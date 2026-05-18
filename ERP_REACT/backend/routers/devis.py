"""
ERP React - Devis Router
Devis/soumissions + lignes + validation publique + HTML generation + public acceptance.
Based on devis.py (13,819 lines) + devis_manager.py (1,557 lines).
"""

import base64
import html as html_mod
import json
import logging
import os
import re
import secrets
import smtplib
import sys
import time

import psycopg2  # pour psycopg2.Binary (BYTEA inserts conversation_documents)
from psycopg2 import sql as psql  # pour schema-qualifier ALTER/UPDATE en _ensure_devis_pct_columns
from datetime import datetime
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.utils import formataddr, formatdate
from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from pydantic import BaseModel, Field, field_validator
from typing import Optional, List

from ..erp_auth import get_current_user, require_role, ErpUser
from .. import erp_database as db
from .ai import check_ai_guard, _check_credits, _deduct_credits, track_ai_usage, _today_prompt_line, _ensure_conversations_last_updated_at, _detect_media_type_from_bytes

# Import Anthropic client for AI features
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
try:
    import anthropic
    _anthropic_client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY", ""))
except ImportError:
    _anthropic_client = None

try:
    import httpx
except ImportError:
    httpx = None

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/devis", tags=["Devis"])


def _call_claude(*, model="claude-opus-4-7", max_tokens=32000, system=None, messages, betas=None, tools=None, tool_choice=None, temperature: float = 0.1, stop_sequences: Optional[list] = None):
    """Call Claude via streaming to support long Opus requests (>10 min).

    Wrapped with retry on transient connection errors. Render's edge proxy or
    Anthropic intermediates can close the chunked stream after ~150s on long
    Opus responses, surfacing as `RemoteProtocolError: peer closed connection
    without sending complete message body`. These are not subclasses of
    `anthropic.APIError`, so they bypass the upstream caller's typed handler
    and turn into generic 500s. One retry is enough to survive the flake.

    OPT-2: si `betas` est fourni (ex: ["files-api-2025-04-14"]), route via
    `client.beta.messages.stream` qui supporte le parametre betas. Sinon,
    route via `client.messages.stream` (chemin standard, retro-compat).

    temperature: default 0.1 pour reduire variance non-deterministe sur les
    chiffres CAD (montants, quantites, prix unitaires). Override possible
    si besoin de creativite (ex: chat libre).

    extra_headers: ajoute systematiquement le beta header
    `extended-cache-ttl-2025-04-11` requis pour cache_control ttl="1h".
    """
    # FIX (2026-05-17): Opus 4.7 a deprecie le parametre `temperature` (le
    # modele utilise du extended thinking implicite qui force temperature=1.0).
    # L'API retourne 400 "temperature is deprecated for this model" si on
    # passe ce param. On le filtre ici pour les modeles qui ne le supportent
    # plus. Les autres modeles (Sonnet, Haiku, legacy) gardent le param.
    _TEMPERATURE_DEPRECATED_PREFIXES = ("claude-opus-4-7",)
    kwargs = {"model": model, "max_tokens": max_tokens, "messages": messages, "timeout": 600.0}
    if not any(model.startswith(p) for p in _TEMPERATURE_DEPRECATED_PREFIXES):
        kwargs["temperature"] = temperature
    elif temperature != 0.1:
        # Log si le caller specifie une valeur != default (info pour audit).
        logger.info(
            "_call_claude: temperature=%s ignore (deprecated par %s)",
            temperature, model,
        )
    if system:
        kwargs["system"] = system
    if stop_sequences:
        kwargs["stop_sequences"] = list(stop_sequences)
    if tools:
        kwargs["tools"] = list(tools)
    if tool_choice:
        kwargs["tool_choice"] = tool_choice

    # Header beta requis pour cache_control TTL 1h (ephemeral 1h au lieu de 5min default)
    kwargs["extra_headers"] = {"anthropic-beta": "extended-cache-ttl-2025-04-11"}

    transient_excs = [anthropic.APIConnectionError]
    if httpx is not None:
        # Couvre toute la famille des erreurs réseau transient httpx :
        # - Read* : lecture interrompue ou timeout côté serveur
        # - Write* : écriture interrompue (rare mais possible)
        # - Connect* : impossible d'établir la connexion (DNS, refused, timeout)
        # - PoolTimeout : épuisement du pool httpx (charge anormale)
        # - RemoteProtocolError : peer closed connection (chunked stream coupé)
        # Toutes méritent une 2e tentative — c'est l'esprit du retry interne.
        transient_excs.extend([
            httpx.RemoteProtocolError,
            httpx.ReadError,
            httpx.ReadTimeout,
            httpx.WriteError,
            httpx.ConnectError,
            httpx.ConnectTimeout,
            httpx.PoolTimeout,
        ])
    transient_excs = tuple(transient_excs)

    # OPT-2: route via client.beta.messages.stream si betas fourni. Le SDK
    # 0.67+ supporte `betas=` sur beta.messages.stream — verifie via inspect
    # au runtime pour eviter erreur si SDK trop ancien.
    if betas:
        kwargs["betas"] = list(betas)
        stream_ctx = lambda: _anthropic_client.beta.messages.stream(**kwargs)
    else:
        stream_ctx = lambda: _anthropic_client.messages.stream(**kwargs)

    last_exc = None
    for attempt in range(2):
        try:
            with stream_ctx() as stream:
                return stream.get_final_message()
        except transient_excs as exc:
            last_exc = exc
            logger.warning(
                "_call_claude transient error on attempt %d: %s — retrying",
                attempt + 1, exc,
            )
            time.sleep(0.5 * (attempt + 1))

    # Both attempts failed. If it's already an anthropic exception, re-raise so
    # callers' typed handlers (`except anthropic.APIError`) match it. Raw httpx
    # errors are NOT subclasses of anthropic.APIError — wrap them so the
    # caller's typed handler still routes them to a 503 rather than letting
    # them fall through to the generic `except Exception` branch (500).
    if isinstance(last_exc, anthropic.APIError):
        raise last_exc
    # Resolve the underlying httpx Request defensively. `httpx.RequestError.request`
    # is a *property* that raises RuntimeError when the request was never set —
    # plain `getattr(..., None)` does NOT swallow that, so we need a real try.
    try:
        req = getattr(last_exc, "request", None)
    except Exception:
        req = None
    try:
        wrapped = anthropic.APIConnectionError(
            message=f"Anthropic connection error: {last_exc}",
            request=req,
        )
    except Exception:
        # SDK signature mismatch — fall back to raising the original error
        raise last_exc
    raise wrapped from last_exc


def _extract_text_from_response(message) -> str:
    """Extrait le texte de la reponse Claude de maniere safe.

    Claude peut retourner plusieurs blocks (thinking, text, tool_use). Le
    premier n'est pas garanti d'etre un text block — si thinking est active
    ou si le SDK change, `message.content[0].text` peut AttributeError.
    Cette fonction iteratr tous les blocks et concatene les text blocks.
    """
    try:
        texts = []
        for block in message.content:
            # Bloc text: attribut .text + type="text"
            btype = getattr(block, "type", None)
            if btype == "text":
                t = getattr(block, "text", None)
                if t:
                    texts.append(t)
        if texts:
            return "\n".join(texts)
        # Pas de fallback non-safe: si Claude retourne uniquement des blocks
        # thinking/tool_use sans text, on retourne chaine vide plutot que de
        # risquer de retourner le contenu d'un thinking block. Le caller gere.
    except Exception as exc:
        logger.warning("_extract_text_from_response failed: %s", exc)
    return ""


def _extract_citations_from_response(message) -> list:
    """Extrait les citations Claude d'une reponse multi-block (Sprint 3 #10).

    Format Anthropic PDF citations:
        block.type == "text"
        block.citations == [{type: page_location, cited_text, document_index,
                             document_title, start_page_number, ...}, ...]

    Returns liste flat de dict normalises.
    """
    citations = []
    try:
        for block in (getattr(message, "content", None) or []):
            if getattr(block, "type", None) != "text":
                continue
            block_text = getattr(block, "text", "") or ""
            raw_cits = getattr(block, "citations", None) or []
            for cit in raw_cits:
                page_start = (
                    getattr(cit, "start_page_number", None)
                    if not isinstance(cit, dict)
                    else cit.get("start_page_number")
                )
                cited_text = (
                    getattr(cit, "cited_text", None)
                    if not isinstance(cit, dict)
                    else cit.get("cited_text")
                ) or ""
                doc_idx = (
                    getattr(cit, "document_index", 0)
                    if not isinstance(cit, dict)
                    else cit.get("document_index", 0)
                )
                doc_title = (
                    getattr(cit, "document_title", None)
                    if not isinstance(cit, dict)
                    else cit.get("document_title")
                )
                citations.append({
                    "text": block_text[:500],
                    "cited_text": (cited_text or "")[:1000],
                    "document_index": doc_idx or 0,
                    "page_number": page_start,
                    "document_title": doc_title,
                })
    except Exception as exc:
        logger.warning("_extract_citations_from_response failed: %s", exc)
    return citations


def _format_response_with_footnotes(message, document_titles=None) -> str:
    """Concatene les text blocks Claude en sortie propre (sans footnotes affichees).

    UX fix (2026-05-17): la section "Sources citees" et les marqueurs [^N] inline
    affichaient des snippets PDF bruts (OCR ratees, cotes melangees, jargon
    "P. 30\" x 80\" 4-8X12 2'-04..."). Le bruit confond les utilisateurs sans
    apporter de valeur — Claude cite deja explicitement les pages dans son texte
    (ex: "page A101 dit \"REZ-DE-CHAUSSEE = 1467,4 pi2\"").

    On garde le parsing pour ne pas casser citations_list (utilise par
    _extract_citations_from_response pour audit/billing), mais on n'injecte
    ni les [^N] dans le texte affiche ni la section bibliographique.
    """
    if not message or not getattr(message, "content", None):
        return ""

    parts = []
    try:
        for block in message.content:
            if getattr(block, "type", None) != "text":
                continue
            block_text = getattr(block, "text", "") or ""
            # Ignore les citations: on ne les injecte plus dans le texte affiche
            # (citations_list reste disponible via _extract_citations_from_response)
            parts.append(block_text)
    except Exception as exc:
        logger.warning("_format_response_with_footnotes parse failed: %s", exc)

    return "\n".join(parts).strip()


def _build_estimate_pass1_observation_prompt(devis_data: dict, additional_context: str = "") -> str:
    """Prompt Pass 1 multipass: observation pure du plan (vision exhaustive).

    Sprint 3 #5 - Decouple Reading vs Interpretation pour reduire les
    hallucinations Vision (cf memoire multipass_pattern.md).
    """
    ctx_extra = f"\n\nContexte utilisateur:\n{additional_context}" if additional_context else ""
    nom_projet = devis_data.get("nom_projet", "N/A") if devis_data else "N/A"
    return (
        f"TACHE: SCAN EXHAUSTIF DU PLAN. Tu observes le plan suivant pour "
        f"projet '{nom_projet}'.{ctx_extra}\n\n"
        f"REGLES STRICTES:\n"
        f"1. LIS CHAQUE ANNOTATION LITTERALEMENT (mot-pour-mot, sans interpretation)\n"
        f"2. NE PAS INVENTER de dimensions/quantites non visibles\n"
        f"3. NE PAS faire d'estimation de couts (Pass 2 s'en chargera)\n"
        f"4. NE PAS donner de recommandations (Pass 2 s'en chargera)\n"
        f"5. Pour chaque element observe, fournir page_number et snippet textuel\n\n"
        f"Reponds UNIQUEMENT en JSON strict (pas de markdown autour):\n"
        f"{{\n"
        f'  "zones_detectees": [\n'
        f'    {{"nom": "BUREAU PRINCIPAL", "dimensions": "12pi x 10pi", "page": 1, "snippet": "..."}}\n'
        f'  ],\n'
        f'  "elements_techniques": [\n'
        f'    {{"type": "porte", "specs": "30pox80po", "quantite": 1, "page": 1, "snippet": "..."}},\n'
        f'    {{"type": "fenetre", "specs": "PVC 36x48", "quantite": 2, "page": 1, "snippet": "..."}}\n'
        f'  ],\n'
        f'  "materiaux_indiques": [\n'
        f'    {{"description": "Gypse 5/8 type X", "page": 2, "snippet": "..."}}\n'
        f'  ],\n'
        f'  "surfaces_pi2": {{"total_plancher": 240, "total_murs": 460}},\n'
        f'  "annotations_techniques": [\n'
        f'    {{"texte": "CONFORMITE RBQ", "page": 1, "snippet": "..."}}\n'
        f'  ],\n'
        f'  "summary": "Plan d\'agrandissement 12x20 avec 1 bureau + 1 salle"\n'
        f"}}"
    )


def _parse_estimate_pass1_json(response_text: str) -> dict:
    """Parse robuste du JSON Pass 1. Fallback en HTTPException 502 si invalide."""
    if not response_text or not response_text.strip():
        raise HTTPException(status_code=502, detail="Reponse IA Pass 1 vide")
    text = response_text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        text = "\n".join(lines)
    try:
        data = json.loads(text)
    except json.JSONDecodeError as exc:
        import re as _re
        m = _re.search(r"\{[\s\S]*\}", text)
        if m:
            try:
                data = json.loads(m.group(0))
            except json.JSONDecodeError:
                logger.warning("Pass 1 JSON parse failed: %s. Preview=%r", exc, text[:200])
                raise HTTPException(
                    status_code=502,
                    detail="Reponse IA Pass 1 non parsable (JSON invalide)",
                )
        else:
            logger.warning("Pass 1 JSON parse failed: %s. Preview=%r", exc, text[:200])
            raise HTTPException(
                status_code=502,
                detail="Reponse IA Pass 1 non parsable",
            )
    if not isinstance(data, dict):
        raise HTTPException(status_code=502, detail="Pass 1 JSON doit etre un object")
    return data


def _build_estimate_pass2_analysis_prompt(
    pass1_data: dict,
    devis_data: dict,
    lignes_text: str,
    citations_list_pass1: list,
    additional_context: str = "",
) -> str:
    """Prompt Pass 2 multipass: analyse business + recommandations sur la base Pass 1.

    Sprint 3 #5 - Aucune vision dans ce Pass, juste raisonnement sur le JSON
    structure du Pass 1. Evite contamination visuelle, focus sur calcul.
    """
    pass1_json_str = json.dumps(pass1_data, ensure_ascii=False, indent=2, default=str)
    ctx_extra = f"\n\nContexte utilisateur supplementaire:\n{additional_context}" if additional_context else ""
    nom_projet = devis_data.get("nom_projet", "N/A") if devis_data else "N/A"
    description = devis_data.get("description", "Aucune description") if devis_data else "N/A"

    citations_hint = ""
    if citations_list_pass1:
        pages_seen = sorted(set(
            c.get("page_number") for c in citations_list_pass1
            if c.get("page_number") is not None
        ))
        if pages_seen:
            citations_hint = (
                f"\nReferences plan (citations PDF du Pass 1): "
                f"pages {', '.join(str(p) for p in pages_seen)}. "
                f"Lorsque tu mentionnes un element du plan, indique entre "
                f"parentheses la page (ex: 'voir page 1')."
            )

    return (
        f"Tu es estimateur professionnel construction Quebec. Un scan exhaustif "
        f"du plan a deja ete fait (Pass 1 ci-dessous). Ton role: analyser et estimer.\n\n"
        f"PROJET: {nom_projet}\n"
        f"Description: {description}\n\n"
        f"LIGNES EXISTANTES DU DEVIS:\n{lignes_text}\n\n"
        f"DONNEES DU PLAN (Pass 1, scan exhaustif):\n{pass1_json_str}{ctx_extra}\n\n"
        f"LIVRABLES (en markdown structure):\n"
        f"1. VERIFICATION DES LIGNES existantes (sont-elles coherentes avec le plan?)\n"
        f"2. ITEMS MANQUANTS evidents (selon les elements observes au Pass 1)\n"
        f"3. RECOMMANDATIONS concretes (prix CAD 2026 Quebec, optimisations)\n"
        f"4. ESTIMATION GLOBALE: fourchette de prix totale TTC + ratio approximatif $/pi2{citations_hint}\n\n"
        f"Format: markdown structure, en francais Quebec, avec tableaux pour comparaisons prix. "
        f"NE PAS reprendre integralement le JSON Pass 1, juste y faire reference."
    )


def _run_estimate_with_plan_multipass(
    pdf_b64: str,
    media_type: str,
    devis_data: dict,
    lignes_text: str,
    additional_context: str,
    filename: str,
    is_pdf: bool,
    system_blocks: list,
    precision_mode: bool,
) -> dict:
    """Orchestre Pass 1 (vision pure JSON) + Pass 2 (analyse texte markdown).

    Sprint 3 #5 MVP multipass anti-hallucination.

    Returns dict avec:
      - response_text: markdown final (Pass 2)
      - citations_list: collectees depuis Pass 1 (PDF only)
      - tokens_in, tokens_out: cumules sur les 2 passes
      - cache_creation, cache_read: cumules
      - thinking_tokens: cumules
      - pass1_summary: brief du JSON Pass 1 pour metadata
    """
    if not _anthropic_client:
        raise HTTPException(status_code=503, detail="Service IA non configure")

    # PASS 1: Vision pure - observation exhaustive en JSON
    pass1_prompt = _build_estimate_pass1_observation_prompt(devis_data, additional_context)
    pass1_plan_block = {
        "type": "document" if is_pdf else "image",
        "source": {"type": "base64", "media_type": media_type, "data": pdf_b64},
        "cache_control": {"type": "ephemeral"},
    }
    if is_pdf:
        pass1_plan_block["citations"] = {"enabled": True}
    pass1_user_content = [
        pass1_plan_block,
        {"type": "text", "text": pass1_prompt},
    ]

    try:
        pass1_message = _call_claude_with_thinking(
            model="claude-opus-4-7",
            thinking_budget=10000,
            max_response_tokens=60000,
            system=system_blocks,
            messages=[{"role": "user", "content": pass1_user_content}],
            effort="high",
        )
    except Exception as exc:
        logger.exception("Multipass Pass 1 Claude call failed: %s", exc)
        raise HTTPException(status_code=502, detail="Erreur Pass 1 (observation plan)")

    pass1_text = _extract_text_blocks(pass1_message) or ""
    citations_list_pass1 = _extract_citations_from_response(pass1_message) if is_pdf else []
    pass1_in = getattr(pass1_message.usage, "input_tokens", 0) or 0
    pass1_out = getattr(pass1_message.usage, "output_tokens", 0) or 0
    pass1_cache_create = getattr(pass1_message.usage, "cache_creation_input_tokens", 0) or 0
    pass1_cache_read = getattr(pass1_message.usage, "cache_read_input_tokens", 0) or 0
    pass1_thinking = _count_thinking_tokens_safe(pass1_message)

    try:
        pass1_data = _parse_estimate_pass1_json(pass1_text)
    except HTTPException:
        # Pass 1 invalide - on remonte
        raise

    # PASS 2: Texte pur - analyse markdown avec contexte Pass 1
    pass2_prompt = _build_estimate_pass2_analysis_prompt(
        pass1_data=pass1_data,
        devis_data=devis_data,
        lignes_text=lignes_text,
        citations_list_pass1=citations_list_pass1,
        additional_context=additional_context,
    )

    try:
        if precision_mode:
            pass2_message = _call_claude_with_thinking(
                model="claude-opus-4-7",
                thinking_budget=10000,
                max_response_tokens=80000,
                system=system_blocks,
                messages=[{"role": "user", "content": pass2_prompt}],
                effort="high",
            )
        else:
            pass2_message = _call_claude(
                model="claude-opus-4-7",
                max_tokens=32000,
                system=system_blocks,
                messages=[{"role": "user", "content": pass2_prompt}],
            )
    except Exception as exc:
        logger.exception("Multipass Pass 2 Claude call failed: %s", exc)
        raise HTTPException(status_code=502, detail="Erreur Pass 2 (analyse business)")

    if precision_mode:
        response_text = _extract_text_blocks(pass2_message) or ""
        pass2_thinking = _count_thinking_tokens_safe(pass2_message)
    else:
        response_text = _extract_text_from_response(pass2_message) or ""
        pass2_thinking = 0

    pass2_in = getattr(pass2_message.usage, "input_tokens", 0) or 0
    pass2_out = getattr(pass2_message.usage, "output_tokens", 0) or 0
    pass2_cache_create = getattr(pass2_message.usage, "cache_creation_input_tokens", 0) or 0
    pass2_cache_read = getattr(pass2_message.usage, "cache_read_input_tokens", 0) or 0

    # UX fix (2026-05-17): plus d'injection de section "Sources citees" en fin
    # de markdown. Les snippets OCR bruts confondaient les utilisateurs (cotes
    # melangees, jargon technique). citations_list_pass1 reste disponible pour
    # audit interne mais n'est plus affichee.

    return {
        "response_text": response_text,
        "citations_list": citations_list_pass1,
        "tokens_in": pass1_in + pass2_in,
        "tokens_out": pass1_out + pass2_out,
        "cache_creation": pass1_cache_create + pass2_cache_create,
        "cache_read": pass1_cache_read + pass2_cache_read,
        "thinking_tokens": pass1_thinking + pass2_thinking,
        "pass1_summary": (pass1_data.get("summary") or "")[:500],
        "pass1_zones_count": len(pass1_data.get("zones_detectees", []) or []),
        "pass1_elements_count": len(pass1_data.get("elements_techniques", []) or []),
    }


def _extract_tool_use_input(message, tool_name: str):
    """Extrait le bloc tool_use du nom donne et retourne son .input dict.

    Avec tool_choice force, Claude doit produire un tool_use de nom X.
    Si manquant (refus safety ou troncature max_tokens), retourne None.
    Le caller doit alors raise une HTTPException 502.
    """
    if not message or not getattr(message, "content", None):
        return None
    for block in message.content:
        if getattr(block, "type", None) == "tool_use" and getattr(block, "name", None) == tool_name:
            return getattr(block, "input", None)
    return None


def _lookup_catalogue_produits(
    schema: str,
    search: str,
    categorie: Optional[str] = None,
    limit: int = 20,
) -> list[dict]:
    """Recherche dans le catalogue produits du tenant.

    Utilise par l'outil Claude `recherche_catalogue_produits` pour grounder
    les prix de l'estimation sur les produits reels du tenant au lieu
    d'halluciner des tarifs marche.

    SECURITE TENANT: requiert `schema` non-None et set_tenant() obligatoire.
    NE JAMAIS appeler cette fonction sans contexte tenant valide.
    """
    if not schema:
        logger.warning("_lookup_catalogue_produits: schema manquant - refus")
        return []
    if not search or len(search.strip()) < 2:
        return []
    safe_limit = max(1, min(int(limit or 20), 50))
    search_pattern = f"%{search.strip().lower()}%"

    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, schema)
        # Defense en profondeur: transaction READ ONLY + statement_timeout 10s.
        # Aligne avec _execute_tenant_select (ai.py) — protection DoS
        # (table corrompue ou geante ne peut pas monopoliser la connexion pool)
        # + empeche toute mutation accidentelle malgre que la query soit figee.
        conn.autocommit = False
        cursor = conn.cursor()
        cursor.execute("SET TRANSACTION READ ONLY")
        cursor.execute("SET LOCAL statement_timeout TO '10000'")

        cursor.execute(
            "SELECT EXISTS (SELECT 1 FROM information_schema.tables "
            "WHERE table_schema = current_schema() AND table_name = 'produits') AS ex"
        )
        row = cursor.fetchone()
        # RealDictCursor garantit toujours un dict — simplification cosmetique
        if not row or not row.get("ex"):
            return []

        where_clauses = [
            "(LOWER(COALESCE(nom,'')) LIKE %s "
            "OR LOWER(COALESCE(description,'')) LIKE %s "
            "OR LOWER(COALESCE(code_produit,'')) LIKE %s)"
        ]
        params: list = [search_pattern, search_pattern, search_pattern]

        if categorie and categorie.strip():
            where_clauses.append("LOWER(COALESCE(categorie,'')) = %s")
            params.append(categorie.strip().lower())

        where_clauses.append("COALESCE(active, TRUE) = TRUE")
        where_sql = " AND ".join(where_clauses)

        sql_main = (
            "SELECT id, "
            "COALESCE(code_produit, '') AS code_produit, "
            "COALESCE(nom, '') AS nom, "
            "COALESCE(description, '') AS description, "
            "COALESCE(categorie, '') AS categorie, "
            "COALESCE(unite_vente, '') AS unite_vente, "
            "COALESCE(prix_unitaire, 0)::float AS prix_unitaire, "
            "COALESCE(stock_disponible, 0)::float AS stock_disponible "
            f"FROM produits WHERE {where_sql} "
            "ORDER BY nom ASC LIMIT %s"
        )
        params_full = params + [safe_limit]

        try:
            cursor.execute(sql_main, params_full)
            rows = cursor.fetchall()
        except Exception as exc:
            logger.info(
                "_lookup_catalogue_produits: fallback vers colonne 'unite' (schema=%s): %s",
                schema, exc,
            )
            try:
                conn.rollback()
            except Exception:
                pass
            sql_fallback = (
                "SELECT id, "
                "COALESCE(code_produit, '') AS code_produit, "
                "COALESCE(nom, '') AS nom, "
                "COALESCE(description, '') AS description, "
                "COALESCE(categorie, '') AS categorie, "
                "COALESCE(unite, '') AS unite_vente, "
                "COALESCE(prix_unitaire, 0)::float AS prix_unitaire, "
                "COALESCE(stock_disponible, 0)::float AS stock_disponible "
                f"FROM produits WHERE {where_sql} "
                "ORDER BY nom ASC LIMIT %s"
            )
            cursor.execute(sql_fallback, params_full)
            rows = cursor.fetchall()

        results = []
        for r in rows:
            d = dict(r)
            results.append({
                "id": d.get("id"),
                "code_produit": d.get("code_produit") or "",
                "nom": d.get("nom") or "",
                "description": d.get("description") or "",
                "categorie": d.get("categorie") or "",
                "unite_vente": d.get("unite_vente") or "",
                "prix_unitaire": float(d.get("prix_unitaire") or 0),
                "stock_disponible": float(d.get("stock_disponible") or 0),
            })

        logger.info(
            "AI catalog search schema=%s search=%r cat=%r -> %d rows",
            schema, search, categorie, len(results),
        )
        return results
    except Exception as exc:
        logger.warning(
            "_lookup_catalogue_produits failed (schema=%s, search=%r): %s",
            schema, search, exc,
        )
        return []
    finally:
        # Rollback explicite pour fermer la transaction READ ONLY proprement
        try:
            conn.rollback()
        except Exception:
            pass
        if cursor:
            try:
                cursor.close()
            except Exception:
                pass
        # Restore autocommit avant de retourner la connexion au pool
        try:
            conn.autocommit = True
        except Exception:
            pass
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        try:
            conn.close()
        except Exception:
            pass


def _extract_tool_use_blocks(message) -> list:
    """Extrait tous les blocs tool_use d'une reponse Claude.

    Complement de `_extract_tool_use_input` (singulier). Pour la boucle
    multi-tours d'ai_estimate_devis, Claude peut emettre PLUSIEURS tool_use
    dans la meme reponse - on doit tous les executer avant de renvoyer
    le batch de tool_result.

    Returns: liste de dicts {id, name, input}. Vide si aucun.
    """
    if not message or not getattr(message, "content", None):
        return []
    blocks = []
    for block in message.content:
        if getattr(block, "type", None) == "tool_use":
            blocks.append({
                "id": getattr(block, "id", "") or "",
                "name": getattr(block, "name", "") or "",
                "input": getattr(block, "input", {}) or {},
            })
    return blocks


_ESTIMATION_TOOL_NAMES = {"recherche_catalogue_produits"}


def _execute_estimation_tool(tool_name: str, tool_input: dict, schema: str) -> dict:
    """Execute un outil d'estimation et retourne le resultat serialisable JSON.

    Whitelist stricte: refuse tout tool_name absent de _ESTIMATION_TOOL_NAMES
    pour eviter qu'un patch futur ajoute un outil sensible sans audit.
    """
    if tool_name not in _ESTIMATION_TOOL_NAMES:
        logger.warning("_execute_estimation_tool: outil non whiteliste: %s", tool_name)
        return {"error": f"Outil inconnu ou non autorise: {tool_name}"}

    if not schema:
        return {"error": "Contexte tenant manquant - outil indisponible"}

    try:
        if tool_name == "recherche_catalogue_produits":
            search = (tool_input or {}).get("search", "")
            categorie = (tool_input or {}).get("categorie")
            limit = (tool_input or {}).get("limit", 20)
            produits = _lookup_catalogue_produits(
                schema=schema,
                search=search,
                categorie=categorie,
                limit=limit,
            )
            return {
                "result": {
                    "count": len(produits),
                    "produits": produits,
                    "search_used": search,
                    "categorie_used": categorie,
                }
            }
        return {"error": f"Dispatcher non implemente pour: {tool_name}"}
    except Exception as exc:
        logger.error(
            "_execute_estimation_tool error tool=%s: %s",
            tool_name, exc,
        )
        return {"error": f"Erreur execution outil: {exc}"}


def _call_claude_with_thinking(
    *,
    model: str = "claude-opus-4-7",
    thinking_budget: int = 10000,  # DEPRECATED on Opus 4.7 - ignored, kept for backward compat
    max_response_tokens: int = 100000,
    system: Optional[str] = None,
    messages: list,
    effort: str = "high",  # adaptive thinking effort: low / medium / high (default) / xhigh / max
    temperature: float = 0.1,
    tools: Optional[list] = None,
    tool_choice: Optional[dict] = None,
):
    """Call Claude with Adaptive Thinking on Opus 4.7.

    Sur Claude Opus 4.7, manual thinking (`{"type": "enabled", "budget_tokens": N}`)
    est REJETE par l'API avec 400 invalid_request_error. La syntaxe officielle est
    `thinking={"type": "adaptive"}` + `output_config={"effort": "..."}`.
    Cf. https://platform.claude.com/docs/en/build-with-claude/adaptive-thinking

    EFFORT par defaut "high" = comportement equivalent a l'ancien `enabled` +
    `budget_tokens=10000`: d'apres la doc Anthropic "At high and max effort levels,
    Claude almost always thinks. Provides deep reasoning on complex tasks." C'est
    le drop-in pour les analyses de plans / estimations multi-section.

    Pour des taches plus exigeantes (Vision plans manuscrits, agentic complexe),
    passer effort="xhigh" (consomme plus de tokens mais raisonnement plus profond).

    Le parametre `thinking_budget` est conserve dans la signature pour ne pas
    casser les callers existants qui le passent (devis.py:4532, devis.py:4803,
    metre_pdf.py:4798), mais n'a plus aucun effet sur Opus 4.7. Le caller
    Pass 2 (metre_pdf.py:4829) a deja ete migre vers `effort="medium"`.

    Thinking tokens sont factures comme output tokens (deja inclus dans
    usage.output_tokens - NE PAS additionner separement, cf. Round 22A).

    STREAMING REQUIRED: cf. commit 4a952f3.
    """
    if not _anthropic_client:
        raise RuntimeError("Anthropic client not initialized")
    # max_tokens fixed at 32000 (decision Sylvain 2026-05-17): uniformite
    # entre tous les call sites. max_response_tokens reste accepte dans la
    # signature pour compatibilite caller mais est ignore ici (cap hard 32k
    # couvre thinking + response COMBINES sur Opus 4.7).
    max_tokens = 32000
    kwargs = {
        "model": model,
        "max_tokens": max_tokens,
        "thinking": {"type": "adaptive"},
        "extra_body": {"output_config": {"effort": effort}},
        "messages": messages,
        "timeout": 600.0,
    }
    # NOTE: Adaptive thinking requiert temperature=1.0 sur Anthropic API.
    # Le param `temperature` est accepte dans la signature pour uniformite avec
    # _call_claude mais IGNORE ici si thinking est actif. Si caller passe une
    # valeur != 1.0 et != 0.1 (default), on log un warning.
    if temperature not in (0.1, 1.0):
        logger.info(
            "_call_claude_with_thinking: temperature=%s requested but adaptive thinking "
            "requires temperature=1.0 — using default 1.0",
            temperature,
        )
    if system is not None:
        kwargs["system"] = system
    if tools:
        kwargs["tools"] = list(tools)
    if tool_choice:
        kwargs["tool_choice"] = tool_choice
    # Header beta requis pour cache_control TTL 1h
    kwargs["extra_headers"] = {"anthropic-beta": "extended-cache-ttl-2025-04-11"}
    with _anthropic_client.messages.stream(**kwargs) as stream:
        return stream.get_final_message()


def _extract_text_blocks(message) -> str:
    """Extract only 'text' blocks from a Claude response (skip 'thinking' blocks).

    Avec Extended Thinking enabled, content peut inclure des blocks
    'thinking' (raisonnement interne) en plus des blocks 'text' (la
    reponse finale). On ne veut que la reponse texte pour l'affichage.
    """
    text = ""
    for block in message.content:
        if getattr(block, 'type', None) == 'text' and hasattr(block, 'text'):
            text += block.text
    return text


def _count_thinking_tokens_safe(message) -> int:
    """Estimate thinking tokens by inspecting content blocks (informational).

    Le SDK Anthropic n'expose PAS `usage.thinking_tokens` comme champ separe -
    les thinking tokens sont bundles dans `output_tokens`. La seule facon
    fiable de detecter si Extended Thinking a vraiment tourne est de compter
    les caracteres dans les blocks de type 'thinking' du content.

    Returns 0 si aucun block thinking present. NE PAS additionner au cost
    calc - les thinking tokens sont deja dans usage.output_tokens (cf.
    Round 22A lesson).

    Note: estimation 1 token = 4 chars (approximation, Claude tokenize plus
    finement) mais permet de distinguer "thinking absent" de "thinking
    present" avec un ordre de grandeur correct.
    """
    if not message or not getattr(message, 'content', None):
        return 0
    total_chars = 0
    for block in message.content:
        if getattr(block, 'type', None) == 'thinking':
            text = getattr(block, 'thinking', '') or ''
            total_chars += len(text)
    return total_chars // 4 if total_chars > 0 else 0


DEVIS_STATUSES = ["Brouillon", "Valide", "Envoye", "En attente", "Accepte", "Refuse", "Termine", "Annule", "Expire"]

# Basic email format check — rejects obvious junk like "aaaaa" before we
# update devis.statut='Envoye' and attempt an SMTP send that's guaranteed to fail.
_EMAIL_RE = re.compile(r"^[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}$")

# Default percentages when columns don't exist yet in DB
_DEFAULT_ADM_PCT = 3.0
_DEFAULT_CON_PCT = 12.0
_DEFAULT_PRO_PCT = 15.0

# Columns that devis→project conversion paths may INSERT into the `projects`
# table. Old tenants are missing some of these (migration gap: CREATE TABLE IF
# NOT EXISTS does not add columns to pre-existing tables). Apply defensive
# ADD COLUMN IF NOT EXISTS before any INSERT to prevent UndefinedColumn errors.
_PROJECTS_INSERT_DEFENSIVE_COLS = [
    ("client_company_id", "INTEGER"),
    ("client_contact_id", "INTEGER"),
    ("client_nom_cache", "TEXT"),
    ("po_client", "TEXT"),
    ("priorite", "TEXT"),
    ("type_projet", "TEXT"),
    ("tache", "TEXT"),
    ("budget_total", "NUMERIC(14,2)"),
    ("date_debut_reel", "DATE"),
    ("date_fin_reel", "DATE"),
    ("date_soumis", "DATE"),
    ("date_prevu", "DATE"),
    ("bd_ft_estime", "NUMERIC"),
    ("prix_estime", "NUMERIC"),
    ("description", "TEXT"),
    ("devis_id", "INTEGER"),
    ("devis_source_id", "INTEGER"),
    ("numero_devis", "TEXT"),
]


# Memoization cache for _ensure_projects_insert_columns. Only written on FULL
# success (all 18 ALTERs completed without exception) to avoid poisoning the
# cache if a transient ALTER failure would otherwise mask a missing column
# permanently for the lifetime of the process. Thread-safe under asyncio
# (single-threaded event loop per uvicorn worker); would need locking if moved
# to a threaded ASGI server.
_projects_cols_ensured_for: set = set()


def _ensure_projects_insert_columns(cursor, conn, schema: str) -> None:
    """Defensively ADD COLUMN IF NOT EXISTS for columns inserted by
    _create_project_from_devis / convert_devis_to_project / accept_public_devis.
    Memoized per tenant schema — subsequent calls are free no-ops. Caller must
    be in autocommit mode (the default for `db.get_conn()` connections)."""
    if schema and schema in _projects_cols_ensured_for:
        return
    all_succeeded = True
    for col, ctype in _PROJECTS_INSERT_DEFENSIVE_COLS:
        try:
            cursor.execute(
                f"ALTER TABLE projects ADD COLUMN IF NOT EXISTS {col} {ctype}"
            )
        except Exception as exc:
            all_succeeded = False
            logger.warning("ALTER projects ADD %s failed: %s", col, exc)
            try:
                conn.rollback()
                if schema:
                    db.set_tenant(conn, schema)
            except Exception:
                pass
    if schema and all_succeeded:
        _projects_cols_ensured_for.add(schema)


# ============================================
# PUBLIC TOKEN HELPERS
# ============================================

import re
import unicodedata


def _slugify(text: str, max_length: int = 30) -> str:
    """Convert text to URL-friendly slug."""
    text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('ascii')
    text = re.sub(r'[^\w\s-]', '', text.lower())
    text = re.sub(r'[-\s]+', '-', text).strip('-')
    return text[:max_length]


def _generate_readable_token(nom_projet: str) -> str:
    """Generate a human-readable token like 'garage-dupont-2025-a3f9'."""
    slug = _slugify(nom_projet, 30)
    year = datetime.now().year
    unique = secrets.token_urlsafe(6)
    return f"{slug}-{year}-{unique}" if slug else f"devis-{year}-{unique}"


# ============================================
# MO / MAT RATIOS — Quebec construction industry
# Sources: CCQ conventions collectives, APCHQ guides de couts,
#          SCHL/CMHC residential cost data, RS Means Canada
# Format: (keywords, mo_pct, mat_pct)
# ============================================

_MO_MAT_RULES = [
    # Pure labor lines emitted by the Métré module (e.g. "Main-d'œuvre — Charpentier").
    # Listed FIRST so the keyword match wins before a trade keyword (e.g. "charpentier")
    # downgrades the line to a 45/55 charpente split. Description is normalised by
    # `_normalize_for_match` which strips accents and lowercases (œ → oe).
    (["main-d'oeuvre", "main d'oeuvre", "main-doeuvre", "maindoeuvre"], 100, 0),
    # Peinture — labor-dominant par excellence (CCQ)
    (["peinture", "teinture", "vernis", "laque"], 70, 30),
    # Demolition / deconstruction — labor + conteneurs/disposition
    (["demolition", "demontage", "deconstruction"], 65, 35),
    # Gypse / platrage — tirage de joints tres labor-intensive
    (["gypse", "platrage", "platre", "tirage de joint"], 60, 40),
    # Electricite (CCQ) — taux CCQ parmi les plus eleves
    (["electricite", "electrique", "cablage", "eclairage", "panneau electrique", "filage"], 55, 45),
    # Ceramique / carrelage — pose labor-intensive
    (["ceramique", "carrelage", "tuile"], 55, 45),
    # Maconnerie — hautement labor-intensive (brique, pierre, mortier)
    (["maconnerie", "brique", "pierre naturelle", "bloc de beton"], 55, 45),
    # Metaux ouvres / soudure — metier specialise
    (["soudure", "metal ouvre", "acier", "fer forge"], 55, 45),
    # Finitions interieures (hors portes/fenetres)
    (["finition", "finitions interieures"], 55, 45),
    # Plomberie — equilibre (appareils sanitaires vs pose)
    (["plomberie", "tuyauterie", "drain", "robinet"], 50, 50),
    # Coffrage
    (["coffrage"], 50, 50),
    # Revetement exterieur
    (["revetement exterieur", "bardage", "parement", "canexel", "vinyle"], 45, 55),
    # Charpente / structure bois — bois d'oeuvre couteux
    (["charpente", "ossature", "structure bois", "colombage"], 45, 55),
    # Toiture / couverture — membranes, bardeaux, isolant rigide
    (["toiture", "couverture", "bardeaux", "membrane", "toit"], 45, 55),
    # Fondations / beton — materiaux lourds (beton, armature, coffrage)
    (["beton", "fondation", "dalle", "structure portante", "semelle", "pilier"], 40, 60),
    # CVAC — equipements couteux (thermopompes, unites)
    (["cvac", "cvca", "chauffage", "ventilation", "climatisation", "thermopompe", "echangeur"], 40, 60),
    # Isolation / enveloppe — materiau-dominant
    (["isolation", "enveloppe thermique", "laine", "styromousse", "urethan", "pare-vapeur"], 35, 65),
    # Amenagement paysager — pave, terre, vegetaux, murets
    (["amenagement paysager", "paysagement", "plantation", "pave"], 35, 65),
    # Excavation / terrassement — equipement + materiaux (gravier, remblai)
    (["excavation", "terrassement", "deblai", "remblai"], 30, 70),
    # Armoires / ebenisterie — fabrication + materiaux couteux
    (["armoire", "ebenisterie", "cuisine", "vanite", "comptoir"], 30, 70),
    # Portes et fenetres — produits couteux
    (["porte", "fenetre", "vitrerie", "vitrage", "moustiquaire"], 30, 70),
]
_DEFAULT_MO_MAT = (50, 50)


def _normalize_for_match(text: str) -> str:
    """Remove accents and lowercase for keyword matching.

    Important: NFKD does NOT decompose the œ ligature (U+0153) nor æ (U+00E6)
    — these are atomic Unicode codepoints, not combining sequences. We expand
    them explicitly so keywords like "main-d'oeuvre" can match descriptions
    that use the typographically correct œ (e.g. "Main-d'œuvre — Charpentier"
    emitted by the Métré module).

    Apostrophe variants are also normalised to ASCII straight apostrophe (U+0027)
    so a description edited via Word/macOS autocorrect (which would convert
    "main-d'oeuvre" → "main-d’oeuvre") still matches our keywords.
    """
    pre = (text.lower()
        .replace("œ", "oe")  # œ ligature
        .replace("æ", "ae")  # æ ligature
        .replace("’", "'")    # right single quotation mark
        .replace("ʼ", "'")    # modifier letter apostrophe
        .replace("‘", "'"))   # left single quotation mark
    nfkd = unicodedata.normalize("NFKD", pre)
    return "".join(c for c in nfkd if not unicodedata.combining(c))


def _get_mo_mat_ratio(description: str) -> tuple:
    """Return (mo_pct, mat_pct) based on trade keywords in description."""
    desc = _normalize_for_match(description)
    for keywords, mo, mat in _MO_MAT_RULES:
        for kw in keywords:
            if kw in desc:
                return (mo, mat)
    return _DEFAULT_MO_MAT


_public_tokens_table_ensured = False


def _ensure_public_tokens_table(conn):
    """Create devis_public_tokens table in public schema if it doesn't exist. Idempotent, runs once per process."""
    global _public_tokens_table_ensured
    if _public_tokens_table_ensured:
        return
    prev_autocommit = conn.autocommit
    try:
        conn.autocommit = True
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS public.devis_public_tokens (
                token TEXT PRIMARY KEY,
                tenant_schema TEXT NOT NULL,
                devis_id INTEGER NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                expires_at TIMESTAMP
            )
        """)
        cursor.execute("""
            CREATE INDEX IF NOT EXISTS idx_dpt_tenant_devis
            ON public.devis_public_tokens(tenant_schema, devis_id)
        """)
        cursor.close()
        _public_tokens_table_ensured = True
    except Exception:
        pass
    finally:
        conn.autocommit = prev_autocommit


def _register_public_token(conn, token: str, tenant_schema: str, devis_id: int, expires_days: int = 90):
    """Register a token in the public lookup table with expiration (default 90 days)."""
    _ensure_public_tokens_table(conn)
    prev_autocommit = conn.autocommit
    try:
        conn.autocommit = True
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO public.devis_public_tokens (token, tenant_schema, devis_id, expires_at) "
            "VALUES (%s, %s, %s, CURRENT_TIMESTAMP + make_interval(days => %s)) "
            "ON CONFLICT (token) DO UPDATE SET "
            "tenant_schema = EXCLUDED.tenant_schema, devis_id = EXCLUDED.devis_id, "
            "expires_at = EXCLUDED.expires_at",
            (token, tenant_schema, devis_id, expires_days),
        )
        cursor.close()
    except Exception:
        pass
    finally:
        conn.autocommit = prev_autocommit


def _lookup_token(conn, token: str):
    """Look up a token from the public tokens table. Returns (tenant_schema, devis_id) or None."""
    _ensure_public_tokens_table(conn)
    cursor = conn.cursor()
    try:
        cursor.execute(
            "SELECT tenant_schema, devis_id FROM public.devis_public_tokens "
            "WHERE token = %s AND (expires_at IS NULL OR expires_at > CURRENT_TIMESTAMP)",
            (token,),
        )
        row = cursor.fetchone()
        if row:
            return row["tenant_schema"], row["devis_id"]
        return None
    except Exception:
        return None
    finally:
        cursor.close()


def _find_devis_by_token_fallback(conn, token: str):
    """Legacy fallback: scan all tenant schemas for a devis with this validation_token."""
    cursor = conn.cursor()
    try:
        cursor.execute(
            "SELECT schema_name FROM information_schema.schemata WHERE schema_name LIKE 'tenant_%%'"
        )
        schemas = [r["schema_name"] for r in cursor.fetchall()]
    finally:
        cursor.close()

    for schema in schemas:
        cur = None
        try:
            db.set_tenant(conn, schema)
            cur = conn.cursor()
            cur.execute("SELECT id FROM devis WHERE validation_token = %s", (token,))
            row = cur.fetchone()
            if row:
                cur.close()
                cur = None
                # Backfill the public tokens table for next time
                _register_public_token(conn, token, schema, row["id"])
                return schema, row["id"]
        except Exception:
            continue
        finally:
            if cur:
                try:
                    cur.close()
                except Exception:
                    pass
    return None


_ai_profiles_ensured_for: set = set()


def _ensure_ai_profiles_tables(conn, schema: str = ""):
    """Create ai_profiles + ai_profile_documents tables if missing. Per-tenant."""
    if schema in _ai_profiles_ensured_for:
        return
    prev_autocommit = conn.autocommit
    try:
        conn.autocommit = True
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS ai_profiles (
                id SERIAL PRIMARY KEY,
                name TEXT NOT NULL,
                instructions TEXT NOT NULL DEFAULT '',
                is_active BOOLEAN DEFAULT TRUE,
                created_by INTEGER,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS ai_profile_documents (
                id SERIAL PRIMARY KEY,
                profile_id INTEGER NOT NULL REFERENCES ai_profiles(id) ON DELETE CASCADE,
                original_name TEXT NOT NULL,
                content_type TEXT,
                file_size INTEGER,
                extracted_text TEXT,
                uploaded_by INTEGER,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        cursor.close()
        _ai_profiles_ensured_for.add(schema)
    except Exception as exc:
        logger.warning("_ensure_ai_profiles_tables: %s", exc)
    finally:
        conn.autocommit = prev_autocommit


_conversation_docs_ensured_for: set = set()

def _ensure_conversation_documents_table(conn, schema: str = ""):
    """Create conversation_documents table for persistent file storage per
    conversation. Used by Estimation IA to keep uploaded plans/docs accessible
    across messages (indefinite continuation) without re-uploading.

    Storage: BYTEA (aligned with ai_profile_documents pattern). Extracted text
    stored for non-binary formats (xlsx, docx, csv) so Claude can reconsult
    content without re-running OCR.
    """
    if schema in _conversation_docs_ensured_for:
        return
    prev_autocommit = conn.autocommit
    try:
        conn.autocommit = True
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS conversation_documents (
                id SERIAL PRIMARY KEY,
                conversation_id INTEGER,
                filename TEXT NOT NULL,
                media_type TEXT,
                file_size INTEGER,
                content BYTEA,
                extracted_text TEXT,
                summary TEXT,
                category_detected TEXT,
                subcategory_detected TEXT,
                superficie_pi2 INTEGER,
                superficie_renovation_pi2 INTEGER,
                superficie_agrandissement_pi2 INTEGER,
                superficie_existant_conserve_pi2 INTEGER,
                is_active_context BOOLEAN DEFAULT TRUE,
                uploaded_by INTEGER,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        # Defensive migrations (tenants provisioned before columns were added)
        for col_name, col_type in (
            ("conversation_id", "INTEGER"),
            ("media_type", "TEXT"),
            ("file_size", "INTEGER"),
            ("content", "BYTEA"),
            ("extracted_text", "TEXT"),
            ("summary", "TEXT"),
            ("category_detected", "TEXT"),
            ("subcategory_detected", "TEXT"),
            ("superficie_pi2", "INTEGER"),
            ("superficie_renovation_pi2", "INTEGER"),
            ("superficie_agrandissement_pi2", "INTEGER"),
            ("superficie_existant_conserve_pi2", "INTEGER"),
            ("is_active_context", "BOOLEAN DEFAULT TRUE"),
            ("uploaded_by", "INTEGER"),
            ("updated_at", "TIMESTAMP DEFAULT CURRENT_TIMESTAMP"),
            # OPT-2: Files API beta — anthropic_file_id permet de reutiliser un
            # fichier deja uploade chez Anthropic plutot que de re-encoder en
            # base64 a chaque tour multi-tour. Gros gain pour /ai-chat-with-files
            # (moins de bytes envoyes, moins de tokens factures, latence reduite).
            ("anthropic_file_id", "TEXT"),
        ):
            try:
                cursor.execute(f"ALTER TABLE conversation_documents ADD COLUMN IF NOT EXISTS {col_name} {col_type}")
            except Exception as exc:
                logger.debug("ALTER conversation_documents.%s skipped: %s", col_name, exc)
        # Index pour lookup rapide par conversation
        try:
            cursor.execute(
                "CREATE INDEX IF NOT EXISTS idx_conv_docs_conversation "
                "ON conversation_documents(conversation_id, is_active_context)"
            )
        except Exception as exc:
            logger.debug("CREATE INDEX idx_conv_docs_conversation skipped: %s", exc)
        cursor.close()
        _conversation_docs_ensured_for.add(schema)
    except Exception as exc:
        logger.warning("_ensure_conversation_documents_table: %s", exc)
    finally:
        conn.autocommit = prev_autocommit


_devis_ai_estimations_ensured_for: set = set()


def _ensure_devis_ai_estimations_table(conn, schema: str = ""):
    """Cree la table devis_ai_estimations + index si manquants. Per-tenant.

    Persiste chaque estimation IA generee par:
      - ai_estimate_devis (type_estimation='simple')
      - ai_estimate_with_plan (type_estimation='with_plan')
      - ai_generate_soumission (type_estimation='soumission')
    """
    if schema in _devis_ai_estimations_ensured_for:
        return
    prev_autocommit = conn.autocommit
    try:
        conn.autocommit = True
        cursor = conn.cursor()
        cursor.execute(
            """
            CREATE TABLE IF NOT EXISTS devis_ai_estimations (
                id SERIAL PRIMARY KEY,
                devis_id INTEGER NOT NULL,
                user_id INTEGER NOT NULL,
                type_estimation TEXT NOT NULL,
                ai_text TEXT NOT NULL,
                metadata_json JSONB NOT NULL DEFAULT '{}'::jsonb,
                tokens_in INTEGER NOT NULL DEFAULT 0,
                tokens_out INTEGER NOT NULL DEFAULT 0,
                cache_creation_tokens INTEGER NOT NULL DEFAULT 0,
                cache_read_tokens INTEGER NOT NULL DEFAULT 0,
                cost_usd NUMERIC(10,4) NOT NULL DEFAULT 0,
                precision_mode BOOLEAN NOT NULL DEFAULT FALSE,
                thinking_tokens INTEGER NOT NULL DEFAULT 0,
                claude_model TEXT NOT NULL DEFAULT 'claude-opus-4-7',
                created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
                archived BOOLEAN NOT NULL DEFAULT FALSE
            )
            """
        )
        for col_name, col_type in (
            ("metadata_json", "JSONB NOT NULL DEFAULT '{}'::jsonb"),
            ("cache_creation_tokens", "INTEGER NOT NULL DEFAULT 0"),
            ("cache_read_tokens", "INTEGER NOT NULL DEFAULT 0"),
            ("cost_usd", "NUMERIC(10,4) NOT NULL DEFAULT 0"),
            ("precision_mode", "BOOLEAN NOT NULL DEFAULT FALSE"),
            ("thinking_tokens", "INTEGER NOT NULL DEFAULT 0"),
            ("claude_model", "TEXT NOT NULL DEFAULT 'claude-opus-4-7'"),
            ("archived", "BOOLEAN NOT NULL DEFAULT FALSE"),
        ):
            try:
                cursor.execute(
                    f"ALTER TABLE devis_ai_estimations "
                    f"ADD COLUMN IF NOT EXISTS {col_name} {col_type}"
                )
            except Exception as exc:
                logger.debug("ALTER devis_ai_estimations.%s skipped: %s", col_name, exc)
        for idx_sql in (
            "CREATE INDEX IF NOT EXISTS idx_devis_ai_estimations_devis_id ON devis_ai_estimations(devis_id)",
            "CREATE INDEX IF NOT EXISTS idx_devis_ai_estimations_created_at ON devis_ai_estimations(created_at DESC)",
            "CREATE INDEX IF NOT EXISTS idx_devis_ai_estimations_user_id ON devis_ai_estimations(user_id)",
            "CREATE INDEX IF NOT EXISTS idx_devis_ai_estimations_devis_active ON devis_ai_estimations(devis_id, created_at DESC) WHERE archived = FALSE",
        ):
            try:
                cursor.execute(idx_sql)
            except Exception as exc:
                logger.debug("CREATE INDEX skipped: %s", exc)
        cursor.close()
        _devis_ai_estimations_ensured_for.add(schema)
    except Exception as exc:
        logger.warning("_ensure_devis_ai_estimations_table: %s", exc)
    finally:
        conn.autocommit = prev_autocommit


def _persist_ai_estimation(
    schema: str,
    devis_id: int,
    user_id: int,
    type_estimation: str,
    ai_text: str,
    metadata: dict,
    tokens_in: int,
    tokens_out: int,
    cache_creation: int = 0,
    cache_read: int = 0,
    cost_usd: float = 0.0,
    precision_mode: bool = False,
    thinking_tokens: int = 0,
    claude_model: str = "claude-opus-4-7",
) -> Optional[int]:
    """Persiste une estimation IA dans devis_ai_estimations (best-effort).

    Retourne l'id si succes, None si echec. NE DOIT JAMAIS raise pour ne
    pas casser le flow IA principal.

    type_estimation: 'simple' | 'with_plan' | 'soumission'
    """
    if not schema:
        logger.warning("_persist_ai_estimation called without schema")
        return None
    if not devis_id or not user_id:
        logger.warning(
            "_persist_ai_estimation invalid devis_id=%s user_id=%s",
            devis_id, user_id,
        )
        return None
    if type_estimation not in ("simple", "with_plan", "soumission"):
        logger.warning(
            "_persist_ai_estimation: type_estimation invalide '%s'",
            type_estimation,
        )
        return None
    if ai_text is None:
        ai_text = ""

    try:
        metadata_str = json.dumps(metadata or {}, ensure_ascii=False, default=str)
    except (TypeError, ValueError) as ser_exc:
        logger.warning(
            "_persist_ai_estimation metadata non serialisable (%s), fallback {}",
            ser_exc,
        )
        metadata_str = "{}"

    conn = None
    cursor = None
    try:
        conn = db.get_conn()
        db.set_tenant(conn, schema)
        _ensure_devis_ai_estimations_table(conn, schema)
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO devis_ai_estimations ("
            "devis_id, user_id, type_estimation, ai_text, metadata_json, "
            "tokens_in, tokens_out, cache_creation_tokens, cache_read_tokens, "
            "cost_usd, precision_mode, thinking_tokens, claude_model"
            ") VALUES (%s, %s, %s, %s, %s::jsonb, %s, %s, %s, %s, %s, %s, %s, %s) "
            "RETURNING id",
            (
                devis_id, user_id, type_estimation, ai_text, metadata_str,
                int(tokens_in or 0), int(tokens_out or 0),
                int(cache_creation or 0), int(cache_read or 0),
                float(cost_usd or 0.0),
                bool(precision_mode),
                int(thinking_tokens or 0),
                claude_model or "claude-opus-4-7",
            ),
        )
        row = cursor.fetchone()
        conn.commit()
        new_id = row["id"] if row else None
        logger.info(
            "AI estimation persisted: tenant=%s devis_id=%s user_id=%s "
            "type=%s id=%s cost=%.4f",
            schema, devis_id, user_id, type_estimation, new_id, float(cost_usd or 0.0),
        )
        return new_id
    except Exception as exc:
        logger.warning(
            "_persist_ai_estimation FAILED (best-effort): tenant=%s devis_id=%s type=%s err=%s",
            schema, devis_id, type_estimation, exc,
        )
        try:
            if conn:
                conn.rollback()
        except Exception:
            pass
        return None
    finally:
        if cursor:
            try:
                cursor.close()
            except Exception:
                pass
        if conn:
            try:
                db.reset_tenant(conn)
            except Exception:
                pass
            try:
                conn.close()
            except Exception:
                pass


_devis_pct_ensured_for: set = set()

_DEVIS_PCT_COLUMNS = [
    "administration_pct NUMERIC DEFAULT 3.0",
    "contingences_pct NUMERIC DEFAULT 12.0",
    "profit_pct NUMERIC DEFAULT 15.0",
    "date_fin DATE",
    "project_id INTEGER",
    # 'Budgétaire' (estimation approximative) | 'Détaillée' (soumission ferme).
    # Existing rows default to 'Détaillée' per product decision — hands-on
    # soumissions have historically been the detailed/ventilated kind.
    "type_soumission VARCHAR(20) DEFAULT 'Détaillée'",
]


def _ensure_devis_pct_columns(conn, schema: str = ""):
    """Add frequently-missing columns to devis if drift exists. Per-tenant.

    Covers pct fields, date_fin, and `project_id` (needed by the orphan
    recovery path in update_devis — old tenants whose devis table predates
    the devis→project link have no project_id column and would crash with
    UndefinedColumn on the SELECT).

    WARNING for future contributors: when the caller invokes this with
    conn.autocommit=False (tx mode), the fallback path issues conn.rollback()
    on the caller's connection if the pool is exhausted. This is safe ONLY
    if the caller has executed nothing more than db.set_tenant() before this
    call. Do NOT call this after SELECT FOR UPDATE / INSERT / UPDATE on
    caller's conn in tx mode — the rollback would silently release locks
    and discard work. (Audited callers in tx mode: only crm.py:1045.)
    """
    # Defense-in-depth: refuse to operate without an explicit tenant schema.
    # All known callers pass user.schema (validated upstream), but an empty
    # schema would cause the ALTER to land on whatever search_path is set —
    # potentially `public.devis` if the caller forgot set_tenant or a rollback
    # wiped it. Refuse explicitly rather than relying on search_path.
    if not schema:
        logger.warning("_ensure_devis_pct_columns called without schema; skipping (defense-in-depth)")
        return
    if schema in _devis_pct_ensured_for:
        return

    # Caller may have flipped to tx mode (conn.autocommit=False) and executed
    # SET search_path via set_tenant, opening an implicit transaction. Toggling
    # autocommit while a tx is open raises "set_session cannot be used inside a
    # transaction" — and the finally restore re-raises, propagating a 500 to
    # the caller (e.g. crm.py:create_devis_from_opportunity). Strategy:
    # 1) Preferred: use a separate pool connection so we never touch caller's tx
    # 2) Fallback (pool exhausted, etc.): rollback caller's tx, ALTER in
    #    autocommit on caller's conn, restore caller's tx state. Safe IIF the
    #    only work the caller did before this call is set_tenant (audited:
    #    crm.py:1045 fits this — rollback only undoes the implicit BEGIN
    #    opened by set_tenant, no user data is lost).
    #
    # Both paths use schema-qualified DDL ("ALTER TABLE {schema}.devis ...")
    # to eliminate any dependency on search_path — even if a rollback wipes
    # search_path mid-fallback, the ALTER targets the correct tenant table
    # and never accidentally writes to public.devis (which exists in this
    # repo, see erp_database.py:4121,8521).
    qualified_devis = psql.SQL("{}.{}").format(psql.Identifier(schema), psql.Identifier("devis"))

    if not conn.autocommit:
        # _devis_pct_ensured_for is a global set; CPython GIL makes set.add()
        # atomic for hashable elements (string schema names) — no lock needed.
        # ProcessPoolExecutor / no-GIL builds would require a threading.Lock here.
        migration_conn = None
        separate_conn_done = False
        try:
            migration_conn = db.get_conn()
            # Pool returns connections in autocommit mode (database_config.py
            # ISOLATION_LEVEL_AUTOCOMMIT) — no transaction to manage here.
            db.set_tenant(migration_conn, schema)
            cursor = migration_conn.cursor()
            for col in _DEVIS_PCT_COLUMNS:
                # Schema-qualified ALTER: never falls back to public.devis
                # even if search_path is somehow wrong. col is a constant
                # literal (hardcoded above), safe to concatenate as SQL.
                cursor.execute(
                    psql.SQL("ALTER TABLE {} ADD COLUMN IF NOT EXISTS " + col).format(qualified_devis)
                )
            cursor.execute(
                psql.SQL("UPDATE {} SET type_soumission = 'Détaillée' WHERE type_soumission IS NULL").format(qualified_devis)
            )
            cursor.close()
            _devis_pct_ensured_for.add(schema)
            separate_conn_done = True
        except Exception as exc:
            # Includes psycopg2.pool.PoolError when pool is exhausted (max=10
            # in prod via ENVIRONMENT=production override). Fall through to
            # caller-conn fallback rather than letting the caller's INSERT
            # fail later with UndefinedColumn (silent failure mode).
            logger.warning("_ensure_devis_pct_columns (separate conn): %s", exc)
        finally:
            if migration_conn is not None:
                try:
                    migration_conn.close()
                except Exception:
                    pass

        if separate_conn_done:
            return

        # Fallback: separate-conn unavailable (pool drained). Use caller's conn.
        # Rollback closes the implicit tx opened by set_tenant (destroys
        # search_path — lesson #14), then we toggle autocommit, ALTER (with
        # schema-qualified DDL, immune to search_path state), and restore the
        # caller's original tx state with search_path re-set.
        try:
            try:
                conn.rollback()
            except Exception:
                pass
            conn.autocommit = True
            cursor = conn.cursor()
            for col in _DEVIS_PCT_COLUMNS:
                # Schema-qualified — ALTER lands on tenant.devis even if
                # search_path was wiped by rollback above and set_tenant
                # has not been re-issued yet.
                cursor.execute(
                    psql.SQL("ALTER TABLE {} ADD COLUMN IF NOT EXISTS " + col).format(qualified_devis)
                )
            cursor.execute(
                psql.SQL("UPDATE {} SET type_soumission = 'Détaillée' WHERE type_soumission IS NULL").format(qualified_devis)
            )
            cursor.close()
            _devis_pct_ensured_for.add(schema)
        except Exception as exc:
            logger.warning("_ensure_devis_pct_columns (caller-conn fallback): %s", exc)
        finally:
            # Restore caller's tx mode (autocommit=False) and search_path so
            # the caller's subsequent statements see the same context they had
            # before this call. If autocommit restore fails (rare — would
            # require a tx to be open after our autocommit ALTERs, which can
            # happen if a CREATE/ALTER triggered an implicit tx in some pg
            # versions), force a rollback then retry once before giving up
            # silently — leaving the caller in autocommit=True would cause
            # silent atomicity loss on subsequent multi-statement ops.
            try:
                conn.autocommit = False
            except Exception:
                try:
                    conn.rollback()
                    conn.autocommit = False
                except Exception as exc2:
                    logger.warning(
                        "_ensure_devis_pct_columns (fallback): cannot restore caller tx mode: %s",
                        exc2,
                    )
            try:
                db.set_tenant(conn, schema)
            except Exception:
                # If we can't restore search_path, the caller's next statement
                # will fail visibly (UndefinedTable on tenant table). Better
                # than silently writing to public.* with a dangling search_path.
                logger.warning(
                    "_ensure_devis_pct_columns (fallback): cannot re-set search_path for %s",
                    schema,
                )
        return

    # Caller's conn is already in autocommit mode — original flow, use directly.
    # Use schema-qualified DDL here too for parity and defense-in-depth.
    prev_autocommit = conn.autocommit
    try:
        conn.autocommit = True
        cursor = conn.cursor()
        for col in _DEVIS_PCT_COLUMNS:
            cursor.execute(
                psql.SQL("ALTER TABLE {} ADD COLUMN IF NOT EXISTS " + col).format(qualified_devis)
            )
        # Backfill NULLs that may exist on tenants that added the column
        # before the DEFAULT was introduced.
        cursor.execute(
            psql.SQL("UPDATE {} SET type_soumission = 'Détaillée' WHERE type_soumission IS NULL").format(qualified_devis)
        )
        cursor.close()
        _devis_pct_ensured_for.add(schema)
    except Exception as exc:
        logger.warning("_ensure_devis_pct_columns: %s", exc)
    finally:
        try:
            conn.autocommit = prev_autocommit
        except Exception as exc:
            # Defensive: if a tx was opened by the loop above and cursor.close
            # didn't end it, restoring autocommit fails with the same set_session
            # error. Rollback then retry once before logging.
            try:
                conn.rollback()
                conn.autocommit = prev_autocommit
            except Exception:
                logger.warning("_ensure_devis_pct_columns: cannot restore autocommit: %s", exc)


def _read_devis_pct(cursor, devis_id: int) -> tuple:
    """Read administration/contingences/profit percentages from devis.
    _ensure_devis_pct_columns guarantees columns exist; fallback is last resort."""
    try:
        cursor.execute(
            "SELECT administration_pct, contingences_pct, profit_pct FROM devis WHERE id = %s",
            (devis_id,),
        )
        row = cursor.fetchone()
        if row:
            adm = float(row["administration_pct"]) if row.get("administration_pct") is not None else _DEFAULT_ADM_PCT
            con = float(row["contingences_pct"]) if row.get("contingences_pct") is not None else _DEFAULT_CON_PCT
            pro = float(row["profit_pct"]) if row.get("profit_pct") is not None else _DEFAULT_PRO_PCT
            return adm, con, pro
    except Exception:
        # Do NOT rollback — destroys search_path (lesson #14).
        # _ensure_devis_pct_columns should have added columns via autocommit.
        pass
    return _DEFAULT_ADM_PCT, _DEFAULT_CON_PCT, _DEFAULT_PRO_PCT


def _read_devis_pct_with_amounts(cursor, devis_id: int) -> dict:
    """Read full devis financials including pct and amounts.
    _ensure_devis_pct_columns guarantees columns exist."""
    try:
        cursor.execute(
            "SELECT total_travaux, administration_pct, contingences_pct, profit_pct, "
            "administration, contingences, profit FROM devis WHERE id = %s",
            (devis_id,),
        )
        return cursor.fetchone()
    except Exception:
        # Do NOT rollback — destroys search_path (lesson #14).
        # Return None; caller's outer except will handle the 500.
        return None


class DevisAssignmentCreate(BaseModel):
    employee_id: int
    role: Optional[str] = None


def _empty_to_none(v: Optional[str]) -> Optional[str]:
    """Convert empty strings to None for optional date fields."""
    return None if v is not None and v.strip() == "" else v


def _strip_non_empty(v):
    """Strip whitespace and reject empty strings. Passes None through.
    Used by required-name validators to block `""` and `"   "` inputs."""
    if v is None:
        return v
    v = str(v).strip()
    if not v:
        raise ValueError("Ne peut pas etre vide")
    return v


class DevisCreate(BaseModel):
    nom_projet: str
    client_company_id: Optional[int] = None
    client_contact_id: Optional[int] = None
    client_nom_direct: Optional[str] = None
    project_id: Optional[str] = None
    description: Optional[str] = None
    date_prevu: Optional[str] = None
    date_soumis: Optional[str] = None
    date_fin: Optional[str] = None
    po_client: Optional[str] = None
    priorite: Optional[str] = None
    tache: Optional[str] = None
    prix_estime: Optional[float] = None
    notes: Optional[str] = None
    administration_pct: Optional[float] = None
    contingences_pct: Optional[float] = None
    profit_pct: Optional[float] = None
    # 'Budgétaire' or 'Détaillée'. None on create → DB default ('Détaillée')
    type_soumission: Optional[str] = None

    _nom_projet_validator = field_validator("nom_projet", mode="before")(_strip_non_empty)

    @field_validator("type_soumission", mode="before")
    @classmethod
    def _clean_type_soumission(cls, v):
        if v is None:
            return None
        v = str(v).strip()
        if not v:
            return None
        if v not in ("Budgétaire", "Détaillée"):
            raise ValueError("type_soumission doit être 'Budgétaire' ou 'Détaillée'")
        return v

    @field_validator("date_prevu", "date_soumis", "date_fin", mode="before")
    @classmethod
    def _clean_date(cls, v):
        return _empty_to_none(v)


class DevisUpdate(BaseModel):
    nom_projet: Optional[str] = None
    statut: Optional[str] = None
    description: Optional[str] = None
    date_prevu: Optional[str] = None
    date_soumis: Optional[str] = None
    date_fin: Optional[str] = None
    notes: Optional[str] = None
    total_travaux: Optional[float] = None
    client_company_id: Optional[int] = None
    client_contact_id: Optional[int] = None
    client_nom_direct: Optional[str] = None
    po_client: Optional[str] = None
    priorite: Optional[str] = None
    tache: Optional[str] = None
    prix_estime: Optional[float] = None
    administration_pct: Optional[float] = None
    contingences_pct: Optional[float] = None
    profit_pct: Optional[float] = None
    administration: Optional[float] = None
    contingences: Optional[float] = None
    profit: Optional[float] = None
    show_administration: Optional[bool] = None
    show_contingences: Optional[bool] = None
    show_profit: Optional[bool] = None
    show_unite: Optional[bool] = None
    show_quantite: Optional[bool] = None
    show_prix_unitaire: Optional[bool] = None
    show_montant_ligne: Optional[bool] = None
    show_mo_mat: Optional[bool] = None
    administration_label: Optional[str] = Field(default=None, max_length=50)
    contingences_label: Optional[str] = Field(default=None, max_length=50)
    profit_label: Optional[str] = Field(default=None, max_length=50)
    conditions_text: Optional[str] = Field(default=None, max_length=10000)
    exclusions_text: Optional[str] = Field(default=None, max_length=10000)
    show_conditions: Optional[bool] = None
    show_exclusions: Optional[bool] = None
    type_soumission: Optional[str] = None

    _nom_projet_validator = field_validator("nom_projet", mode="before")(_strip_non_empty)

    @field_validator("date_prevu", "date_soumis", "date_fin", mode="before")
    @classmethod
    def _clean_date(cls, v):
        return _empty_to_none(v)

    @field_validator("type_soumission", mode="before")
    @classmethod
    def _clean_type_soumission_update(cls, v):
        if v is None:
            return None
        v = str(v).strip()
        if not v:
            return None
        if v not in ("Budgétaire", "Détaillée"):
            raise ValueError("type_soumission doit être 'Budgétaire' ou 'Détaillée'")
        return v

    @field_validator("conditions_text", "exclusions_text", mode="before")
    @classmethod
    def _limit_lines(cls, v):
        # DoS guard: 200 lignes max — généreusement au-dessus du besoin réel (~15 items)
        # tout en empêchant un payload adversarial de 10k caractères de newlines seuls.
        if isinstance(v, str) and v.count("\n") > 200:
            raise ValueError("Maximum 200 lignes par champ")
        return v


class DevisLigneCreate(BaseModel):
    description: str
    quantite: float = Field(default=1, gt=0)
    unite: str = "unite"
    prix_unitaire: float = Field(default=0, ge=0)
    categorie: Optional[str] = None
    notes_ligne: Optional[str] = None
    sequence_ligne: int = 0
    code_article: Optional[str] = None
    # Custom MO/MAT ratios per line (0-100, sum should = 100).
    # If both None, auto-detection via _get_mo_mat_ratio(description) is used.
    mo_pct: Optional[float] = Field(default=None, ge=0, le=100)
    mat_pct: Optional[float] = Field(default=None, ge=0, le=100)
    # Per-line markup overrides (0-100). NULL = inherit the devis-level
    # administration_pct / contingences_pct / profit_pct (default behaviour).
    # When non-NULL, these % are applied to THIS line's HT amount only,
    # leaving other lines untouched. Used for granular pricing decisions
    # (e.g. higher profit on rare items, lower on negotiated ones).
    admin_pct_ligne: Optional[float] = Field(default=None, ge=0, le=100)
    contingence_pct_ligne: Optional[float] = Field(default=None, ge=0, le=100)
    profit_pct_ligne: Optional[float] = Field(default=None, ge=0, le=100)


class CCQCalculation(BaseModel):
    montant_main_oeuvre: float
    metiers: List[str] = []


class CNESSTCalculation(BaseModel):
    montant_main_oeuvre: float
    taux_unite: float = 1.80


# ============ AI Profile Models ============

class AiProfileCreate(BaseModel):
    name: str
    instructions: str = ""


class AiProfileUpdate(BaseModel):
    name: Optional[str] = None
    instructions: Optional[str] = None
    is_active: Optional[bool] = None


# ============ Estimation IA Models ============

class AiChatRequest(BaseModel):
    messages: List[dict]
    profile_id: Optional[str] = None
    devis_id: Optional[int] = None
    conversation_id: Optional[int] = None


class AiGenerateSoumissionRequest(BaseModel):
    messages: List[dict]
    projet_type: Optional[str] = None
    superficie: Optional[float] = None
    profile_id: Optional[str] = None
    devis_id: Optional[int] = None  # Sprint 3 #7: lien optionnel pour persistance


class ConversationSave(BaseModel):
    name: str
    devis_id: Optional[int] = None
    subject: Optional[str] = None
    messages: List[dict]
    expert_profile: Optional[str] = None
    metadata: Optional[str] = None


def _get_profiles_dir() -> Optional[str]:
    """Find the profiles directory relative to the repo root."""
    base = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
    for sub in ("profiles", os.path.join("EXPERTS_AI", "profiles")):
        d = os.path.join(base, sub)
        if os.path.isdir(d):
            return d
    return None


# ============================================
# DEVIS DEFAULTS (CONDITIONS + EXCLUSIONS per entreprise)
# ============================================

class DevisDefaultsUpdate(BaseModel):
    """Body for PUT /devis/defaults. Empty string clears the default (falls back
    to hardcoded DEVIS_CONDITIONS/DEVIS_EXCLUSIONS constants)."""
    conditions: Optional[str] = Field(default=None, max_length=10000)
    exclusions: Optional[str] = Field(default=None, max_length=10000)

    @field_validator("conditions", "exclusions", mode="before")
    @classmethod
    def _limit_lines(cls, v):
        if isinstance(v, str) and v.count("\n") > 200:
            raise ValueError("Maximum 200 lignes par champ")
        return v


@router.get("/defaults")
async def get_devis_defaults(user: ErpUser = Depends(get_current_user)):
    """Return the entreprise-wide default conditions + exclusions for new devis.
    Empty string when unset — caller shows hardcoded placeholder."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        cursor = conn.cursor()
        defaults = _get_entreprise_devis_defaults(cursor)
        return {
            "conditions": defaults.get("conditions") or "",
            "exclusions": defaults.get("exclusions") or "",
            "conditions_fallback": "\n".join(DEVIS_CONDITIONS),
            "exclusions_fallback": "\n".join(DEVIS_EXCLUSIONS),
        }
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("get_devis_defaults error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors du chargement des defauts")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.put("/defaults")
async def update_devis_defaults(body: DevisDefaultsUpdate, user: ErpUser = Depends(require_role("admin"))):
    """Update the entreprise-wide default conditions + exclusions (admin only —
    same surface as /config/entreprise endpoints).
    Writes to entreprise_config.config_data as JSON keys:
      - devis_conditions_default
      - devis_exclusions_default
    Empty string clears the key (falls back to hardcoded constants)."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    prev_autocommit = None
    try:
        db.set_tenant(conn, user.schema)
        cursor = conn.cursor()
        # Ensure the entreprise_config table exists (old tenants may not have it).
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS entreprise_config (
                id SERIAL PRIMARY KEY,
                config_data JSONB,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # FIX P1 (round 8): forcer autocommit=False pour que SELECT FOR UPDATE
        # tienne son lock pendant le merge config + UPDATE/INSERT. En autocommit,
        # 2 saves concurrents peuvent perdre des modifications (last-write-wins).
        try:
            prev_autocommit = conn.autocommit
        except Exception:
            prev_autocommit = None
        try:
            conn.autocommit = False
        except Exception:
            pass

        cursor.execute("SELECT config_data FROM entreprise_config WHERE id = 1 FOR UPDATE")
        row = cursor.fetchone()
        if row:
            raw = row.get("config_data")
            if isinstance(raw, dict):
                current = raw
            elif isinstance(raw, str) and raw.strip():
                try:
                    current = json.loads(raw)
                except Exception:
                    current = {}
            else:
                current = {}
        else:
            current = {}

        # Only update fields that were sent (empty string is meaningful — "clear")
        fields = body.model_dump(exclude_unset=True)
        if "conditions" in fields:
            val = (fields["conditions"] or "").strip()
            if val:
                current["devis_conditions_default"] = val
            else:
                current.pop("devis_conditions_default", None)
        if "exclusions" in fields:
            val = (fields["exclusions"] or "").strip()
            if val:
                current["devis_exclusions_default"] = val
            else:
                current.pop("devis_exclusions_default", None)

        if row:
            cursor.execute(
                "UPDATE entreprise_config SET config_data = %s::jsonb, "
                "updated_at = CURRENT_TIMESTAMP WHERE id = 1",
                (json.dumps(current),),
            )
        else:
            cursor.execute(
                "INSERT INTO entreprise_config (id, config_data, created_at, updated_at) "
                "VALUES (1, %s::jsonb, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)",
                (json.dumps(current),),
            )
        conn.commit()
        return {
            "message": "Defauts enregistres",
            "conditions": current.get("devis_conditions_default", ""),
            "exclusions": current.get("devis_exclusions_default", ""),
        }
    except HTTPException:
        try:
            conn.rollback()
        except Exception:
            pass
        raise
    except Exception as exc:
        logger.error("update_devis_defaults error: %s", exc)
        try:
            conn.rollback()
            db.set_tenant(conn, user.schema)
        except Exception:
            pass
        raise HTTPException(status_code=500, detail="Erreur lors de la sauvegarde des defauts")
    finally:
        # Restaurer l'autocommit avant retour au pool.
        if prev_autocommit is not None:
            try:
                conn.autocommit = prev_autocommit
            except Exception:
                pass
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.get("")
async def list_devis(
    user: ErpUser = Depends(get_current_user),
    page: int = Query(1, ge=1),
    per_page: int = Query(20, ge=1, le=100),
    search: Optional[str] = None,
    statut: Optional[str] = None,
    # The Axios request interceptor already converts camelCase query params
    # (`typeSoumission`) to snake_case (`type_soumission`) before the request
    # hits the backend, so no FastAPI alias is needed. Using an alias would
    # actually BREAK the filter — FastAPI would only accept the aliased name
    # and silently drop the snake_case form sent by the interceptor.
    type_soumission: Optional[str] = Query(None),
):
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        _ensure_devis_pct_columns(conn, user.schema)
        cursor = conn.cursor()
        wheres, params = [], []
        if search:
            wheres.append("(LOWER(d.nom_projet) LIKE %s OR LOWER(d.numero_devis) LIKE %s)")
            s = f"%{search.lower()}%"
            params.extend([s, s])
        if statut:
            # Qualify with alias — `companies` also has a `statut` column so
            # the JOIN in the SELECT below makes this reference ambiguous.
            wheres.append("d.statut = %s")
            params.append(statut)
        if type_soumission and type_soumission in ("Budgétaire", "Détaillée"):
            wheres.append("d.type_soumission = %s")
            params.append(type_soumission)
        w = " AND ".join(wheres) if wheres else "TRUE"
        cursor.execute(f"SELECT COUNT(*) as total FROM devis d WHERE {w}", params)
        total = cursor.fetchone()["total"]
        offset = (page - 1) * per_page
        cursor.execute(
            f"SELECT d.id, d.numero_devis, d.nom_projet, d.description, d.statut, "
            f"d.client_company_id, d.client_contact_id, d.client_nom_direct, d.project_id, "
            f"d.total_travaux, d.tps, d.tvq, d.investissement_total, "
            f"d.created_at, d.date_prevu, d.date_soumis, d.date_fin, "
            f"d.po_client, d.priorite, d.tache, d.prix_estime, "
            f"d.type_soumission, "
            f"COALESCE(d.client_nom_cache, c.nom, "
            f"NULLIF(TRIM(COALESCE(ct.prenom,'') || ' ' || COALESCE(ct.nom_famille,'')), ''), "
            f"d.client_nom_direct) as client_nom_cache "
            f"FROM devis d "
            f"LEFT JOIN companies c ON d.client_company_id = c.id "
            f"LEFT JOIN contacts ct ON d.client_contact_id = ct.id "
            f"WHERE {w} ORDER BY d.created_at DESC LIMIT %s OFFSET %s",
            params + [per_page, offset],
        )
        items = []
        for row in cursor.fetchall():
            d = dict(row)
            for k in ("created_at", "date_prevu", "date_soumis", "date_fin", "date_decision"):
                if d.get(k) is not None:
                    d[k] = str(d[k])
            for k in ("total_travaux", "tps", "tvq", "investissement_total", "prix_estime"):
                if d.get(k) is not None:
                    d[k] = float(d[k])
            items.append(d)
        return {"items": items, "total": total, "page": page, "per_page": per_page}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("list_devis error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


# ============================================
# FIXED-PATH ENDPOINTS (must be before /{devis_id})
# ============================================

@router.get("/statistics")
async def get_devis_statistics(user: ErpUser = Depends(get_current_user)):
    """Get devis statistics: totals, acceptance rate, amounts by status."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        _ensure_devis_pct_columns(conn, user.schema)
        cursor = conn.cursor()
        cursor.execute(
            "SELECT statut, COUNT(*) as count, COALESCE(SUM(investissement_total), 0) as montant "
            "FROM devis GROUP BY statut"
        )
        rows = cursor.fetchall()
        par_statut = []
        total = 0
        montant_total = 0.0
        acceptes = 0
        refuses = 0
        en_attente = 0
        brouillons = 0
        envoyes = 0
        for row in rows:
            r = dict(row)
            s = r["statut"] or "Inconnu"
            c = int(r["count"])
            m = float(r["montant"])
            par_statut.append({"statut": s, "count": c, "montant": round(m, 2)})
            total += c
            montant_total += m
            if s == "Accepte":
                acceptes = c
            elif s == "Refuse":
                refuses = c
            elif s in ("En attente", "Envoye"):
                if s == "En attente":
                    en_attente += c
                else:
                    envoyes += c
            elif s == "Brouillon":
                brouillons = c
        # en_attente also includes Envoye for the summary field
        en_attente_total = en_attente + envoyes
        denominator = acceptes + refuses
        taux_acceptation = round((acceptes / denominator) * 100, 1) if denominator > 0 else 0.0

        # Per-type counts (Budgétaire vs Détaillée). Legacy rows with NULL
        # type_soumission fall back to Détaillée via the _ensure defensive
        # migration above, so every row contributes to exactly one bucket.
        budgetaires = 0
        detaillees = 0
        try:
            cursor.execute(
                "SELECT COALESCE(type_soumission, 'Détaillée') as t, COUNT(*) as count "
                "FROM devis GROUP BY COALESCE(type_soumission, 'Détaillée')"
            )
            for tr in cursor.fetchall():
                tdict = dict(tr)
                if tdict["t"] == "Budgétaire":
                    budgetaires = int(tdict["count"])
                elif tdict["t"] == "Détaillée":
                    detaillees = int(tdict["count"])
        except Exception as exc:
            logger.warning("devis statistics type_soumission group: %s", exc)

        return {
            "total": total,
            "taux_acceptation": taux_acceptation,
            "montant_total": round(montant_total, 2),
            "en_attente": en_attente_total,
            "brouillons": brouillons,
            "envoyes": envoyes,
            "budgetaires": budgetaires,
            "detaillees": detaillees,
            "par_statut": par_statut,
        }
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("get_devis_statistics error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.post("/calculate-ccq")
async def calculate_ccq(body: CCQCalculation):
    """Calculate CCQ employer costs for construction."""
    # CCQ rates by trade (2026)
    CCQ_RATES = {
        "general": 12.5,
        "electricien": 11.8,
        "plombier": 11.8,
        "charpentier": 12.5,
        "menuisier": 12.5,
        "ferblantier": 11.8,
        "peintre": 12.5,
        "calorifugeur": 12.5,
        "frigoriste": 11.8,
        "mecanicien": 11.8,
        "soudeur": 12.5,
        "manoeuvre": 12.5,
        "operateur": 12.5,
        "grutier": 12.5,
    }
    default_rate = 12.5
    details_metiers = []
    if body.metiers:
        total_cotisation = 0.0
        part_per_metier = body.montant_main_oeuvre / len(body.metiers) if len(body.metiers) > 0 else 0
        for metier in body.metiers:
            metier_lower = metier.lower().strip()
            taux = CCQ_RATES.get(metier_lower, default_rate)
            cotisation_metier = round(part_per_metier * taux / 100, 2)
            total_cotisation += cotisation_metier
            details_metiers.append({
                "metier": metier,
                "montant_main_oeuvre": round(part_per_metier, 2),
                "taux_ccq": taux,
                "cotisation_ccq": cotisation_metier,
            })
        taux_ccq = round(total_cotisation / body.montant_main_oeuvre * 100, 2) if body.montant_main_oeuvre > 0 else default_rate
    else:
        taux_ccq = default_rate
        total_cotisation = round(body.montant_main_oeuvre * taux_ccq / 100, 2)
    total_avec_ccq = round(body.montant_main_oeuvre + total_cotisation, 2)
    return {
        "montant_main_oeuvre": body.montant_main_oeuvre,
        "taux_ccq": taux_ccq,
        "cotisation_ccq": round(total_cotisation, 2),
        "total_avec_ccq": total_avec_ccq,
        "details_metiers": details_metiers,
    }


@router.post("/calculate-cnesst")
async def calculate_cnesst(body: CNESSTCalculation):
    """Calculate CNESST employer costs."""
    cotisation_cnesst = round(body.montant_main_oeuvre * body.taux_unite / 100, 2)
    total_avec_cnesst = round(body.montant_main_oeuvre + cotisation_cnesst, 2)
    return {
        "montant_main_oeuvre": body.montant_main_oeuvre,
        "taux_unite": body.taux_unite,
        "cotisation_cnesst": cotisation_cnesst,
        "total_avec_cnesst": total_avec_cnesst,
    }


# ------------------------------------------------------------------
# DIAGNOSTIC PREMIERE ANALYSE — prompt de categorisation base sur le
# profil Entrepreneur general (5 categories + sous-categorie gamme)
# ------------------------------------------------------------------
_DIAGNOSTIC_PROMPT = """En tant qu'Entrepreneur general chevronne (40 ans d'experience au Quebec),
analyse ce document avec RIGUEUR TECHNIQUE de lecteur de plans professionnel.

═══════════════════════════════════════════════════════════════════════════
METHODOLOGIE OBLIGATOIRE DE LECTURE DU PLAN PDF
═══════════════════════════════════════════════════════════════════════════

Avant de produire le diagnostic, applique cette methodologie rigoureuse:

**A. Inventaire des vues** (LIS TOUTES LES PAGES)
Un plan professionnel comprend plusieurs vues. Identifie-les une par une:
- **Plan d'implantation** (vu du dessus, position sur le terrain, orientation
  N/S/E/O, marges de recul)
- **Plan en plan** (vue du dessus de chaque etage: fondation, RDC, etage,
  sous-sol). Permet de lire la FORME DU BATIMENT et les dimensions.
- **Elevations** (facades avant/arriere/cotes). Permet de lire les hauteurs,
  revetements exterieurs, types d'ouvertures, pentes de toit.
- **Coupes/sections** (tranche verticale du batiment). Permet de lire les
  fondations, hauteurs sous-plafond, structure, isolation.
- **Details d'assemblage** (zoom techniques, murs, toitures, parapets).

**B. Lecture des cotations** (dimensions marquees)
- Identifie l'**echelle** du plan (ex: "1:50", "1/4" = 1'-0\\"", "1:100").
- Extrais les **cotes exterieures** (longueur × largeur totale du batiment).
- Pour les formes complexes (L, T, U, avec decroches): decompose en
  rectangles et somme les superficies.
- Calcul superficie = longueur × largeur pour chaque rectangle, puis somme.
- **Privilegier TOUJOURS les cotations ecrites** aux estimations visuelles.
- Si cotations absentes mais echelle presente: utilise l'echelle pour
  estimer avec marge d'erreur notee.
- Si cotations ET echelle absentes: note "Plan non cote — dimensions a
  confirmer par le client" plutot que d'inventer des chiffres.

**C. Detection du contour exterieur du batiment**
Trace mentalement le perimetre du mur exterieur:
- Identifie les **decroches** (changements d'angle du contour).
- Repere les **saillies**: balcons, galeries, auvents, avant-toits
  (pas dans la superficie habitable mais mentionner).
- Si renovation: **distingue les traits** existant vs agrandissement:
  * Trait plein epais + hachure = contour existant conserve
  * Trait plein fin = nouveau construit
  * Pointilles / traits barres = elements a demolir
  * Legende du plan donne la convention exacte — LIS-LA.

**D. Conventions de dessin techniques quebecoises**
- Cartouche (coin inferieur droit): nom concepteur, dossier, revision, date
- Symboles OTPQ/OAQ/OIQ: indiquent qualification du concepteur
- Murs porteurs: trait plus epais + souvent hachures
- Cloisons: trait plus fin
- Portes: arc de cercle 90° indiquant sens d'ouverture
- Fenetres: double trait parallele avec parfois indication type (G=guillotine,
  A=auvent, B=battant, F=fixe)
- Escaliers: fleche de direction + ligne de coupe
- Cotes en pieds-pouces (ex: 24'-6\\") ou metrique (7315 mm)
- Hauteurs marquees sous "H=" ou cote elevation (TOS, TOP, FF)

**E. Lecture des coupes pour fondations et structure**
- Type de fondation visible: pieux, dalle, mur beton coule, blocs, pilotis
- Hauteur fondation vs sol (hors-sol, semi-enterre, complet)
- Type de structure: 2x4 vs 2x6, espacement @ 16\\" c/c ou 24\\" c/c
- Poutres: LVL, Microlam, bois lamelle-colle, acier, avec dimensions
- Toiture: charpente conventionnelle vs fermes prefabriquees, pentes

**F. Synthese**
Apres cet inventaire, produis le diagnostic des 8 sections ci-dessous
en citant EXPLICITEMENT les pages/vues d'ou chaque information provient.
Si tu ne peux pas determiner une information, dis-le clairement plutot
que d'inventer.

═══════════════════════════════════════════════════════════════════════════

Produis maintenant le DIAGNOSTIC PRELIMINAIRE structure avec les
sections OBLIGATOIRES suivantes en markdown.

### 0. Inventaire des vues du plan
Liste les vues detectees page par page (si le plan fait plusieurs pages):
- Page 1: <type de vue, ex: Plan d'implantation + cartouche>
- Page 2: <ex: Plan fondation>
- Page 3: <ex: Plan rez-de-chaussee>
- Page 4: <ex: Plan etage>
- Page 5: <ex: Elevation avant + elevation arriere>
- Page 6: <ex: Elevation cotes droit et gauche>
- Page 7: <ex: Coupe A-A + coupe B-B>
- Page 8: <ex: Plan electrique>
- etc.

Mentionne l'**echelle** principale (ex: "1:50" ou "1/4\\" = 1'-0\\"") et
l'**orientation** (nord sur le plan).

### 1. Categorie detectee
Choisis UNE categorie parmi ces 5 (correspondant aux sections du profil Entrepreneur general):
- **Residentiel neuf** — construction residentielle neuve (unifamiliale, multifamiliale)
- **Renovation residentielle** — renovation ou agrandissement d'une residence existante
- **Commercial neuf** — construction commerciale neuve (bureaux, commerces, restaurants)
- **Commercial renovation** — renovation ou transformation d'un espace commercial
- **Institutionnel / public** — ecoles, CHSLD, hopitaux, edifices gouvernementaux

Justifie en 1-2 phrases les indices qui t'ont permis de categoriser
(type de plan, superficie, usage, materiaux, etc.).

### 2. Niveau de gamme — TOUJOURS ECONOMIQUE PAR DEFAUT
**REGLE ABSOLUE**: Au diagnostic preliminaire, utilise TOUJOURS la gamme
**Economique** (prix plancher) de la categorie detectee, peu importe la
qualite apparente des finitions visibles au plan.

Raison metier: les finitions interieures et exterieures sont les elements
les plus volatils du budget (revetement, toiture, planchers, armoires,
salles de bain, etc.). Le client peut decider de garder les finitions
premium du plan original OU de les remplacer par des options economiques.
On part donc TOUJOURS de la base la plus basse et on monte ensuite selon
les demandes explicites du client au fil de la conversation.

Format de la section (OBLIGATOIRE):
- **Gamme appliquee**: Economique (fourchette $/pi² de la categorie)
- **Indices de gamme superieure detectes au plan** (liste a titre
  indicatif seulement — NON APPLIQUES au calcul par defaut):
  * liste des finitions/equipements premium identifies sur le plan
  * ex: Revetement en lattes aluminium imitation bois
  * ex: Tôle à joints pincés sans vis apparentes
  * ex: Planchers chauffants sur dalle
  * ex: VRC EnergyStar haute efficacite
- **Note au client**: "L'estimation par defaut utilise la gamme
  Economique. Si vous souhaitez conserver les finitions premium
  identifiees ci-dessus, precisez-le et je recalculerai en gamme
  De base, Moyenne ou Haut de gamme selon vos finitions interieures
  et exterieures souhaitees."

Echelle de reference (pour escalade sur demande client):
- **Economique** — finition minimale, materiaux de base (PAR DEFAUT)
- **De base** — finition standard
- **Moyenne** — finition moderne avec quelques elements haut de gamme
- **Haut de gamme** — finitions nobles, equipements premium

### 3. Superficies par zone de travaux — DECOMPOSITION OBLIGATOIRE
**METHODE DE CALCUL** (obligatoire):
1. Lis les cotations exterieures du plan en plan (longueur × largeur).
2. Pour forme rectangulaire simple: superficie = longueur × largeur.
3. Pour forme complexe (L, T, U, decroches): decompose en rectangles
   elementaires, calcule chaque rectangle, puis additionne.
4. Indique la source des dimensions: "selon cotations page 3" ou
   "selon echelle 1:50, dimensions visuelles" ou "cotations absentes —
   a confirmer par le client".
5. Convertis au besoin: 1 m² = 10.764 pi² / 1 pi² = 0.0929 m².

**REGLE METIER**: Pour une renovation + agrandissement, tu DOIS decomposer
le plan en 3 zones distinctes et n'estimer QUE les zones touchees
(renovation + agrandissement). L'existant conserve est EXCLU du calcul.

Indices visuels du plan:
- **Zone A — Existant tel quel (NON ESTIMEE)**: mentions "EXISTANT TEL QUEL",
  "A CONSERVER", murs/planchers sans trait de demolition ni modification.
  Partie du sous-sol non-excavee.
- **Zone B — Renovation (existant modifie)**: murs/portes/fenetres barres
  (legende demolition), nouveau revetement sur mur existant, remplacement
  cuisine/salle de bain, reamenagement interieur.
- **Zone C — Agrandissement / Ajout (neuf)**: mentions "AGRANDISSEMENT",
  "AJOUT", "NOUVEAU", nouvelle fondation (pieux, dalle neuve), structure
  entierement neuve.

═══════════════════════════════════════════════════════════════════════════
METHODOLOGIE DE CALCUL — OBLIGATOIRE SI PAS DE SUPERFICIE TEXTUELLE
═══════════════════════════════════════════════════════════════════════════

Certains plans (dessinateurs en batiment, plans de renovation) n'inscrivent
PAS de superficie totale dans leur cartouche. Tu vois seulement des
**cotations dimensionnelles** (ex: 52'-2", 16'-0", 22'-0").

**REGLE ABSOLUE**: tu n'inventes JAMAIS une superficie. Si pas de chiffre
textuel "X pi²" / "X P.C." / "X m²" lisible, tu APPLIQUES cette methodologie
de calcul deterministe et tu MONTRES tes calculs visiblement:

**ETAPE 1 — Identification du contour exterieur**
Lis les cotations de la facade ET de la profondeur du batiment.
Decompose la forme globale en rectangles elementaires (R1, R2, R3...).

**ETAPE 2 — Calcul detaille rectangle par rectangle**
Pour chaque rectangle, applique aire = longueur × largeur en pieds.
Convertis pieds-pouces en decimal: 52'-2" = 52 + 2/12 = 52.17 pi
                                   16'-6" = 16 + 6/12 = 16.5 pi

**ETAPE 3 — Affiche les calculs en bloc code**
OBLIGATOIRE: utilise un bloc markdown ```text ... ``` pour montrer
l'arithmetique de maniere auditable. Exemple type pour une forme en L:

```text
DECOMPOSITION DU CONTOUR (page A300, cotations exterieures):

  R1 — Corps principal     : 52'-2" × 16'-0"  = 52.17 × 16.0  = 834.7 pi²
  R2 — Extension cuisine   : 10'-0" × 22'-0"  = 10.0  × 22.0  = 220.0 pi²
  R3 — Avancee facade      : 16'-0" × 6'-0"   = 16.0  × 6.0   =  96.0 pi²

  TOTAL EMPREINTE RDC       = 834.7 + 220.0 + 96.0 = 1 150.7 pi²

  Sources cotations:
    - 52'-2" : facade avant (page A300)
    - 16'-0" : profondeur principale (page A300)
    - 22'-0" : profondeur extension (page A301)
    - 10'-0" : largeur extension (page A300)

  Si projet 2 etages identiques:
    SUPERFICIE HABITABLE TOTALE = 1 150.7 × 2 = 2 301 pi²
```

**ETAPE 4 — Auto-validation**
Apres calcul, verifie la coherence:
- Resultat ≈ surface visible sur le plan a l'echelle? (sanity check)
- Toutes les cotations utilisees correspondent a celles du plan?
- Y-a-t-il un retour, decroche ou annexe que j'ai oublie?

**ETAPE 5 — Communication transparente au client**
Termine la section superficies par:
> *Note: superficies calculees a partir des cotations exterieures car
> aucune superficie totale n'est inscrite au cartouche. Veuillez confirmer
> ou corriger ces valeurs avant chiffrage final.*

**INTERDIT**:
- Donner un chiffre rond sans montrer le calcul (ex: "environ 1500 pi²" sans
  decomposition rectangulaire)
- Approximer "~1 600 pi²" si tu peux calculer 1 597.3 pi² depuis les cotations
- Mentionner une superficie sans citer la page source des cotations utilisees

═══════════════════════════════════════════════════════════════════════════

Format OBLIGATOIRE de la section (utilise les valeurs reelles extraites du plan):
- **Zone A — Existant conserve** (NON ESTIMEE): X pi² [liste des pieces/zones]
- **Zone B — Renovation** (existant modifie): X pi² [liste]
- **Zone C — Agrandissement / Neuf**: X pi² [liste, separer etages si applicable]
- **Garage attache (si present)**: X pi² (tarifie separement, ~150 $/pi² economique)
- **Sous-sol** (si visible):
  * Existant non touche: X pi² (non estime)
  * Neuf/fini/renove: X pi² (estime selon classification B ou C)
- **SUPERFICIE TOTALE A ESTIMER** = Zone B + Zone C + Garage + Sous-sol touche = **X pi²**

Cas particuliers:
- **Projet 100% neuf** (construction neuve): Zone A = 0, Zone B = 0,
  Zone C = totalite = superficie a estimer
- **Projet 100% renovation sans agrandissement**: Zone C = 0, estimation
  sur Zone B uniquement au tarif renovation
- **Superficies non lisibles au plan**: DEMANDE explicitement au client
  "Pouvez-vous preciser la superficie de l'agrandissement / renovation?"
  plutot que d'utiliser la superficie habitable totale.

**Tarification differenciee** (appliquer aux bonnes zones):
- Zone B (Renovation): tarif "Renovation residentielle" de la gamme
- Zone C (Agrandissement): tarif "Residentiel neuf" de la gamme
- Garage: tarif reduit (~150 $/pi² economique, +30% par gamme superieure)

### 4. Particularites techniques (extraites des coupes et elevations)
Source obligatoire: mentionne la page/vue pour chaque element.

- **Contour exterieur du batiment**: forme globale (rectangulaire, L, T, U,
  avec decroches), perimetre approximatif en pi. lin.
- **Fondation** (coupes): type (pieux visses, dalle beton, mur coule,
  blocs, pilotis), hauteur hors-sol, isolation perimetrique R-X
- **Structure** (plan + coupes): 2x4 vs 2x6, espacement c/c, type de
  poutres (LVL, microlam, acier), portees, murs porteurs vs cloisons
- **Toiture** (elevations + coupes): type (pente / plate / combinee),
  pentes cotees (ex: 4/12, 8/12), materiau (bardeaux, tole, membrane,
  ardoise), debords de toit
- **Enveloppe** (elevations): revetement principal par facade (brique,
  vinyle, bois, pierre, crepi, lattes aluminium), materiaux differents
  par facade si applicable
- **Ouvertures** (plan + elevations): compte approximatif portes et
  fenetres par facade, types (guillotine, auvent, battant, fixe, patio)
- **Hauteur totale du batiment** (elevations): hauteur faitage / sol,
  verifier conformite vs max permis municipal
- **Orientation** (plan d'implantation): facade principale vers quel
  point cardinal
- **Autres elements notables**: galerie, balcon, avant-toit, garage
  attache/detache, piscine, cabanon

### 5. Corps de metier identifies
Liste les corps de metier necessaires **applicables specifiquement a la categorie detectee**
(ex: si Renovation, mentionne demolition; si Commercial, mentionne systemes speciaux).

### 6. Contingences recommandees
- **Neuf**: 12% (meilleure previsibilite)
- **Renovation**: 15% (imprevus chantier existant)
- **Institutionnel / commercial complexe**: 15-20%
Precise le taux applicable.

### 7. Points d'attention
- Permis municipal / zonage a verifier
- Normes specifiques (Novoclimat, LEED, CNB, CCQ, RBQ)
- Risques particuliers (decontamination, structure existante, sol, etc.)
- Documents complementaires requis

### 8. Categorie JSON (pour traitement automatique)
Termine OBLIGATOIREMENT ta reponse par un bloc JSON sur une seule ligne,
commencant par `<CATEGORY>` et finissant par `</CATEGORY>`. La gamme
DOIT TOUJOURS etre "Economique" au diagnostic initial (regle absolue).

`superficie_pi2` = **SUPERFICIE A ESTIMER** uniquement (Zone B + Zone C +
Garage + Sous-sol touche). Pour un projet mixte renovation + agrandissement,
decompose aussi dans les champs separes.

Exemples:
- Construction 100% neuve (unifamiliale 2500 pi²):
  `<CATEGORY>{"categorie":"Residentiel neuf","gamme":"Economique","superficie_pi2":2500,"superficie_renovation_pi2":0,"superficie_agrandissement_pi2":2500,"superficie_existant_conserve_pi2":0}</CATEGORY>`
- Renovation + agrandissement (comme projet Papineau — total habitable 3196,
  existant conserve 725, renovation 1057, agrandissement 1414, ne pas
  compter le garage dans superficie_pi2 si tarife separement):
  `<CATEGORY>{"categorie":"Renovation residentielle","gamme":"Economique","superficie_pi2":2471,"superficie_renovation_pi2":1057,"superficie_agrandissement_pi2":1414,"superficie_existant_conserve_pi2":725}</CATEGORY>`
- Renovation pure (reno cuisine 200 pi² dans une maison 2000 pi²):
  `<CATEGORY>{"categorie":"Renovation residentielle","gamme":"Economique","superficie_pi2":200,"superficie_renovation_pi2":200,"superficie_agrandissement_pi2":0,"superficie_existant_conserve_pi2":1800}</CATEGORY>`

A la fin de ta reponse (avant le bloc CATEGORY), pose la question:
**"Voulez-vous que je prepare une estimation detaillee par corps de metier
au prix Economique par defaut, ou preferez-vous d'abord ajuster la gamme
(De base / Moyenne / Haut de gamme) selon les finitions interieures et
exterieures souhaitees?"**
"""


def _parse_category_tag(text: str) -> dict:
    """Extract the <CATEGORY>{...}</CATEGORY> JSON tag from Claude output."""
    m = re.search(r"<CATEGORY>\s*(\{.*?\})\s*</CATEGORY>", text, re.DOTALL)
    if not m:
        return {}
    try:
        return json.loads(m.group(1))
    except (json.JSONDecodeError, TypeError):
        return {}


def _strip_category_tag(text: str) -> str:
    """Remove the <CATEGORY>...</CATEGORY> tag from displayed text."""
    return re.sub(r"\s*<CATEGORY>.*?</CATEGORY>\s*", "\n", text, flags=re.DOTALL).strip()


# ============================================================
# BUG A FIX (2026-05-17): anti-hallucination superficies
# ----------------------------------------------------------------
# Observation: le meme plan Maison Papineau-Leduc analyse 8 fois en 8 jours
# a produit des superficies variant de 2087 a 3552 pi² (variation 70%) alors
# que les chiffres ecrits NOIRS sur BLANCS sur le plan sont stables.
#
# Cause: Claude lit les chiffres visuellement et hallucine. Le texte OCR du
# PDF (pdfplumber) contient pourtant les chiffres exacts en clair.
#
# Fix: extraire deterministiquement les superficies du texte OCR via regex
# whitelist (labels connus du domaine), puis les injecter dans le prompt
# comme GROUND TRUTH avant l'analyse. Claude doit privilegier ces valeurs
# sur son estimation visuelle et flag les contradictions inter-pages
# explicitement.
# ============================================================

# Capture une valeur de superficie avec ses unites quebecoises courantes:
#   - "1234,5 pi²" / "1234 pi2"           (standard moderne)
#   - "1497 P.C."  / "1497 p.c."           (pieds carres, abbreviation tres
#                                           courante chez les dessinateurs QC)
#   - "1497P.C."   / "1497p.c."            (colle, ex. Plan2 "TOTAL = 2994P.C.")
#   - "1497 pi.ca." / "1497 pi ca"         (variante rare)
# Filtre value < 50 ou > 50000 pour ecarter les ID, les hauteurs (pieds
# lineaires sans pi²) et les artefacts OCR. (?i) = case insensitive sur
# tout le pattern.
_SUPERFICIE_VALUE_RE = re.compile(
    r"(?i)"
    r"(\d{2,5}(?:[\s,\.]\d{1,3})?)"  # capture numerique avec decimale optionnelle
    r"\s*"
    r"(?:"
    r"pi\s*[²²2]"                    # pi² ou pi2 ou pi 2
    r"|"
    r"p\s*\.?\s*c\s*\.?"             # P.C. / p.c. / P C / pc / p . c .
    r"|"
    r"pi\s*\.\s*ca\.?"               # pi.ca. (rare)
    r")"
    r"(?=\s|$|[^A-Za-z])"           # boundary: evite de matcher "pcQ" / "pcontact"
)
# Labels metier qu'on whitelist pour identifier ce a quoi la valeur se rapporte.
# IMPORTANT: ordre n'a pas d'importance — on prend le PLUS PROCHE de la valeur
# dans le contexte (60 chars avant).
#
# Note typo: certains plans QC ecrivent "SUPPERFICIE" (double P) — on tolere
# via SUPP?ERFICIE. Observe sur Plan2.pdf (Danielle Poitras, dessinatrice).
_SUPERFICIE_LABEL_RE = re.compile(
    r"(?i)\b("
    r"REZ[\s\-]?DE[\s\-]?CHAUSS[ÉéE]E|"  # rez-de-chaussee
    r"RDC|RC|"
    r"[ÉéE]TAGE|"  # etage
    r"EXISTANT|AJOUT|AGRANDISSEMENT|GARAGE|"
    r"SOUS[\s\-]?SOL|TOTAL|"
    r"AIRE\s+HABITABLE|SUPP?ERFICIE|"  # tolere typo SUPPERFICIE
    r"NEUF|PROPOS[ÉéE]E?|CONSERV[ÉéE]E?|HABITABLE"
    r")\b"
)


def _extract_known_superficies(text: str) -> list:
    """Extract `[{label, value_pi2, context}, ...]` from PDF OCR text.

    Strategy: scan all "XXXX pi²" matches, then walk back ~80 chars to find
    a whitelisted label (RC, Étage, Existant, Ajout, etc.). Dedup by
    (label, value) to avoid 4x duplication when the same chiffre apparait
    dans plusieurs vues.

    Conservative: ne retourne RIEN si pas de label valide a proximite
    (eviter de polluer le prompt avec des chiffres ambigus).
    """
    if not text:
        return []

    results = []
    seen = set()
    for m in _SUPERFICIE_VALUE_RE.finditer(text):
        raw_value = m.group(1).replace(",", ".").replace(" ", "")
        try:
            value = float(raw_value)
        except ValueError:
            continue
        # Filter aberrantes (cotes en pieds lineaires, ID, etc.)
        if value < 50 or value > 50000:
            continue

        # Walk back ~80 chars to find a known label
        start_ctx = max(0, m.start() - 80)
        ctx = text[start_ctx:m.start()].replace("\r", " ").replace("\n", " ")
        label_matches = list(_SUPERFICIE_LABEL_RE.finditer(ctx))
        if not label_matches:
            continue
        # Best label = closest to the value (last in ctx)
        best = label_matches[-1]
        label_norm = best.group(0).upper()

        # Capture phrase de contexte (label + 15 chars before pour qualificatifs
        # type "Existant RC", "Étage existant"). Trim a 50 chars max.
        ext_start = max(0, best.start() - 20)
        phrase = ctx[ext_start:].strip()
        if len(phrase) > 50:
            phrase = phrase[-50:].strip()

        key = (label_norm, round(value, 1))
        if key in seen:
            continue
        seen.add(key)

        results.append({
            "label": label_norm,
            "value_pi2": value,
            "context": phrase or label_norm,
        })

    return results


def _format_superficies_for_prompt(superficies: list) -> str:
    """Format extracted superficies as a markdown block for the Claude prompt.

    Vide si la liste est vide — on ne pollue pas le prompt avec un bloc inutile.
    """
    if not superficies:
        return ""
    lines = [
        "=== SUPERFICIES EXTRAITES TEXTUELLEMENT DU PDF (ground truth deterministe) ===",
    ]
    for s in superficies:
        lines.append(f"- \"{s['context']}\" => **{s['value_pi2']:.1f} pi²**")
    lines.append("")
    lines.append(
        "IMPORTANT: Ces valeurs sont extraites par regex sur le texte OCR du PDF "
        "(deterministe et reproductible). UTILISE-LES exactement plutot que "
        "d'estimer visuellement. Si tu detectes des contradictions inter-pages "
        "(ex: une meme zone avec 2 valeurs differentes), FLAG-LES explicitement "
        "et indique la page source de chaque chiffre."
    )
    lines.append("=== FIN GROUND TRUTH SUPERFICIES ===")
    return "\n".join(lines)


@router.post("/ai-analyze-document")
async def ai_analyze_document(
    file: UploadFile = File(...),
    conversation_id: Optional[int] = Form(None),
    user: ErpUser = Depends(get_current_user),
):
    """Analyze an uploaded document (plan, devis, cahier des charges) as an
    Entrepreneur general: detect category (Residentiel neuf / Renovation / Commercial /
    Institutionnel), estimate gamme, extract superficies, list corps de metier.

    Le document est persiste dans `conversation_documents` pour consultation
    ulterieure (continuation indefinie de la conversation). Prompt caching
    Anthropic active (5 min TTL) pour reduire les couts sur les messages suivants.
    """
    if not _anthropic_client:
        raise HTTPException(status_code=503, detail="Service IA non disponible (client Anthropic non configure)")

    # AI billing: guard + credit check
    allowed, error_msg = check_ai_guard(user)
    if not allowed:
        raise HTTPException(status_code=403, detail=error_msg or "Acces IA refuse")
    credits_ok, balance = _check_credits(user)
    if not credits_ok:
        raise HTTPException(status_code=402, detail="Credits IA epuises. Veuillez recharger votre solde.")

    try:
        file_bytes = await file.read()
        # FIX: garde explicit fichier vide pour eviter envoi b64="" a Anthropic
        # qui retourne un 400 obscur "image specified but appears to be invalid".
        if not file_bytes:
            raise HTTPException(status_code=400, detail="Fichier vide ou impossible a lire")
        if len(file_bytes) > 32 * 1024 * 1024:  # 32 MB — limite Anthropic payload
            raise HTTPException(status_code=413, detail="Fichier trop volumineux (max 32 Mo — limite API Anthropic)")
        filename = file.filename or "document"
        file_ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        media_type_map = {
            "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
            "gif": "image/gif", "webp": "image/webp", "pdf": "application/pdf",
        }
        media_type = media_type_map.get(file_ext, None)

        # Pre-validation PDF: count pages avant envoi Anthropic pour eviter 413
        # request_too_large + facturation IA partielle (le client paye le payload
        # parti meme si Anthropic rejette). Au-dela de 100 pages on tronque
        # frequemment la reponse + le cout monte vite (chaque page PDF = ~1500
        # tokens output Vision sur Opus 4.7). Le user obtient un message UX
        # clair plutot qu'un 413 obscur a la fin du round-trip.
        if media_type == "application/pdf":
            try:
                import fitz  # PyMuPDF, deja installe pour metre_pdf.py
                pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
                page_count = pdf_doc.page_count
                pdf_doc.close()
                if page_count > 100:
                    raise HTTPException(
                        status_code=413,
                        detail=(
                            f"PDF trop volumineux: {page_count} pages (max 100). "
                            "Reduisez le nombre de pages ou envoyez en plusieurs parties."
                        ),
                    )
            except ImportError:
                logger.warning("PyMuPDF (fitz) not installed - PDF page count skipped")
            except HTTPException:
                raise
            except Exception as pdf_err:
                # PDF corrompu / chiffre / format inattendu: laisser passer et
                # qu'Anthropic le rejette avec son propre message (deja gere
                # par le catch APIError plus bas).
                logger.warning("PDF page count failed for %s: %s", filename, pdf_err)

        # Compression image si > 4.5 MB (limit Claude 5 MB)
        # Opus 4.7 native res = 2576 px sur le long edge (4784 tokens). Compresser
        # plus bas (ex: 2048) gaspille de la resolution alors que le modele peut
        # accepter 2576 sans resize interne. Plus de pixels = meilleure analyse
        # de plans/devis. Ref: doc Anthropic Vision (Opus 4.7).
        if media_type and media_type.startswith("image/") and len(file_bytes) > 4.5 * 1024 * 1024:
            try:
                from PIL import Image
                import io
                img = Image.open(io.BytesIO(file_bytes))
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                max_dim = 2576
                if max(img.size) > max_dim:
                    img.thumbnail((max_dim, max_dim), Image.LANCZOS)
                buf = io.BytesIO()
                img.save(buf, format='JPEG', quality=80, optimize=True)
                file_bytes = buf.getvalue()
                media_type = "image/jpeg"
            except Exception as resize_err:
                logger.warning("Image resize failed: %s — sending original", resize_err)
        if media_type and media_type.startswith("image/") and len(file_bytes) > 5 * 1024 * 1024:
            raise HTTPException(
                status_code=400,
                detail=f"Image trop volumineuse (max 5 Mo apres compression). Fichier: {filename}, taille: {len(file_bytes) / (1024*1024):.1f} Mo",
            )

        # Charge le profil Entrepreneur general comme system prompt
        # (le profil contient 40 ans de connaissance + les 5 categories + baremes $/pi²)
        # Split today_line (non cacheable) du profil (cacheable TTL 1h) pour eviter
        # cache miss quotidien a minuit.
        entrepreneur_prompt = _load_profile_system_prompt("ENTREPRENEUR_GENERAL", user.schema)
        system_blocks = [
            # Bloc 1: today_line seul (change chaque jour, non cacheable)
            {"type": "text", "text": _today_prompt_line()},
            # Bloc 2: profil Entrepreneur general + markdown rules (cacheable TTL 1h)
            # Le profil fait ~130 KB donc > 4096 tokens min cacheable pour Opus 4.7.
            {
                "type": "text",
                "text": entrepreneur_prompt + _MARKDOWN_TABLE_RULES,
                "cache_control": {"type": "ephemeral", "ttl": "1h"},
            },
        ]

        # Construit le content block selon le format
        if media_type and media_type.startswith("image/"):
            # Detection magic bytes: si l'extension/Content-Type ment, utiliser
            # le type reel des bytes pour eviter erreur 400 Anthropic.
            actual_media_type = _detect_media_type_from_bytes(file_bytes) or media_type
            b64_data = base64.standard_b64encode(file_bytes).decode("utf-8")
            user_content = [
                {
                    "type": "image",
                    "source": {"type": "base64", "media_type": actual_media_type, "data": b64_data},
                    "cache_control": {"type": "ephemeral"},
                },
                {"type": "text", "text": _DIAGNOSTIC_PROMPT},
            ]
            extracted_text = ""  # pas d'extraction pour images
        elif media_type == "application/pdf":
            b64_data = base64.standard_b64encode(file_bytes).decode("utf-8")
            # Extraction texte du PDF pour stockage (evite re-OCR plus tard) ET
            # ground truth deterministe des superficies (BUG A fix anti-hallucination).
            extracted_text = _extract_text_from_file(file_bytes, "application/pdf", filename)
            sup_block = _format_superficies_for_prompt(
                _extract_known_superficies(extracted_text or "")
            )
            user_content = [
                {
                    "type": "document",
                    "source": {"type": "base64", "media_type": "application/pdf", "data": b64_data},
                    "cache_control": {"type": "ephemeral"},
                    "citations": {"enabled": True},
                },
            ]
            if sup_block:
                # Injecte le ground truth AVANT le prompt pour que Claude le voie
                # comme contexte autoritaire. Le PDF reste cacheable; ce bloc texte
                # change peu pour un meme PDF (regex deterministe).
                user_content.append({"type": "text", "text": sup_block})
                logger.info(
                    "BUG A fix: %d superficie(s) extracted deterministically for %s",
                    sup_block.count("\n- "), filename,
                )
            user_content.append({"type": "text", "text": _DIAGNOSTIC_PROMPT})
        else:
            # xlsx / docx / txt / csv — extraction texte + envoi texte
            extracted_text = _extract_text_from_file(file_bytes, "", filename)
            if not extracted_text.strip():
                raise HTTPException(
                    status_code=400,
                    detail=f"Format non supporte ou contenu vide: {filename}. Formats supportes: PDF, PNG, JPG, XLSX, DOCX, CSV, TXT.",
                )
            # BUG A fix: si le doc textuel mentionne des superficies, ground truth.
            sup_block = _format_superficies_for_prompt(
                _extract_known_superficies(extracted_text)
            )
            prompt_prefix = f"{sup_block}\n\n" if sup_block else ""
            user_content = [{
                "type": "text",
                "text": f"{prompt_prefix}{_DIAGNOSTIC_PROMPT}\n\n=== Contenu du document ({filename}) ===\n{extracted_text}",
            }]

        message = _call_claude(
            model="claude-opus-4-7", max_tokens=32000,
            system=system_blocks,
            messages=[{"role": "user", "content": user_content}],
        )

        # Sprint 3 #10 - Format reponse avec footnotes citations PDF
        ai_text = _format_response_with_footnotes(message, document_titles=[filename])
        citations_list = _extract_citations_from_response(message)

        # Extraction du tag <CATEGORY> pour metadata BD
        category_data = _parse_category_tag(ai_text)
        display_text = _strip_category_tag(ai_text)

        # AI billing: track usage + deduct credits AVANT la persistence doc
        # pour atomicite: si billing crashe, le doc n'est pas persiste et
        # l'user recoit une erreur nette plutot qu'un doc orphelin non-facture.
        tokens_in = message.usage.input_tokens
        tokens_out = message.usage.output_tokens
        cache_creation = getattr(message.usage, "cache_creation_input_tokens", 0) or 0
        cache_read = getattr(message.usage, "cache_read_input_tokens", 0) or 0
        # Opus 4.7: $15/M input, $75/M output, $18.75/M cache write, $1.50/M cache read
        cost = (
            tokens_in * 15 / 1_000_000
            + tokens_out * 75 / 1_000_000
            + cache_creation * 18.75 / 1_000_000
            + cache_read * 1.50 / 1_000_000
        ) * 1.30  # 30% markup
        # BUG G fix (2026-05-17): passe model explicitement. Sans ca, track_ai_usage
        # logge AI_MODEL (claude-sonnet-4-6) alors que le call _call_claude utilise
        # claude-opus-4-7 — desync entre tracking et realite (analyse couts erronee).
        track_ai_usage(user, "devis_analyze_document", tokens_in, tokens_out, cost, 0, True, model="claude-opus-4-7")
        _deduct_credits(user, cost)

        # Helper pour caster les superficies en INT de maniere safe
        def _safe_int(v):
            try:
                if v is None:
                    return None
                return int(float(v))
            except (TypeError, ValueError):
                return None

        # Persistence du document dans conversation_documents (apres billing)
        document_id = None
        if user.schema:
            conn_doc = db.get_conn()
            try:
                db.set_tenant(conn_doc, user.schema)
                _ensure_conversation_documents_table(conn_doc, user.schema)
                cur_doc = conn_doc.cursor()
                cur_doc.execute(
                    "INSERT INTO conversation_documents "
                    "(conversation_id, filename, media_type, file_size, content, "
                    "extracted_text, summary, category_detected, subcategory_detected, "
                    "superficie_pi2, superficie_renovation_pi2, superficie_agrandissement_pi2, "
                    "superficie_existant_conserve_pi2, "
                    "is_active_context, uploaded_by, created_at, updated_at) "
                    "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,TRUE,%s,NOW(),NOW()) RETURNING id",
                    (
                        conversation_id,
                        filename,
                        media_type or "application/octet-stream",
                        len(file_bytes),
                        psycopg2.Binary(file_bytes) if (media_type and (media_type.startswith("image/") or media_type == "application/pdf")) else None,
                        extracted_text[:500000] if extracted_text else None,  # Cap 500K chars
                        display_text[:50000],  # Cap summary a 50K chars
                        category_data.get("categorie"),
                        category_data.get("gamme"),
                        _safe_int(category_data.get("superficie_pi2")),
                        _safe_int(category_data.get("superficie_renovation_pi2")),
                        _safe_int(category_data.get("superficie_agrandissement_pi2")),
                        _safe_int(category_data.get("superficie_existant_conserve_pi2")),
                        user.user_id,
                    ),
                )
                row = cur_doc.fetchone()
                document_id = row["id"] if row else None
                conn_doc.commit()
                cur_doc.close()

                # OPT-2: upload best-effort vers Files API beta pour reutilisation
                # multi-tour. N'echoue pas le endpoint si l'upload casse — la row
                # est deja persistee avec content (BYTEA), le fallback base64
                # fonctionnera. On met a jour anthropic_file_id apres coup pour
                # que les tours suivants utilisent la source file_id (gros gain
                # bytes/tokens). Seulement images et PDF — pas les xlsx/docx/csv
                # qui sont stockes en extracted_text.
                # BUG F fix (2026-05-17): retire la condition `conversation_id`. Le
                # premier upload (avant /conversations) avait conversation_id=NULL,
                # ce qui skippait l'upload Files API. Resultat: 8 analyses du PDF
                # Maison Papineau-Leduc en 8 jours avec anthropic_file_id NULL,
                # forcant la reupload base64 4.9 MB a chaque tour. Maintenant on
                # upload le file_id meme si la row est orpheline; il sera lie a
                # la conv plus tard via save_conversation (UPDATE orphan window 1h).
                if document_id and media_type and (
                    media_type.startswith("image/") or media_type == "application/pdf"
                ):
                    file_id = _upload_to_anthropic_files(file_bytes, filename, media_type)
                    if file_id:
                        try:
                            cur_up = conn_doc.cursor()
                            cur_up.execute(
                                "UPDATE conversation_documents SET anthropic_file_id = %s WHERE id = %s",
                                (file_id, document_id),
                            )
                            conn_doc.commit()
                            cur_up.close()
                        except Exception as up_exc:
                            logger.warning("UPDATE anthropic_file_id failed for doc %s: %s", document_id, up_exc)
                            try:
                                conn_doc.rollback()
                            except Exception:
                                pass
            except Exception as persist_exc:
                logger.warning("persist conversation_document failed: %s", persist_exc)
                try:
                    conn_doc.rollback()
                except Exception:
                    pass
            finally:
                try:
                    db.reset_tenant(conn_doc)
                except Exception:
                    pass
                conn_doc.close()

        return {
            "summary": display_text,
            "filename": filename,
            "document_id": document_id,
            "category": category_data.get("categorie"),
            "subcategory": category_data.get("gamme"),
            "superficie_pi2": _safe_int(category_data.get("superficie_pi2")),
            "superficie_renovation_pi2": _safe_int(category_data.get("superficie_renovation_pi2")),
            "superficie_agrandissement_pi2": _safe_int(category_data.get("superficie_agrandissement_pi2")),
            "superficie_existant_conserve_pi2": _safe_int(category_data.get("superficie_existant_conserve_pi2")),
            "citations": citations_list,
            "usage": {
                "input_tokens": tokens_in,
                "output_tokens": tokens_out,
                "cache_creation_input_tokens": cache_creation,
                "cache_read_input_tokens": cache_read,
                "cost_usd": round(cost, 6),
            },
        }

    except HTTPException:
        raise
    except anthropic.APIError as exc:
        logger.error("ai_analyze_document API error: %s", exc)
        status = getattr(exc, "status_code", 0)
        exc_str = str(exc).lower()
        if status == 413 or (status == 400 and "too_large" in exc_str):
            raise HTTPException(
                status_code=413,
                detail="Document trop volumineux pour l'analyse IA. Essayez un fichier plus petit ou un PDF avec moins de pages.",
            )
        if status == 400 and "image exceeds" in exc_str:
            raise HTTPException(
                status_code=413,
                detail="Image trop volumineuse (max 5 Mo). Reduisez la taille de l'image ou utilisez un format compresse (JPEG).",
            )
        if status == 529 or "overloaded" in exc_str:
            raise HTTPException(status_code=503, detail="Le service IA est temporairement surcharge. Veuillez reessayer dans quelques secondes.")
        if status == 429 or "rate_limit" in exc_str:
            raise HTTPException(status_code=429, detail="Trop de demandes vers l'IA. Attendez quelques instants avant de reessayer.")
        logger.error("Erreur API Claude: %s", exc)
        raise HTTPException(status_code=503, detail="Service IA temporairement indisponible")
    except Exception as exc:
        logger.error("ai_analyze_document error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de l'analyse du document")


# ============================================================
# AI PROFILES — Custom profiles with knowledge base
# ============================================================


# Protection ZIP bomb sur fichiers Office (xlsx, docx — conteneurs ZIP).
# Limites: 200 MB d'expansion totale, ratio max 100x (compressé→décompressé),
# 1000 entrées max dans le ZIP.
_ZIP_MAX_TOTAL_UNCOMPRESSED = 200 * 1024 * 1024  # 200 MB
_ZIP_MAX_COMPRESSION_RATIO = 100
_ZIP_MAX_ENTRIES = 1000


def _validate_office_zip_safe(content: bytes) -> None:
    """Vérifie qu'un fichier Office (xlsx/docx) n'est pas une ZIP bomb.

    Lève ValueError si suspect. Appelle AVANT de passer le contenu à
    openpyxl/python-docx pour éviter une expansion catastrophique.
    """
    import io
    import zipfile

    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            entries = zf.infolist()
            if len(entries) > _ZIP_MAX_ENTRIES:
                raise ValueError(f"Trop d'entrées dans le ZIP ({len(entries)} > {_ZIP_MAX_ENTRIES})")
            total_uncompressed = sum(e.file_size for e in entries)
            if total_uncompressed > _ZIP_MAX_TOTAL_UNCOMPRESSED:
                raise ValueError(
                    f"Décompression > {_ZIP_MAX_TOTAL_UNCOMPRESSED // (1024*1024)} MB refusée"
                )
            total_compressed = sum(e.compress_size for e in entries) or 1
            ratio = total_uncompressed / total_compressed
            if ratio > _ZIP_MAX_COMPRESSION_RATIO:
                raise ValueError(
                    f"Ratio de compression suspect ({ratio:.0f}x > {_ZIP_MAX_COMPRESSION_RATIO}x)"
                )
    except zipfile.BadZipFile:
        raise ValueError("Fichier ZIP invalide")


def _extract_text_from_file(content: bytes, content_type: str, filename: str) -> str:
    """Extract plain text from uploaded file for knowledge base injection."""
    lower = filename.lower()

    # Plain text / CSV
    if lower.endswith((".txt", ".csv", ".tsv", ".md")):
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return content.decode(enc)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="replace")

    # PDF via pdfplumber
    if lower.endswith(".pdf"):
        try:
            import pdfplumber
            import io
            text_parts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            if text_parts:
                return "\n\n".join(text_parts)
        except Exception as exc:
            logger.warning("pdfplumber extraction failed for %s: %s", filename, exc)
        return ""

    # XLSX via openpyxl
    if lower.endswith(".xlsx"):
        try:
            _validate_office_zip_safe(content)
            import openpyxl
            import io
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            lines = []
            for ws in wb.worksheets:
                lines.append(f"--- {ws.title} ---")
                for row in ws.iter_rows(values_only=True):
                    vals = [str(c) if c is not None else "" for c in row]
                    lines.append("\t".join(vals))
            wb.close()
            return "\n".join(lines)
        except ValueError as exc:
            logger.warning("xlsx rejected (zip-bomb guard) for %s: %s", filename, exc)
        except Exception as exc:
            logger.warning("openpyxl extraction failed for %s: %s", filename, exc)
        return ""

    # DOCX via python-docx
    if lower.endswith(".docx"):
        try:
            _validate_office_zip_safe(content)
            import docx
            import io
            doc = docx.Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text)
        except ValueError as exc:
            logger.warning("docx rejected (zip-bomb guard) for %s: %s", filename, exc)
        except Exception as exc:
            logger.warning("docx extraction failed for %s: %s", filename, exc)
        return ""

    return ""


@router.get("/ai-profiles")
async def list_ai_profiles(user: ErpUser = Depends(get_current_user)):
    """List custom AI profiles for the tenant."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Schema requis")
    conn = db.get_conn()
    try:
        db.set_tenant(conn, user.schema)
        _ensure_ai_profiles_tables(conn, user.schema)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT p.id, p.name, p.instructions, p.is_active, p.created_by,
                   p.created_at, p.updated_at,
                   COUNT(d.id) AS document_count
            FROM ai_profiles p
            LEFT JOIN ai_profile_documents d ON d.profile_id = p.id
            GROUP BY p.id
            ORDER BY p.name
        """)
        items = [dict(r) for r in cursor.fetchall()]
        cursor.close()
        return {"items": items}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("list_ai_profiles error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors du chargement des profils")
    finally:
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.post("/ai-profiles")
async def create_ai_profile(body: AiProfileCreate, user: ErpUser = Depends(get_current_user)):
    """Create a custom AI profile."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Schema requis")
    conn = db.get_conn()
    try:
        db.set_tenant(conn, user.schema)
        _ensure_ai_profiles_tables(conn, user.schema)
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO ai_profiles (name, instructions, created_by)
            VALUES (%s, %s, %s) RETURNING id
        """, (body.name, body.instructions, user.user_id))
        row = cursor.fetchone()
        conn.commit()
        cursor.close()
        return {"id": row["id"], "message": "Profil cree"}
    except HTTPException:
        raise
    except Exception as exc:
        conn.rollback()
        logger.error("create_ai_profile error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de la creation du profil")
    finally:
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.get("/ai-profiles/{profile_id}")
async def get_ai_profile(profile_id: int, user: ErpUser = Depends(get_current_user)):
    """Get a custom AI profile with its documents."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Schema requis")
    conn = db.get_conn()
    try:
        db.set_tenant(conn, user.schema)
        _ensure_ai_profiles_tables(conn, user.schema)
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM ai_profiles WHERE id = %s", (profile_id,))
        profile = cursor.fetchone()
        if not profile:
            raise HTTPException(status_code=404, detail="Profil non trouve")
        profile = dict(profile)
        cursor.execute("""
            SELECT id, original_name, content_type, file_size,
                   LENGTH(extracted_text) AS extracted_text_length,
                   created_at
            FROM ai_profile_documents WHERE profile_id = %s ORDER BY created_at
        """, (profile_id,))
        profile["documents"] = [dict(r) for r in cursor.fetchall()]
        cursor.close()
        return profile
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("get_ai_profile error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors du chargement du profil")
    finally:
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.put("/ai-profiles/{profile_id}")
async def update_ai_profile(profile_id: int, body: AiProfileUpdate, user: ErpUser = Depends(get_current_user)):
    """Update a custom AI profile."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Schema requis")
    conn = db.get_conn()
    try:
        db.set_tenant(conn, user.schema)
        _ensure_ai_profiles_tables(conn, user.schema)
        cursor = conn.cursor()
        sets, vals = [], []
        if body.name is not None:
            sets.append("name = %s")
            vals.append(body.name)
        if body.instructions is not None:
            sets.append("instructions = %s")
            vals.append(body.instructions)
        if body.is_active is not None:
            sets.append("is_active = %s")
            vals.append(body.is_active)
        if not sets:
            raise HTTPException(status_code=400, detail="Aucun champ a modifier")
        sets.append("updated_at = CURRENT_TIMESTAMP")
        vals.append(profile_id)
        cursor.execute(f"UPDATE ai_profiles SET {', '.join(sets)} WHERE id = %s", vals)
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="Profil non trouve")
        conn.commit()
        cursor.close()
        return {"message": "Profil mis a jour"}
    except HTTPException:
        raise
    except Exception as exc:
        conn.rollback()
        logger.error("update_ai_profile error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de la mise a jour du profil")
    finally:
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.delete("/ai-profiles/{profile_id}")
async def delete_ai_profile(profile_id: int, user: ErpUser = Depends(get_current_user)):
    """Delete a custom AI profile (cascade deletes documents)."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Schema requis")
    conn = db.get_conn()
    try:
        db.set_tenant(conn, user.schema)
        _ensure_ai_profiles_tables(conn, user.schema)
        cursor = conn.cursor()
        cursor.execute("DELETE FROM ai_profiles WHERE id = %s", (profile_id,))
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="Profil non trouve")
        conn.commit()
        cursor.close()
        return {"message": "Profil supprime"}
    except HTTPException:
        raise
    except Exception as exc:
        conn.rollback()
        logger.error("delete_ai_profile error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de la suppression du profil")
    finally:
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.post("/ai-profiles/{profile_id}/documents")
async def upload_profile_document(
    profile_id: int,
    file: UploadFile = File(...),
    user: ErpUser = Depends(get_current_user),
):
    """Upload a knowledge base document to a custom AI profile."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Schema requis")
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="Fichier trop volumineux (max 20 Mo)")

    extracted = _extract_text_from_file(content, file.content_type or "", file.filename or "doc")
    if not extracted:
        logger.warning("No text extracted from %s — storing file without text", file.filename)

    conn = db.get_conn()
    try:
        db.set_tenant(conn, user.schema)
        _ensure_ai_profiles_tables(conn, user.schema)
        cursor = conn.cursor()
        # Verify profile exists
        cursor.execute("SELECT id FROM ai_profiles WHERE id = %s", (profile_id,))
        if not cursor.fetchone():
            raise HTTPException(status_code=404, detail="Profil non trouve")
        cursor.execute("""
            INSERT INTO ai_profile_documents
                (profile_id, original_name, content_type, file_size, extracted_text, uploaded_by)
            VALUES (%s, %s, %s, %s, %s, %s) RETURNING id
        """, (profile_id, file.filename, file.content_type, len(content), extracted, user.user_id))
        doc_id = cursor.fetchone()["id"]
        conn.commit()
        cursor.close()
        return {
            "id": doc_id,
            "original_name": file.filename,
            "file_size": len(content),
            "extracted_text_length": len(extracted) if extracted else 0,
            "message": "Document ajoute",
        }
    except HTTPException:
        raise
    except Exception as exc:
        conn.rollback()
        logger.error("upload_profile_document error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de l'upload du document")
    finally:
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.delete("/ai-profiles/{profile_id}/documents/{doc_id}")
async def delete_profile_document(
    profile_id: int, doc_id: int, user: ErpUser = Depends(get_current_user)
):
    """Remove a knowledge base document from a custom AI profile."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Schema requis")
    conn = db.get_conn()
    try:
        db.set_tenant(conn, user.schema)
        _ensure_ai_profiles_tables(conn, user.schema)
        cursor = conn.cursor()
        cursor.execute(
            "DELETE FROM ai_profile_documents WHERE id = %s AND profile_id = %s",
            (doc_id, profile_id),
        )
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="Document non trouve")
        conn.commit()
        cursor.close()
        return {"message": "Document supprime"}
    except HTTPException:
        raise
    except Exception as exc:
        conn.rollback()
        logger.error("delete_profile_document error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de la suppression du document")
    finally:
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


_MARKDOWN_TABLE_RULES = """

=== REGLES STRICTES POUR LES TABLEAUX MARKDOWN ===
Quand tu inclus un tableau markdown dans ta reponse, tu DOIS respecter ces regles absolues:

1. Chaque ligne du tableau (header, separator, data, totaux) DOIT contenir EXACTEMENT le meme nombre de |.
2. Les lignes TOTAL, SOUS-TOTAL et GRAND TOTAL DOIVENT avoir le meme nombre de colonnes que le header — ne jamais retirer de colonnes.
3. INTERDIT de laisser une cellule vide au milieu du tableau : ecris "-" ou "n/a" si une valeur n'est pas applicable.
4. Separateur standard obligatoire: |---|---|---| avec autant de |---| que de colonnes.
5. Pas de colspan, pas de fusion, pas de lignes manquantes.
6. Exemples valides:
   | # | CORPS DE METIER | $/pi² | TOTAL | M.O. | MAT. |
   |---|---|---|---|---|---|
   | 1 | Excavation | 37,36 | 53 055 $ | 19 059 $ | 34 006 $ |
   | TOTAL | - | 206,43 | 293 138 $ | 138 319 $ | 189 752 $ |

Le rendu frontend depend de ces regles: toute deviation casse l'alignement visuel des colonnes.
"""


def _load_profile_system_prompt(profile_id: Optional[str], schema: Optional[str] = None) -> str:
    """Load system prompt for a profile ID. Handles system (disk) and custom (DB) profiles."""
    default = "Tu es un expert en construction au Quebec. Reponds en francais."
    if not profile_id:
        return default

    # Custom profile: load from tenant DB
    if profile_id.startswith("custom_"):
        try:
            db_id = int(profile_id.split("_", 1)[1])
        except (ValueError, IndexError):
            return default
        if not schema:
            return default
        conn = db.get_conn()
        try:
            db.set_tenant(conn, schema)
            _ensure_ai_profiles_tables(conn, schema)
            cursor = conn.cursor()
            cursor.execute("SELECT instructions FROM ai_profiles WHERE id = %s AND is_active = TRUE", (db_id,))
            row = cursor.fetchone()
            if not row:
                cursor.close()
                return default
            prompt = row["instructions"] or default
            # Load knowledge base documents
            cursor.execute(
                "SELECT original_name, extracted_text FROM ai_profile_documents "
                "WHERE profile_id = %s AND extracted_text IS NOT NULL AND extracted_text != '' "
                "ORDER BY created_at",
                (db_id,),
            )
            docs = cursor.fetchall()
            if docs:
                prompt += "\n\n=== BASE DE CONNAISSANCES ===\n"
                for doc in docs:
                    prompt += f"\n--- {doc['original_name']} ---\n{doc['extracted_text']}\n"
            cursor.close()
            return prompt
        except Exception as exc:
            logger.warning("_load_profile_system_prompt custom error: %s", exc)
            return default
        finally:
            try:
                db.reset_tenant(conn)
            except Exception:
                pass
            conn.close()

    # System profile: load from disk
    profiles_dir = _get_profiles_dir()
    if profiles_dir:
        path = os.path.join(profiles_dir, f"{profile_id}_profil.txt")
        if os.path.isfile(path):
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
    return default


# ============================================================
# ESTIMATION IA — Expert profiles, chat, soumission, conversations
# ============================================================

@router.get("/expert-profiles")
async def list_expert_profiles(user: ErpUser = Depends(get_current_user)):
    """List available expert profiles: system (from disk) + custom (from tenant DB)."""
    profiles = []

    # System profiles from disk
    # Mapping for proper French accented display names
    _ACCENT_NAMES: dict[str, str] = {
        "AGENT_IMMOBILIER": "Agent Immobilier",
        "AMENAGEMENT_PAYSAGER": "Aménagement Paysager",
        "ANALYSEUR_DOCUMENTS": "Analyseur Documents",
        "APRES_SINISTRE": "Après Sinistre",
        "ARCHITECTE": "Architecte",
        "ARMOIRES_CUISINE": "Armoires Cuisine",
        "ARPENTEUR_GEOMETRE": "Arpenteur Géomètre",
        "ASCENSEURS": "Ascenseurs",
        "AVOCAT_EN_DROIT": "Avocat en Droit",
        "CALCULS_COLONNES": "Calculs Colonnes",
        "CALCULS_HAUTEUR_PLANCHER": "Calculs Hauteur Plancher",
        "CALCULS_HEEL_HEIGHT": "Calculs Heel Height",
        "CALCULS_JACKS": "Calculs Jacks",
        "CALCULS_LINTEAUX": "Calculs Linteaux",
        "CALCULS_POUTRES": "Calculs Poutres",
        "CALCULS_TRAJETS": "Calculs Trajets",
        "CALULS_PENTES": "Calculs Pentes",
        "COMPTABLE_CONSTRUCTION": "Comptable Construction",
        "COURTIER_HYPOTHECAIRE": "Courtier Hypothécaire",
        "CVCA_HVAC": "CVCA / HVAC",
        "DECONTAMINATION": "Décontamination",
        "DEMOLITION": "Démolition",
        "DESIGNER_INTERIEUR": "Designer Intérieur",
        "EBENISTE": "Ébéniste",
        "ELECTRICIEN": "Électricien",
        "ELECTRO_MECANICIEN": "Électro Mécanicien",
        "ENTREPRENEUR_GENERAL": "Entrepreneur Général",
        "ENTREPRENEUR_PAVAGE": "Entrepreneur Pavage",
        "EPOXY": "Époxy",
        "ESCALIERS": "Escaliers",
        "ETRIERS_MITEK": "Étriers Mitek",
        "ETRIER_SIMPSON": "Étrier Simpson",
        "EXCAVATION": "Excavation",
        "EXPERT_CAMELEON": "Expert Caméléon",
        "FINITION_INTERIEURE": "Finition Intérieure",
        "FONDATIONS": "Fondations",
        "FOND_PREVOYANCE": "Fond Prévoyance",
        "GYPSE": "Gypse",
        "INGENIEUR_EN_METALLURGIE": "Ingénieur en Métallurgie",
        "INGENIEUR_STRUCTURE": "Ingénieur Structure",
        "INSPECTEUR": "Inspecteur",
        "LEED": "LEED",
        "LOGISTICIEN_CONSTRUCTION": "Logisticien Construction",
        "MACONNERIE": "Maçonnerie",
        "MAGASINIER": "Magasinier",
        "METALLURGIE": "Métallurgie",
        "MODIFIER_DEVIS": "Modifier Devis",
        "PEINTRE": "Peintre",
        "PLOMBIER": "Plombier",
        "PORTEFEUILLE_IMMOBILIER": "Portefeuille Immobilier",
        "PORTES_FENETRES": "Portes Fenêtres",
        "PROGRAMMEUR_CNC": "Programmeur CNC",
        "RBQ_ET_CCQ": "RBQ et CCQ",
        "REPRESENTANT": "Représentant",
        "REVETEMENT_EXTERIEUR": "Revêtement Extérieur",
        "REVETEMENT_PLANCHER": "Revêtement Plancher",
        "SIGNALISATION": "Signalisation",
        "SOUDEUR_MONTEUR": "Soudeur Monteur",
        "STRUCTURE_ACIER": "Structure Acier",
        "STRUCTURE_DE_BOIS": "Structure de Bois",
        "SUBVENTIONS": "Subventions",
        "TECHNOLOGUE": "Technologue",
        "TOITURE": "Toiture",
        "URBANISTE": "Urbaniste",
    }

    profiles_dir = _get_profiles_dir()
    if profiles_dir:
        for fname in sorted(os.listdir(profiles_dir)):
            if fname.endswith("_profil.txt"):
                pid = fname.replace("_profil.txt", "")
                name = _ACCENT_NAMES.get(pid, pid.replace("_", " ").title())
                profiles.append({"id": pid, "name": name, "filename": fname, "source": "system"})

    # Custom profiles from tenant DB
    if user.schema:
        conn = db.get_conn()
        try:
            db.set_tenant(conn, user.schema)
            _ensure_ai_profiles_tables(conn, user.schema)
            cursor = conn.cursor()
            cursor.execute(
                "SELECT id, name FROM ai_profiles WHERE is_active = TRUE ORDER BY name"
            )
            for row in cursor.fetchall():
                profiles.append({
                    "id": f"custom_{row['id']}",
                    "name": row["name"],
                    "filename": "",
                    "source": "custom",
                })
            cursor.close()
        except Exception as exc:
            logger.warning("list_expert_profiles custom error: %s", exc)
        finally:
            try:
                db.reset_tenant(conn)
            except Exception:
                pass
            conn.close()

    return {"profiles": profiles}


def _build_system_blocks_with_cache(profile_id: Optional[str], schema: Optional[str], devis_id: Optional[int]) -> list:
    """Build Claude system as a list of content blocks with prompt caching.

    Structure:
      Block 1: today_prompt_line SEUL (NOT cached — change chaque jour, casserait le cache a minuit)
      Block 2: profile_system_prompt + markdown rules (CACHED — TTL 1h via beta header)
      Block 3: devis context (NOT cached — varies per devis)

    Le prompt du profil (ex: Entrepreneur general = ~130 KB) est stable entre
    les messages d'une meme conversation, donc cache hit = 10% du cout input.
    En isolant `today_line` dans son propre bloc non-cache, on evite le cache miss
    quotidien a 00:00 et on monte le TTL a 1h (au lieu de 5 min ephemeral default).
    """
    profile_prompt = _load_profile_system_prompt(profile_id, schema)
    cacheable_text = profile_prompt + _MARKDOWN_TABLE_RULES
    blocks: list = [
        # Bloc 1: today_line seul (non cacheable)
        {"type": "text", "text": _today_prompt_line()},
        # Bloc 2: profil + markdown rules (cacheable TTL 1h)
        {"type": "text", "text": cacheable_text, "cache_control": {"type": "ephemeral", "ttl": "1h"}},
    ]
    if devis_id and schema:
        conn = db.get_conn()
        cur = None
        devis_ctx = ""
        try:
            db.set_tenant(conn, schema)
            cur = conn.cursor()
            cur.execute("SELECT * FROM devis WHERE id = %s", (devis_id,))
            d = cur.fetchone()
            if d:
                d = dict(d)
                devis_ctx = (
                    f"\n\nContexte du devis:\nProjet: {d.get('nom_projet','N/A')}\n"
                    f"Description: {d.get('description','N/A')}\n"
                    f"Total travaux: {d.get('total_travaux',0)}$\n"
                )
                cur.execute(
                    "SELECT description, quantite, unite, prix_unitaire, montant_ligne "
                    "FROM devis_lignes WHERE devis_id = %s ORDER BY sequence_ligne", (devis_id,))
                for i, l in enumerate(cur.fetchall(), 1):
                    l = dict(l)
                    devis_ctx += (
                        f"  {i}. {l.get('description','')} "
                        f"- {l.get('quantite',0)} {l.get('unite','')} x {l.get('prix_unitaire',0)}$\n"
                    )
        except Exception as e:
            logger.warning("Could not load devis context: %s", e)
        finally:
            if cur:
                cur.close()
            try:
                db.reset_tenant(conn)
            except Exception:
                pass
            conn.close()
        if devis_ctx:
            blocks.append({"type": "text", "text": devis_ctx})
    return blocks


# OPT-2: Files API beta — beta header constant
_ANTHROPIC_FILES_API_BETA = "files-api-2025-04-14"


def _upload_to_anthropic_files(content_bytes: bytes, filename: str, media_type: str) -> Optional[str]:
    """Upload un fichier (image/PDF) vers l'API Files Anthropic et retourne le file_id.

    Best-effort: retourne None si l'upload echoue (le caller fallback vers base64).
    Le file_id retourne peut etre reutilise sur plusieurs tours de conversation
    sans re-uploader les bytes — gros gain pour /ai-chat-with-files (multi-tour).

    Doc Anthropic: *"each request resends the full conversation history. If
    images are base64-encoded, the full image bytes are included in the
    payload on every turn..."* — Files API resout ce probleme.
    """
    if not _anthropic_client or not content_bytes:
        return None
    try:
        import io
        # SDK accepte un tuple (filename, file_obj, content_type) — le content_type
        # est requis pour que l'API valide le fichier comme image/PDF/etc.
        file_obj = io.BytesIO(content_bytes)
        result = _anthropic_client.beta.files.upload(
            file=(filename or "document", file_obj, media_type),
            betas=[_ANTHROPIC_FILES_API_BETA],
        )
        # Le SDK retourne un FileMetadata avec un attribut .id
        file_id = getattr(result, "id", None)
        if file_id:
            logger.info("Files API upload OK: filename=%s media_type=%s file_id=%s", filename, media_type, file_id)
        return file_id
    except Exception as exc:
        logger.warning("Files API upload failed for %s (%s): %s — fallback base64", filename, media_type, exc)
        return None


def _delete_from_anthropic_files(file_id: Optional[str]) -> bool:
    """Best-effort delete d'un fichier Files API. Retourne True si succes ou si rien a faire.

    Utilise pour cleanup quand un document est supprime de la conversation.
    Erreurs sont loggees mais n'echouent pas l'operation principale.
    """
    if not file_id or not _anthropic_client:
        return True
    try:
        _anthropic_client.beta.files.delete(file_id=file_id, betas=[_ANTHROPIC_FILES_API_BETA])
        return True
    except Exception as exc:
        logger.warning("Files API delete failed for %s: %s", file_id, exc)
        return False


def _load_active_conversation_documents(schema: Optional[str], conversation_id: Optional[int], user_id: Optional[int] = None) -> list:
    """Load documents persistes lies a une conversation avec is_active_context=TRUE.

    SECURITE: si user_id est fourni, un check ownership via JOIN sur
    conversations.user_id previent qu'un user injecte le conv_id d'un
    autre user du meme tenant pour recevoir ses documents confidentiels.

    Retourne une liste de content blocks Claude (image/document/text). Le
    cache_control ephemeral n'est applique qu'au DERNIER bloc image/PDF
    pour respecter la limite Anthropic de 4 breakpoints par requete
    (1 deja utilise sur system + max 3 pour les docs).

    OPT-2: si une row a `anthropic_file_id` non-null, on utilise une source
    de type "file" (file_id reference) au lieu de re-encoder les bytes en
    base64 a chaque tour. Sinon, fallback base64 (retro-compat).
    """
    if not conversation_id or not schema:
        return []
    conn = db.get_conn()
    cur = None
    blocks: list = []
    try:
        db.set_tenant(conn, schema)
        _ensure_conversation_documents_table(conn, schema)
        cur = conn.cursor()
        # Ownership check: JOIN conversations pour s'assurer que la conv
        # appartient bien au user courant (ou est orpheline user_id IS NULL
        # pour backward compat avec legacy convs).
        # OPT-2: SELECT anthropic_file_id pour reutilisation Files API multi-tour.
        if user_id is not None:
            cur.execute(
                "SELECT cd.id, cd.filename, cd.media_type, cd.content, cd.extracted_text, cd.anthropic_file_id "
                "FROM conversation_documents cd "
                "JOIN conversations c ON c.id = cd.conversation_id "
                "WHERE cd.conversation_id = %s AND cd.is_active_context = TRUE "
                "  AND (c.user_id = %s OR c.user_id IS NULL) "
                "ORDER BY cd.created_at",
                (conversation_id, user_id),
            )
        else:
            cur.execute(
                "SELECT id, filename, media_type, content, extracted_text, anthropic_file_id "
                "FROM conversation_documents "
                "WHERE conversation_id = %s AND is_active_context = TRUE "
                "ORDER BY created_at",
                (conversation_id,),
            )
        rows = cur.fetchall()
        for row in rows:
            r = dict(row)
            mt = r.get("media_type") or ""
            content_bytes = r.get("content")
            file_id = r.get("anthropic_file_id")
            # OPT-2: si file_id present, utiliser source de type "file" — pas
            # de re-encoding base64, pas de re-upload. Reutilise l'objet stocke
            # cote Anthropic. Fallback base64 si file_id absent.
            if mt.startswith("image/"):
                if file_id:
                    blocks.append({
                        "type": "image",
                        "source": {"type": "file", "file_id": file_id},
                    })
                elif content_bytes:
                    raw_bytes = bytes(content_bytes)
                    # media_type stocke en BD peut mentir (ex: PNG-renomme JPEG).
                    # Re-detecter via magic bytes pour eviter erreur 400 Anthropic.
                    actual_mt = _detect_media_type_from_bytes(raw_bytes) or mt
                    b64_data = base64.standard_b64encode(raw_bytes).decode("utf-8")
                    blocks.append({
                        "type": "image",
                        "source": {"type": "base64", "media_type": actual_mt, "data": b64_data},
                    })
            elif mt == "application/pdf":
                if file_id:
                    blocks.append({
                        "type": "document",
                        "source": {"type": "file", "file_id": file_id},
                    })
                elif content_bytes:
                    b64_data = base64.standard_b64encode(bytes(content_bytes)).decode("utf-8")
                    blocks.append({
                        "type": "document",
                        "source": {"type": "base64", "media_type": "application/pdf", "data": b64_data},
                    })
            elif r.get("extracted_text"):
                # xlsx/docx/csv/txt — envoi en texte
                blocks.append({
                    "type": "text",
                    "text": f"\n=== Document: {r['filename']} ===\n{r['extracted_text']}\n",
                })
        # Applique cache_control uniquement sur le DERNIER bloc cacheable
        # (image/document) pour consolider en 1 breakpoint. Text blocks ne
        # declenchent pas de breakpoint separe.
        for blk in reversed(blocks):
            if blk.get("type") in ("image", "document"):
                blk["cache_control"] = {"type": "ephemeral"}
                break
    except Exception as exc:
        logger.warning("_load_active_conversation_documents failed: %s", exc)
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()
    # OPT-2: tag la liste avec l'info has_file_source pour le caller (pattern
    # leger via attribut sur la liste — Python autorise pas les attrs sur list,
    # donc on retourne (blocks, has_file_source) via un wrapper.).
    # Pour eviter de casser tous les call sites existants, on garde la signature
    # qui retourne list. Le caller qui veut l'info peut appeler le helper
    # _conversation_docs_uses_file_api(blocks) ci-dessous.
    return blocks


def _conversation_docs_uses_file_api(blocks: list) -> bool:
    """Detecte si au moins un block dans la liste utilise une source file_id
    (Files API beta). Le caller doit alors passer betas=["files-api-2025-04-14"]
    dans l'appel Claude.
    """
    for blk in blocks or []:
        src = blk.get("source") if isinstance(blk, dict) else None
        if isinstance(src, dict) and src.get("type") == "file":
            return True
    return False


@router.post("/ai-chat")
async def ai_chat(body: AiChatRequest, user: ErpUser = Depends(get_current_user)):
    """Chat with an AI expert using Claude Opus (30 000 tokens) + prompt caching.

    Si body.conversation_id est fourni, les documents persistes de la conversation
    sont automatiquement injectes dans le premier user message (avec cache).
    """
    if not _anthropic_client:
        raise HTTPException(status_code=503, detail="Service IA non disponible")

    # AI billing: guard + credit check
    allowed, error_msg = check_ai_guard(user)
    if not allowed:
        raise HTTPException(status_code=403, detail=error_msg or "Acces IA refuse")
    credits_ok, balance = _check_credits(user)
    if not credits_ok:
        raise HTTPException(status_code=402, detail="Credits IA epuises. Veuillez recharger votre solde.")

    # System blocks avec cache (profile + markdown rules + devis context)
    system_blocks = _build_system_blocks_with_cache(body.profile_id, user.schema, body.devis_id)

    # Build Claude messages
    claude_msgs = [
        {"role": m["role"], "content": m["content"]}
        for m in body.messages
        if m.get("role") in ("user", "assistant") and m.get("content")
    ]
    if not claude_msgs:
        raise HTTPException(status_code=400, detail="Aucun message fourni")

    # Injecte les documents persistes de la conversation dans le premier user msg
    # Ownership check: user.user_id passe pour empêcher la fuite cross-user intra-tenant
    conv_id = getattr(body, "conversation_id", None)
    doc_blocks = _load_active_conversation_documents(user.schema, conv_id, user.user_id)
    if doc_blocks:
        first_user_idx = None
        for i, m in enumerate(claude_msgs):
            if m["role"] == "user":
                first_user_idx = i
                break
        if first_user_idx is not None:
            orig = claude_msgs[first_user_idx]["content"]
            orig_text = orig if isinstance(orig, str) else ""
            claude_msgs[first_user_idx]["content"] = doc_blocks + [{"type": "text", "text": orig_text}]

    # OPT-2: detecte si Files API beta est requis (sources file_id en jeu)
    needs_files_beta = _conversation_docs_uses_file_api(doc_blocks)

    try:
        resp = _call_claude(
            model="claude-opus-4-7",
            max_tokens=32000,
            system=system_blocks,
            messages=claude_msgs,
            betas=([_ANTHROPIC_FILES_API_BETA] if needs_files_beta else None),
        )

        # AI billing: track usage + deduct credits (inclut cost cache)
        tokens_in = resp.usage.input_tokens
        tokens_out = resp.usage.output_tokens
        cache_creation = getattr(resp.usage, "cache_creation_input_tokens", 0) or 0
        cache_read = getattr(resp.usage, "cache_read_input_tokens", 0) or 0
        # Opus 4.7: $15/M input, $75/M output, $18.75/M cache write, $1.50/M cache read
        cost = (
            tokens_in * 15 / 1_000_000
            + tokens_out * 75 / 1_000_000
            + cache_creation * 18.75 / 1_000_000
            + cache_read * 1.50 / 1_000_000
        ) * 1.30  # 30% markup
        track_ai_usage(user, "devis_ai_chat", tokens_in, tokens_out, cost, 0, True, model="claude-opus-4-7")
        _deduct_credits(user, cost)

        return {
            "response": _extract_text_from_response(resp),
            "usage": {
                "input_tokens": tokens_in,
                "output_tokens": tokens_out,
                "cache_creation_input_tokens": cache_creation,
                "cache_read_input_tokens": cache_read,
                "cost_usd": round(cost, 6),
            },
        }
    except HTTPException:
        raise
    except anthropic.NotFoundError as exc:
        # FIX-2: file_id expire ou supprime cote Anthropic (Files API beta).
        # Clear les anthropic_file_id stales en DB pour qu'au prochain tour le
        # _load_active_conversation_documents reencode en base64. On ne fait pas
        # de retry automatique — le user reessaie 1 click.
        if needs_files_beta and conv_id and user.schema:
            logger.warning("ai_chat Anthropic 404 (file_id expire), clearing stale file_ids: %s", exc)
            _conn_fix = None
            try:
                _conn_fix = db.get_conn()
                db.set_tenant(_conn_fix, user.schema)
                _cur_fix = _conn_fix.cursor()
                _cur_fix.execute(
                    "UPDATE conversation_documents SET anthropic_file_id = NULL "
                    "WHERE conversation_id = %s AND anthropic_file_id IS NOT NULL",
                    (conv_id,),
                )
                _conn_fix.commit()
                _cur_fix.close()
            except Exception:
                logger.exception("Failed to clear stale anthropic_file_id (ai_chat)")
                try:
                    if _conn_fix is not None:
                        _conn_fix.rollback()
                except Exception:
                    pass
            finally:
                try:
                    if _conn_fix is not None:
                        db.reset_tenant(_conn_fix)
                except Exception:
                    pass
                try:
                    if _conn_fix is not None:
                        _conn_fix.close()
                except Exception:
                    pass
            raise HTTPException(
                status_code=503,
                detail="Pieces jointes expirees, recharge ta conversation et reessaie.",
            )
        # 404 non lie aux fichiers — laisse le handler APIError ci-dessous traiter
        logger.error("ai_chat NotFoundError (non-file): %s", exc)
        raise HTTPException(status_code=503, detail="Service IA temporairement indisponible")
    except anthropic.APIError as exc:
        logger.error("ai_chat API error: %s", exc)
        status = getattr(exc, "status_code", 0)
        exc_str = str(exc).lower()
        if status == 413 or (status == 400 and "too_large" in exc_str):
            raise HTTPException(
                status_code=413,
                detail="Conversation trop longue pour l'IA. Essayez de demarrer une nouvelle conversation.",
            )
        if status == 400 and "image exceeds" in exc_str:
            raise HTTPException(
                status_code=413,
                detail="Image trop volumineuse (max 5 Mo). Reduisez la taille de l'image ou utilisez un format compresse (JPEG).",
            )
        if status == 529 or "overloaded" in exc_str:
            raise HTTPException(status_code=503, detail="Le service IA est temporairement surcharge. Veuillez reessayer dans quelques secondes.")
        if status == 429 or "rate_limit" in exc_str:
            raise HTTPException(status_code=429, detail="Trop de demandes vers l'IA. Attendez quelques instants avant de reessayer.")
        logger.error("Erreur API Claude: %s", exc)
        raise HTTPException(status_code=503, detail="Service IA temporairement indisponible")
    except Exception as exc:
        # Connection errors from httpx can bypass _call_claude's wrapping when
        # anthropic.APIConnectionError construction fails (e.g. SDK signature
        # drift). Observed in prod 2026-05-02: "peer closed connection without
        # sending complete message body (incomplete chunked read)" landed
        # here as a generic 500 instead of a transient 503. Catch the known
        # httpx transient classes explicitly as a secondary safety net so
        # the user gets a clear "retry" signal rather than a hard failure.
        # Couvre aussi WriteError / ConnectError / TimeoutException / PoolTimeout
        # qui sont également transient et méritent un 503 retry-friendly.
        if httpx is not None and isinstance(
            exc,
            (
                httpx.RemoteProtocolError,
                httpx.ReadError,
                httpx.ReadTimeout,
                httpx.WriteError,
                httpx.ConnectError,
                httpx.ConnectTimeout,
                httpx.PoolTimeout,
                httpx.TimeoutException,
            ),
        ):
            logger.warning("ai_chat httpx transient: %s", exc)
            raise HTTPException(
                status_code=503,
                detail="La connexion au service IA a ete interrompue. Veuillez reessayer.",
            )
        logger.error("ai_chat error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors du chat IA")


@router.post("/ai-chat-with-files")
async def ai_chat_with_files(
    messages_json: str = Form(...),
    profile_id: Optional[str] = Form(None),
    devis_id: Optional[int] = Form(None),
    conversation_id: Optional[int] = Form(None),
    files: List[UploadFile] = File(default=[]),
    user: ErpUser = Depends(get_current_user),
):
    """Chat with an AI expert using Claude Opus + cache + persistent documents.

    - `conversation_id`: si fourni, les documents persistes de la conversation
      sont injectes automatiquement (cache Anthropic 5 min TTL).
    - `files`: nouveaux fichiers attaches a ce message (max 5, 10 Mo chacun).
      Persistes dans conversation_documents si conversation_id fourni.
    """
    if not _anthropic_client:
        raise HTTPException(status_code=503, detail="Service IA non disponible")

    # Validate file count
    if len(files) > 5:
        raise HTTPException(status_code=413, detail="Maximum 5 fichiers par message")

    # Parse messages from JSON string
    try:
        messages = json.loads(messages_json)
    except (json.JSONDecodeError, TypeError):
        raise HTTPException(status_code=400, detail="messages_json invalide")

    # AI billing: guard + credit check
    allowed, error_msg = check_ai_guard(user)
    if not allowed:
        raise HTTPException(status_code=403, detail=error_msg or "Acces IA refuse")
    credits_ok, balance = _check_credits(user)
    if not credits_ok:
        raise HTTPException(status_code=402, detail="Credits IA epuises. Veuillez recharger votre solde.")

    # System blocks avec cache (profile + markdown rules + devis context)
    system_blocks = _build_system_blocks_with_cache(profile_id, user.schema, devis_id)

    # Build Claude messages
    claude_msgs = [
        {"role": m["role"], "content": m["content"]}
        for m in messages
        if m.get("role") in ("user", "assistant") and m.get("content")
    ]
    if not claude_msgs:
        raise HTTPException(status_code=400, detail="Aucun message fourni")

    # Injecte les documents persistes de la conversation dans le premier user msg
    # Ownership check via user_id pour prevenir fuite intra-tenant
    persistent_doc_blocks = _load_active_conversation_documents(user.schema, conversation_id, user.user_id)
    if persistent_doc_blocks:
        first_user_idx = None
        for i, m in enumerate(claude_msgs):
            if m["role"] == "user":
                first_user_idx = i
                break
        if first_user_idx is not None:
            orig = claude_msgs[first_user_idx]["content"]
            orig_text = orig if isinstance(orig, str) else ""
            claude_msgs[first_user_idx]["content"] = persistent_doc_blocks + [{"type": "text", "text": orig_text}]

    # If files are attached, build content blocks for the last user message
    # + persist files in conversation_documents for indefinite conversation continuation
    newly_persisted_files: list = []  # pour tracking / reponse
    if files:
        # Find the last user message
        last_user_idx = None
        for idx in range(len(claude_msgs) - 1, -1, -1):
            if claude_msgs[idx]["role"] == "user":
                last_user_idx = idx
                break

        if last_user_idx is not None:
            last_user_msg_text = claude_msgs[last_user_idx]["content"]
            content_blocks = []
            for f_upload in files:
                file_bytes = await f_upload.read()
                # FIX: garde explicit fichier vide (cf. ai_analyze_document)
                if not file_bytes:
                    raise HTTPException(status_code=400, detail=f"Fichier {f_upload.filename or 'sans nom'} vide ou impossible a lire")
                if len(file_bytes) > 10 * 1024 * 1024:
                    raise HTTPException(status_code=413, detail=f"Fichier {f_upload.filename} trop volumineux (max 10 Mo)")
                b64 = base64.standard_b64encode(file_bytes).decode("utf-8")
                ext = (f_upload.filename or "").rsplit(".", 1)[-1].lower()
                media_map = {
                    "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                    "gif": "image/gif", "webp": "image/webp", "pdf": "application/pdf",
                }
                media_type = media_map.get(ext)
                extracted_text_for_persist = ""
                if media_type and media_type.startswith("image/"):
                    # Compression si > 4.5 MB. Opus 4.7 native res = 2576 px sur le
                    # long edge (vs 1568 pour Sonnet) — compresser plus bas que ca
                    # gaspille de la resolution. Voir OPT-1 dans /ai-analyze-document.
                    if len(file_bytes) > 4.5 * 1024 * 1024:
                        try:
                            from PIL import Image
                            import io
                            img = Image.open(io.BytesIO(file_bytes))
                            if img.mode != 'RGB':
                                img = img.convert('RGB')
                            max_dim = 2576
                            if max(img.size) > max_dim:
                                img.thumbnail((max_dim, max_dim), Image.LANCZOS)
                            buf = io.BytesIO()
                            img.save(buf, format='JPEG', quality=80, optimize=True)
                            file_bytes = buf.getvalue()
                            b64 = base64.standard_b64encode(file_bytes).decode("utf-8")
                            media_type = "image/jpeg"
                        except Exception as resize_err:
                            logger.warning("Image resize failed: %s — sending original", resize_err)
                    if len(file_bytes) > 5 * 1024 * 1024:
                        raise HTTPException(
                            status_code=400,
                            detail=f"Image trop volumineuse (max 5 Mo apres compression). Fichier: {f_upload.filename}",
                        )
                    # Validation finale: si l'extension ment ou que PIL a echoue,
                    # detecter le vrai media_type des bytes pour eviter erreur 400 Anthropic.
                    actual_media_type = _detect_media_type_from_bytes(file_bytes) or media_type
                    content_blocks.append({
                        "type": "image",
                        "source": {"type": "base64", "media_type": actual_media_type, "data": b64},
                    })
                elif media_type == "application/pdf":
                    content_blocks.append({
                        "type": "document",
                        "source": {"type": "base64", "media_type": "application/pdf", "data": b64},
                    })
                    # Extraction texte pour stockage (evite re-OCR plus tard)
                    try:
                        extracted_text_for_persist = _extract_text_from_file(file_bytes, "application/pdf", f_upload.filename or "")
                    except Exception:
                        pass
                else:
                    # xlsx/docx/csv/txt — extraction + envoi en texte
                    extracted_text_for_persist = _extract_text_from_file(file_bytes, "", f_upload.filename or "")
                    if extracted_text_for_persist.strip():
                        content_blocks.append({
                            "type": "text",
                            "text": f"\n=== Fichier joint: {f_upload.filename} ===\n{extracted_text_for_persist}\n",
                        })

                # Persiste dans conversation_documents si conversation_id fourni
                if conversation_id and user.schema:
                    inserted_doc_id = None
                    try:
                        conn_p = db.get_conn()
                        db.set_tenant(conn_p, user.schema)
                        _ensure_conversation_documents_table(conn_p, user.schema)
                        cur_p = conn_p.cursor()
                        cur_p.execute(
                            "INSERT INTO conversation_documents "
                            "(conversation_id, filename, media_type, file_size, content, "
                            "extracted_text, is_active_context, uploaded_by, created_at, updated_at) "
                            "VALUES (%s,%s,%s,%s,%s,%s,TRUE,%s,NOW(),NOW()) RETURNING id",
                            (
                                conversation_id, f_upload.filename or "file",
                                media_type or "application/octet-stream",
                                len(file_bytes),
                                psycopg2.Binary(file_bytes) if (media_type and (media_type.startswith("image/") or media_type == "application/pdf")) else None,
                                extracted_text_for_persist[:500000] if extracted_text_for_persist else None,
                                user.user_id,
                            ),
                        )
                        row_p = cur_p.fetchone()
                        if row_p:
                            inserted_doc_id = row_p["id"]
                            newly_persisted_files.append({"id": inserted_doc_id, "filename": f_upload.filename})
                        conn_p.commit()
                        cur_p.close()

                        # OPT-2: upload best-effort vers Files API pour reutilisation
                        # multi-tour. Cf. /ai-analyze-document pour le rationale.
                        if inserted_doc_id and media_type and (
                            media_type.startswith("image/") or media_type == "application/pdf"
                        ):
                            file_id = _upload_to_anthropic_files(
                                file_bytes, f_upload.filename or "file", media_type
                            )
                            if file_id:
                                try:
                                    cur_up = conn_p.cursor()
                                    cur_up.execute(
                                        "UPDATE conversation_documents SET anthropic_file_id = %s WHERE id = %s",
                                        (file_id, inserted_doc_id),
                                    )
                                    conn_p.commit()
                                    cur_up.close()
                                except Exception as up_exc:
                                    logger.warning("UPDATE anthropic_file_id failed for doc %s: %s", inserted_doc_id, up_exc)
                                    try:
                                        conn_p.rollback()
                                    except Exception:
                                        pass
                    except Exception as persist_exc:
                        logger.warning("persist file in conversation_documents failed: %s", persist_exc)
                        try:
                            conn_p.rollback()
                        except Exception:
                            pass
                    finally:
                        try:
                            db.reset_tenant(conn_p)
                        except Exception:
                            pass
                        try:
                            conn_p.close()
                        except Exception:
                            pass

            # Applique cache_control uniquement sur le DERNIER bloc cacheable
            # (image/document) pour respecter la limite Anthropic de 4 breakpoints.
            for blk in reversed(content_blocks):
                if blk.get("type") in ("image", "document"):
                    blk["cache_control"] = {"type": "ephemeral"}
                    break
            content_blocks.append({"type": "text", "text": last_user_msg_text})
            claude_msgs[last_user_idx]["content"] = content_blocks

    # OPT-2: detecte si au moins un block de message utilise une source file_id
    # (Files API beta). Si oui, on doit passer betas=[files-api-2025-04-14] dans
    # l'appel Claude, sinon l'API ignore la source file et retourne une 400.
    # On scanne persistent_doc_blocks ET le content du last_user_msg car les
    # nouveaux files persistes sur ce tour peuvent avoir un file_id, mais ne
    # sont injectes qu'au tour SUIVANT via _load_active_conversation_documents.
    # Sur ce tour-ci, les nouveaux files sont en base64 (content_blocks). Le
    # vrai gain Files API est sur les tours suivants.
    needs_files_beta = False
    for msg in claude_msgs:
        c = msg.get("content")
        if isinstance(c, list):
            if _conversation_docs_uses_file_api(c):
                needs_files_beta = True
                break

    try:
        resp = _call_claude(
            model="claude-opus-4-7",
            max_tokens=32000,
            system=system_blocks,
            messages=claude_msgs,
            betas=([_ANTHROPIC_FILES_API_BETA] if needs_files_beta else None),
        )

        # AI billing: track usage + deduct credits (inclut cost cache)
        tokens_in = resp.usage.input_tokens
        tokens_out = resp.usage.output_tokens
        cache_creation = getattr(resp.usage, "cache_creation_input_tokens", 0) or 0
        cache_read = getattr(resp.usage, "cache_read_input_tokens", 0) or 0
        cost = (
            tokens_in * 15 / 1_000_000
            + tokens_out * 75 / 1_000_000
            + cache_creation * 18.75 / 1_000_000
            + cache_read * 1.50 / 1_000_000
        ) * 1.30
        track_ai_usage(user, "devis_ai_chat", tokens_in, tokens_out, cost, 0, True, model="claude-opus-4-7")
        _deduct_credits(user, cost)

        return {
            "response": _extract_text_from_response(resp),
            "persisted_files": newly_persisted_files,
            "usage": {
                "input_tokens": tokens_in,
                "output_tokens": tokens_out,
                "cache_creation_input_tokens": cache_creation,
                "cache_read_input_tokens": cache_read,
                "cost_usd": round(cost, 6),
            },
        }
    except HTTPException:
        raise
    except anthropic.NotFoundError as exc:
        # FIX-2: file_id expire ou supprime cote Anthropic (Files API beta).
        # Clear les anthropic_file_id stales en DB pour qu'au prochain tour le
        # _load_active_conversation_documents reencode en base64. Pas de retry
        # automatique — user reessaie 1 click.
        if needs_files_beta and conversation_id and user.schema:
            logger.warning("ai_chat_with_files Anthropic 404 (file_id expire), clearing stale file_ids: %s", exc)
            _conn_fix = None
            try:
                _conn_fix = db.get_conn()
                db.set_tenant(_conn_fix, user.schema)
                _cur_fix = _conn_fix.cursor()
                _cur_fix.execute(
                    "UPDATE conversation_documents SET anthropic_file_id = NULL "
                    "WHERE conversation_id = %s AND anthropic_file_id IS NOT NULL",
                    (conversation_id,),
                )
                _conn_fix.commit()
                _cur_fix.close()
            except Exception:
                logger.exception("Failed to clear stale anthropic_file_id (ai_chat_with_files)")
                try:
                    if _conn_fix is not None:
                        _conn_fix.rollback()
                except Exception:
                    pass
            finally:
                try:
                    if _conn_fix is not None:
                        db.reset_tenant(_conn_fix)
                except Exception:
                    pass
                try:
                    if _conn_fix is not None:
                        _conn_fix.close()
                except Exception:
                    pass
            raise HTTPException(
                status_code=503,
                detail="Pieces jointes expirees, recharge ta conversation et reessaie.",
            )
        # 404 non lie aux fichiers — fallback 503 generique
        logger.error("ai_chat_with_files NotFoundError (non-file): %s", exc)
        raise HTTPException(status_code=503, detail="Service IA temporairement indisponible")
    except anthropic.APIError as exc:
        logger.error("ai_chat_with_files API error: %s", exc)
        status = getattr(exc, "status_code", 0)
        exc_str = str(exc).lower()
        if status == 413 or (status == 400 and "too_large" in exc_str):
            raise HTTPException(
                status_code=413,
                detail="Conversation trop longue pour l'IA. Essayez de demarrer une nouvelle conversation.",
            )
        if status == 400 and "image exceeds" in exc_str:
            raise HTTPException(
                status_code=413,
                detail="Image trop volumineuse (max 5 Mo). Reduisez la taille de l'image ou utilisez un format compresse (JPEG).",
            )
        if status == 529 or "overloaded" in exc_str:
            raise HTTPException(status_code=503, detail="Le service IA est temporairement surcharge. Veuillez reessayer dans quelques secondes.")
        if status == 429 or "rate_limit" in exc_str:
            raise HTTPException(status_code=429, detail="Trop de demandes vers l'IA. Attendez quelques instants avant de reessayer.")
        logger.error("Erreur API Claude: %s", exc)
        raise HTTPException(status_code=503, detail="Service IA temporairement indisponible")
    except Exception as exc:
        # Same secondary safety net as ai_chat — couvre tous les types httpx
        # transient (Read/Write/Connect/Timeout/Pool) plutôt que les seuls
        # 3 initiaux. Si une variante échappe au wrapper de _call_claude,
        # on route vers 503 (retry-friendly) au lieu du 500 générique.
        if httpx is not None and isinstance(
            exc,
            (
                httpx.RemoteProtocolError,
                httpx.ReadError,
                httpx.ReadTimeout,
                httpx.WriteError,
                httpx.ConnectError,
                httpx.ConnectTimeout,
                httpx.PoolTimeout,
                httpx.TimeoutException,
            ),
        ):
            logger.warning("ai_chat_with_files httpx transient: %s", exc)
            raise HTTPException(
                status_code=503,
                detail="La connexion au service IA a ete interrompue. Veuillez reessayer.",
            )
        logger.error("ai_chat_with_files error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors du chat IA")


# ============================================
# SOUMISSION GENERATION - tool_use schema (Sprint 2.A)
# ============================================
# Enums fermes pour forcer Claude a produire des categories et unites
# normalisees (eliminer les variantes 'Plomb.' / 'Plomberie residentielle' qui
# cassent le groupement par corps de metier dans le UI).

_SOUMISSION_CATEGORIES = [
    "Fondation", "Charpente", "Toiture", "Revetement exterieur",
    "Portes et fenetres", "Plomberie", "Electricite", "Ventilation HVAC",
    "Isolation", "Finition interieure", "Finition exterieure", "Demolition",
    "Excavation", "Beton", "Maconnerie", "Peinture", "Plancher",
    "Armoires cuisine", "Salle de bain", "Terrain amenagement", "Permis et frais",
]

_SOUMISSION_UNITES = [
    "UN", "H", "JR", "SEM", "MOIS", "ML", "M2", "M3",
    "KG", "T", "LOT", "FORF", "PI", "PI2", "PI3", "V",
]

_SOUMISSION_TOOL = {
    "name": "creer_lignes_devis",
    "description": (
        "Cree les lignes detaillees d'une soumission de construction "
        "organisees par corps de metier. Chaque ligne represente un travail "
        "concret (materiaux + MO + equipement) avec quantite, unite et prix "
        "unitaire CAD 2026 Quebec. NE PAS inclure de lignes pour "
        "Administration, Contingences, Profit, Gestion de projet ou Frais "
        "generaux ces pourcentages sont calcules automatiquement par le "
        "systeme de devis Constructo AI. Regroupe les items par categorie "
        "(corps de metier) pour faciliter la totalisation."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "items": {
                "type": "array",
                "minItems": 1,
                "description": "Liste des lignes de la soumission, groupees logiquement par corps de metier.",
                "items": {
                    "type": "object",
                    "properties": {
                        "description": {"type": "string", "minLength": 3, "maxLength": 500},
                        "quantite": {"type": "number", "minimum": 0.001, "maximum": 100000},
                        "unite": {"type": "string", "enum": _SOUMISSION_UNITES},
                        "prix_unitaire": {"type": "number", "minimum": 0, "maximum": 1000000},
                        "categorie": {"type": "string", "enum": _SOUMISSION_CATEGORIES},
                        "notes_ligne": {"type": "string", "maxLength": 1000},
                    },
                    "required": ["description", "quantite", "unite", "prix_unitaire", "categorie"],
                    "additionalProperties": False,
                },
            },
        },
        "required": ["items"],
        "additionalProperties": False,
    },
}


_CATALOGUE_LOOKUP_TOOL = {
    "name": "recherche_catalogue_produits",
    "description": (
        "Recherche dans le catalogue produits du tenant (table `produits`) pour grounder "
        "les estimations sur les prix REELS de l'entreprise au lieu d'inventer des tarifs marche. "
        "Utilise systematiquement cet outil avant de proposer un prix unitaire pour: materiaux "
        "(bois d'oeuvre, gypse, isolant, ceramique, peinture), produits manufactures (portes, "
        "fenetres, robinetterie, luminaires), et composantes catalogues. "
        "Privilegier 2-3 mots-cles courts (ex: 'gypse 5/8', 'fenetre PVC', 'thermopompe 18000'). "
        "Si aucun resultat, retenter avec un terme plus generique. "
        "Le champ `prix_unitaire` est en CAD (avant taxes), `stock_disponible` indique la "
        "quantite en inventaire (ne pas confondre avec le prix de marche). "
        "NE PAS utiliser pour: main-d'oeuvre (taux CCQ), permis, frais administratifs, "
        "contingences (calcules automatiquement). Max 50 produits par appel."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "search": {
                "type": "string",
                "minLength": 2,
                "maxLength": 100,
                "description": (
                    "Termes de recherche (matche nom, description, code_produit en ILIKE). "
                    "Ex: 'gypse 5/8', '2x4x8 SPF', 'fenetre PVC 36x48'."
                ),
            },
            "categorie": {
                "type": "string",
                "maxLength": 100,
                "description": (
                    "Filtre optionnel sur la categorie produit (egalite stricte, insensible "
                    "a la casse). Ex: 'Charpente', 'Plomberie', 'Electricite'. "
                    "Omettre si l'on veut chercher dans toutes les categories."
                ),
            },
            "limit": {
                "type": "integer",
                "minimum": 1,
                "maximum": 50,
                "default": 20,
                "description": "Nombre max de produits retournes (defaut 20, cap 50).",
            },
        },
        "required": ["search"],
        "additionalProperties": False,
    },
}


# ============================================
# SOUMISSION VALIDATION POST-HOC (Sprint 3 #9)
# ============================================
# Fourchettes APCHQ Quebec 2026 en $/pi2 par type de projet.
# Sources publiees (Association des Professionnels de la Construction et de
# l'Habitation du Quebec) + ajustements inflation 2024-2026. Marge de tolerance
# volontairement large pour eviter les faux positifs sur les marches regionaux.

_APCHQ_FOURCHETTES_PI2_2026 = {
    "cuisine_renovation": {"min": 200, "max": 1000},
    "salle_de_bain_renovation": {"min": 250, "max": 1200},
    "agrandissement_residentiel": {"min": 250, "max": 500},
    "construction_neuve_residentiel": {"min": 200, "max": 450},
    "renovation_majeure_residentiel": {"min": 150, "max": 400},
    "sous_sol_finition": {"min": 80, "max": 200},
    "toiture_refection": {"min": 8, "max": 25},
    "garage_construction": {"min": 100, "max": 250},
    "commercial_leger": {"min": 150, "max": 350},
    "commercial_lourd": {"min": 250, "max": 600},
    "industriel": {"min": 100, "max": 300},
}

_APCHQ_FOURCHETTES_DEFAUT = {"min": 50, "max": 2000}
_PRIX_UNITAIRE_MIN = 0.01
_PRIX_UNITAIRE_MAX = 1_000_000.0
_QUANTITE_MIN = 0.001
_QUANTITE_MAX = 100_000.0
_TOTAL_AGREGE_MAX = 10_000_000.0
_RATIO_CATEGORIE_MAX = 0.40

_KEYWORDS_CATEGORIE_MAP = {
    "beton": {"Beton", "Fondation"},
    "fondation": {"Fondation", "Beton", "Excavation"},
    "charpente": {"Charpente"},
    "toiture": {"Toiture"},
    "bardeau": {"Toiture"},
    "couvre-plancher": {"Plancher"},
    "plancher": {"Plancher"},
    "ceramique": {"Plancher", "Salle de bain", "Finition interieure"},
    "gypse": {"Finition interieure"},
    "peinture": {"Peinture"},
    "armoire": {"Armoires cuisine", "Finition interieure"},
    "plomberie": {"Plomberie"},
    "toilette": {"Plomberie", "Salle de bain"},
    "robinet": {"Plomberie"},
    "lavabo": {"Plomberie", "Salle de bain"},
    "electricite": {"Electricite"},
    "luminaire": {"Electricite"},
    "prise": {"Electricite"},
    "ventilation": {"Ventilation HVAC"},
    "thermopompe": {"Ventilation HVAC"},
    "fenetre": {"Portes et fenetres"},
    "porte": {"Portes et fenetres"},
    "isolation": {"Isolation"},
    "uretane": {"Isolation"},
    "demolition": {"Demolition"},
    "excavation": {"Excavation"},
    "permis": {"Permis et frais"},
}

_MONTANT_PATTERN = re.compile(
    r"(?:(?<![\w.])\$\s?([0-9]{1,3}(?:[ ,][0-9]{3})*(?:[.,][0-9]{1,2})?|[0-9]+(?:[.,][0-9]{1,2})?))"
    r"|(?:(?<![\w.])([0-9]{1,3}(?:[ ,][0-9]{3})*(?:[.,][0-9]{1,2})?|[0-9]+(?:[.,][0-9]{1,2})?)\s?\$)"
)


def _normalize_projet_type(projet_type):
    """Normalise un projet_type pour matcher les clefs APCHQ."""
    if not projet_type:
        return None
    normalized = projet_type.strip().lower()
    accents = {"é": "e", "è": "e", "ê": "e", "à": "a", "â": "a", "ô": "o", "î": "i", "ï": "i", "ç": "c", "û": "u", "ù": "u"}
    for src, dst in accents.items():
        normalized = normalized.replace(src, dst)
    normalized = normalized.replace("-", "_").replace(" ", "_")
    while "__" in normalized:
        normalized = normalized.replace("__", "_")
    if normalized in _APCHQ_FOURCHETTES_PI2_2026:
        return normalized
    aliases = {
        "cuisine": "cuisine_renovation",
        "renovation_cuisine": "cuisine_renovation",
        "salle_de_bain": "salle_de_bain_renovation",
        "sdb": "salle_de_bain_renovation",
        "renovation_sdb": "salle_de_bain_renovation",
        "agrandissement": "agrandissement_residentiel",
        "neuf": "construction_neuve_residentiel",
        "neuf_residentiel": "construction_neuve_residentiel",
        "construction_neuve": "construction_neuve_residentiel",
        "sous_sol": "sous_sol_finition",
        "toiture": "toiture_refection",
        "garage": "garage_construction",
        "commercial": "commercial_leger",
    }
    return aliases.get(normalized)


def _parse_montant_quebecois(raw):
    """Parse une valeur monetaire quebecoise: '1 234,56' -> 1234.56."""
    s = (raw or "").strip().replace(" ", "")
    if not s:
        return None
    if "," in s and "." in s:
        if s.rindex(",") > s.rindex("."):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    elif "," in s:
        parts = s.split(",")
        if len(parts) == 2 and 1 <= len(parts[1]) <= 2:
            s = s.replace(",", ".")
        else:
            s = s.replace(",", "")
    try:
        return float(s)
    except ValueError:
        return None


def _validate_estimation_items(items, superficie_pi2=None, type_projet=None):
    """Valide les items d'une soumission generee par IA et flag les aberrations."""
    warnings_list = []
    critical_count = 0

    if not items:
        return {
            "valid": True,
            "warnings": ["Aucun item a valider (liste vide)."],
            "critical_count": 0,
            "stats": {"total": 0.0, "total_par_categorie": {}, "ratio_par_pi2": None, "nb_items": 0},
        }

    total_agrege = 0.0
    total_par_categorie = {}
    dedup_keys = {}

    for idx, it in enumerate(items, start=1):
        try:
            qte = float(it.get("quantite", 0) or 0)
        except (TypeError, ValueError):
            qte = 0.0
        try:
            prix = float(it.get("prix_unitaire", 0) or 0)
        except (TypeError, ValueError):
            prix = 0.0

        montant = it.get("montant_ligne")
        if montant is None:
            montant = round(qte * prix, 2)
        else:
            try:
                montant = float(montant)
            except (TypeError, ValueError):
                montant = round(qte * prix, 2)

        description = (it.get("description") or "").strip()
        categorie = (it.get("categorie") or "Sans categorie").strip()
        label = description[:60] if description else f"item #{idx}"

        if prix < _PRIX_UNITAIRE_MIN:
            warnings_list.append(
                f"[CRITIQUE] Ligne {idx} ({label}): prix unitaire {prix:.4f}$ "
                f"sous le seuil minimum ({_PRIX_UNITAIRE_MIN}$)."
            )
            critical_count += 1
        elif prix > _PRIX_UNITAIRE_MAX:
            warnings_list.append(
                f"[CRITIQUE] Ligne {idx} ({label}): prix unitaire {prix:,.2f}$ "
                f"depasse le maximum tolere ({_PRIX_UNITAIRE_MAX:,.0f}$)."
            )
            critical_count += 1

        if qte <= 0:
            warnings_list.append(
                f"[CRITIQUE] Ligne {idx} ({label}): quantite {qte} <= 0."
            )
            critical_count += 1
        elif qte > _QUANTITE_MAX:
            warnings_list.append(
                f"[CRITIQUE] Ligne {idx} ({label}): quantite {qte:,.0f} "
                f"depasse le maximum tolere ({_QUANTITE_MAX:,.0f})."
            )
            critical_count += 1

        desc_lower = description.lower()
        for keyword, expected_cats in _KEYWORDS_CATEGORIE_MAP.items():
            if keyword in desc_lower and categorie not in expected_cats:
                warnings_list.append(
                    f"[SOFT] Ligne {idx} ({label}): description contient "
                    f"\"{keyword}\" mais categorie est \"{categorie}\" "
                    f"(attendu: {', '.join(sorted(expected_cats))})."
                )
                break

        dedup_key = (desc_lower, categorie, round(prix, 2))
        if dedup_key in dedup_keys:
            warnings_list.append(
                f"[SOFT] Ligne {idx} ({label}): doublon probable avec la "
                f"ligne {dedup_keys[dedup_key]} (meme description, categorie "
                f"et prix unitaire)."
            )
        else:
            dedup_keys[dedup_key] = idx

        total_agrege += montant
        total_par_categorie[categorie] = total_par_categorie.get(categorie, 0.0) + montant

    if total_agrege > _TOTAL_AGREGE_MAX:
        warnings_list.append(
            f"[CRITIQUE] Total agrege {total_agrege:,.2f}$ depasse "
            f"{_TOTAL_AGREGE_MAX:,.0f}$. Verifier qu'il s'agit bien d'une "
            f"soumission residentielle ou commercial leger."
        )
        critical_count += 1

    ratio_par_pi2 = None
    if superficie_pi2 and superficie_pi2 > 0:
        ratio_par_pi2 = round(total_agrege / superficie_pi2, 2)
        type_norm = _normalize_projet_type(type_projet)
        fourchette = _APCHQ_FOURCHETTES_PI2_2026.get(type_norm) if type_norm else None
        if fourchette is None:
            fourchette = _APCHQ_FOURCHETTES_DEFAUT
            label_fourchette = "fourchette par defaut (type inconnu)"
        else:
            label_fourchette = f"APCHQ {type_norm}"

        if ratio_par_pi2 < fourchette["min"]:
            warnings_list.append(
                f"[SOFT] Ratio {ratio_par_pi2:,.0f} $/pi2 sous la fourchette "
                f"{label_fourchette} ({fourchette['min']}-{fourchette['max']} $/pi2). "
                f"Soumission peut etre incomplete."
            )
        elif ratio_par_pi2 > fourchette["max"]:
            warnings_list.append(
                f"[CRITIQUE] Ratio {ratio_par_pi2:,.0f} $/pi2 depasse la "
                f"fourchette {label_fourchette} ({fourchette['min']}-{fourchette['max']} "
                f"$/pi2). Verifier les prix unitaires et quantites."
            )
            critical_count += 1

    if total_agrege > 0:
        for cat, montant_cat in total_par_categorie.items():
            ratio_cat = montant_cat / total_agrege
            if ratio_cat > _RATIO_CATEGORIE_MAX:
                warnings_list.append(
                    f"[SOFT] Categorie \"{cat}\" represente "
                    f"{ratio_cat * 100:.0f}% du total ({montant_cat:,.2f}$). "
                    f"Concentration superieure a {_RATIO_CATEGORIE_MAX * 100:.0f}%, "
                    f"verifier l'absence de doublons ou de quantites multipliees."
                )

    return {
        "valid": critical_count == 0,
        "warnings": warnings_list,
        "critical_count": critical_count,
        "stats": {
            "total": round(total_agrege, 2),
            "total_par_categorie": {k: round(v, 2) for k, v in total_par_categorie.items()},
            "ratio_par_pi2": ratio_par_pi2,
            "nb_items": len(items),
        },
    }


def _format_validation_warnings(validation):
    """Formatte les warnings en markdown UI-ready."""
    warnings_list = validation.get("warnings") or []
    if not warnings_list:
        return ""

    critical = [w for w in warnings_list if w.startswith("[CRITIQUE]")]
    soft = [w for w in warnings_list if w.startswith("[SOFT]")]
    info = [w for w in warnings_list if not w.startswith("[CRITIQUE]") and not w.startswith("[SOFT]")]

    parts = ["", "---", "", "## Validation automatique des montants", ""]
    if critical:
        parts.append("### Anomalies critiques a verifier")
        parts.extend(f"- {w[len('[CRITIQUE] '):]}" for w in critical)
        parts.append("")
    if soft:
        parts.append("### Points d'attention")
        parts.extend(f"- {w[len('[SOFT] '):]}" for w in soft)
        parts.append("")
    if info:
        parts.append("### Informations")
        parts.extend(f"- {w}" for w in info)
        parts.append("")

    stats = validation.get("stats") or {}
    parts.append("### Statistiques")
    parts.append(f"- Total agrege: {stats.get('total', 0):,.2f}$")
    if stats.get("ratio_par_pi2") is not None:
        parts.append(f"- Ratio: {stats['ratio_par_pi2']:,.2f} $/pi2")
    parts.append(f"- Nombre d'items: {stats.get('nb_items', 0)}")
    parts.append("")
    return "\n".join(parts)


def _validate_estimation_text(text, superficie_pi2=None):
    """Valide les montants mentionnes dans une estimation textuelle (markdown)."""
    warnings_list = []
    critical_count = 0

    if not text:
        return {
            "valid": True,
            "warnings": ["Texte vide, aucune validation possible."],
            "critical_count": 0,
            "stats": {"montants_detectes": 0, "montant_max": None, "ratio_par_pi2": None},
        }

    montants = []
    for m in _MONTANT_PATTERN.finditer(text):
        raw = m.group(1) or m.group(2) or ""
        val = _parse_montant_quebecois(raw)
        if val is not None:
            montants.append(val)

    if not montants:
        return {
            "valid": True,
            "warnings": ["Aucun montant detecte dans le texte."],
            "critical_count": 0,
            "stats": {"montants_detectes": 0, "montant_max": None, "ratio_par_pi2": None},
        }

    montants_aberrants_bas = [v for v in montants if 0 < v < 1.0]
    montants_aberrants_haut = [v for v in montants if v > _TOTAL_AGREGE_MAX]
    if montants_aberrants_bas:
        warnings_list.append(
            f"[SOFT] {len(montants_aberrants_bas)} montant(s) detecte(s) sous "
            f"1.00$ (peut etre erreurs typographiques): "
            f"{', '.join(f'{v:.4f}$' for v in montants_aberrants_bas[:3])}."
        )
    if montants_aberrants_haut:
        warnings_list.append(
            f"[CRITIQUE] {len(montants_aberrants_haut)} montant(s) detecte(s) "
            f"depassant {_TOTAL_AGREGE_MAX:,.0f}$: "
            f"{', '.join(f'{v:,.2f}$' for v in montants_aberrants_haut[:3])}."
        )
        critical_count += 1

    montant_max = max(montants)
    ratio_par_pi2 = None
    if superficie_pi2 and superficie_pi2 > 0 and montant_max > 0:
        ratio_par_pi2 = round(montant_max / superficie_pi2, 2)
        fourchette = _APCHQ_FOURCHETTES_DEFAUT
        if ratio_par_pi2 > fourchette["max"] * 1.5:
            warnings_list.append(
                f"[SOFT] Plus gros montant detecte ({montant_max:,.2f}$) / "
                f"superficie ({superficie_pi2:,.0f} pi2) = "
                f"{ratio_par_pi2:,.0f} $/pi2 — au-dela des fourchettes "
                f"connues. Verifier si ce montant est bien le total."
            )

    return {
        "valid": critical_count == 0,
        "warnings": warnings_list,
        "critical_count": critical_count,
        "stats": {
            "montants_detectes": len(montants),
            "montant_max": round(montant_max, 2),
            "ratio_par_pi2": ratio_par_pi2,
        },
    }


@router.post("/ai-generate-soumission")
async def ai_generate_soumission(body: AiGenerateSoumissionRequest, user: ErpUser = Depends(get_current_user)):
    """Generate structured soumission items from conversation context using Opus."""
    if not _anthropic_client:
        raise HTTPException(status_code=503, detail="Service IA non disponible")

    # AI billing: guard + credit check
    allowed, error_msg = check_ai_guard(user)
    if not allowed:
        raise HTTPException(status_code=403, detail=error_msg or "Acces IA refuse")
    credits_ok, balance = _check_credits(user)
    if not credits_ok:
        raise HTTPException(status_code=402, detail="Credits IA epuises. Veuillez recharger votre solde.")

    conv_text = ""
    for m in body.messages:
        role = "Client" if m.get("role") == "user" else "Expert"
        conv_text += f"{role}: {m.get('content','')}\n\n"

    # Sprint 2.A: utilise le system block avec cache 1h via _build_system_blocks_with_cache.
    # Si pas de profile_id, on retombe sur "ENTREPRENEUR_GENERAL" comme defaut sensible
    # pour beneficier quand meme du cache hit (le prompt est stable inter-appels).
    effective_profile_id = body.profile_id or "ENTREPRENEUR_GENERAL"
    system_blocks = _build_system_blocks_with_cache(
        profile_id=effective_profile_id,
        schema=user.schema,
        devis_id=None,
    )

    # Scope directive (non cached) — restreint la soumission au metier du profil
    # custom (charge dynamiquement le nom depuis ai_profiles). Pour les profils
    # systeme on derive le nom via title-case (legacy comportement).
    scope_directive = ""
    if body.profile_id:
        if body.profile_id.startswith("custom_"):
            try:
                db_id = int(body.profile_id.split("_", 1)[1])
                if user.schema:
                    conn_p = db.get_conn()
                    try:
                        db.set_tenant(conn_p, user.schema)
                        _ensure_ai_profiles_tables(conn_p, user.schema)
                        cur_p = conn_p.cursor()
                        cur_p.execute("SELECT name FROM ai_profiles WHERE id = %s", (db_id,))
                        prow = cur_p.fetchone()
                        cur_p.close()
                    finally:
                        try:
                            db.reset_tenant(conn_p)
                        except Exception:
                            pass
                        conn_p.close()
                    if prow:
                        pname = prow["name"]
                        scope_directive = (
                            f"L'expert consulte est: {pname}. La soumission doit UNIQUEMENT "
                            f"couvrir les travaux relevant de cette specialite. Ne PAS inclure les "
                            f"corps de metier qui ne relevent pas de ce specialiste.\n\n"
                        )
            except (ValueError, IndexError):
                pass
        else:
            pname = body.profile_id.replace("_", " ").title()
            scope_directive = (
                f"L'expert consulte est un {pname}. La soumission doit UNIQUEMENT "
                f"couvrir les travaux relevant de cette specialite. Ne PAS inclure les "
                f"corps de metier qui ne relevent pas de ce specialiste.\n\n"
            )

    # Sprint 2.A: prompt court qui delegue la structuration au tool_use. Les
    # contraintes (categories, unites, "pas de pourcentages") sont desormais
    # encodees dans le JSON Schema de _SOUMISSION_TOOL — pas besoin de les
    # repeter en texte libre. Claude doit appeler le tool, pas produire du JSON.
    prompt = (
        "Base sur cette conversation entre un client et un expert en construction au Quebec, "
        "genere une soumission detaillee organisee par corps de metier en appelant le tool "
        "`creer_lignes_devis`.\n\n"
        + scope_directive
        + f"Conversation:\n{conv_text}\n"
        + (f"Type de projet: {body.projet_type}\n" if body.projet_type else "")
        + (f"Superficie: {body.superficie} pi2\n" if body.superficie else "")
        + "\nProduis chaque ligne avec quantite estimee, unite normalisee, prix unitaire "
        "CAD 2026 Quebec (jamais 0) et categorie correspondant a un corps de metier. "
        "Regroupe les items par categorie.\n"
    )

    try:
        resp = _call_claude(
            model="claude-opus-4-7",
            max_tokens=32000,  # tool_use input ~ qq lignes JSON structure (decision Sylvain: uniformite 32k)
            temperature=0.1,
            system=system_blocks,
            messages=[{"role": "user", "content": prompt}],
            tools=[_SOUMISSION_TOOL],
            tool_choice={"type": "tool", "name": "creer_lignes_devis"},
        )

        tool_input = _extract_tool_use_input(resp, "creer_lignes_devis")
        if tool_input is None:
            # tool_choice etait force — l'absence d'un tool_use indique soit un
            # refus safety, soit une troncature max_tokens. Retourner 502 plutot
            # qu'une liste vide silencieuse (qui afficherait "0 ligne" sans erreur).
            logger.warning(
                "ai_generate_soumission: tool_use creer_lignes_devis absent — "
                "stop_reason=%s",
                getattr(resp, "stop_reason", "?"),
            )
            raise HTTPException(
                status_code=502,
                detail="Le service IA n'a pas pu structurer la soumission. Reessayez.",
            )

        items = list(tool_input.get("items", []) or [])
        for it in items:
            qte = float(it.get("quantite", 0) or 0)
            prix = float(it.get("prix_unitaire", 0) or 0)
            it["montant_ligne"] = round(qte * prix, 2)

        # Sprint 3 #9: validation post-hoc des montants generes par l'IA.
        try:
            validation = _validate_estimation_items(
                items=items,
                superficie_pi2=body.superficie,
                type_projet=body.projet_type,
            )
        except Exception as val_exc:
            logger.exception("ai_generate_soumission validation failed: %s", val_exc)
            validation = {
                "valid": True,
                "warnings": [f"Validation indisponible: {val_exc}"],
                "critical_count": 0,
                "stats": {},
            }

        if validation.get("critical_count", 0) > 0:
            logger.warning(
                "ai_generate_soumission validation CRITICAL tenant=%s "
                "nb_items=%d total=%.2f ratio_pi2=%s warnings=%d",
                user.schema, len(items),
                validation.get("stats", {}).get("total", 0),
                validation.get("stats", {}).get("ratio_par_pi2"),
                validation.get("critical_count"),
            )

        # AI billing: track usage + deduct credits
        tokens_in = resp.usage.input_tokens
        tokens_out = resp.usage.output_tokens
        cache_creation = getattr(resp.usage, "cache_creation_input_tokens", 0) or 0
        cache_read = getattr(resp.usage, "cache_read_input_tokens", 0) or 0
        # Opus 4.7: $15/M input, $75/M output, $18.75/M cache write, $1.50/M cache read
        cost = (
            tokens_in * 15 / 1_000_000
            + tokens_out * 75 / 1_000_000
            + cache_creation * 18.75 / 1_000_000
            + cache_read * 1.50 / 1_000_000
        ) * 1.30  # 30% markup
        track_ai_usage(user, "devis_generate_soumission", tokens_in, tokens_out, cost, 0, True, model="claude-opus-4-7")
        _deduct_credits(user, cost)

        # Sprint 3 #7: persistance best-effort si devis_id fourni
        persisted_id = None
        if body.devis_id and user.schema:
            try:
                items_dump = json.dumps(items, ensure_ascii=False, default=str)
            except Exception:
                items_dump = "[]"
            persisted_id = _persist_ai_estimation(
                schema=user.schema,
                devis_id=body.devis_id,
                user_id=user.user_id,
                type_estimation="soumission",
                ai_text=items_dump,
                metadata={
                    "profile_id": effective_profile_id,
                    "projet_type": body.projet_type,
                    "superficie": body.superficie,
                    "nb_items": len(items),
                    "validation": validation,
                },
                tokens_in=tokens_in,
                tokens_out=tokens_out,
                cache_creation=cache_creation,
                cache_read=cache_read,
                cost_usd=cost,
                precision_mode=False,
                thinking_tokens=0,
                claude_model="claude-opus-4-7",
            )

        return {
            "items": items,
            # raw_response retire — avec tool_use force, soit items est rempli
            # soit on a deja raise un 502 plus haut. Pas de texte brut a renvoyer.
            "raw_response": None,
            "validation": validation,
            "estimation_id": persisted_id,
            "usage": {
                "input_tokens": tokens_in,
                "output_tokens": tokens_out,
                "cache_creation_input_tokens": cache_creation,
                "cache_read_input_tokens": cache_read,
                "cost_usd": round(cost, 6),
            },
        }
    except HTTPException:
        raise
    except anthropic.APIError as exc:
        logger.error("ai_generate_soumission API error: %s", exc)
        status = getattr(exc, "status_code", 0)
        exc_str = str(exc).lower()
        if status == 413 or (status == 400 and "too_large" in exc_str):
            raise HTTPException(
                status_code=413,
                detail="Contenu trop volumineux pour la generation. Essayez avec moins de messages dans la conversation.",
            )
        if status == 400 and "image exceeds" in exc_str:
            raise HTTPException(status_code=413, detail="Image trop volumineuse (max 5 Mo). Reduisez la taille de l'image ou utilisez un format compresse (JPEG).")
        if status == 529 or "overloaded" in exc_str:
            raise HTTPException(status_code=503, detail="Le service IA est temporairement surcharge. Veuillez reessayer dans quelques secondes.")
        if status == 429 or "rate_limit" in exc_str:
            raise HTTPException(status_code=429, detail="Trop de demandes vers l'IA. Attendez quelques instants avant de reessayer.")
        logger.error("Erreur API Claude: %s", exc)
        raise HTTPException(status_code=503, detail="Service IA temporairement indisponible")
    except Exception as exc:
        logger.error("ai_generate_soumission error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de la generation")


# ---- Conversations CRUD ----

@router.get("/conversations")
async def list_conversations(user: ErpUser = Depends(get_current_user)):
    """List user's estimation IA conversations."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        cur = conn.cursor()
        cur.execute(
            "SELECT id, name, devis_id, subject, status, created_at, updated_at "
            "FROM conversations WHERE user_id = %s OR user_id IS NULL "
            "ORDER BY updated_at DESC NULLS LAST",
            (user.user_id,),
        )
        items = []
        for row in cur.fetchall():
            r = dict(row)
            items.append({
                "id": r["id"], "name": r.get("name",""), "devisId": r.get("devis_id"),
                "subject": r.get("subject",""), "status": r.get("status","active"),
                "createdAt": str(r.get("created_at","")), "updatedAt": str(r.get("updated_at","")),
            })
        return {"items": items}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("list_conversations error: %s", exc)
        return {"items": []}
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.post("/conversations")
async def save_conversation(body: ConversationSave, user: ErpUser = Depends(get_current_user)):
    """Save a new conversation."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        cur = conn.cursor()
        # Defensif: certains tenants legacy ont une colonne `last_updated_at`
        # NOT NULL sans default — l'INSERT qui ne la specifie pas crashe avec
        # "null value in column last_updated_at" (9x/72h vu dans les logs).
        # Helper memoize par schema (skip ALTER si deja sync) pour eviter le
        # lock ACCESS EXCLUSIVE repete sur chaque save (cf. emails.py:107).
        _ensure_conversations_last_updated_at(cur, user.schema)
        msgs_json = json.dumps(body.messages, ensure_ascii=False)
        # NB: last_updated_at est inclus explicitement avec NOW() pour
        # garantir qu'on satisfait toujours la contrainte (meme si l'ALTER
        # ci-dessus a echoue avec un lock timeout). L'ADD COLUMN IF NOT EXISTS
        # garantit que la colonne existe sur tous les tenants — l'INSERT
        # inconditionnel est donc sur.
        cur.execute(
            "INSERT INTO conversations (name, devis_id, user_id, subject, status, messages_json, messages, metadata, created_at, updated_at, last_updated_at) "
            "VALUES (%s,%s,%s,%s,'active',%s::jsonb,%s,%s,NOW(),NOW(),NOW()) RETURNING id",
            (body.name, body.devis_id, user.user_id, body.subject or body.name, msgs_json, msgs_json, body.metadata),
        )
        row = cur.fetchone()
        new_conv_id = row["id"]
        # Commit la conversation d'abord pour garantir qu'elle existe meme si
        # le link orphan echoue plus bas (la conv est plus importante).
        conn.commit()

        # Lier les documents orphelins (conversation_id=NULL uploades recemment
        # par ce user via /ai-analyze-document avant qu'une conversation soit
        # creee) a cette nouvelle conversation. Window: 1 heure. Isole dans sa
        # propre transaction avec rollback en cas d'echec.
        try:
            _ensure_conversation_documents_table(conn, user.schema)
            cur.execute(
                "UPDATE conversation_documents "
                "SET conversation_id = %s, updated_at = NOW() "
                "WHERE conversation_id IS NULL "
                "  AND uploaded_by = %s "
                "  AND created_at >= NOW() - INTERVAL '1 hour'",
                (new_conv_id, user.user_id),
            )
            conn.commit()
        except Exception as link_exc:
            logger.warning("link orphan documents to conversation failed: %s", link_exc)
            try:
                conn.rollback()
            except Exception:
                pass

        # BUG B fix (2026-05-17): liaison bidirectionnelle devis <-> conversation.
        # Avant: conversation.devis_id pointait vers le devis mais devis.conversation_id
        # restait NULL, rendant impossible la navigation devis -> historique IA.
        # On UPDATE seulement si devis.conversation_id IS NULL pour ne pas ecraser
        # un lien existant (un devis peut avoir plusieurs convs, on garde la plus
        # ancienne — la plus "fondatrice" — comme reference principale).
        if body.devis_id:
            try:
                cur.execute(
                    "UPDATE devis SET conversation_id = %s, updated_at = NOW() "
                    "WHERE id = %s AND conversation_id IS NULL",
                    (new_conv_id, body.devis_id),
                )
                conn.commit()
            except Exception as link_devis_exc:
                logger.warning("link conversation to devis %s failed: %s", body.devis_id, link_devis_exc)
                try:
                    conn.rollback()
                except Exception:
                    pass

        return {"id": new_conv_id, "message": "Conversation sauvegardee"}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("save_conversation error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de la sauvegarde")
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.put("/conversations/{conv_id}")
async def update_conversation(conv_id: int, body: ConversationSave, user: ErpUser = Depends(get_current_user)):
    """Update an existing conversation (messages, name, metadata)."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        cur = conn.cursor()
        # FIX: garantir que la colonne last_updated_at existe sur tenants legacy
        # avant d'y faire reference dans le UPDATE (sinon UndefinedColumn 500).
        # Memoize par schema cote helper, donc cout zero apres le premier appel.
        _ensure_conversations_last_updated_at(cur, user.schema)
        msgs_json = json.dumps(body.messages, ensure_ascii=False)
        cur.execute(
            "UPDATE conversations SET name=%s, messages_json=%s::jsonb, messages=%s, "
            "metadata=%s, subject=%s, updated_at=NOW(), last_updated_at=NOW() "
            "WHERE id=%s AND (user_id=%s OR user_id IS NULL)",
            (body.name, msgs_json, msgs_json, body.metadata,
             body.subject or body.name, conv_id, user.user_id),
        )
        # FIX: rowcount check — coherent avec rename_conversation.
        # Retourne 404 explicit si conversation inexistante ou pas owner,
        # plutot qu'un faux 200 silencieux.
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="Conversation non trouvée")
        # FIX: commit explicit — l'absence de commit etait un bug latent qui
        # empechait les updates de persister en mode non-autocommit.
        conn.commit()
        return {"id": conv_id, "message": "Conversation mise à jour"}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("update_conversation error: %s", exc)
        try:
            conn.rollback()
        except Exception:
            pass
        raise HTTPException(status_code=500, detail="Erreur lors de la mise à jour")
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


class ConversationRename(BaseModel):
    """Body pour PATCH /conversations/{id} — rename seul."""
    name: str


@router.patch("/conversations/{conv_id}")
async def rename_conversation(
    conv_id: int,
    body: ConversationRename,
    user: ErpUser = Depends(get_current_user),
):
    """Rename une conversation sans toucher aux messages/metadata.

    Endpoint dedie au rename (vs PUT qui exige tout le payload ConversationSave).
    Ownership check via WHERE user_id + rowcount 404.
    """
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")

    new_name = (body.name or "").strip()
    if not new_name:
        raise HTTPException(status_code=400, detail="Le nom ne peut pas etre vide")
    if len(new_name) > 200:
        raise HTTPException(status_code=400, detail="Le nom est trop long (max 200 caracteres)")

    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        cur = conn.cursor()
        # FIX: garantir que la colonne last_updated_at existe sur tenants legacy
        # avant d'y faire reference dans le UPDATE (sinon UndefinedColumn 500).
        _ensure_conversations_last_updated_at(cur, user.schema)
        cur.execute(
            "UPDATE conversations SET name = %s, updated_at = NOW(), last_updated_at = NOW() "
            "WHERE id = %s AND (user_id = %s OR user_id IS NULL) "
            "RETURNING id",
            (new_name, conv_id, user.user_id),
        )
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Conversation non trouvee")
        conn.commit()
        return {"id": conv_id, "name": new_name, "message": "Conversation renommee"}
    except HTTPException:
        raise
    except Exception as exc:
        try:
            conn.rollback()
        except Exception:
            pass
        logger.error("rename_conversation error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors du renommage")
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.get("/conversations/{conv_id}")
async def get_conversation(conv_id: int, user: ErpUser = Depends(get_current_user)):
    """Get a single conversation with messages."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        cur = conn.cursor()
        cur.execute("SELECT * FROM conversations WHERE id = %s AND (user_id = %s OR user_id IS NULL)", (conv_id, user.user_id))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Conversation non trouvée")
        r = dict(row)
        messages = []
        if r.get("messages_json"):
            messages = r["messages_json"] if isinstance(r["messages_json"], list) else json.loads(r["messages_json"])
        elif r.get("messages"):
            try:
                messages = json.loads(r["messages"])
            except (json.JSONDecodeError, TypeError):
                pass
        return {
            "id": r["id"], "name": r.get("name",""), "devisId": r.get("devis_id"),
            "subject": r.get("subject",""), "status": r.get("status","active"),
            "messages": messages, "metadata": r.get("metadata"),
            "createdAt": str(r.get("created_at","")), "updatedAt": str(r.get("updated_at","")),
        }
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("get_conversation error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.delete("/conversations/{conv_id}")
async def delete_conversation(conv_id: int, user: ErpUser = Depends(get_current_user)):
    """Delete a conversation."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        cur = conn.cursor()
        # FIX-3: SELECT all anthropic_file_id des documents lies AVANT le DELETE
        # cascade pour cleanup Files API apres le DELETE OK (best-effort).
        file_ids_to_delete: list = []
        try:
            cur.execute(
                "SELECT cd.anthropic_file_id FROM conversation_documents cd "
                "JOIN conversations c ON c.id = cd.conversation_id "
                "WHERE cd.conversation_id = %s "
                "  AND (c.user_id = %s OR c.user_id IS NULL) "
                "  AND cd.anthropic_file_id IS NOT NULL",
                (conv_id, user.user_id),
            )
            for _row_fid in cur.fetchall():
                _fid = _row_fid.get("anthropic_file_id") if isinstance(_row_fid, dict) else _row_fid[0]
                if _fid:
                    file_ids_to_delete.append(_fid)
        except Exception as _sel_exc:
            logger.warning("SELECT anthropic_file_ids pre-delete-conversation failed: %s", _sel_exc)
        cur.execute("DELETE FROM conversations WHERE id = %s AND (user_id = %s OR user_id IS NULL)", (conv_id, user.user_id))
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="Conversation non trouvée")
        # FIX: commit explicit — coherent avec update_conversation/rename_conversation.
        # Sans commit, le DELETE ne persiste pas en mode non-autocommit.
        conn.commit()
        # FIX-3: cleanup Anthropic Files API APRES commit DB OK (best-effort).
        for _fid in file_ids_to_delete:
            try:
                _delete_from_anthropic_files(_fid)
            except Exception:
                logger.warning("Anthropic file delete failed for %s (non-critical)", _fid)
        return {"message": "Conversation supprimée"}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("delete_conversation error: %s", exc)
        try:
            conn.rollback()
        except Exception:
            pass
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


# ---- Conversation Documents CRUD ----
# Permet au frontend de lister/ajouter/retirer/toggle les documents persistes
# d'une conversation (plan PDF, devis Excel, etc. accessibles a Claude via cache).

@router.get("/conversations/{conv_id}/documents")
async def list_conversation_documents(
    conv_id: int, user: ErpUser = Depends(get_current_user)
):
    """Liste les documents persistes d'une conversation (metadata seulement,
    pas le BYTEA pour alleger)."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        _ensure_conversation_documents_table(conn, user.schema)
        cur = conn.cursor()
        # Ownership check: le user doit etre proprietaire de la conversation
        # (ou la conv doit etre orpheline user_id IS NULL pour backward compat)
        cur.execute(
            "SELECT cd.id, cd.filename, cd.media_type, cd.file_size, cd.summary, "
            "cd.category_detected, cd.subcategory_detected, "
            "cd.superficie_pi2, cd.superficie_renovation_pi2, "
            "cd.superficie_agrandissement_pi2, cd.superficie_existant_conserve_pi2, "
            "cd.is_active_context, cd.created_at "
            "FROM conversation_documents cd "
            "JOIN conversations c ON c.id = cd.conversation_id "
            "WHERE cd.conversation_id = %s "
            "  AND (c.user_id = %s OR c.user_id IS NULL) "
            "ORDER BY cd.created_at ASC",
            (conv_id, user.user_id),
        )
        items = []
        for row in cur.fetchall():
            r = dict(row)
            items.append({
                "id": r["id"],
                "filename": r["filename"],
                "mediaType": r.get("media_type"),
                "fileSize": r.get("file_size"),
                "summary": r.get("summary"),
                "category": r.get("category_detected"),
                "subcategory": r.get("subcategory_detected"),
                "superficiePi2": r.get("superficie_pi2"),
                "superficieRenovationPi2": r.get("superficie_renovation_pi2"),
                "superficieAgrandissementPi2": r.get("superficie_agrandissement_pi2"),
                "superficieExistantConservePi2": r.get("superficie_existant_conserve_pi2"),
                "isActiveContext": r.get("is_active_context", True),
                "createdAt": str(r.get("created_at", "")),
            })
        return {"items": items}
    except Exception as exc:
        logger.error("list_conversation_documents error: %s", exc)
        return {"items": []}
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.get("/conversations/{conv_id}/documents/{doc_id}/download")
async def download_conversation_document(
    conv_id: int, doc_id: int, user: ErpUser = Depends(get_current_user)
):
    """Telecharge le binaire d'un document persiste."""
    from fastapi.responses import Response
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        cur = conn.cursor()
        # Ownership check via conversations.user_id
        cur.execute(
            "SELECT cd.filename, cd.media_type, cd.content FROM conversation_documents cd "
            "JOIN conversations c ON c.id = cd.conversation_id "
            "WHERE cd.id = %s AND cd.conversation_id = %s "
            "  AND (c.user_id = %s OR c.user_id IS NULL)",
            (doc_id, conv_id, user.user_id),
        )
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Document non trouve")
        r = dict(row)
        content = r.get("content")
        if not content:
            raise HTTPException(status_code=404, detail="Contenu binaire absent")
        # Escape filename pour eviter HTTP header injection (CRLF, quotes).
        # RFC 5987: filename*=UTF-8''<urlencoded> pour caracteres non-ASCII,
        # + fallback ASCII dans filename=.
        from urllib.parse import quote as _urlquote
        raw_fn = (r["filename"] or "document").replace("\r", "").replace("\n", "")
        ascii_fn = raw_fn.encode("ascii", "replace").decode("ascii").replace('"', "").replace("\\", "")
        utf8_fn = _urlquote(raw_fn, safe="")
        cd_header = f'attachment; filename="{ascii_fn}"; filename*=UTF-8\'\'{utf8_fn}'
        return Response(
            content=bytes(content),
            media_type=r.get("media_type") or "application/octet-stream",
            headers={"Content-Disposition": cd_header},
        )
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("download_conversation_document error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur telechargement")
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.put("/conversations/{conv_id}/documents/{doc_id}/toggle")
async def toggle_document_active(
    conv_id: int, doc_id: int, user: ErpUser = Depends(get_current_user)
):
    """Toggle is_active_context: activer/desactiver un document pour les
    prochains messages (sans le supprimer)."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        cur = conn.cursor()
        # Ownership check via EXISTS subquery (UPDATE cannot use JOIN directly)
        cur.execute(
            "UPDATE conversation_documents "
            "SET is_active_context = NOT COALESCE(is_active_context, TRUE), updated_at = NOW() "
            "WHERE id = %s AND conversation_id = %s "
            "  AND EXISTS (SELECT 1 FROM conversations c WHERE c.id = conversation_id "
            "              AND (c.user_id = %s OR c.user_id IS NULL)) "
            "RETURNING is_active_context",
            (doc_id, conv_id, user.user_id),
        )
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Document non trouve")
        conn.commit()
        return {"id": doc_id, "isActiveContext": row["is_active_context"]}
    except HTTPException:
        raise
    except Exception as exc:
        try:
            conn.rollback()
        except Exception:
            pass
        logger.error("toggle_document_active error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur toggle")
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.delete("/conversations/{conv_id}/documents/{doc_id}")
async def delete_conversation_document(
    conv_id: int, doc_id: int, user: ErpUser = Depends(get_current_user)
):
    """Supprime definitivement un document persiste."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        cur = conn.cursor()
        # FIX-3: SELECT anthropic_file_id avant DELETE pour cleanup Files API
        # apres le DELETE OK (best-effort, n'echoue pas le DELETE).
        file_id_to_delete = None
        try:
            cur.execute(
                "SELECT cd.anthropic_file_id FROM conversation_documents cd "
                "JOIN conversations c ON c.id = cd.conversation_id "
                "WHERE cd.id = %s AND cd.conversation_id = %s "
                "  AND (c.user_id = %s OR c.user_id IS NULL)",
                (doc_id, conv_id, user.user_id),
            )
            _row_pre = cur.fetchone()
            if _row_pre:
                file_id_to_delete = _row_pre.get("anthropic_file_id") if isinstance(_row_pre, dict) else _row_pre[0]
        except Exception as _sel_exc:
            logger.warning("SELECT anthropic_file_id pre-delete failed: %s", _sel_exc)
        # Ownership check via EXISTS subquery
        cur.execute(
            "DELETE FROM conversation_documents "
            "WHERE id = %s AND conversation_id = %s "
            "  AND EXISTS (SELECT 1 FROM conversations c WHERE c.id = conversation_id "
            "              AND (c.user_id = %s OR c.user_id IS NULL))",
            (doc_id, conv_id, user.user_id),
        )
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="Document non trouve")
        conn.commit()
        # FIX-3: cleanup Anthropic Files API APRES commit DB OK (best-effort).
        if file_id_to_delete:
            try:
                _delete_from_anthropic_files(file_id_to_delete)
            except Exception:
                logger.warning("Anthropic file delete failed for %s (non-critical)", file_id_to_delete)
        return {"message": "Document supprime"}
    except HTTPException:
        raise
    except Exception as exc:
        try:
            conn.rollback()
        except Exception:
            pass
        logger.error("delete_conversation_document error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur suppression")
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.get("/{devis_id}")
async def get_devis(devis_id: int, user: ErpUser = Depends(get_current_user)):
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        _ensure_devis_pct_columns(conn, user.schema)
        cursor = conn.cursor()
        cursor.execute(
            "SELECT d.*, c.nom as client_nom "
            "FROM devis d "
            "LEFT JOIN companies c ON d.client_company_id = c.id "
            "WHERE d.id = %s",
            (devis_id,),
        )
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Devis non trouvé")
        d = dict(row)
        for k in ("created_at", "date_prevu", "date_soumis", "date_decision", "signature_date"):
            if d.get(k) is not None:
                d[k] = str(d[k])
        for k in ("total_travaux", "tps", "tvq", "investissement_total",
                   "total_avant_taxes", "prix_estime", "administration", "contingences",
                   "profit", "administration_pct", "contingences_pct", "profit_pct"):
            if d.get(k) is not None:
                d[k] = float(d[k])
        # Get lignes
        _ensure_visibility_columns(cursor)
        cursor.execute(
            "SELECT id, description, quantite, unite, prix_unitaire, montant_ligne, "
            "sequence_ligne, categorie, notes_ligne, COALESCE(visible, TRUE) as visible, "
            "mo_pct, mat_pct, admin_pct_ligne, contingence_pct_ligne, profit_pct_ligne "
            "FROM devis_lignes WHERE devis_id = %s ORDER BY sequence_ligne ASC",
            (devis_id,),
        )
        lignes = []
        for l in cursor.fetchall():
            ld = dict(l)
            for k in ("quantite", "prix_unitaire", "montant_ligne", "mo_pct", "mat_pct",
                     "admin_pct_ligne", "contingence_pct_ligne", "profit_pct_ligne"):
                if ld.get(k) is not None:
                    ld[k] = float(ld[k])
            lignes.append(ld)
        d["lignes"] = lignes
        return d
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("get_devis error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


# ============================================
# DEVIS CONVERSION & AI ENDPOINTS
# ============================================

@router.post("/{devis_id}/convert-to-project")
async def convert_devis_to_project(devis_id: int, user: ErpUser = Depends(get_current_user)):
    """Convert an accepted devis into a project."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        cursor = conn.cursor()

        # Fetch the devis
        cursor.execute("SELECT * FROM devis WHERE id = %s", (devis_id,))
        devis = cursor.fetchone()
        if not devis:
            raise HTTPException(status_code=404, detail="Devis non trouvé")
        devis = dict(devis)

        # Verify status allows conversion
        if devis.get("statut") not in ("Accepte", "Acceptée", "Termine"):
            raise HTTPException(
                status_code=400,
                detail=f"Le devis doit être Accepté ou Terminé pour être converti (statut actuel: {devis.get('statut')})"
            )

        # Idempotent: if already converted, return the existing link instead
        # of raising 400. Avoids confusing errors on double-click / retry and
        # allows the frontend button to serve as a "navigate to project" cta.
        if devis.get("project_id"):
            return {
                "project_id": devis["project_id"],
                "created": False,
                "message": "Un projet est deja lie a ce devis",
            }

        now = datetime.now().isoformat()

        # Ensure all columns we're about to INSERT exist on this tenant (lazy migration)
        _ensure_projects_insert_columns(cursor, conn, user.schema)

        # Get opportunity info if linked
        opportunity_id = devis.get("opportunity_id")

        # Fix sequence if out of sync (can happen after manual INSERTs).
        # Use GREATEST(max, 1) + 3-arg setval to avoid "value 0 is out of bounds"
        # when projects table is empty (PostgreSQL sequence range is 1..2^31-1).
        try:
            cursor.execute(
                "SELECT setval(pg_get_serial_sequence('projects', 'id'), "
                "GREATEST(COALESCE((SELECT MAX(id) FROM projects), 0), 1), "
                "(SELECT COUNT(*) > 0 FROM projects))"
            )
        except Exception as seq_exc:
            logger.warning("setval projects_id_seq failed: %s", seq_exc)
            try:
                conn.rollback()
                db.set_tenant(conn, user.schema)
            except Exception:
                pass

        # Compute budget with same fallback chain as _create_project_from_devis
        # (investissement_total → total_avant_taxes → prix_estime → total_travaux
        # + computed TPS/TVQ). Prevents budget=0 when only the subtotal fields
        # are populated (which is common for devis generated via Metre / Manuel).
        budget = float(devis.get("investissement_total") or 0)
        if devis.get("investissement_total") is None:
            st = float(devis.get("total_avant_taxes") or devis.get("prix_estime") or devis.get("total_travaux") or 0)
            tps_v = float(devis.get("tps") or 0)
            tvq_v = float(devis.get("tvq") or 0)
            if devis.get("tps") is None and st:
                tps_v = round(st * 0.05, 2)
            if devis.get("tvq") is None and st:
                tvq_v = round(st * 0.09975, 2)
            budget = round(st + tps_v + tvq_v, 2)

        # Insert into projects — use RETURNING id
        cursor.execute(
            "INSERT INTO projects (nom_projet, client_company_id, client_contact_id, "
            "client_nom_cache, po_client, statut, priorite, tache, date_soumis, date_prevu, "
            "date_debut_reel, bd_ft_estime, prix_estime, budget_total, description, "
            "devis_id, devis_source_id, numero_devis, "
            "created_at, updated_at) "
            "VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, "
            "%s, %s, %s, %s, %s) RETURNING id",
            (
                devis.get("nom_projet"),
                devis.get("client_company_id"),
                devis.get("client_contact_id"),
                devis.get("client_nom_cache"),
                devis.get("po_client"),
                "En cours",
                devis.get("priorite", "Moyenne"),
                devis.get("tache"),
                devis.get("date_soumis"),
                devis.get("date_prevu"),
                now[:10],  # date_debut_reel = today
                devis.get("bd_ft_estime"),
                devis.get("prix_estime"),
                budget,
                devis.get("description"),
                devis_id,
                devis_id,
                devis.get("numero_devis"),
                now,
                now,
            ),
        )
        project_id = cursor.fetchone()["id"]

        # Generate numero_projet (PROJ-YYYY-NNNNN) — coherent avec
        # projects._generate_numero_projet et _backfill_numero_projet.
        # Try/except simple : la connexion du pool est en autocommit
        # (database_config.py:368,394 ISOLATION_LEVEL_AUTOCOMMIT), donc
        # chaque statement = sa propre transaction. Si l'UPDATE echoue
        # (ex: colonne manquante sur tenant pathologique malgre
        # _ensure_projects_insert_columns appele plus haut), l'erreur
        # est isolee — les statements suivants restent OK. Le _backfill
        # paresseux dans projects.py:189 rattrapera au prochain list/get.
        # IMPORTANT: PAS de SAVEPOINT — en autocommit, SAVEPOINT raise
        # InvalidSavepointSpecification (no transaction block), no-op silencieux.
        try:
            cursor.execute(
                "UPDATE projects "
                "SET numero_projet = 'PROJ-' || EXTRACT(YEAR FROM COALESCE(created_at, CURRENT_TIMESTAMP))::int "
                "|| '-' || LPAD(id::text, 5, '0') "
                "WHERE id = %s AND (numero_projet IS NULL OR numero_projet = '')",
                (project_id,),
            )
        except Exception as numproj_exc:
            logger.error("Could not set numero_projet for project %s: %s — _backfill_numero_projet rattrapera au prochain list_projects", project_id, numproj_exc)

        # Atomic race-safe link: only set if project_id is still NULL. Two
        # concurrent clicks on "Convertir en projet" both INSERT a project;
        # the loser (rowcount=0) must delete its duplicate and return the
        # winning project_id instead. Mirrors the pattern in
        # _create_project_from_devis:3710.
        cursor.execute(
            "UPDATE devis SET project_id = %s WHERE id = %s AND project_id IS NULL",
            (project_id, devis_id),
        )
        if cursor.rowcount == 0:
            try:
                cursor.execute("DELETE FROM projects WHERE id = %s", (project_id,))
            except Exception as del_exc:
                logger.warning(
                    "convert_devis race: failed to delete duplicate project %s: %s",
                    project_id, del_exc,
                )
            cursor.execute("SELECT project_id FROM devis WHERE id = %s", (devis_id,))
            winner = cursor.fetchone()
            if winner and winner.get("project_id"):
                logger.info(
                    "convert_devis race: duplicate %s discarded, returning winner %s",
                    project_id, winner["project_id"],
                )
                return {
                    "project_id": winner["project_id"],
                    "devis_id": devis_id,
                    "opportunity_updated": False,
                    "created": False,
                    "message": "Un projet est deja lie a ce devis",
                }
            raise HTTPException(status_code=500, detail="Echec de liaison devis/projet")

        # Update linked opportunity if exists. Defensive ALTER for old tenants
        # whose opportunities table predates projet_id / converted_at /
        # updated_at. Mirrors accept_public_devis:4340.
        if opportunity_id:
            for col, ctype in (
                ("projet_id", "INTEGER"),
                ("converted_at", "TIMESTAMP"),
                ("updated_at", "TIMESTAMP"),
            ):
                try:
                    cursor.execute(
                        f"ALTER TABLE opportunities ADD COLUMN IF NOT EXISTS {col} {ctype}"
                    )
                except Exception as opp_alter_exc:
                    logger.warning("ALTER opportunities ADD %s failed: %s", col, opp_alter_exc)
            try:
                cursor.execute(
                    "UPDATE opportunities SET projet_id = %s, statut = 'GAGNE', "
                    "converted_at = CURRENT_TIMESTAMP, updated_at = CURRENT_TIMESTAMP "
                    "WHERE id = %s",
                    (project_id, opportunity_id),
                )
            except Exception as opp_exc:
                logger.warning("UPDATE opportunities failed for devis %s: %s", devis_id, opp_exc)

        # Copy attachments (devis_attachments → project_attachments). Best-effort;
        # failure does not roll back the conversion.
        try:
            cursor.execute(
                "INSERT INTO project_attachments "
                "(project_id, filename, original_filename, file_size, file_type, file_extension, "
                "category, description, file_path, file_data, file_hash, upload_date, "
                "uploaded_by, is_active) "
                "SELECT %s, filename, original_filename, file_size, file_type, file_extension, "
                "category, description, file_path, file_data, file_hash, upload_date, "
                "uploaded_by, is_active "
                "FROM devis_attachments WHERE devis_id = %s AND is_active = TRUE",
                (project_id, devis_id),
            )
            if cursor.rowcount > 0:
                logger.info("Copied %d attachments from devis %s to project %s", cursor.rowcount, devis_id, project_id)
        except Exception as att_exc:
            logger.warning("Could not copy devis attachments on convert: %s", att_exc)

        return {
            "project_id": project_id,
            "devis_id": devis_id,
            "opportunity_updated": opportunity_id is not None,
            "created": True,
            "message": "Projet créé depuis le devis",
        }

    except HTTPException:
        raise
    except Exception as exc:
        logger.error("convert_devis_to_project error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de la conversion du devis en projet")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


class AiEstimatePayload(BaseModel):
    """Body parametre optionnel pour /ai-estimate.

    `precision_mode` = True (defaut) active Adaptive Thinking sur Opus 4.7
    avec effort="high" (Claude raisonne quasi systematiquement avant de produire
    la reponse). Plus precis (recommande, defaut). ~3x plus cher (~60-90s vs 30s).
    """
    precision_mode: bool = True


@router.post("/{devis_id}/ai-estimate")
async def ai_estimate_devis(
    devis_id: int,
    payload: Optional[AiEstimatePayload] = None,
    user: ErpUser = Depends(get_current_user),
):
    """Get AI estimation for a devis using Claude.

    Mode standard (default): _call_claude (streaming, max_tokens=32000).
    Mode precision (`precision_mode=True`): _call_claude_with_thinking
    (Adaptive Thinking effort="high", max_response_tokens=100000). Sylvain
    veut le mode "Precision etendue" applique a l'Estimation IA.
    """
    if not _anthropic_client:
        raise HTTPException(status_code=503, detail="Service IA non disponible (client Anthropic non configure)")

    # AI billing: guard + credit check
    allowed, error_msg = check_ai_guard(user)
    if not allowed:
        raise HTTPException(status_code=403, detail=error_msg or "Acces IA refuse")
    credits_ok, balance = _check_credits(user)
    if not credits_ok:
        raise HTTPException(status_code=402, detail="Credits IA epuises. Veuillez recharger votre solde.")

    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")

    precision_mode = bool(payload.precision_mode) if payload else True

    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        cursor = conn.cursor()

        # Fetch devis
        cursor.execute("SELECT * FROM devis WHERE id = %s", (devis_id,))
        devis = cursor.fetchone()
        if not devis:
            raise HTTPException(status_code=404, detail="Devis non trouvé")
        devis = dict(devis)

        # Fetch lignes
        cursor.execute(
            "SELECT description, quantite, unite, prix_unitaire, montant_ligne, categorie "
            "FROM devis_lignes WHERE devis_id = %s ORDER BY sequence_ligne ASC",
            (devis_id,),
        )
        lignes = [dict(r) for r in cursor.fetchall()]

        # Build prompt
        lignes_text = ""
        for i, l in enumerate(lignes, 1):
            qte = float(l.get("quantite", 0) or 0)
            prix = float(l.get("prix_unitaire", 0) or 0)
            montant = float(l.get("montant_ligne", 0) or 0)
            lignes_text += (
                f"  {i}. {l.get('description', '')} - Qte: {qte} {l.get('unite', '')} "
                f"x {prix:.2f}$ = {montant:.2f}$ [{l.get('categorie', 'General')}]\n"
            )
        if not lignes_text:
            lignes_text = "  (Aucune ligne existante)\n"

        total_travaux = float(devis.get("total_travaux", 0) or 0)
        investissement = float(devis.get("investissement_total", 0) or 0)

        # Sprint 2.B: prompt trim — on retire "SCAN SYSTEMATIQUE OBLIGATOIRE",
        # "RIGUEUR" et "RECOMMENCE 2eme fois" (instructions meta-cognitives qui
        # ballonnent le prompt sans gain mesurable car Adaptive Thinking effort=high
        # fait deja le raisonnement multi-pass en interne). Le profil ENTREPRENEUR_GENERAL
        # injecte en system block contient deja les tarifs CCQ 2026 et la checklist
        # corps de metier — pas besoin de les repeter en user prompt.
        prompt = (
            f"Analyse ce devis et fournis une estimation detaillee.\n\n"
            f"Projet: {devis.get('nom_projet', 'Non specifie')}\n"
            f"Description: {devis.get('description', 'Aucune description')}\n"
            f"Total travaux actuel: {total_travaux:.2f}$\n"
            f"Investissement total (TTC): {investissement:.2f}$\n\n"
            f"Lignes existantes:\n{lignes_text}\n"
            f"Livrables attendus:\n"
            f"1. Une analyse des prix (sont-ils realistes pour le Quebec en 2026?)\n"
            f"2. Des items potentiellement manquants\n"
            f"3. Des recommandations pour optimiser les couts\n"
            f"4. Une estimation des couts CCQ et CNESST si applicable\n"
            f"5. Un resume avec fourchette de prix estimee\n\n"
            "GROUNDING CATALOGUE: utilise l'outil `recherche_catalogue_produits` pour "
            "valider les prix unitaires des materiaux mentionnes dans les lignes "
            "existantes (ex: 'gypse', 'fenetre PVC') et detecter les ecarts vs le "
            "catalogue reel du tenant. Si un produit n'est pas au catalogue, signaler "
            "explicitement qu'on utilise un prix marche estime.\n\n"
            f"Format: markdown structure, en francais, avec sections claires.\n"
        )

        # Sprint 2.B: system block + cache TTL 1h via _build_system_blocks_with_cache.
        # profile_id="ENTREPRENEUR_GENERAL" sert de profil par defaut pour beneficier
        # du cache hit cross-conversation (~130 KB de prompt expert reutilise).
        # devis_id transmis pour permettre au helper d'enrichir le contexte (bloc non
        # cache car varie par devis).
        system_blocks = _build_system_blocks_with_cache(
            profile_id="ENTREPRENEUR_GENERAL",
            schema=user.schema,
            devis_id=devis_id,
        )

        thinking_tokens = 0
        # Cumul des usages sur tous les appels Claude (boucle tool-use).
        tokens_in_total = 0
        tokens_out_total = 0
        cache_creation_total = 0
        cache_read_total = 0

        # Boucle tool-use: Claude peut appeler `recherche_catalogue_produits`
        # jusqu'a max_iterations fois pour grounder les prix sur le catalogue
        # tenant avant de produire l'estimation finale en texte.
        max_iterations = 5
        messages_loop = [{"role": "user", "content": prompt}]
        message = None
        last_stop_reason = None

        for iteration in range(max_iterations):
            if precision_mode:
                message = _call_claude_with_thinking(
                    model="claude-opus-4-7",
                    thinking_budget=10000,
                    max_response_tokens=60000,
                    system=system_blocks,
                    messages=messages_loop,
                    tools=[_CATALOGUE_LOOKUP_TOOL],
                )
                thinking_tokens += _count_thinking_tokens_safe(message)
            else:
                message = _call_claude(
                    model="claude-opus-4-7",
                    max_tokens=32000,
                    system=system_blocks,
                    messages=messages_loop,
                    tools=[_CATALOGUE_LOOKUP_TOOL],
                )

            tokens_in_total += getattr(message.usage, "input_tokens", 0) or 0
            tokens_out_total += getattr(message.usage, "output_tokens", 0) or 0
            cache_creation_total += getattr(message.usage, "cache_creation_input_tokens", 0) or 0
            cache_read_total += getattr(message.usage, "cache_read_input_tokens", 0) or 0

            last_stop_reason = getattr(message, "stop_reason", None)
            if last_stop_reason != "tool_use":
                break

            tool_blocks = _extract_tool_use_blocks(message)
            if not tool_blocks:
                logger.warning(
                    "ai_estimate_devis: stop_reason=tool_use mais aucun bloc extrait (devis_id=%s)",
                    devis_id,
                )
                break

            tool_results_blocks = []
            for tb in tool_blocks:
                result = _execute_estimation_tool(
                    tool_name=tb["name"],
                    tool_input=tb["input"],
                    schema=user.schema,
                )
                result_json = json.dumps(result, ensure_ascii=False, default=str)
                if len(result_json) > 12000:
                    logger.info(
                        "tool_result truncated tool=%s len=%d -> 12000",
                        tb["name"], len(result_json),
                    )
                    result_json = result_json[:12000] + "...[TRUNCATED]"
                tool_results_blocks.append({
                    "type": "tool_result",
                    "tool_use_id": tb["id"],
                    "content": result_json,
                })

            messages_loop.append({"role": "assistant", "content": message.content})
            messages_loop.append({"role": "user", "content": tool_results_blocks})
        else:
            logger.warning(
                "ai_estimate_devis: max_iterations=%d atteint pour devis_id=%s "
                "(stop_reason=%s) - estimation peut etre incomplete",
                max_iterations, devis_id, last_stop_reason,
            )

        if precision_mode:
            ai_text = _extract_text_blocks(message) if message else ""
        else:
            ai_text = _extract_text_from_response(message) if message else ""

        if not ai_text and last_stop_reason == "tool_use":
            ai_text = (
                "L'estimation a necessite plusieurs recherches dans le catalogue "
                "mais n'a pas pu etre finalisee dans la limite imposee. "
                "Veuillez relancer ou affiner le contexte du devis."
            )

        # Sprint 3 #9: validation post-hoc du texte d'estimation.
        try:
            superficie_validation = devis.get("superficie_pi2") or devis.get("superficie")
            if superficie_validation:
                try:
                    superficie_validation = float(superficie_validation)
                except (TypeError, ValueError):
                    superficie_validation = None
            validation_text = _validate_estimation_text(
                text=ai_text,
                superficie_pi2=superficie_validation,
            )
        except Exception as val_exc:
            logger.exception("ai_estimate_devis validation failed: %s", val_exc)
            validation_text = {
                "valid": True,
                "warnings": [f"Validation indisponible: {val_exc}"],
                "critical_count": 0,
                "stats": {},
            }

        ai_text = ai_text + _format_validation_warnings(validation_text)

        if validation_text.get("critical_count", 0) > 0:
            logger.warning(
                "ai_estimate_devis validation CRITICAL devis_id=%s tenant=%s "
                "montants=%d max=%s warnings=%d",
                devis_id, user.schema,
                validation_text.get("stats", {}).get("montants_detectes"),
                validation_text.get("stats", {}).get("montant_max"),
                validation_text.get("critical_count"),
            )

        # AI billing: track usage + deduct credits (cumul TOUS les appels)
        tokens_in = tokens_in_total
        tokens_out = tokens_out_total
        cache_creation = cache_creation_total
        cache_read = cache_read_total
        # Opus 4.7: $15/M input, $75/M output, $18.75/M cache write, $1.50/M cache read
        cost = (
            tokens_in * 15 / 1_000_000
            + tokens_out * 75 / 1_000_000
            + cache_creation * 18.75 / 1_000_000
            + cache_read * 1.50 / 1_000_000
        ) * 1.30  # 30% markup
        feature_name = "devis_ai_estimate_precision" if precision_mode else "devis_ai_estimate"
        track_ai_usage(user, feature_name, tokens_in, tokens_out, cost, 0, True, model="claude-opus-4-7")
        _deduct_credits(user, cost)

        # Sprint 3 #7: persistance best-effort
        persisted_id = _persist_ai_estimation(
            schema=user.schema,
            devis_id=devis_id,
            user_id=user.user_id,
            type_estimation="simple",
            ai_text=ai_text,
            metadata={
                "nom_projet": devis.get("nom_projet"),
                "validation": validation_text,
                "feature_name": feature_name,
                "stop_reason": last_stop_reason,
            },
            tokens_in=tokens_in,
            tokens_out=tokens_out,
            cache_creation=cache_creation,
            cache_read=cache_read,
            cost_usd=cost,
            precision_mode=precision_mode,
            thinking_tokens=thinking_tokens,
            claude_model="claude-opus-4-7",
        )

        return {
            "estimation": ai_text,
            "devis_id": devis_id,
            "nom_projet": devis.get("nom_projet"),
            "precision_mode_used": precision_mode,
            "thinking_tokens": thinking_tokens,
            "validation": validation_text,
            "estimation_id": persisted_id,
            "usage": {
                "input_tokens": tokens_in,
                "output_tokens": tokens_out,
                "cache_creation_input_tokens": cache_creation,
                "cache_read_input_tokens": cache_read,
                "cost_usd": round(cost, 6),
            },
        }

    except HTTPException:
        raise
    except anthropic.APIError as exc:
        logger.error("ai_estimate_devis API error: %s", exc)
        status = getattr(exc, "status_code", 0)
        exc_str = str(exc).lower()
        if status == 413 or (status == 400 and "too_large" in exc_str):
            raise HTTPException(
                status_code=413,
                detail="Devis trop volumineux pour l'estimation IA. Essayez avec moins de lignes.",
            )
        if status == 400 and "image exceeds" in exc_str:
            raise HTTPException(status_code=413, detail="Image trop volumineuse (max 5 Mo). Reduisez la taille de l'image ou utilisez un format compresse (JPEG).")
        if status == 529 or "overloaded" in exc_str:
            raise HTTPException(status_code=503, detail="Le service IA est temporairement surcharge. Veuillez reessayer dans quelques secondes.")
        if status == 429 or "rate_limit" in exc_str:
            raise HTTPException(status_code=429, detail="Trop de demandes vers l'IA. Attendez quelques instants avant de reessayer.")
        logger.error("Erreur API Claude: %s", exc)
        raise HTTPException(status_code=503, detail="Service IA temporairement indisponible")
    except Exception as exc:
        logger.error("ai_estimate_devis error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de l'estimation IA")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.post("/{devis_id}/ai-estimate-with-plan")
async def ai_estimate_with_plan(
    devis_id: int,
    file: UploadFile = File(...),
    precision_mode: bool = Form(default=True),
    multipass_mode: bool = Form(default=False),  # Sprint 3 #5 MVP
    additional_context: Optional[str] = Form(default=None),
    user: ErpUser = Depends(get_current_user),
):
    """Estimation IA avec analyse Vision du plan PDF/image.

    Combine:
    - Vision Claude Opus 4.7 sur le plan upload
    - Contexte des lignes existantes du devis
    - Analyse CCQ Quebec 2026 (categories, prix, tarifs)
    - Optionnel Adaptive Thinking effort="high" (precision_mode=True, ~3x cost)

    Retourne estimation textuelle structuree avec recommandations.
    """
    if not _anthropic_client:
        raise HTTPException(status_code=503, detail="Service IA non disponible")

    # AI billing guards
    allowed, error_msg = check_ai_guard(user)
    if not allowed:
        raise HTTPException(status_code=403, detail=error_msg or "Acces IA refuse")
    credits_ok, balance = _check_credits(user)
    if not credits_ok:
        raise HTTPException(status_code=402, detail="Credits IA epuises")

    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")

    try:
        # 1. Validate file
        file_bytes = await file.read()
        if not file_bytes:
            raise HTTPException(status_code=400, detail="Fichier vide")
        if len(file_bytes) > 32 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="Fichier trop volumineux (max 32 Mo)")
        filename = file.filename or "plan"
        file_ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        media_type_map = {
            "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
            "gif": "image/gif", "webp": "image/webp", "pdf": "application/pdf",
        }
        media_type = media_type_map.get(file_ext)
        if not media_type:
            raise HTTPException(
                status_code=400,
                detail=f"Format non supporte: {filename}. Formats acceptes: PDF, PNG, JPG, GIF, WebP."
            )

        # 2. Compress image si > 4.5 MB (limite Claude 5 MB)
        # Opus 4.7 native res = 2576 px sur le long edge.
        if media_type.startswith("image/") and len(file_bytes) > 4.5 * 1024 * 1024:
            try:
                from PIL import Image
                import io as _io
                img = Image.open(_io.BytesIO(file_bytes))
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                if max(img.size) > 2576:
                    img.thumbnail((2576, 2576), Image.LANCZOS)
                buf = _io.BytesIO()
                img.save(buf, format='JPEG', quality=85, optimize=True)
                file_bytes = buf.getvalue()
                media_type = "image/jpeg"
            except Exception as exc:
                logger.warning("Image compression failed: %s", exc)

        if media_type.startswith("image/") and len(file_bytes) > 5 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="Image trop volumineuse apres compression")

        # 3. Verifier media type via magic bytes (corrige extensions menteuses)
        actual_media_type = _detect_media_type_from_bytes(file_bytes) or media_type
        b64_data = base64.standard_b64encode(file_bytes).decode("utf-8")

        # 4. Recuperer contexte devis existant
        conn = db.get_conn()
        devis_data = None
        existing_lignes = []
        try:
            db.set_tenant(conn, user.schema)
            cur = conn.cursor()
            cur.execute("SELECT * FROM devis WHERE id = %s", (devis_id,))
            devis_data = cur.fetchone()
            if not devis_data:
                cur.close()
                raise HTTPException(status_code=404, detail="Devis non trouve")
            devis_data = dict(devis_data)

            cur.execute(
                "SELECT description, quantite, unite, prix_unitaire, montant_ligne, categorie "
                "FROM devis_lignes WHERE devis_id = %s ORDER BY sequence_ligne",
                (devis_id,)
            )
            existing_lignes = [dict(r) for r in cur.fetchall()]
            cur.close()
        finally:
            try:
                db.reset_tenant(conn)
            except Exception:
                pass
            conn.close()

        # 5. Charger profil ENTREPRENEUR_GENERAL
        # Split today_line (non cacheable) du profil (cacheable TTL 1h) pour eviter
        # cache miss quotidien a minuit.
        entrepreneur_prompt = _load_profile_system_prompt("ENTREPRENEUR_GENERAL", user.schema)
        system_blocks = [
            # Bloc 1: today_line seul (change chaque jour, non cacheable)
            {"type": "text", "text": _today_prompt_line()},
            # Bloc 2: profil + markdown rules (cacheable TTL 1h)
            {
                "type": "text",
                "text": entrepreneur_prompt + _MARKDOWN_TABLE_RULES,
                "cache_control": {"type": "ephemeral", "ttl": "1h"},
            },
        ]

        # 6. Construire le prompt avec contexte devis + plan + scan systematique
        def _safe_float(v):
            try:
                return float(v) if v is not None else 0.0
            except (TypeError, ValueError):
                return 0.0

        lignes_text = "\n".join([
            f"- {l.get('description', '')}: {_safe_float(l.get('quantite'))} {l.get('unite', '')} "
            f"x {_safe_float(l.get('prix_unitaire')):.2f}$ = {_safe_float(l.get('montant_ligne')):.2f}$"
            f" [{l.get('categorie') or 'Sans categorie'}]"
            for l in existing_lignes
        ]) or "(aucune ligne existante)"

        ctx_block = f"\n\nCONTEXTE ADDITIONNEL UTILISATEUR: {additional_context}" if additional_context else ""

        total_travaux = _safe_float(devis_data.get("total_travaux"))
        investissement = _safe_float(devis_data.get("investissement_total"))

        prompt = f"""Tu es expert estimateur en construction au Quebec/CCQ.

CONTEXTE DEVIS EXISTANT:
- Numero: {devis_data.get('numero_devis') or 'N/A'}
- Nom projet: {devis_data.get('nom_projet') or 'Non specifie'}
- Description projet: {devis_data.get('description') or 'Non specifiee'}
- Type projet: {devis_data.get('type_projet') or 'Non specifie'}
- Total travaux: {total_travaux:.2f}$
- Investissement total (TTC): {investissement:.2f}$

LIGNES EXISTANTES DU DEVIS:
{lignes_text}{ctx_block}

PLAN PDF/IMAGE FOURNI (analyser via Vision):
- Lis ATTENTIVEMENT le plan: dimensions, annotations, pieces, materiaux indiques
- Extrais surfaces (pi.ca), longueurs (pi.li), comptages (portes, fenetres, prises, etc.)

ANALYSE A PRODUIRE:

1. **VERIFICATION DES LIGNES EXISTANTES** (coherence avec le plan):
   - Pour chaque ligne du devis: verifie si quantite/dimensions correspondent au plan
   - Identifie les ECARTS (sous-estimations, sur-estimations)
   - Verifie les prix unitaires vs marche Quebec 2026

2. **ITEMS MANQUANTS** (CCQ checklist):
   - Permis municipaux + plans
   - Excavation / preparation terrain (si applicable)
   - Demolition existant (si renovation)
   - Disposition dechets / location conteneur
   - Echafaudage / equipement special
   - Inspections (electricite RBQ, plomberie, gaz, batiment)
   - Nettoyage final
   - Garanties (1 an / 5 ans)
   - CNESST / cotisations sociales
   - Profit / contingences (15-25% standard)
   - Imprevus chantier

3. **RECOMMANDATIONS**:
   - Lignes a ajouter avec quantite estimee et fourchette prix
   - Lignes a ajuster (avec justification)
   - Risques techniques visibles sur le plan

4. **ESTIMATION GLOBALE**:
   - Fourchette finale revisee (min - max - probable)
   - Justification synthetique

SCAN SYSTEMATIQUE OBLIGATOIRE:
1. Identifie chaque PIECE/ZONE du plan
2. Pour CHAQUE zone, liste ce que tu vois (dimensions, surfaces, finitions)
3. Compare aux lignes existantes du devis
4. Repere les oublis evidents
5. RECOMMENCE l'analyse une 2eme fois pour verifier

Reponds en MARKDOWN structure (titres ##, listes, tableaux). Sois precis avec des chiffres.
Inclus une SECTION "RESUME EXECUTIF" en haut (3-5 bullet points)."""

        # Sprint 3 #10 - Citations PDF (anti-hallucination, tracabilite)
        is_pdf = (media_type == "application/pdf")

        # 7. Call Claude (single-pass ou multipass Sprint 3 #5)
        thinking_tokens = 0
        citations_list: list = []
        multipass_result = None  # Sprint 3 #5 - reference pour metadata persistance
        if multipass_mode:
            # Sprint 3 #5 MVP - multipass anti-hallucination
            # Pass 1 = vision exhaustive JSON; Pass 2 = analyse business markdown
            try:
                multipass_result = _run_estimate_with_plan_multipass(
                    pdf_b64=b64_data,
                    media_type=actual_media_type,
                    devis_data=devis_data,
                    lignes_text=lignes_text,
                    additional_context=additional_context or "",
                    filename=filename,
                    is_pdf=is_pdf,
                    system_blocks=system_blocks,
                    precision_mode=precision_mode,
                )
            except HTTPException:
                raise
            except anthropic.APIError as exc:
                logger.error("ai_estimate_with_plan multipass API error: %s", exc)
                status = getattr(exc, "status_code", 0)
                exc_str = str(exc).lower()
                if status == 413 or (status == 400 and "too_large" in exc_str):
                    raise HTTPException(status_code=413, detail="Plan trop volumineux pour l'estimation IA.")
                if status == 400 and "image exceeds" in exc_str:
                    raise HTTPException(status_code=413, detail="Image trop volumineuse (max 5 Mo).")
                if status == 529 or "overloaded" in exc_str:
                    raise HTTPException(status_code=503, detail="Le service IA est temporairement surcharge.")
                if status == 429 or "rate_limit" in exc_str:
                    raise HTTPException(status_code=429, detail="Trop de demandes vers l'IA.")
                raise HTTPException(status_code=502, detail="Erreur lors de l'appel IA. Reessaie plus tard.")
            except Exception:
                logger.exception("Claude Vision multipass call failed")
                raise HTTPException(status_code=502, detail="Erreur lors de l'appel IA. Reessaie plus tard.")

            response_text = multipass_result["response_text"]
            citations_list = multipass_result["citations_list"]
            tokens_in = multipass_result["tokens_in"]
            tokens_out = multipass_result["tokens_out"]
            cache_creation = multipass_result["cache_creation"]
            cache_read = multipass_result["cache_read"]
            thinking_tokens = multipass_result["thinking_tokens"]
        else:
            # Mode single-pass existant (inchange)
            plan_block = {
                "type": "document" if is_pdf else "image",
                "source": {"type": "base64", "media_type": actual_media_type, "data": b64_data},
                "cache_control": {"type": "ephemeral"},
            }
            if is_pdf:
                plan_block["citations"] = {"enabled": True}

            user_content = [
                plan_block,
                {"type": "text", "text": prompt},
            ]

            try:
                if precision_mode:
                    # Mode precision: Adaptive Thinking effort="high" (default).
                    # thinking_budget IGNORE sur Opus 4.7, kept for backward compat.
                    message = _call_claude_with_thinking(
                        model="claude-opus-4-7",
                        thinking_budget=10000,
                        max_response_tokens=100000,
                        system=system_blocks,
                        messages=[{"role": "user", "content": user_content}],
                    )
                    # Sprint 3 #10 - Format reponse avec footnotes citations PDF
                    if is_pdf:
                        response_text = _format_response_with_footnotes(message, document_titles=[filename])
                        citations_list = _extract_citations_from_response(message)
                    else:
                        response_text = _extract_text_blocks(message)
                    thinking_tokens = _count_thinking_tokens_safe(message)
                else:
                    message = _call_claude(
                        model="claude-opus-4-7",
                        max_tokens=32000,
                        system=system_blocks,
                        messages=[{"role": "user", "content": user_content}],
                    )
                    # Sprint 3 #10 - Format reponse avec footnotes citations PDF
                    if is_pdf:
                        response_text = _format_response_with_footnotes(message, document_titles=[filename])
                        citations_list = _extract_citations_from_response(message)
                    else:
                        response_text = _extract_text_from_response(message)
            except anthropic.APIError as exc:
                logger.error("ai_estimate_with_plan API error: %s", exc)
                status = getattr(exc, "status_code", 0)
                exc_str = str(exc).lower()
                if status == 413 or (status == 400 and "too_large" in exc_str):
                    raise HTTPException(status_code=413, detail="Plan trop volumineux pour l'estimation IA.")
                if status == 400 and "image exceeds" in exc_str:
                    raise HTTPException(status_code=413, detail="Image trop volumineuse (max 5 Mo).")
                if status == 529 or "overloaded" in exc_str:
                    raise HTTPException(status_code=503, detail="Le service IA est temporairement surcharge.")
                if status == 429 or "rate_limit" in exc_str:
                    raise HTTPException(status_code=429, detail="Trop de demandes vers l'IA.")
                raise HTTPException(status_code=502, detail="Erreur lors de l'appel IA. Reessaie plus tard.")
            except Exception:
                logger.exception("Claude Vision estimation call failed")
                raise HTTPException(status_code=502, detail="Erreur lors de l'appel IA. Reessaie plus tard.")

            # 8. Billing
            # Round 22A lecon: thinking_tokens deja inclus dans output_tokens, NE PAS double-billing.
            usage = getattr(message, 'usage', None)
            tokens_in = getattr(usage, 'input_tokens', 0) if usage else 0
            tokens_out = getattr(usage, 'output_tokens', 0) if usage else 0
            tokens_in = tokens_in or 0
            tokens_out = tokens_out or 0
            cache_creation = getattr(usage, "cache_creation_input_tokens", 0) if usage else 0
            cache_read = getattr(usage, "cache_read_input_tokens", 0) if usage else 0
            cache_creation = cache_creation or 0
            cache_read = cache_read or 0
        # Opus 4.7: $15/M input, $75/M output, $18.75/M cache write, $1.50/M cache read
        cost_usd = (
            tokens_in * 15 / 1_000_000
            + tokens_out * 75 / 1_000_000
            + cache_creation * 18.75 / 1_000_000
            + cache_read * 1.50 / 1_000_000
        ) * 1.30  # 30% markup

        feature_label = "devis_ai_estimate_with_plan_precision" if precision_mode else "devis_ai_estimate_with_plan"
        try:
            track_ai_usage(user, feature_label, tokens_in, tokens_out, cost_usd, 0, True, model="claude-opus-4-7")
        except Exception:
            logger.exception("track_ai_usage failed")
        try:
            _deduct_credits(user, cost_usd)
        except Exception:
            logger.error(
                "CRITICAL: _deduct_credits FAILED tenant=%s cost=%s - REVENUE LEAK",
                user.schema, cost_usd,
            )

        # Sprint 3 #9: validation post-hoc du texte d'estimation Vision.
        try:
            superficie_validation = None
            if additional_context:
                m_super = re.search(
                    r"([0-9]{1,5}(?:[.,][0-9]+)?)\s*(?:pi\.?2|pi2|pi\.?ca|p\.c\.|sqft|sf)",
                    additional_context.lower(),
                )
                if m_super:
                    val = _parse_montant_quebecois(m_super.group(1))
                    if val and 10 < val < 1_000_000:
                        superficie_validation = val
            if superficie_validation is None:
                superficie_validation = devis_data.get("superficie_pi2") or devis_data.get("superficie")
                if superficie_validation:
                    try:
                        superficie_validation = float(superficie_validation)
                    except (TypeError, ValueError):
                        superficie_validation = None

            validation_text = _validate_estimation_text(
                text=response_text,
                superficie_pi2=superficie_validation,
            )
        except Exception as val_exc:
            logger.exception("ai_estimate_with_plan validation failed: %s", val_exc)
            validation_text = {
                "valid": True,
                "warnings": [f"Validation indisponible: {val_exc}"],
                "critical_count": 0,
                "stats": {},
            }

        response_text = response_text + _format_validation_warnings(validation_text)

        if validation_text.get("critical_count", 0) > 0:
            logger.warning(
                "ai_estimate_with_plan validation CRITICAL devis_id=%s tenant=%s "
                "filename=%s warnings=%d",
                devis_id, user.schema, filename,
                validation_text.get("critical_count"),
            )

        # Sprint 3 #7: persistance best-effort
        persisted_id = _persist_ai_estimation(
            schema=user.schema,
            devis_id=devis_id,
            user_id=user.user_id,
            type_estimation="with_plan",
            ai_text=response_text,
            metadata={
                "filename": filename,
                "additional_context": additional_context,
                "validation": validation_text,
                "superficie_pi2_used": superficie_validation,
                # Sprint 3 #5 MVP multipass
                "multipass_mode": multipass_mode,
                "pass1_summary": multipass_result.get("pass1_summary", "") if multipass_result else None,
                "pass1_zones_count": multipass_result.get("pass1_zones_count", 0) if multipass_result else None,
                "pass1_elements_count": multipass_result.get("pass1_elements_count", 0) if multipass_result else None,
            },
            tokens_in=tokens_in,
            tokens_out=tokens_out,
            cache_creation=cache_creation,
            cache_read=cache_read,
            cost_usd=cost_usd,
            precision_mode=precision_mode,
            thinking_tokens=thinking_tokens,
            claude_model="claude-opus-4-7",
        )

        return {
            "estimation": response_text,
            "devis_id": devis_id,
            "filename": filename,
            "precision_mode_used": precision_mode,
            "thinking_tokens": thinking_tokens,
            "validation": validation_text,
            "estimation_id": persisted_id,
            "citations": citations_list,
            "cost_usd": round(cost_usd, 6),
            "tokens_in": tokens_in,
            "tokens_out": tokens_out,
            "claude_model": "claude-opus-4-7",
        }

    except HTTPException:
        raise
    except Exception:
        logger.exception("ai_estimate_with_plan failed")
        raise HTTPException(status_code=500, detail="Erreur interne IA")


@router.get("/{devis_id}/ai-estimations")
async def list_devis_ai_estimations(
    devis_id: int,
    limit: int = 20,
    include_archived: bool = False,
    user: ErpUser = Depends(get_current_user),
):
    """Liste les estimations IA d'un devis (audit + recovery)."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    limit = max(1, min(int(limit or 20), 100))

    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        _ensure_devis_ai_estimations_table(conn, user.schema)
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM devis WHERE id = %s", (devis_id,))
        if not cursor.fetchone():
            raise HTTPException(status_code=404, detail="Devis non trouve")

        if include_archived:
            cursor.execute(
                "SELECT id, devis_id, user_id, type_estimation, "
                "LEFT(ai_text, 500) AS ai_text_preview, "
                "LENGTH(ai_text) AS ai_text_length, "
                "metadata_json, tokens_in, tokens_out, "
                "cache_creation_tokens, cache_read_tokens, "
                "cost_usd, precision_mode, thinking_tokens, "
                "claude_model, created_at, archived "
                "FROM devis_ai_estimations "
                "WHERE devis_id = %s "
                "ORDER BY created_at DESC LIMIT %s",
                (devis_id, limit),
            )
        else:
            cursor.execute(
                "SELECT id, devis_id, user_id, type_estimation, "
                "LEFT(ai_text, 500) AS ai_text_preview, "
                "LENGTH(ai_text) AS ai_text_length, "
                "metadata_json, tokens_in, tokens_out, "
                "cache_creation_tokens, cache_read_tokens, "
                "cost_usd, precision_mode, thinking_tokens, "
                "claude_model, created_at, archived "
                "FROM devis_ai_estimations "
                "WHERE devis_id = %s AND archived = FALSE "
                "ORDER BY created_at DESC LIMIT %s",
                (devis_id, limit),
            )
        items = []
        for row in cursor.fetchall():
            r = dict(row)
            if r.get("created_at") is not None:
                r["created_at"] = str(r["created_at"])
            if r.get("cost_usd") is not None:
                r["cost_usd"] = float(r["cost_usd"])
            items.append(r)
        return {"items": items, "count": len(items), "limit": limit}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("list_devis_ai_estimations error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.get("/{devis_id}/ai-estimations/{estimation_id}")
async def get_devis_ai_estimation(
    devis_id: int,
    estimation_id: int,
    user: ErpUser = Depends(get_current_user),
):
    """Recupere une estimation IA specifique (markdown complet)."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        _ensure_devis_ai_estimations_table(conn, user.schema)
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, devis_id, user_id, type_estimation, ai_text, "
            "metadata_json, tokens_in, tokens_out, "
            "cache_creation_tokens, cache_read_tokens, "
            "cost_usd, precision_mode, thinking_tokens, "
            "claude_model, created_at, archived "
            "FROM devis_ai_estimations "
            "WHERE id = %s AND devis_id = %s",
            (estimation_id, devis_id),
        )
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Estimation non trouvee")
        r = dict(row)
        if r.get("created_at") is not None:
            r["created_at"] = str(r["created_at"])
        if r.get("cost_usd") is not None:
            r["cost_usd"] = float(r["cost_usd"])
        return r
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("get_devis_ai_estimation error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.delete("/{devis_id}/ai-estimations/{estimation_id}")
async def archive_devis_ai_estimation(
    devis_id: int,
    estimation_id: int,
    user: ErpUser = Depends(get_current_user),
):
    """Soft delete (archived=TRUE) d'une estimation IA."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        _ensure_devis_ai_estimations_table(conn, user.schema)
        cursor = conn.cursor()
        cursor.execute(
            "UPDATE devis_ai_estimations "
            "SET archived = TRUE "
            "WHERE id = %s AND devis_id = %s AND archived = FALSE",
            (estimation_id, devis_id),
        )
        if cursor.rowcount == 0:
            raise HTTPException(
                status_code=404,
                detail="Estimation non trouvee ou deja archivee",
            )
        conn.commit()
        logger.info(
            "AI estimation archived: tenant=%s devis_id=%s estimation_id=%s user_id=%s",
            user.schema, devis_id, estimation_id, user.user_id,
        )
        return {"message": "Estimation archivee", "id": estimation_id}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("archive_devis_ai_estimation error: %s", exc)
        try:
            conn.rollback()
        except Exception:
            pass
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


class DevisBatchUpdate(BaseModel):
    devis_ids: List[int]
    statut: Optional[str] = None


@router.post("/batch-update")
async def batch_update_devis(body: DevisBatchUpdate, user: ErpUser = Depends(get_current_user)):
    """Batch update statut for multiple devis."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    if not body.devis_ids:
        raise HTTPException(status_code=400, detail="Aucun devis selectionne")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        cursor = conn.cursor()
        updated = 0
        if body.statut:
            cursor.execute(
                "UPDATE devis SET statut = %s, updated_at = CURRENT_TIMESTAMP "
                "WHERE id = ANY(%s)",
                (body.statut, body.devis_ids),
            )
            updated = cursor.rowcount
        conn.commit()
        return {"updated": updated, "message": f"{updated} soumission(s) mise(s) a jour"}
    except Exception as exc:
        logger.error("batch_update_devis error: %s", exc)
        conn.rollback()
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.post("")
async def create_devis(body: DevisCreate, user: ErpUser = Depends(get_current_user)):
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        _ensure_devis_pct_columns(conn, user.schema)
        cursor = conn.cursor()
        # Generate devis number: TEMP-then-UPDATE pattern (race-safe)
        year = datetime.now().year
        token = _generate_readable_token(body.nom_projet)
        # Resolve client name for cache
        client_nom_cache = body.client_nom_direct
        if body.client_company_id and not client_nom_cache:
            cursor.execute("SELECT nom FROM companies WHERE id = %s", (body.client_company_id,))
            crow = cursor.fetchone()
            if crow:
                client_nom_cache = crow["nom"]
        adm_pct = body.administration_pct if body.administration_pct is not None else _DEFAULT_ADM_PCT
        con_pct = body.contingences_pct if body.contingences_pct is not None else _DEFAULT_CON_PCT
        pro_pct = body.profit_pct if body.profit_pct is not None else _DEFAULT_PRO_PCT
        type_soum = body.type_soumission or "Détaillée"
        # Si prix_estime fourni a la creation, calculer la cascade HT->TTC
        # pour que la colonne Montant / panel Total TTC affichent des valeurs
        # correctes des la creation (sans attendre un PUT subsequent).
        ht_init = float(body.prix_estime) if body.prix_estime is not None else 0.0
        if ht_init > 0:
            admin_init = round(ht_init * float(adm_pct) / 100, 2)
            conting_init = round(ht_init * float(con_pct) / 100, 2)
            profit_init = round(ht_init * float(pro_pct) / 100, 2)
            total_at_init = round(ht_init + admin_init + conting_init + profit_init, 2)
            tps_init = round(total_at_init * 0.05, 2)
            tvq_init = round(total_at_init * 0.09975, 2)
            inv_total_init = round(total_at_init + tps_init + tvq_init, 2)
        else:
            admin_init = conting_init = profit_init = 0.0
            total_at_init = tps_init = tvq_init = inv_total_init = 0.0
        cursor.execute(
            "INSERT INTO devis (numero_devis, nom_projet, description, client_company_id, client_contact_id, "
            "client_nom_direct, client_nom_cache, project_id, statut, total_travaux, tps, tvq, "
            "investissement_total, created_at, date_prevu, date_soumis, date_fin, notes, "
            "po_client, priorite, tache, prix_estime, "
            "administration, contingences, profit, total_avant_taxes, "
            "administration_pct, contingences_pct, profit_pct, "
            "type_soumission, validation_token) "
            "VALUES ('TEMP',%s,%s,%s,%s,%s,%s,%s,'Brouillon',%s,%s,%s,%s,"
            "CURRENT_TIMESTAMP,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) RETURNING id",
            (body.nom_projet, body.description, body.client_company_id,
             body.client_contact_id, body.client_nom_direct, client_nom_cache, body.project_id,
             ht_init, tps_init, tvq_init, inv_total_init,
             body.date_prevu, body.date_soumis, body.date_fin, body.notes,
             body.po_client, body.priorite, body.tache, body.prix_estime,
             admin_init, conting_init, profit_init, total_at_init,
             adm_pct, con_pct, pro_pct, type_soum, token),
        )
        row = cursor.fetchone()
        devis_id = row["id"]
        numero = f"DEV-{year}-{devis_id:03d}"
        cursor.execute("UPDATE devis SET numero_devis = %s WHERE id = %s", (numero, devis_id))
        # Register token in public lookup table
        _register_public_token(conn, token, user.schema, devis_id)
        return {"id": devis_id, "numero_devis": numero, "message": "Devis créé"}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("create_devis error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.put("/{devis_id}")
async def update_devis(devis_id: int, body: DevisUpdate, user: ErpUser = Depends(get_current_user)):
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    ALLOWED = {"nom_projet", "statut", "description", "date_prevu", "date_soumis", "date_fin",
               "notes", "total_travaux", "tps", "tvq", "investissement_total",
               "client_company_id", "client_contact_id", "client_nom_direct",
               "po_client", "priorite", "tache", "prix_estime",
               "administration_pct", "contingences_pct", "profit_pct",
               "administration", "contingences", "profit",
               "show_administration", "show_contingences", "show_profit",
               "show_unite", "show_quantite", "show_prix_unitaire", "show_montant_ligne", "show_mo_mat",
               "administration_label", "contingences_label", "profit_label",
               "conditions_text", "exclusions_text", "show_conditions", "show_exclusions",
               "type_soumission"}
    fields = {k: v for k, v in body.model_dump().items() if v is not None and k in ALLOWED}
    if not fields:
        raise HTTPException(status_code=400, detail="Aucun champ")
    # Aligne total_travaux sur prix_estime quand l'admin saisit un montant
    # dans la modal "Modifier la soumission" (champ "Prix ($)"). Sans cet
    # alignement, prix_estime est persiste mais n'alimente PAS le recalcul
    # cascade (total_avant_taxes, tps, tvq, investissement_total restent a 0)
    # -> la colonne Montant et le panel Total TTC affichent 0.
    if "prix_estime" in fields and "total_travaux" not in fields:
        fields["total_travaux"] = float(fields["prix_estime"])
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        _ensure_devis_pct_columns(conn, user.schema)
        cursor = conn.cursor()
        _ensure_visibility_columns(cursor)
        # Recalculate if total_travaux, percentages, or dollar amounts changed
        pct_changed = any(k in fields for k in ("administration_pct", "contingences_pct", "profit_pct"))
        amt_changed = any(k in fields for k in ("administration", "contingences", "profit"))
        if "total_travaux" in fields or pct_changed or amt_changed:
            # Read current devis to get existing values (with fallback if pct columns missing)
            current = _read_devis_pct_with_amounts(cursor, devis_id)
            if not current:
                raise HTTPException(status_code=404, detail="Devis non trouvé")
            # Detect if pct columns exist in DB (fallback sets them to None)
            pct_cols_exist = current.get("administration_pct") is not None or current.get("contingences_pct") is not None or current.get("profit_pct") is not None
            ht = float(fields.get("total_travaux", current["total_travaux"] or 0))

            # Dollar amounts take priority: if a $ amount is provided, back-calculate the %
            if "administration" in fields:
                administration = round(float(fields["administration"]), 2)
                adm_pct_val = round(administration / ht * 100, 4) if ht > 0 else 0
                if pct_cols_exist:
                    fields["administration_pct"] = adm_pct_val
            else:
                cur_adm = current.get("administration_pct")
                adm_pct_val = float(fields.get("administration_pct", cur_adm if cur_adm is not None else 3.0))
                administration = round(ht * adm_pct_val / 100, 2)

            if "contingences" in fields:
                contingences = round(float(fields["contingences"]), 2)
                con_pct_val = round(contingences / ht * 100, 4) if ht > 0 else 0
                if pct_cols_exist:
                    fields["contingences_pct"] = con_pct_val
            else:
                cur_con = current.get("contingences_pct")
                con_pct_val = float(fields.get("contingences_pct", cur_con if cur_con is not None else 12.0))
                contingences = round(ht * con_pct_val / 100, 2)

            if "profit" in fields:
                profit = round(float(fields["profit"]), 2)
                pro_pct_val = round(profit / ht * 100, 4) if ht > 0 else 0
                if pct_cols_exist:
                    fields["profit_pct"] = pro_pct_val
            else:
                cur_pro = current.get("profit_pct")
                pro_pct_val = float(fields.get("profit_pct", cur_pro if cur_pro is not None else 15.0))
                profit = round(ht * pro_pct_val / 100, 2)

            # If pct columns don't exist yet, strip them from fields to avoid UPDATE failure
            if not pct_cols_exist:
                fields.pop("administration_pct", None)
                fields.pop("contingences_pct", None)
                fields.pop("profit_pct", None)

            total_avant_taxes = round(ht + administration + contingences + profit, 2)
            tps = round(total_avant_taxes * 0.05, 2)
            tvq = round(total_avant_taxes * 0.09975, 2)
            investissement_total = round(total_avant_taxes + tps + tvq, 2)
            fields["total_travaux"] = ht
            fields["administration"] = administration
            fields["contingences"] = contingences
            fields["profit"] = profit
            fields["total_avant_taxes"] = total_avant_taxes
            fields["tps"] = tps
            fields["tvq"] = tvq
            fields["investissement_total"] = investissement_total
        set_parts = [f"{k} = %s" for k in fields]
        set_parts.append("updated_at = CURRENT_TIMESTAMP")
        values = list(fields.values()) + [devis_id]
        cursor.execute(f"UPDATE devis SET {', '.join(set_parts)} WHERE id = %s", values)
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="Devis non trouvé")

        # If the global %s (or amounts) changed, the totals computed above
        # apply them uniformly to total_ht — but per-line overrides on
        # admin_pct_ligne / contingence_pct_ligne / profit_pct_ligne would
        # be flattened out. Re-aggregate via the helper so each line uses
        # COALESCE(line_override, devis_global_just_updated). The helper
        # writes the final totals back into the devis row.
        if pct_changed or amt_changed or "total_travaux" in fields:
            _recompute_devis_totals(cursor, devis_id)

        # Auto-create project in two cases:
        #   1. statut transitioning Envoye/Brouillon → Accepte (normal flow)
        #   2. devis already in Accepte state but has no project_id (orphan
        #      recovery from the pre-a001128 setval bug — the old code
        #      committed statut='Accepte' then crashed on setval, leaving
        #      devis stranded without a project). Any save on the devis now
        #      triggers retroactive project creation.
        project_id = None
        trigger_creation = False
        if fields.get("statut") == "Accepte":
            trigger_creation = True
        else:
            cursor.execute(
                "SELECT statut, project_id FROM devis WHERE id = %s",
                (devis_id,),
            )
            cur_row = cursor.fetchone()
            if cur_row and cur_row.get("statut") == "Accepte" and not cur_row.get("project_id"):
                trigger_creation = True

        if trigger_creation:
            try:
                db.set_tenant(conn, user.schema)
                project_id = _create_project_from_devis(conn, devis_id, schema=user.schema)
            except Exception as exc_proj:
                logger.warning("Auto-create project from devis failed: %s", exc_proj)

        result = {"message": "Devis mis à jour"}
        if project_id:
            result["project_id"] = project_id
            result["message"] = "Devis mis à jour — Projet créé automatiquement"
        return result
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("update_devis error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


# ============================================
# DEVIS LIGNES
# ============================================


def _recompute_devis_totals(cursor, devis_id: int) -> dict:
    """Recompute total_ht, admin/contingence/profit, taxes and grand total
    for the devis, applying per-line markup overrides where present.

    Each line contributes:
        admin       = montant_ligne × COALESCE(admin_pct_ligne, devis.admin_pct) / 100
        contingence = montant_ligne × COALESCE(contingence_pct_ligne, devis.contingence_pct) / 100
        profit      = montant_ligne × COALESCE(profit_pct_ligne, devis.profit_pct) / 100

    A NULL on a per-line override falls back to the devis-level percentage,
    so lines that don't override behave exactly like before — full backward
    compatibility for existing data. The summed admin/contingence/profit
    persisted on the `devis` row stays a flat aggregate (no schema change).

    Persists the recomputed totals into the `devis` row and returns the dict.
    """
    adm_pct, con_pct, pro_pct = _read_devis_pct(cursor, devis_id)
    cursor.execute(
        """
        SELECT
            COALESCE(SUM(montant_ligne), 0) AS total_ht,
            COALESCE(SUM(montant_ligne * COALESCE(admin_pct_ligne, %s) / 100.0), 0) AS administration,
            COALESCE(SUM(montant_ligne * COALESCE(contingence_pct_ligne, %s) / 100.0), 0) AS contingences,
            COALESCE(SUM(montant_ligne * COALESCE(profit_pct_ligne, %s) / 100.0), 0) AS profit
        FROM devis_lignes
        WHERE devis_id = %s
        """,
        (adm_pct, con_pct, pro_pct, devis_id),
    )
    row = cursor.fetchone()
    total_ht = float(row["total_ht"] or 0)
    administration = round(float(row["administration"] or 0), 2)
    contingences = round(float(row["contingences"] or 0), 2)
    profit = round(float(row["profit"] or 0), 2)
    total_avant_taxes = round(total_ht + administration + contingences + profit, 2)
    tps = round(total_avant_taxes * 0.05, 2)
    tvq = round(total_avant_taxes * 0.09975, 2)
    total = round(total_avant_taxes + tps + tvq, 2)
    cursor.execute(
        "UPDATE devis SET total_travaux = %s, administration = %s, contingences = %s, "
        "profit = %s, total_avant_taxes = %s, tps = %s, tvq = %s, "
        "investissement_total = %s WHERE id = %s",
        (total_ht, administration, contingences, profit, total_avant_taxes, tps, tvq, total, devis_id),
    )
    return {
        "total_ht": total_ht,
        "administration": administration,
        "contingences": contingences,
        "profit": profit,
        "total_avant_taxes": total_avant_taxes,
        "tps": tps,
        "tvq": tvq,
        "total": total,
    }

@router.post("/{devis_id}/lignes")
async def add_devis_ligne(devis_id: int, body: DevisLigneCreate, user: ErpUser = Depends(get_current_user)):
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        _ensure_devis_pct_columns(conn, user.schema)
        cursor = conn.cursor()
        _ensure_visibility_columns(cursor)
        montant = round(body.quantite * body.prix_unitaire, 2)
        if body.sequence_ligne in (None, 0):
            cursor.execute(
                "SELECT COALESCE(MAX(sequence_ligne), 0) + 1 as next_seq "
                "FROM devis_lignes WHERE devis_id = %s", (devis_id,)
            )
            seq = cursor.fetchone()["next_seq"]
        else:
            seq = body.sequence_ligne
        cursor.execute(
            "INSERT INTO devis_lignes (devis_id, description, quantite, unite, "
            "prix_unitaire, montant_ligne, sequence_ligne, categorie, notes_ligne, "
            "mo_pct, mat_pct, admin_pct_ligne, contingence_pct_ligne, profit_pct_ligne) "
            "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) RETURNING id",
            (devis_id, body.description, body.quantite, body.unite,
             body.prix_unitaire, montant, seq, body.categorie, body.notes_ligne,
             body.mo_pct, body.mat_pct,
             body.admin_pct_ligne, body.contingence_pct_ligne, body.profit_pct_ligne),
        )
        ligne_id = cursor.fetchone()["id"]
        _recompute_devis_totals(cursor, devis_id)
        return {"id": ligne_id, "montant": montant, "message": "Ligne ajoutee"}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("add_devis_ligne error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.post("/{devis_id}/lignes/batch")
async def add_devis_lignes_batch(devis_id: int, body: List[DevisLigneCreate], user: ErpUser = Depends(get_current_user)):
    """Add multiple lines at once (avoids rate limiting)."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        _ensure_devis_pct_columns(conn, user.schema)
        cursor = conn.cursor()
        _ensure_visibility_columns(cursor)
        cursor.execute(
            "SELECT COALESCE(MAX(sequence_ligne), 0) as max_seq FROM devis_lignes WHERE devis_id = %s",
            (devis_id,),
        )
        seq = cursor.fetchone()["max_seq"]
        ids = []
        for item in body:
            seq += 1
            montant = round(item.quantite * item.prix_unitaire, 2)
            cursor.execute(
                "INSERT INTO devis_lignes (devis_id, description, quantite, unite, "
                "prix_unitaire, montant_ligne, sequence_ligne, categorie, notes_ligne, "
                "mo_pct, mat_pct, admin_pct_ligne, contingence_pct_ligne, profit_pct_ligne) "
                "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) RETURNING id",
                (devis_id, item.description, item.quantite, item.unite,
                 item.prix_unitaire, montant, seq, item.categorie, item.notes_ligne,
                 item.mo_pct, item.mat_pct,
                 item.admin_pct_ligne, item.contingence_pct_ligne, item.profit_pct_ligne),
            )
            ids.append(cursor.fetchone()["id"])
        _recompute_devis_totals(cursor, devis_id)
        return {"ids": ids, "count": len(ids), "message": f"{len(ids)} lignes ajoutees"}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("add_devis_lignes_batch error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.put("/{devis_id}/lignes/{ligne_id}")
async def update_devis_ligne(devis_id: int, ligne_id: int, body: DevisLigneCreate, user: ErpUser = Depends(get_current_user)):
    """Update a devis line item and recalculate totals."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        _ensure_devis_pct_columns(conn, user.schema)
        cursor = conn.cursor()
        _ensure_visibility_columns(cursor)
        montant = round((body.quantite if body.quantite is not None else 1) * (body.prix_unitaire if body.prix_unitaire is not None else 0), 2)
        cursor.execute(
            "UPDATE devis_lignes SET description = %s, quantite = %s, unite = %s, "
            "prix_unitaire = %s, montant_ligne = %s, categorie = %s, notes_ligne = %s, "
            "sequence_ligne = %s, code_article = %s, "
            "mo_pct = %s, mat_pct = %s, "
            "admin_pct_ligne = %s, contingence_pct_ligne = %s, profit_pct_ligne = %s "
            "WHERE id = %s AND devis_id = %s",
            (body.description, body.quantite, body.unite,
             body.prix_unitaire, montant, body.categorie, body.notes_ligne,
             body.sequence_ligne, body.code_article,
             body.mo_pct, body.mat_pct,
             body.admin_pct_ligne, body.contingence_pct_ligne, body.profit_pct_ligne,
             ligne_id, devis_id),
        )
        _recompute_devis_totals(cursor, devis_id)
        return {"id": ligne_id, "montant_ligne": montant, "message": "Ligne mise à jour"}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("update_devis_ligne error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


def _ensure_visibility_columns(cursor):
    """Add visibility + label columns if they don't exist yet."""
    for stmt in [
        "ALTER TABLE devis_lignes ADD COLUMN IF NOT EXISTS visible BOOLEAN DEFAULT TRUE",
        "ALTER TABLE devis_lignes ADD COLUMN IF NOT EXISTS mo_pct NUMERIC(5,2)",
        "ALTER TABLE devis_lignes ADD COLUMN IF NOT EXISTS mat_pct NUMERIC(5,2)",
        # Per-line markup overrides: NULL = inherit the devis-level
        # administration_pct / contingences_pct / profit_pct (default behaviour).
        # Non-null = override only this line's markup, leaving siblings untouched.
        # Used for granular pricing (e.g. higher profit on rare items, lower on
        # negotiated ones). Internal: NOT shown in the client-facing PDF/HTML.
        "ALTER TABLE devis_lignes ADD COLUMN IF NOT EXISTS admin_pct_ligne NUMERIC(5,2)",
        "ALTER TABLE devis_lignes ADD COLUMN IF NOT EXISTS contingence_pct_ligne NUMERIC(5,2)",
        "ALTER TABLE devis_lignes ADD COLUMN IF NOT EXISTS profit_pct_ligne NUMERIC(5,2)",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS show_administration BOOLEAN DEFAULT TRUE",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS show_contingences BOOLEAN DEFAULT TRUE",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS show_profit BOOLEAN DEFAULT TRUE",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS show_unite BOOLEAN DEFAULT TRUE",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS show_quantite BOOLEAN DEFAULT TRUE",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS show_prix_unitaire BOOLEAN DEFAULT TRUE",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS show_montant_ligne BOOLEAN DEFAULT TRUE",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS show_mo_mat BOOLEAN DEFAULT FALSE",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS administration_label TEXT DEFAULT 'Administration'",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS contingences_label TEXT DEFAULT 'Contingences'",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS profit_label TEXT DEFAULT 'Profit'",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS conditions_text TEXT",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS exclusions_text TEXT",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS show_conditions BOOLEAN DEFAULT TRUE",
        "ALTER TABLE devis ADD COLUMN IF NOT EXISTS show_exclusions BOOLEAN DEFAULT TRUE",
    ]:
        try:
            cursor.execute(stmt)
        except Exception:
            pass


@router.patch("/{devis_id}/lignes/{ligne_id}/visibility")
async def toggle_ligne_visibility(devis_id: int, ligne_id: int, body: dict, user: ErpUser = Depends(get_current_user)):
    """Toggle visibility of a devis line item for HTML export."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    visible = body.get("visible", True)
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        cursor = conn.cursor()
        _ensure_visibility_columns(cursor)
        cursor.execute(
            "UPDATE devis_lignes SET visible = %s WHERE id = %s AND devis_id = %s",
            (visible, ligne_id, devis_id),
        )
        return {"message": "ok", "visible": visible}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("toggle_ligne_visibility error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            try:
                cursor.close()
            except Exception:
                pass
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.delete("/{devis_id}/lignes/{ligne_id}")
async def delete_devis_ligne(devis_id: int, ligne_id: int, user: ErpUser = Depends(get_current_user)):
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        _ensure_devis_pct_columns(conn, user.schema)
        cursor = conn.cursor()
        # Required for legacy tenants that haven't yet hit any line-write
        # endpoint after the per-line markup migration: `_recompute_devis_totals`
        # reads `admin_pct_ligne / contingence_pct_ligne / profit_pct_ligne`
        # which only exist after `_ensure_visibility_columns` has run.
        _ensure_visibility_columns(cursor)
        cursor.execute("DELETE FROM devis_lignes WHERE id = %s AND devis_id = %s", (ligne_id, devis_id))
        _recompute_devis_totals(cursor, devis_id)
        return {"message": "Ligne supprimée"}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("delete_devis_ligne error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()



@router.delete("/{devis_id}")
async def delete_devis(devis_id: int, user: ErpUser = Depends(get_current_user)):
    """Delete a devis and all associated data."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        cursor = conn.cursor()
        # Verify devis exists and check status
        cursor.execute("SELECT id, statut FROM devis WHERE id = %s", (devis_id,))
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Devis non trouvé")
        statut = (row.get("statut") or "").upper()
        if statut in ("ACCEPTE", "TERMINE"):
            raise HTTPException(status_code=400, detail="Impossible de supprimer un devis accepte ou termine")
        # Clean up all related tables (no FK constraints, but avoid orphans)
        for table in ("devis_lignes", "devis_assignations", "devis_attachments",
                       "devis_dependencies", "devis_envois", "dossier_devis",
                       "conversations", "b2b_avancement_commentaires"):
            try:
                cursor.execute(f"DELETE FROM {table} WHERE devis_id = %s", (devis_id,))
            except Exception:
                db.set_tenant(conn, user.schema)
        # Detach factures and projects (SET NULL, not delete)
        for table in ("factures", "projects"):
            try:
                cursor.execute(f"UPDATE {table} SET devis_id = NULL WHERE devis_id = %s", (devis_id,))
            except Exception:
                db.set_tenant(conn, user.schema)
        # Detach opportunities
        try:
            cursor.execute("UPDATE opportunities SET devis_id = NULL WHERE devis_id = %s", (devis_id,))
        except Exception:
            db.set_tenant(conn, user.schema)
        # Cleanup Gantt dependencies attached to this devis (any direction)
        try:
            cursor.execute(
                "DELETE FROM gantt_dependencies "
                "WHERE (source_type = 'devis' AND source_id = %s) "
                "   OR (target_type = 'devis' AND target_id = %s)",
                (str(devis_id), str(devis_id)),
            )
        except Exception:
            db.set_tenant(conn, user.schema)
        # Delete the devis
        cursor.execute("DELETE FROM devis WHERE id = %s", (devis_id,))
        # Bug pre-existant fixe : sans conn.commit() la transaction etait
        # rollback a conn.close() (psycopg2 non-autocommit par defaut) et
        # le DELETE n'etait jamais persiste.
        conn.commit()
        return {"message": "Devis supprime"}
    except HTTPException:
        try:
            conn.rollback()
        except Exception:
            pass
        raise
    except Exception as exc:
        logger.error("delete_devis error: %s", exc)
        try:
            conn.rollback()
        except Exception:
            pass
        raise HTTPException(status_code=500, detail="Erreur lors de la suppression")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()

# ============================================
# DEVIS HTML GENERATION
# ============================================

DEVIS_EXCLUSIONS = [
    "Travaux de demolition non mentionnes dans la description des travaux",
    "Reparation ou remplacement des fondations existantes",
    "Travaux de decontamination des sols ou des materiaux",
    "Travaux d'amenagement paysager et de plantation",
    "Permis et frais de ville (a la charge du client)",
    "Branchements temporaires (electricite, eau, egout)",
    "Frais de stationnement et de signalisation de chantier",
    "Protection hivernale et chauffage temporaire",
    "Etudes geotechniques et releves d'arpentage",
    "Honoraires professionnels (architecte, ingenieur)",
    "Mobilier, equipements et accessoires non mentionnes",
    "Tests et essais en laboratoire",
    "Assurance chantier et cautionnement",
    "Travaux de toiture non specifies",
    "Raccordements definitifs aux services municipaux",
]

# Default conditions for new devis (one item per line).
# Override per-devis via devis.conditions_text, or globally per-entreprise via
# entreprise_config.devis_conditions_default.
DEVIS_CONDITIONS = [
    "Ce devis est valide pour une periode de 30 jours a compter de la date d'emission.",
    "Calendrier de paiement: 30% a la signature, 40% en cours de travaux, 30% a la fin des travaux.",
    "Garantie de 1 an sur la main-d'oeuvre et materiaux, selon les normes RBQ.",
    "Les travaux debuteront dans un delai convenu apres l'acceptation du devis.",
    "Toute modification au devis fera l'objet d'un avenant signe par les deux parties.",
]


def _get_entreprise_devis_defaults(cursor) -> dict:
    """Read default conditions/exclusions from entreprise_config (tenant-wide).
    Returns {'conditions': str|None, 'exclusions': str|None}. Silent fallback
    to None if table missing or keys unset — caller uses hardcoded constants."""
    result = {"conditions": None, "exclusions": None}
    try:
        cursor.execute("SELECT config_data FROM entreprise_config WHERE id = 1")
        row = cursor.fetchone()
        if not row:
            return result
        raw = row.get("config_data")
        if isinstance(raw, dict):
            data = raw
        elif isinstance(raw, str) and raw.strip():
            try:
                data = json.loads(raw)
            except Exception as exc:
                logger.warning("_get_entreprise_devis_defaults JSON parse failed: %s", exc)
                return result
        else:
            return result
        cond = data.get("devis_conditions_default")
        excl = data.get("devis_exclusions_default")
        if isinstance(cond, str) and cond.strip():
            result["conditions"] = cond
        if isinstance(excl, str) and excl.strip():
            result["exclusions"] = excl
    except Exception as exc:
        logger.warning("_get_entreprise_devis_defaults SELECT failed: %s", exc)
    return result


def _fmt_money(val: float) -> str:
    """Format a float as Canadian currency string."""
    if val is None:
        val = 0
    s = f"{val:,.2f}"
    # Replace comma thousand separator with space, and period with comma
    parts = s.split(".")
    integer_part = parts[0].replace(",", " ")
    return f"{integer_part},{parts[1]} $"


def _generate_devis_html(
    devis: dict,
    lignes: list,
    client_company: Optional[dict],
    enterprise: Optional[dict],
    client_contact: Optional[dict] = None,
    entreprise_defaults: Optional[dict] = None,
    theme: Optional[dict] = None,
) -> str:
    """Generate a professional HTML document for a devis.

    Conditions/exclusions resolution order:
      1. devis.conditions_text / devis.exclusions_text (per-devis override)
      2. entreprise_defaults['conditions'] / ['exclusions'] (tenant-wide default)
      3. DEVIS_CONDITIONS / DEVIS_EXCLUSIONS (hardcoded fallback)

    `theme` is an optional tenant color palette from get_document_theme(). When
    omitted, DEFAULT_DOCUMENT_THEME is used so rendering never breaks.
    """
    from .html_utils import DEFAULT_DOCUMENT_THEME, THEME_KEYS
    _t = dict(DEFAULT_DOCUMENT_THEME)
    if isinstance(theme, dict):
        for k in THEME_KEYS:
            v = theme.get(k)
            if isinstance(v, str) and v.strip():
                _t[k] = v

    _raw_esc = html_mod.escape  # shorthand for HTML escaping
    _esc = lambda v: _raw_esc(str(v) if v is not None else "")  # None-safe wrapper

    # Enterprise info (from normalized config or fallback to enterprise dict)
    if enterprise:
        ent_name = _esc(enterprise.get("nom", "") or enterprise.get("nom_entreprise", "") or "Entreprise")
        ent_address = _esc(enterprise.get("adresse", ""))
        ent_ville = _esc(enterprise.get("ville", ""))
        ent_province = _esc(enterprise.get("province", ""))
        ent_cp = _esc(enterprise.get("code_postal", ""))
        ent_phone = _esc(enterprise.get("telephone", "") or enterprise.get("telephone_bureau", ""))
        ent_email = _esc(enterprise.get("courriel", "") or enterprise.get("email", ""))
        ent_rbq = _esc(enterprise.get("rbq", "") or enterprise.get("numero_rbq", ""))
        ent_neq = _esc(enterprise.get("neq", "") or enterprise.get("numero_neq", ""))
        ent_tps = _esc(enterprise.get("tps", "") or enterprise.get("numero_tps", ""))
        ent_tvq = _esc(enterprise.get("tvq", "") or enterprise.get("numero_tvq", ""))
        _raw_logo = enterprise.get("logo_base64", "") or ""
        ent_logo = _raw_logo if _raw_logo.startswith("data:image/") else ""
    else:
        ent_name = "Entreprise"
        ent_address = ent_ville = ent_province = ent_cp = ""
        ent_phone = ent_email = ent_rbq = ent_neq = ent_tps = ent_tvq = ""
        ent_logo = ""

    # Logo HTML (escaped to prevent XSS via malicious data URL)
    ent_logo_html = f'<img src="{_esc(ent_logo)}" alt="{ent_name}" class="company-logo">' if ent_logo else ""

    # Client info — use company first, then contact, then devis fields
    def _build_address(src: dict) -> str:
        """Build full address from adresse, ville, province, code_postal fields."""
        parts = []
        addr = src.get("adresse", "") or ""
        if addr:
            parts.append(addr)
        city_parts = []
        for f in ("ville", "province", "code_postal"):
            v = src.get(f, "") or ""
            if v:
                city_parts.append(v)
        if city_parts:
            parts.append(", ".join(city_parts))
        return "\n".join(parts)

    if client_company:
        cli_name = _esc(client_company.get("nom", "Client"))
        cli_address = _esc(_build_address(client_company))
        cli_phone = _esc(client_company.get("telephone", ""))
        cli_email = _esc(client_company.get("email", ""))
    elif client_contact:
        contact_name = f"{client_contact.get('prenom', '')} {client_contact.get('nom_famille', '') or client_contact.get('nom', '')}".strip()
        cli_name = _esc(contact_name or "Client")
        cli_address = _esc(_build_address(client_contact))
        cli_phone = _esc(client_contact.get("telephone", "") or client_contact.get("mobile", "") or "")
        cli_email = _esc(client_contact.get("email", "") or "")
    else:
        cli_name = _esc(devis.get("client_nom") or devis.get("client_nom_cache") or devis.get("client_nom_direct") or "Client")
        cli_address = ""
        cli_phone = ""
        cli_email = ""
    # Add contact person line if both company and contact exist
    cli_contact_line = ""
    if client_company and client_contact:
        contact_name = f"{client_contact.get('prenom', '')} {client_contact.get('nom_famille', '') or client_contact.get('nom', '')}".strip()
        if contact_name:
            cli_contact_line = _esc(contact_name)

    # Devis info
    numero = _esc(devis.get("numero_devis", ""))
    titre = _esc(devis.get("nom_projet", ""))
    date_creation = str(devis.get("created_at") or "")[:10]
    date_validite = str(devis.get("date_prevu") or "")[:10] if devis.get("date_prevu") else ""
    # Type de soumission — affiché comme badge dans le header pour que le
    # client voie clairement si la soumission est ferme (Détaillée) ou
    # indicative (Budgétaire). Seul le type Budgétaire déclenche un
    # avertissement visuel; Détaillée reste muet (comportement par défaut).
    type_soumission = (devis.get("type_soumission") or "Détaillée").strip()
    is_budgetaire = type_soumission == "Budgétaire"
    type_badge_html = (
        '<div style="background:#FEF3C7;color:#92400E;padding:4px 10px;border-radius:4px;'
        'font-size:11px;font-weight:700;letter-spacing:0.5px;margin-top:6px;'
        'display:inline-block;">BUDGÉTAIRE</div>'
        if is_budgetaire else ''
    )
    budgetaire_banner_html = (
        '<div style="background:#FEF3C7;border-left:4px solid #F59E0B;color:#92400E;'
        'padding:10px 14px;border-radius:4px;margin:16px 0;font-size:12px;">'
        '<strong>Soumission budgétaire</strong> — Ce document est une estimation '
        'approximative à titre indicatif. Les montants peuvent varier selon le '
        'relevé final des mesures et l\'analyse détaillée du projet.</div>'
        if is_budgetaire else ''
    )

    # Calculate totals — from lines if available, otherwise from devis record
    lignes_total = sum(float(l.get("montant_ligne", 0) or 0) for l in lignes)

    if lignes_total > 0:
        # Calculate from line items, honouring per-line markup overrides so
        # the rendered "Sous-total + Administration + Contingences + Profit"
        # block matches the sum of the line amounts (which use _line_markup).
        # Without this, a line with an override would inflate the line price
        # but the summary block would still apply the global %, causing a
        # visible mismatch between line totals and the grand subtotal.
        sous_total_ht = lignes_total
        admin_pct = float(devis.get("administration_pct") if devis.get("administration_pct") is not None else 3.0) / 100
        contingences_pct = float(devis.get("contingences_pct") if devis.get("contingences_pct") is not None else 12.0) / 100
        profit_pct = float(devis.get("profit_pct") if devis.get("profit_pct") is not None else 15.0) / 100
        admin = 0.0
        contingences = 0.0
        profit = 0.0
        for _ln in lignes:
            _m = float(_ln.get("montant_ligne", 0) or 0)
            _a = _ln.get("admin_pct_ligne")
            _c = _ln.get("contingence_pct_ligne")
            _p = _ln.get("profit_pct_ligne")
            admin += _m * ((float(_a) / 100) if _a is not None else admin_pct)
            contingences += _m * ((float(_c) / 100) if _c is not None else contingences_pct)
            profit += _m * ((float(_p) / 100) if _p is not None else profit_pct)
        admin = round(admin, 2)
        contingences = round(contingences, 2)
        profit = round(profit, 2)
        sous_total_avant_taxes = round(sous_total_ht + admin + contingences + profit, 2)
        tps = round(sous_total_avant_taxes * 0.05, 2)
        tvq = round(sous_total_avant_taxes * 0.09975, 2)
        total_ttc = round(sous_total_avant_taxes + tps + tvq, 2)
    else:
        # Fallback: use stored values from the devis record
        def _fv(key):
            v = devis.get(key)
            if v is None:
                return 0.0
            try:
                return float(v)
            except (ValueError, TypeError):
                return 0.0

        _tv = _fv("total_travaux")
        sous_total_ht = _tv if devis.get("total_travaux") is not None else (_fv("prix_estime") if devis.get("prix_estime") is not None else 0.0)
        admin = _fv("administration")
        contingences = _fv("contingences")
        profit = _fv("profit")
        sous_total_avant_taxes = _fv("total_avant_taxes") if devis.get("total_avant_taxes") is not None else round(sous_total_ht + admin + contingences + profit, 2)
        tps = _fv("tps") if devis.get("tps") is not None else round(sous_total_avant_taxes * 0.05, 2)
        tvq = _fv("tvq") if devis.get("tvq") is not None else round(sous_total_avant_taxes * 0.09975, 2)
        total_ttc = _fv("investissement_total") if devis.get("investissement_total") is not None else round(sous_total_avant_taxes + tps + tvq, 2)
        # Define pct variables for HTML display
        admin_pct = float(devis.get("administration_pct") if devis.get("administration_pct") is not None else 3.0) / 100
        contingences_pct = float(devis.get("contingences_pct") if devis.get("contingences_pct") is not None else 12.0) / 100
        profit_pct = float(devis.get("profit_pct") if devis.get("profit_pct") is not None else 15.0) / 100

    # Filter out hidden lines for HTML export (they still count in totals above)
    visible_lignes = [l for l in lignes if l.get("visible", True) is not False]

    # Column visibility flags
    show_unite = devis.get("show_unite", True) is not False
    show_quantite = devis.get("show_quantite", True) is not False
    show_prix_unitaire = devis.get("show_prix_unitaire", True) is not False
    show_montant_ligne = devis.get("show_montant_ligne", True) is not False
    show_mo_mat = devis.get("show_mo_mat", False) is True
    col_count = 1 + (1 if show_unite else 0) + (1 if show_quantite else 0) + (1 if show_prix_unitaire else 0) + (1 if show_montant_ligne else 0) + (1 if show_mo_mat else 0)

    # Majoration visibility flags (included in line prices, but optionally shown in summary breakdown)
    show_admin = devis.get("show_administration", True) is not False
    show_cont = devis.get("show_contingences", True) is not False
    show_prof = devis.get("show_profit", True) is not False
    admin_label = _esc(devis.get("administration_label") or "Administration")
    cont_label = _esc(devis.get("contingences_label") or "Contingences")
    profit_label = _esc(devis.get("profit_label") or "Profit")
    # Display percentages (integer-formatted when whole, else with up to 2 decimals)
    def _fmt_pct(p: float) -> str:
        # p is a ratio (0.03 for 3%) — convert to percentage value
        v = p * 100
        if abs(v - round(v)) < 0.005:
            return f"{v:.0f}"
        return f"{v:.2f}".rstrip("0").rstrip(".")
    admin_pct_str = _fmt_pct(admin_pct)
    cont_pct_str = _fmt_pct(contingences_pct)
    profit_pct_str = _fmt_pct(profit_pct)
    # Only show each majoration line if flag is on AND amount is non-zero (avoid "Administration (0%) 0,00 $")
    show_admin_line = show_admin and admin > 0
    show_cont_line = show_cont and contingences > 0
    show_prof_line = show_prof and profit > 0
    any_majoration = show_admin_line or show_cont_line or show_prof_line

    # Group visible lines by category
    categories = {}
    for l in visible_lignes:
        cat = l.get("categorie") or "General"
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(l)

    # Build lines HTML
    lines_html = ""
    # Per-line markup: each line's price is silently inflated by its own
    # admin/contingence/profit %s — overrides if set on the line, else the
    # devis-level default. Per-line overrides are NEVER displayed to the
    # client (they remain an internal pricing tool); only the resulting
    # unit price and amount are shown, identical to the legacy behaviour
    # for lines that don't override.
    def _line_markup(l: dict) -> float:
        a = l.get("admin_pct_ligne")
        c = l.get("contingence_pct_ligne")
        p = l.get("profit_pct_ligne")
        a_ratio = (float(a) / 100) if a is not None else admin_pct
        c_ratio = (float(c) / 100) if c is not None else contingences_pct
        p_ratio = (float(p) / 100) if p is not None else profit_pct
        return 1 + a_ratio + c_ratio + p_ratio
    total_mo = 0.0
    total_mat = 0.0
    for cat_name, cat_lines in categories.items():
        # Compute category subtotal (always, needed when montant_ligne hidden)
        cat_subtotal = sum(round(float(l.get("montant_ligne", 0) or 0) * _line_markup(l), 2) for l in cat_lines)
        if len(categories) > 1 or not show_montant_ligne:
            # Show category header — with subtotal when per-line montant is hidden
            subtotal_span = ""
            if not show_montant_ligne:
                subtotal_span = f'<span style="float:right;white-space:nowrap;">{_fmt_money(cat_subtotal)}</span>'
            lines_html += f"""
            <tr class="cat-row">
              <td colspan="{col_count}" style="background:#f0f4f8;font-weight:700;padding:10px 12px;
                  font-size:13px;color:{_t['primary']};border-top:2px solid #cbd5e0;">
                {_esc(cat_name)}{subtotal_span}
              </td>
            </tr>"""
        for l in cat_lines:
            desc = _esc(l.get("description", ""))
            unite = l.get("unite", "")
            qte = float(l.get("quantite", 0) or 0)
            line_markup = _line_markup(l)
            prix = round(float(l.get("prix_unitaire", 0) or 0) * line_markup, 2)
            montant = round(float(l.get("montant_ligne", 0) or 0) * line_markup, 2)
            # MO / MAT calculation (use custom per-line if set, else auto-detect via keywords)
            mo_mat_cell = ""
            if show_mo_mat:
                custom_mo = l.get("mo_pct")
                custom_mat = l.get("mat_pct")
                if custom_mo is not None or custom_mat is not None:
                    # User override: if only one is set, derive the other
                    if custom_mo is None:
                        mo_pct = max(0.0, 100.0 - float(custom_mat))
                        mat_pct = float(custom_mat)
                    elif custom_mat is None:
                        mo_pct = float(custom_mo)
                        mat_pct = max(0.0, 100.0 - float(custom_mo))
                    else:
                        mo_pct = float(custom_mo)
                        mat_pct = float(custom_mat)
                else:
                    mo_pct, mat_pct = _get_mo_mat_ratio(l.get("description", ""))
                mo_val = round(montant * mo_pct / 100, 2)
                mat_val = round(montant * mat_pct / 100, 2)
                # Format pct without trailing ".0" for cosmetic consistency
                # (auto-detection returns int; custom values stored as Decimal → float).
                mo_pct_str = f"{int(mo_pct)}" if float(mo_pct).is_integer() else f"{mo_pct:g}"
                mat_pct_str = f"{int(mat_pct)}" if float(mat_pct).is_integer() else f"{mat_pct:g}"
                total_mo += mo_val
                total_mat += mat_val
                mo_mat_cell = (
                    f'<td style="padding:6px 8px;border-bottom:1px solid #e2e8f0;text-align:right;min-width:160px;vertical-align:middle;">'
                    f'<div style="display:inline-flex;flex-direction:column;gap:3px;align-items:flex-end;">'
                    f'<div style="background:#eff6ff;color:#1d4ed8;font-size:12px;font-weight:600;padding:2px 8px;border-radius:4px;white-space:nowrap;">'
                    f'MO {mo_pct_str}% : {_fmt_money(mo_val)}</div>'
                    f'<div style="background:#fffbeb;color:#b45309;font-size:12px;font-weight:600;padding:2px 8px;border-radius:4px;white-space:nowrap;">'
                    f'MAT {mat_pct_str}% : {_fmt_money(mat_val)}</div>'
                    f'</div></td>'
                )
            lines_html += f"""
            <tr>
              <td style="padding:8px 12px;border-bottom:1px solid #e2e8f0;">{desc}</td>
              {'<td style="padding:8px 12px;border-bottom:1px solid #e2e8f0;text-align:center;">' + _esc(unite) + '</td>' if show_unite else ''}
              {'<td style="padding:8px 12px;border-bottom:1px solid #e2e8f0;text-align:right;">' + f'{qte:,.2f}' + '</td>' if show_quantite else ''}
              {'<td style="padding:8px 12px;border-bottom:1px solid #e2e8f0;text-align:right;">' + _fmt_money(prix) + '</td>' if show_prix_unitaire else ''}
              {'<td style="padding:8px 12px;border-bottom:1px solid #e2e8f0;text-align:right;font-weight:600;">' + _fmt_money(montant) + '</td>' if show_montant_ligne else ''}
              {mo_mat_cell}
            </tr>"""

    # Resolve conditions + exclusions (per-devis → entreprise default → constants)
    defaults = entreprise_defaults or {}
    cond_raw = devis.get("conditions_text")
    if not (isinstance(cond_raw, str) and cond_raw.strip()):
        cond_raw = defaults.get("conditions") or "\n".join(DEVIS_CONDITIONS)
    excl_raw = devis.get("exclusions_text")
    if not (isinstance(excl_raw, str) and excl_raw.strip()):
        excl_raw = defaults.get("exclusions") or "\n".join(DEVIS_EXCLUSIONS)

    # Split to lines, strip, and drop empties. Leading "- " or "* " bullet chars
    # are also stripped so users can paste markdown-style lists transparently.
    def _parse_items(text: str) -> list:
        items = []
        for raw in (text or "").replace("\r\n", "\n").split("\n"):
            line = raw.strip().lstrip("-•*").strip()
            if line:
                items.append(line)
        return items

    conditions_items = _parse_items(cond_raw)
    exclusions_items = _parse_items(excl_raw)

    show_conditions = devis.get("show_conditions", True) is not False
    show_exclusions = devis.get("show_exclusions", True) is not False

    conditions_html = "\n".join(f"<li>{_esc(c)}</li>" for c in conditions_items)
    exclusions_html = "\n".join(f"<li>{_esc(e)}</li>" for e in exclusions_items)

    html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Devis {numero} - {titre}</title>
  <style>
    @page {{ size: letter; margin: 15mm 18mm; }}
    * {{ margin: 0; padding: 0; box-sizing: border-box; }}
    body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; color: #2d3748; line-height: 1.5; background: #fff; }}
    .page {{ max-width: 8.5in; margin: 0 auto; padding: 30px 36px; }}
    .header {{ display: flex; justify-content: space-between; align-items: stretch; margin-bottom: 24px; }}
    .header-left {{ display: flex; align-items: center; gap: 18px; max-width: 65%; }}
    .company-logo {{ height: 90px; width: auto; object-fit: contain; }}
    .company-details {{ }}
    .enterprise-name {{ font-size: 22px; font-weight: 800; color: {_t['primary']}; margin-bottom: 2px; }}
    .enterprise-info {{ font-size: 11px; color: #64748b; line-height: 1.5; }}
    .enterprise-info .ent-nums {{ color: #94a3b8; font-size: 10px; margin-top: 2px; }}
    .header-right {{ background: {_t['primary']}; color: {_t['header_text']}; padding: 20px 28px; border-radius: 6px; text-align: center; display: flex; flex-direction: column; justify-content: center; min-width: 160px; }}
    .devis-label {{ font-size: 24px; font-weight: 800; letter-spacing: 2px; color: {_t['header_text']}; }}
    .devis-numero {{ font-size: 14px; color: {_t['accent_light']}; margin-top: 4px; font-weight: 600; }}
    .header-separator {{ height: 4px; background: linear-gradient(90deg, {_t['primary']} 0%, {_t['accent']} 50%, {_t['primary']} 100%); border-radius: 2px; margin-bottom: 20px; }}
    .info-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 20px; margin-bottom: 24px; }}
    .info-box {{ background: {_t['info_bg']}; border-radius: 6px; padding: 14px 18px; border-left: 4px solid {_t['accent']}; }}
    .info-box h4 {{ font-size: 11px; text-transform: uppercase; color: {_t['accent']}; letter-spacing: 1px; margin-bottom: 8px; font-weight: 700; }}
    .info-box p {{ font-size: 12px; color: #334155; }}
    .info-box .name {{ font-size: 14px; font-weight: 700; color: {_t['primary']}; margin-bottom: 4px; }}
    table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; table-layout: auto; }}
    thead th {{ background: {_t['primary']}; color: {_t['header_text']}; padding: 9px 12px; font-size: 11px; text-transform: uppercase; letter-spacing: 0.5px; white-space: nowrap; }}
    thead th:first-child {{ text-align: left; width: auto; }}
    tbody td {{ font-size: 12px; }}
    td.cat-subtotal {{ white-space: nowrap; min-width: 110px; }}
    .summary {{ margin-left: auto; width: 300px; margin-bottom: 24px; }}
    .summary-row {{ display: flex; justify-content: space-between; padding: 5px 0; font-size: 12px; }}
    .summary-row.sub {{ border-top: 1px solid #e2e8f0; padding-top: 8px; margin-top: 4px; }}
    .summary-row.total {{ border-top: 3px solid {_t['primary']}; padding-top: 10px; margin-top: 8px; font-size: 16px; font-weight: 800; color: {_t['primary']}; }}
    .conditions {{ margin-bottom: 24px; }}
    .conditions h3 {{ font-size: 13px; font-weight: 700; color: {_t['primary']}; margin-bottom: 8px; text-transform: uppercase; }}
    .conditions ul {{ font-size: 11px; color: #4a5568; padding-left: 20px; }}
    .conditions li {{ margin-bottom: 3px; }}
    .exclusions {{ margin-bottom: 24px; }}
    .exclusions h3 {{ font-size: 13px; font-weight: 700; color: {_t['primary']}; margin-bottom: 8px; text-transform: uppercase; }}
    .exclusions ol {{ font-size: 11px; color: #4a5568; padding-left: 20px; }}
    .exclusions li {{ margin-bottom: 3px; }}
    .signatures {{ display: grid; grid-template-columns: 1fr 1fr; gap: 40px; margin-top: 30px; padding-top: 16px; border-top: 1px solid #e2e8f0; }}
    .sig-block {{ text-align: center; }}
    .sig-block h4 {{ font-size: 12px; font-weight: 700; color: {_t['primary']}; margin-bottom: 30px; }}
    .sig-line {{ border-top: 1px solid #2d3748; padding-top: 8px; font-size: 11px; color: #718096; }}
    .footer {{ margin-top: 30px; padding-top: 12px; border-top: 1px solid #e2e8f0; text-align: center; font-size: 10px; color: #a0aec0; }}
    @media print {{
      .page {{ padding: 0; max-width: 100%; }}
      body {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
    }}
  </style>
</head>
<body>
  <div class="page">
    <!-- Header -->
    <div class="header">
      <div class="header-left">
        {ent_logo_html}
        <div class="company-details">
          <div class="enterprise-name">{ent_name}</div>
          <div class="enterprise-info">
            {f'{ent_address}' if ent_address else ''}{f', {ent_ville}, {ent_province} {ent_cp}' if ent_ville else ''}<br>
            {f'Tel: {ent_phone}' if ent_phone else ''}{f' | {ent_email}' if ent_email else ''}
            {'<div class="ent-nums">' + ' | '.join(filter(None, [f'RBQ: {ent_rbq}' if ent_rbq else '', f'NEQ: {ent_neq}' if ent_neq else '', f'TPS: {ent_tps}' if ent_tps else '', f'TVQ: {ent_tvq}' if ent_tvq else ''])) + '</div>' if any([ent_rbq, ent_neq, ent_tps, ent_tvq]) else ''}
          </div>
        </div>
      </div>
      <div class="header-right">
        <div class="devis-label">DEVIS</div>
        <div class="devis-numero">{numero}</div>
        {type_badge_html}
      </div>
    </div>
    <div class="header-separator"></div>
    {budgetaire_banner_html}

    <!-- Info Grid -->
    <div class="info-grid">
      <div class="info-box">
        <h4>Client</h4>
        <p class="name">{cli_name}</p>
        {f'<p style="font-style:italic;color:#475569">Att: {cli_contact_line}</p>' if cli_contact_line else ''}
        {(chr(10).join(f'<p>{line}</p>' for line in cli_address.split(chr(10)) if line.strip())) if cli_address else ''}
        {f'<p>Tel: {cli_phone}</p>' if cli_phone else ''}
        {f'<p>{cli_email}</p>' if cli_email else ''}
      </div>
      <div class="info-box">
        <h4>Informations du devis</h4>
        <p><strong>Titre:</strong> {titre}</p>
        <p><strong>Date:</strong> {date_creation}</p>
        {f'<p><strong>Validite:</strong> {date_validite}</p>' if date_validite else '<p><strong>Validite:</strong> 30 jours</p>'}
      </div>
    </div>

    <!-- Lines Table -->
    <table>
      <thead>
        <tr>
          <th style="text-align:left;">Description</th>
          {'<th style="text-align:center;">Unite</th>' if show_unite else ''}
          {'<th style="text-align:right;">Quantite</th>' if show_quantite else ''}
          {'<th style="text-align:right;">Prix unitaire</th>' if show_prix_unitaire else ''}
          {'<th style="text-align:right;">Montant</th>' if show_montant_ligne else ''}
          {'<th style="text-align:right;min-width:160px;">MO / MAT</th>' if show_mo_mat else ''}
        </tr>
      </thead>
      <tbody>
        {lines_html}
        {'<tr><td colspan="' + str(col_count) + '" style="padding:20px;text-align:center;color:#a0aec0;font-style:italic;">Aucune ligne</td></tr>' if not visible_lignes else ''}
      </tbody>
    </table>

    <!-- Summary (admin/contingences/profit are distributed in line prices, shown here as breakdown when flags enabled) -->
    <div class="summary">
      <div class="summary-row sub"><span>Sous-total</span><span style="font-weight:600;">{_fmt_money(sous_total_avant_taxes)}</span></div>
      {'<div style="font-size:10px;color:#94a3b8;padding:4px 0 2px 0;font-style:italic;">Dont majoration incluse dans les prix unitaires :</div>' if any_majoration else ''}
      {f'<div class="summary-row" style="font-size:11px;color:#64748b;padding-left:10px;"><span style="overflow:hidden;text-overflow:ellipsis;white-space:nowrap;max-width:180px;">{admin_label} ({admin_pct_str}%)</span><span style="white-space:nowrap;">{_fmt_money(admin)}</span></div>' if show_admin_line else ''}
      {f'<div class="summary-row" style="font-size:11px;color:#64748b;padding-left:10px;"><span style="overflow:hidden;text-overflow:ellipsis;white-space:nowrap;max-width:180px;">{cont_label} ({cont_pct_str}%)</span><span style="white-space:nowrap;">{_fmt_money(contingences)}</span></div>' if show_cont_line else ''}
      {f'<div class="summary-row" style="font-size:11px;color:#64748b;padding-left:10px;"><span style="overflow:hidden;text-overflow:ellipsis;white-space:nowrap;max-width:180px;">{profit_label} ({profit_pct_str}%)</span><span style="white-space:nowrap;">{_fmt_money(profit)}</span></div>' if show_prof_line else ''}
      {f'<div class="summary-row" style="font-size:12px;"><span style="background:#eff6ff;color:#1d4ed8;font-weight:600;padding:2px 8px;border-radius:4px;">Main-d&#39;oeuvre (MO)</span><span style="color:#1d4ed8;font-weight:600;">{_fmt_money(total_mo)}</span></div>' if show_mo_mat else ''}
      {f'<div class="summary-row" style="font-size:12px;"><span style="background:#fffbeb;color:#b45309;font-weight:600;padding:2px 8px;border-radius:4px;">Mat&eacute;riaux (MAT)</span><span style="color:#b45309;font-weight:600;">{_fmt_money(total_mat)}</span></div>' if show_mo_mat else ''}
      <div class="summary-row"><span>TPS (5%)</span><span>{_fmt_money(tps)}</span></div>
      <div class="summary-row"><span>TVQ (9,975%)</span><span>{_fmt_money(tvq)}</span></div>
      <div class="summary-row total"><span>TOTAL TTC</span><span>{_fmt_money(total_ttc)}</span></div>
    </div>

    <!-- Conditions -->
    {f'''<div class="conditions">
      <h3>Conditions</h3>
      <ul>
        {conditions_html}
      </ul>
    </div>''' if show_conditions and conditions_items else ''}

    <!-- Exclusions -->
    {f'''<div class="exclusions">
      <h3>Exclusions</h3>
      <ol>
        {exclusions_html}
      </ol>
    </div>''' if show_exclusions and exclusions_items else ''}

    <!-- Signatures -->
    <div class="signatures">
      <div class="sig-block">
        <h4>Entrepreneur</h4>
        <div class="sig-line">
          Nom: _______________________________<br>
          Date: _______________________________<br>
          Signature: ___________________________
        </div>
      </div>
      <div class="sig-block">
        <h4>Client</h4>
        <div class="sig-line">
          Nom: _______________________________<br>
          Date: _______________________________<br>
          Signature: ___________________________
        </div>
      </div>
    </div>

    <!-- Footer -->
    <div class="footer">
      {ent_name} — Devis {numero} — Genere le {datetime.now().strftime('%Y-%m-%d %H:%M')}
    </div>
  </div>
</body>
</html>"""

    # Post-hoc: swap the hardcoded gray border across all inline cells and
    # sections to the tenant's theme border color (no-op when unchanged).
    # Keeps the inline preview HTML coherent with the themed SHARED_CSS used
    # by exports.py, without having to refactor 12+ inline `border:1px solid`
    # sites to f-string substitutions. Also inject tbody row alternation so
    # the generated HTML matches what the frontend ThemePreview advertises.
    html = html.replace(
        '</style>',
        f"tbody tr:nth-child(even){{background:{_t['table_row_alt']};}}</style>",
        1,
    )
    html = html.replace('#e2e8f0', _t['border'])

    return html


class PreviewLigneItem(BaseModel):
    """Lightweight ligne for preview (separate from DevisLigneCreate to enforce
    payload size caps — prevents CPU DoS via large HTML rendering)."""
    description: str = Field(..., min_length=1, max_length=2000)
    quantite: float = Field(default=1, gt=0, le=1_000_000)
    unite: str = Field(default="unite", max_length=50)
    prix_unitaire: float = Field(default=0, ge=0, le=10_000_000)
    categorie: Optional[str] = Field(default=None, max_length=200)
    notes_ligne: Optional[str] = Field(default=None, max_length=2000)
    sequence_ligne: int = Field(default=0, ge=0)
    code_article: Optional[str] = Field(default=None, max_length=100)


class PreviewHtmlRequest(BaseModel):
    """Body for preview-html-with-items endpoint.

    Items are merged in-memory with the devis existing lignes before calling
    _generate_devis_html. Nothing is persisted — the client chooses whether to
    commit later via /lignes/batch. Frontend sends camelCase `extraItems`;
    the api client interceptor converts it to snake_case `extra_items`.

    Capped at 2000 items to prevent CPU DoS on _generate_devis_html rendering.
    Even large construction projects rarely exceed a few hundred lines.
    """
    extra_items: List[PreviewLigneItem] = Field(default_factory=list, max_length=2000)


@router.post("/{devis_id}/preview-html-with-items")
async def preview_devis_html_with_items(
    devis_id: int,
    body: PreviewHtmlRequest,
    user: ErpUser = Depends(get_current_user),
):
    """Preview HTML for a devis merged with additional items (no persistence).

    Loads the current devis + its persisted lignes, appends `extraItems` in memory,
    then renders the shared _generate_devis_html template. Used by Metre and Manuel
    tabs to show an accurate preview before the user commits the items with
    POST /lignes/batch. Zero database writes.
    """
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")

    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        _ensure_devis_pct_columns(conn, user.schema)
        cursor = conn.cursor()
        _ensure_visibility_columns(cursor)

        cursor.execute("SELECT * FROM devis WHERE id = %s", (devis_id,))
        devis_row = cursor.fetchone()
        if not devis_row:
            raise HTTPException(status_code=404, detail="Devis non trouve")
        devis = dict(devis_row)

        cursor.execute(
            "SELECT * FROM devis_lignes WHERE devis_id = %s ORDER BY sequence_ligne ASC",
            (devis_id,),
        )
        lignes = [dict(r) for r in cursor.fetchall()]

        # Append extraItems as in-memory dicts matching the lignes structure.
        # _generate_devis_html recomputes all totals from `lignes` when lignes_total > 0
        # (see _generate_devis_html ~line 3016), so we don't need to update devis[*] fields.
        if body.extra_items:
            max_seq = max((int(l.get("sequence_ligne") or 0) for l in lignes), default=0)
            for item in body.extra_items:
                max_seq += 1
                montant = round(float(item.quantite) * float(item.prix_unitaire), 2)
                lignes.append({
                    "id": None,
                    "devis_id": devis_id,
                    "description": item.description,
                    "quantite": float(item.quantite),
                    "unite": item.unite or "unite",
                    "prix_unitaire": float(item.prix_unitaire),
                    "montant_ligne": montant,
                    "sequence_ligne": max_seq,
                    "categorie": item.categorie,
                    "notes_ligne": item.notes_ligne,
                    "visible": True,
                })

        client_company = None
        if devis.get("client_company_id"):
            cursor.execute("SELECT * FROM companies WHERE id = %s", (devis["client_company_id"],))
            row = cursor.fetchone()
            if row:
                client_company = dict(row)

        client_contact = None
        if devis.get("client_contact_id"):
            try:
                cursor.execute("SELECT * FROM contacts WHERE id = %s", (devis["client_contact_id"],))
                row = cursor.fetchone()
                if row:
                    client_contact = dict(row)
            except Exception:
                pass

        from .html_utils import get_company_info, get_document_theme
        enterprise = get_company_info(cursor)
        entreprise_defaults = _get_entreprise_devis_defaults(cursor)
        theme = get_document_theme(cursor)

        html = _generate_devis_html(devis, lignes, client_company, enterprise, client_contact, entreprise_defaults, theme=theme)

        return {
            "html": html,
            "devis_id": devis_id,
            "numero": devis.get("numero_devis", ""),
            "extra_items_count": len(body.extra_items),
        }

    except HTTPException:
        raise
    except Exception as exc:
        logger.error("preview_devis_html_with_items error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de la preview HTML")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.post("/{devis_id}/generate-html")
async def generate_devis_html(devis_id: int, user: ErpUser = Depends(get_current_user)):
    """Generate a professional HTML document for a devis."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")

    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        cursor = conn.cursor()
        _ensure_visibility_columns(cursor)

        # Fetch devis
        cursor.execute("SELECT * FROM devis WHERE id = %s", (devis_id,))
        devis = cursor.fetchone()
        if not devis:
            raise HTTPException(status_code=404, detail="Devis non trouvé")
        devis = dict(devis)

        # Fetch lignes
        cursor.execute(
            "SELECT * FROM devis_lignes WHERE devis_id = %s ORDER BY sequence_ligne ASC",
            (devis_id,),
        )
        lignes = [dict(r) for r in cursor.fetchall()]

        # Fetch client company
        client_company = None
        if devis.get("client_company_id"):
            cursor.execute("SELECT * FROM companies WHERE id = %s", (devis["client_company_id"],))
            row = cursor.fetchone()
            if row:
                client_company = dict(row)

        # Fetch client contact
        client_contact = None
        if devis.get("client_contact_id"):
            try:
                cursor.execute("SELECT * FROM contacts WHERE id = %s", (devis["client_contact_id"],))
                row = cursor.fetchone()
                if row:
                    client_contact = dict(row)
            except Exception:
                pass

        # Fetch enterprise config from tenant entreprise_config table
        from .html_utils import get_company_info, get_document_theme
        enterprise = get_company_info(cursor)
        entreprise_defaults = _get_entreprise_devis_defaults(cursor)
        theme = get_document_theme(cursor)

        # Generate HTML
        html = _generate_devis_html(devis, lignes, client_company, enterprise, client_contact, entreprise_defaults, theme=theme)

        # Store HTML in metadonnees_json JSONB field
        try:
            # Try to update metadonnees_json column
            existing_meta = devis.get("metadonnees_json") or {}
            if isinstance(existing_meta, str):
                try:
                    existing_meta = json.loads(existing_meta)
                except (json.JSONDecodeError, TypeError):
                    existing_meta = {}
            existing_meta["html_generated"] = True
            existing_meta["html_generated_at"] = datetime.now().isoformat()
            existing_meta["html_content"] = html

            cursor.execute(
                "UPDATE devis SET metadonnees_json = %s WHERE id = %s",
                (json.dumps(existing_meta), devis_id),
            )
        except Exception as meta_exc:
            logger.warning("Could not store HTML in metadonnees_json: %s", meta_exc)
            # metadonnees_json column may not exist, just return the HTML

        return {"html": html, "devis_id": devis_id, "numero": devis.get("numero_devis", "")}

    except HTTPException:
        raise
    except Exception as exc:
        logger.error("generate_devis_html error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de la generation HTML")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


# ============================================
# EXPORT XLSX (Excel)
# ============================================


@router.get("/{devis_id}/export-xlsx")
async def export_devis_xlsx(devis_id: int, user: ErpUser = Depends(get_current_user)):
    """Export a devis to .xlsx format (Excel) with formatted layout."""
    from fastapi.responses import Response
    import io
    from urllib.parse import quote as _urlquote

    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")

    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        cursor = conn.cursor()
        _ensure_visibility_columns(cursor)

        cursor.execute("SELECT * FROM devis WHERE id = %s", (devis_id,))
        devis_row = cursor.fetchone()
        if not devis_row:
            raise HTTPException(status_code=404, detail="Devis non trouvé")
        devis = dict(devis_row)

        cursor.execute(
            "SELECT * FROM devis_lignes WHERE devis_id = %s ORDER BY sequence_ligne ASC",
            (devis_id,),
        )
        lignes = [dict(r) for r in cursor.fetchall()]

        client_company = None
        if devis.get("client_company_id"):
            cursor.execute("SELECT * FROM companies WHERE id = %s", (devis["client_company_id"],))
            row = cursor.fetchone()
            if row:
                client_company = dict(row)

        client_contact = None
        if devis.get("client_contact_id"):
            try:
                cursor.execute("SELECT * FROM contacts WHERE id = %s", (devis["client_contact_id"],))
                row = cursor.fetchone()
                if row:
                    client_contact = dict(row)
            except Exception:
                pass

        from .html_utils import get_company_info
        enterprise = get_company_info(cursor)
        # Tenant defaults for conditions/exclusions (cohérent avec _generate_devis_html)
        entreprise_defaults = _get_entreprise_devis_defaults(cursor)

        # ==== Helper anti-formula-injection Excel + filtrage caractères de contrôle ====
        # 1) Formula injection : une cellule string commençant par = + - @ \t \r
        #    est interprétée par Excel comme une FORMULE à l'ouverture. Un
        #    attaquant qui met "=WEBSERVICE(...)" dans un nom de client pourrait
        #    exfiltrer des données de la machine du collègue qui ouvre le xlsx.
        # 2) Caractères de contrôle : openpyxl lève IllegalCharacterError pour
        #    \x00-\x08, \x0B, \x0C, \x0E-\x1F. Un champ BD importé depuis Word/
        #    Outlook peut contenir des null bytes ou du \x07 (BEL) invisibles
        #    qui feraient crasher l'export (HTTP 500 pour ce devis).
        # NOTE: _safe() est local au endpoint — tout refactor extrayant une
        # sous-fonction DOIT passer _safe en argument. Ne JAMAIS appeler _safe
        # sur une valeur numérique (un -0.5 deviendrait la string "'-0.5").
        _ILLEGAL_XLSX_CHARS = re.compile(r'[\x00-\x08\x0B\x0C\x0E-\x1F]')
        def _safe(v):
            if v is None:
                return v
            s = _ILLEGAL_XLSX_CHARS.sub('', str(v))
            if s and s[0] in ("=", "+", "-", "@", "\t", "\r"):
                return "'" + s
            return s

        # Read pct (admin/contingences/profit) from devis, fallback defaults.
        # Use `is not None` (not `or`) so 0 % explicite from client is preserved
        # instead of silently replaced by the default (cohérent avec _generate_devis_html).
        _raw_adm = devis.get("administration_pct")
        _raw_con = devis.get("contingences_pct")
        _raw_pro = devis.get("profit_pct")
        adm_pct = float(_raw_adm if _raw_adm is not None else _DEFAULT_ADM_PCT)
        con_pct = float(_raw_con if _raw_con is not None else _DEFAULT_CON_PCT)
        pro_pct = float(_raw_pro if _raw_pro is not None else _DEFAULT_PRO_PCT)

        # Helper local pour l'adresse complète (miroir de _generate_devis_html._build_address)
        def _full_address(src: dict) -> str:
            parts = []
            addr = src.get("adresse") or ""
            if addr:
                parts.append(addr)
            city_parts = []
            for f in ("ville", "province", "code_postal"):
                v = src.get(f) or ""
                if v:
                    city_parts.append(v)
            if city_parts:
                parts.append(", ".join(city_parts))
            return " | ".join(parts)

        # Flags visibilité (cohérents avec _generate_devis_html)
        show_admin = devis.get("show_administration") is not False
        show_cont = devis.get("show_contingences") is not False
        show_profit = devis.get("show_profit") is not False
        show_unite = devis.get("show_unite") is not False
        show_qte = devis.get("show_quantite") is not False
        show_prix = devis.get("show_prix_unitaire") is not False
        show_mntl = devis.get("show_montant_ligne") is not False

        # Build xlsx
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        # Sheet title : xlsx max 31 chars + interdiction de : \ / ? * [ ]
        _raw_title = (devis.get("numero_devis") or "Devis")
        _safe_title = re.sub(r'[:\\/?*\[\]]', '_', _raw_title)[:31] or "Devis"
        ws.title = _safe_title

        navy = "002050"
        header_fill = PatternFill(start_color=navy, end_color=navy, fill_type="solid")
        white_bold = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
        bold = Font(name="Calibri", size=11, bold=True)
        thin = Side(border_style="thin", color="999999")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        row = 1

        # === En-tête entreprise ===
        # get_company_info() retourne des clés FRANÇAISES normalisées (cf. html_utils.py)
        ws.cell(row=row, column=1, value=_safe(enterprise.get("nom") or "Entreprise")).font = Font(
            name="Calibri", size=16, bold=True, color=navy
        )
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
        row += 1
        ent_full_addr = _full_address(enterprise)
        ent_bits = []
        if ent_full_addr:
            ent_bits.append(ent_full_addr)
        for k, fmt in (("telephone", "Tél: {}"), ("courriel", "{}"), ("rbq", "RBQ: {}"), ("neq", "NEQ: {}")):
            v = enterprise.get(k)
            if v:
                ent_bits.append(fmt.format(v))
        ent_addr = " | ".join(ent_bits)
        if ent_addr:
            ws.cell(row=row, column=1, value=_safe(ent_addr)).font = Font(name="Calibri", size=9, color="606770")
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
            row += 1
        row += 1

        # === Titre Soumission ===
        label = "SOUMISSION BUDGÉTAIRE" if str(devis.get("type_soumission") or "").lower().startswith("budg") else "SOUMISSION"
        ws.cell(row=row, column=1, value=label).font = Font(name="Calibri", size=14, bold=True, color=navy)
        ws.cell(row=row, column=5, value="N°").font = bold
        ws.cell(row=row, column=6, value=_safe(devis.get("numero_devis") or "")).font = bold
        row += 1
        ws.cell(row=row, column=5, value="Date").font = bold
        ws.cell(row=row, column=6, value=str(devis.get("date_soumis") or devis.get("created_at") or "")[:10])
        row += 2

        # === Client ===
        ws.cell(row=row, column=1, value="CLIENT").font = bold
        row += 1
        # Construction nom client sans tomber sur "None None" si prénom/nom_famille sont NULL en BD.
        # `or ''` APRÈS .get() neutralise une valeur None explicite (différent de `.get(key, '')`
        # qui ne neutralise que l'absence de clé).
        contact_full_name = None
        if client_contact:
            _p = (client_contact.get("prenom") or "").strip()
            _n = (client_contact.get("nom_famille") or "").strip()
            _combined = (f"{_p} {_n}").strip()
            contact_full_name = _combined or None
        client_name = (
            (client_company or {}).get("nom")
            or contact_full_name
            or devis.get("client_nom_direct")
            or "--"
        )
        ws.cell(row=row, column=1, value=_safe(client_name))
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
        row += 1
        # Adresse client complète (adresse + ville + province + code postal), miroir HTML
        client_addr = (
            _full_address(client_company or {})
            or _full_address(client_contact or {})
        )
        if client_addr:
            ws.cell(row=row, column=1, value=_safe(client_addr)).font = Font(name="Calibri", size=10, color="606770")
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
            row += 1
        if devis.get("nom_projet"):
            ws.cell(row=row, column=1, value="Projet :").font = bold
            ws.cell(row=row, column=2, value=_safe(devis.get("nom_projet")))
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)
            row += 1
        row += 1

        # === En-têtes lignes ===
        headers = ["#", "Description"]
        col_keys = ["_idx", "description"]
        if show_unite:
            headers.append("Unité"); col_keys.append("unite")
        if show_qte:
            headers.append("Qté"); col_keys.append("quantite")
        if show_prix:
            headers.append("Prix unitaire"); col_keys.append("prix_unitaire")
        if show_mntl:
            headers.append("Montant"); col_keys.append("montant_ligne")

        header_row = row
        for i, h in enumerate(headers, start=1):
            c = ws.cell(row=row, column=i, value=h)
            c.font = white_bold
            c.fill = header_fill
            c.alignment = Alignment(horizontal="center", vertical="center")
            c.border = border
        row += 1

        # === Lignes (visibles seulement, mais totaux plus bas incluent tout) ===
        # Per-line markup honours admin_pct_ligne / contingence_pct_ligne /
        # profit_pct_ligne overrides — same logic as `_line_markup` in
        # `_generate_devis_html`. Ensures the XLSX export matches the PDF.
        def _line_markup_xlsx(l):
            _a = l.get("admin_pct_ligne")
            _c = l.get("contingence_pct_ligne")
            _p = l.get("profit_pct_ligne")
            _ar = (float(_a) / 100.0) if _a is not None else (adm_pct / 100.0)
            _cr = (float(_c) / 100.0) if _c is not None else (con_pct / 100.0)
            _pr = (float(_p) / 100.0) if _p is not None else (pro_pct / 100.0)
            return 1.0 + _ar + _cr + _pr
        visible_lignes = [l for l in lignes if l.get("visible") is not False]
        idx = 1
        for l in visible_lignes:
            qte = float(l.get("quantite") or 0)
            _lm = _line_markup_xlsx(l)
            prix = round(float(l.get("prix_unitaire") or 0) * _lm, 2)
            montant = round(float(l.get("montant_ligne") or 0) * _lm, 2)
            values = {
                "_idx": idx,
                "description": _safe(l.get("description") or ""),
                "unite": _safe(l.get("unite") or ""),
                "quantite": qte,
                "prix_unitaire": prix,
                "montant_ligne": montant,
            }
            for ci, key in enumerate(col_keys, start=1):
                c = ws.cell(row=row, column=ci, value=values.get(key))
                c.border = border
                if key in ("quantite",):
                    c.number_format = "#,##0.00"
                    c.alignment = Alignment(horizontal="right")
                elif key in ("prix_unitaire", "montant_ligne"):
                    c.number_format = '#,##0.00 "$"'
                    c.alignment = Alignment(horizontal="right")
                elif key == "_idx":
                    c.alignment = Alignment(horizontal="center")
            row += 1
            idx += 1

        amt_last_col = len(headers)

        # === Totaux (formules arithmétiques alignées EXACTEMENT sur _generate_devis_html lignes 4567-4583) ===
        # Sous-total HT = somme brute des lignes (SANS markup) — toutes les lignes, visibles ou masquées.
        lignes_total = sum(float(l.get("montant_ligne") or 0) for l in lignes)
        adm_decimal = adm_pct / 100.0
        con_decimal = con_pct / 100.0
        pro_decimal = pro_pct / 100.0

        if lignes_total > 0:
            # Honour per-line overrides — keep the XLSX summary in sync with
            # the per-line amounts above (which use _line_markup_xlsx).
            sous_total_ht = lignes_total
            admin_amt = 0.0
            contingences_amt = 0.0
            profit_amt = 0.0
            for _ln in lignes:
                _m = float(_ln.get("montant_ligne") or 0)
                _a = _ln.get("admin_pct_ligne")
                _c = _ln.get("contingence_pct_ligne")
                _p = _ln.get("profit_pct_ligne")
                admin_amt += _m * ((float(_a) / 100.0) if _a is not None else adm_decimal)
                contingences_amt += _m * ((float(_c) / 100.0) if _c is not None else con_decimal)
                profit_amt += _m * ((float(_p) / 100.0) if _p is not None else pro_decimal)
            admin_amt = round(admin_amt, 2)
            contingences_amt = round(contingences_amt, 2)
            profit_amt = round(profit_amt, 2)
            sous_total_avant_taxes = round(sous_total_ht + admin_amt + contingences_amt + profit_amt, 2)
            tps = round(sous_total_avant_taxes * 0.05, 2)
            tvq = round(sous_total_avant_taxes * 0.09975, 2)
            total_ttc = round(sous_total_avant_taxes + tps + tvq, 2)
        else:
            # Fallback identique à HTML (devis.py:4584-4603) : valeurs stockées sur le devis
            def _fv(key):
                v = devis.get(key)
                if v is None:
                    return 0.0
                try:
                    return float(v)
                except (ValueError, TypeError):
                    return 0.0
            sous_total_ht = _fv("total_travaux") if devis.get("total_travaux") is not None else (
                _fv("prix_estime") if devis.get("prix_estime") is not None else 0.0
            )
            admin_amt = _fv("administration")
            contingences_amt = _fv("contingences")
            profit_amt = _fv("profit")
            sous_total_avant_taxes = _fv("total_avant_taxes") if devis.get("total_avant_taxes") is not None else round(
                sous_total_ht + admin_amt + contingences_amt + profit_amt, 2
            )
            tps = _fv("tps") if devis.get("tps") is not None else round(sous_total_avant_taxes * 0.05, 2)
            tvq = _fv("tvq") if devis.get("tvq") is not None else round(sous_total_avant_taxes * 0.09975, 2)
            total_ttc = _fv("investissement_total") if devis.get("investissement_total") is not None else round(
                sous_total_avant_taxes + tps + tvq, 2
            )

        row += 1

        def total_row(label_text, value, bold_total=False):
            nonlocal row
            lc = ws.cell(row=row, column=amt_last_col - 1, value=label_text)
            vc = ws.cell(row=row, column=amt_last_col, value=value)
            lc.alignment = Alignment(horizontal="right")
            vc.number_format = '#,##0.00 "$"'
            vc.alignment = Alignment(horizontal="right")
            if bold_total:
                lc.font = bold
                vc.font = bold
                lc.fill = PatternFill(start_color="EEF2FF", end_color="EEF2FF", fill_type="solid")
                vc.fill = PatternFill(start_color="EEF2FF", end_color="EEF2FF", fill_type="solid")
            row += 1

        total_row("Sous-total HT", sous_total_ht)
        # Admin / Contingences / Profit affichés séparément pour transparence (selon flags)
        if show_admin and adm_pct > 0:
            adm_label = (devis.get("administration_label") or "Administration") + f" ({adm_pct:g}%)"
            total_row(adm_label, admin_amt)
        if show_cont and con_pct > 0:
            con_label = (devis.get("contingences_label") or "Contingences") + f" ({con_pct:g}%)"
            total_row(con_label, contingences_amt)
        if show_profit and pro_pct > 0:
            pro_label = (devis.get("profit_label") or "Profit") + f" ({pro_pct:g}%)"
            total_row(pro_label, profit_amt)
        total_row("Sous-total avant taxes", sous_total_avant_taxes, bold_total=True)
        total_row("TPS (5%)", tps)
        total_row("TVQ (9.975%)", tvq)
        total_row("TOTAL TTC", total_ttc, bold_total=True)

        # === Conditions / Exclusions (fallback devis → entreprise_defaults → constantes) ===
        # Miroir exact de _generate_devis_html:4722-4729
        ed = entreprise_defaults or {}
        cond_raw = devis.get("conditions_text")
        if not (isinstance(cond_raw, str) and cond_raw.strip()):
            cond_raw = ed.get("conditions") or "\n".join(DEVIS_CONDITIONS)
        excl_raw = devis.get("exclusions_text")
        if not (isinstance(excl_raw, str) and excl_raw.strip()):
            excl_raw = ed.get("exclusions") or "\n".join(DEVIS_EXCLUSIONS)

        show_cond = devis.get("show_conditions") is not False
        show_excl = devis.get("show_exclusions") is not False

        def _clean_items(text):
            out = []
            for raw in (text or "").replace("\r\n", "\n").split("\n"):
                line = raw.strip().lstrip("-•*").strip()
                if line:
                    out.append(line)
            return out

        cond_items = _clean_items(cond_raw) if show_cond else []
        excl_items = _clean_items(excl_raw) if show_excl else []

        if cond_items or excl_items:
            row += 2
            if cond_items:
                ws.cell(row=row, column=1, value="CONDITIONS").font = bold
                row += 1
                for line in cond_items:
                    ws.cell(row=row, column=1, value=_safe(line))
                    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=amt_last_col)
                    row += 1
                row += 1
            if excl_items:
                ws.cell(row=row, column=1, value="EXCLUSIONS").font = bold
                row += 1
                for line in excl_items:
                    ws.cell(row=row, column=1, value=_safe(line))
                    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=amt_last_col)
                    row += 1

        # === Largeurs colonnes ===
        widths = [5, 52, 10, 10, 16, 18]
        for i, w in enumerate(widths[:amt_last_col], start=1):
            ws.column_dimensions[get_column_letter(i)].width = w

        # === Freeze ligne en-tête tableau ===
        ws.freeze_panes = ws.cell(row=header_row + 1, column=1)

        # === Sérialiser ===
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        payload = buf.getvalue()

        # Filename RFC 5987 safe
        numero = (devis.get("numero_devis") or f"devis-{devis_id}").replace("\r", "").replace("\n", "")
        raw_fn = f"{numero}.xlsx"
        ascii_fn = raw_fn.encode("ascii", "replace").decode("ascii").replace('"', "").replace("\\", "")
        utf8_fn = _urlquote(raw_fn, safe="")
        cd_header = f'attachment; filename="{ascii_fn}"; filename*=UTF-8\'\'{utf8_fn}'

        return Response(
            content=payload,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": cd_header},
        )

    except HTTPException:
        raise
    except Exception as exc:
        logger.error("export_devis_xlsx error: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail="Erreur export Excel")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


# ============================================
# AUTO-CREATE PROJECT FROM DEVIS
# ============================================


def _create_project_from_devis(conn, devis_id: int, schema: str = "") -> Optional[int]:
    """Create a project from an accepted devis. Returns project_id or None.

    Copies all relevant fields + attachments. Never raises. Pass the tenant
    `schema` so the defensive-migration helper can re-bind search_path if an
    ALTER rolls back.
    """
    cursor = None
    try:
        cursor = conn.cursor()

        # Fetch full devis
        cursor.execute("SELECT * FROM devis WHERE id = %s", (devis_id,))
        devis = cursor.fetchone()
        if not devis:
            return None
        d = dict(devis)

        # Skip if already linked to a project
        if d.get("project_id"):
            return d["project_id"]

        # Compute budget (TTC) — use actual devis values, fallback to stored taxes
        budget = float(d.get("investissement_total") or 0)
        if d.get("investissement_total") is None:
            st = float(d.get("total_avant_taxes") or d.get("prix_estime") or d.get("total_travaux") or 0)
            tps_v = float(d.get("tps") or 0)
            tvq_v = float(d.get("tvq") or 0)
            if d.get("tps") is None and st:
                tps_v = round(st * 0.05, 2)
            if d.get("tvq") is None and st:
                tvq_v = round(st * 0.09975, 2)
            budget = round(st + tps_v + tvq_v, 2)

        # Ensure all columns we're about to INSERT exist on this tenant (lazy migration)
        _ensure_projects_insert_columns(cursor, conn, schema)

        # Fix sequence if out of sync (can happen after manual INSERTs).
        # Use GREATEST(max, 1) + 3-arg setval to avoid "value 0 is out of bounds"
        # when projects table is empty (PostgreSQL sequence range is 1..2^31-1).
        try:
            cursor.execute(
                "SELECT setval(pg_get_serial_sequence('projects', 'id'), "
                "GREATEST(COALESCE((SELECT MAX(id) FROM projects), 0), 1), "
                "(SELECT COUNT(*) > 0 FROM projects))"
            )
        except Exception as seq_exc:
            logger.warning("setval projects_id_seq failed: %s", seq_exc)

        # Insert project — use RETURNING id
        # Map devis dates: date_prevu → date_debut_reel, date_fin → date_fin_reel
        cursor.execute(
            "INSERT INTO projects (nom_projet, client_company_id, client_contact_id, "
            "client_nom_cache, po_client, statut, priorite, type_projet, "
            "budget_total, date_debut_reel, date_fin_reel, description, "
            "devis_id, devis_source_id, numero_devis, "
            "created_at, updated_at) "
            "VALUES (%s,%s,%s,%s,%s,'En cours',%s,'Construction',"
            "%s,%s,%s,%s,%s,%s,%s,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP) RETURNING id",
            (
                d.get("nom_projet") or f"Projet DEV-{devis_id}",
                d.get("client_company_id"),
                d.get("client_contact_id"),
                d.get("client_nom_cache"),
                d.get("po_client"),
                d.get("priorite") or "Moyenne",
                budget,
                d.get("date_prevu"),
                d.get("date_fin"),
                d.get("description") or d.get("notes") or f"Projet créé depuis devis #{devis_id}",
                devis_id,
                devis_id,
                d.get("numero_devis"),
            ),
        )
        project_id = cursor.fetchone()["id"]

        # Generate numero_projet (PROJ-YYYY-NNNNN) — try/except simple,
        # voir commentaire detaille au site #1 (~ligne 3710).
        try:
            cursor.execute(
                "UPDATE projects "
                "SET numero_projet = 'PROJ-' || EXTRACT(YEAR FROM COALESCE(created_at, CURRENT_TIMESTAMP))::int "
                "|| '-' || LPAD(id::text, 5, '0') "
                "WHERE id = %s AND (numero_projet IS NULL OR numero_projet = '')",
                (project_id,),
            )
        except Exception as numproj_exc:
            logger.error("Could not set numero_projet for project %s: %s — _backfill_numero_projet rattrapera au prochain list_projects", project_id, numproj_exc)

        # Atomic race-safe link: only set if project_id is still NULL. Two
        # concurrent orphan-recovery calls both pass the guard at 3644 (both
        # see project_id=NULL before either INSERTs), both create projects
        # here, then race on the UPDATE. The first wins; the loser sees
        # rowcount=0 and must clean up the duplicate project it just created.
        cursor.execute(
            "UPDATE devis SET project_id = %s WHERE id = %s AND project_id IS NULL",
            (project_id, devis_id),
        )
        if cursor.rowcount == 0:
            # Another call won the race — delete our duplicate and return
            # the winning project_id instead.
            try:
                cursor.execute("DELETE FROM projects WHERE id = %s", (project_id,))
            except Exception as del_exc:
                logger.warning(
                    "Orphan recovery race: failed to delete duplicate project %s: %s",
                    project_id, del_exc,
                )
            cursor.execute("SELECT project_id FROM devis WHERE id = %s", (devis_id,))
            winner = cursor.fetchone()
            if winner and winner.get("project_id"):
                logger.info(
                    "Orphan recovery race resolved: duplicate %s discarded, returning winner %s",
                    project_id, winner["project_id"],
                )
                return winner["project_id"]
            return None

        # Update linked opportunity if the devis was tied to one. Mirrors
        # accept_public_devis + convert_devis_to_project so all 3 devis→project
        # paths consistently mark the opportunity GAGNE. Defensive ALTER for
        # old tenants whose opportunities table predates these columns.
        opportunity_id = d.get("opportunity_id")
        if opportunity_id:
            for col, ctype in (
                ("projet_id", "INTEGER"),
                ("converted_at", "TIMESTAMP"),
                ("updated_at", "TIMESTAMP"),
            ):
                try:
                    cursor.execute(
                        f"ALTER TABLE opportunities ADD COLUMN IF NOT EXISTS {col} {ctype}"
                    )
                except Exception as opp_alter_exc:
                    logger.warning("ALTER opportunities ADD %s failed: %s", col, opp_alter_exc)
            try:
                cursor.execute(
                    "UPDATE opportunities SET projet_id = %s, statut = 'GAGNE', "
                    "converted_at = CURRENT_TIMESTAMP, updated_at = CURRENT_TIMESTAMP "
                    "WHERE id = %s",
                    (project_id, opportunity_id),
                )
            except Exception as opp_exc:
                logger.warning(
                    "UPDATE opportunities failed for devis %s: %s", devis_id, opp_exc
                )

        # Copy attachments (try/except to prevent abort if table missing)
        try:
            cursor.execute(
                "INSERT INTO project_attachments "
                "(project_id, filename, original_filename, file_size, file_type, file_extension, "
                "category, description, file_path, file_data, file_hash, upload_date, "
                "uploaded_by, is_active) "
                "SELECT %s, filename, original_filename, file_size, file_type, file_extension, "
                "category, description, file_path, file_data, file_hash, upload_date, "
                "uploaded_by, is_active "
                "FROM devis_attachments WHERE devis_id = %s AND is_active = TRUE",
                (project_id, devis_id),
            )
            copied_count = cursor.rowcount
            if copied_count > 0:
                logger.info("Copied %d attachments from devis %s to project %s", copied_count, devis_id, project_id)
        except Exception as att_exc:
            logger.warning("Could not copy devis attachments: %s", att_exc)
        logger.info("Auto-created project %s from devis %s (budget: %s)", project_id, devis_id, budget)
        return project_id

    except Exception as exc:
        logger.warning("_create_project_from_devis failed for devis %s: %s", devis_id, exc)
        return None
    finally:
        if cursor:
            cursor.close()


# ============================================
# DEVIS SEND
# ============================================


def _send_devis_email(
    to_email: str,
    enterprise_name: str,
    numero_devis: str,
    nom_projet: str,
    total_ttc: str,
    public_url: str,
    theme: Optional[dict] = None,
) -> bool:
    """Send a professional HTML email to the client with the devis link.

    `theme` is an optional tenant color palette; falls back to defaults when
    not provided so callers that don't have a cursor still work.
    Returns True if email was sent, False otherwise. Never raises.
    """
    from ..erp_config import SMTP_HOST, SMTP_PORT, SMTP_USER, SMTP_PASSWORD, SMTP_FROM_NAME, SMTP_USE_SSL, APP_BASE_URL
    from .html_utils import DEFAULT_DOCUMENT_THEME, THEME_KEYS

    if not SMTP_HOST or not SMTP_USER or not SMTP_PASSWORD:
        logger.warning("SMTP not configured -- skipping email for devis %s", numero_devis)
        return False

    _t = dict(DEFAULT_DOCUMENT_THEME)
    if isinstance(theme, dict):
        for k in THEME_KEYS:
            v = theme.get(k)
            if isinstance(v, str) and v.strip():
                _t[k] = v

    full_url = f"{APP_BASE_URL.rstrip('/')}{public_url}"
    esc = html_mod.escape
    safe_url = esc(full_url)

    html_body = f"""<!DOCTYPE html>
<html lang="fr">
<head><meta charset="UTF-8"></head>
<body style="margin:0;padding:0;background:#f4f6f9;font-family:'Segoe UI',Tahoma,Geneva,Verdana,sans-serif;">
  <table width="100%" cellpadding="0" cellspacing="0" style="background:#f4f6f9;padding:40px 0;">
    <tr><td align="center">
      <table width="600" cellpadding="0" cellspacing="0" style="background:#ffffff;border-radius:8px;overflow:hidden;box-shadow:0 2px 8px rgba(0,0,0,0.08);">
        <tr>
          <td style="background:{_t['primary']};padding:28px 40px;">
            <h1 style="margin:0;color:{_t['header_text']};font-size:22px;font-weight:700;">{esc(enterprise_name)}</h1>
          </td>
        </tr>
        <tr>
          <td style="padding:40px;">
            <p style="margin:0 0 20px;font-size:16px;color:#2d3748;line-height:1.6;">Bonjour,</p>
            <p style="margin:0 0 24px;font-size:16px;color:#2d3748;line-height:1.6;">
              Veuillez trouver ci-dessous votre soumission de la part de <strong>{esc(enterprise_name)}</strong>.
            </p>
            <table width="100%" cellpadding="0" cellspacing="0" style="background:#f7fafc;border:1px solid #e2e8f0;border-radius:8px;margin-bottom:32px;">
              <tr><td style="padding:24px;">
                <table width="100%" cellpadding="0" cellspacing="0">
                  <tr>
                    <td style="padding:6px 0;font-size:14px;color:#718096;">Soumission :</td>
                    <td style="padding:6px 0;font-size:14px;color:{_t['primary']};font-weight:600;text-align:right;">{esc(numero_devis)}</td>
                  </tr>
                  <tr>
                    <td style="padding:6px 0;font-size:14px;color:#718096;">Projet :</td>
                    <td style="padding:6px 0;font-size:14px;color:{_t['primary']};font-weight:600;text-align:right;">{esc(nom_projet)}</td>
                  </tr>
                  <tr>
                    <td colspan="2" style="border-top:1px solid #e2e8f0;padding-top:12px;">
                      <table width="100%" cellpadding="0" cellspacing="0">
                        <tr>
                          <td style="font-size:16px;color:{_t['primary']};font-weight:700;">Total TTC</td>
                          <td style="font-size:20px;color:{_t['primary']};font-weight:800;text-align:right;">{esc(total_ttc)}</td>
                        </tr>
                      </table>
                    </td>
                  </tr>
                </table>
              </td></tr>
            </table>
            <table width="100%" cellpadding="0" cellspacing="0">
              <tr>
                <td align="center" style="padding:8px 0 32px;">
                  <a href="{safe_url}" target="_blank"
                     style="display:inline-block;background:{_t['primary']};color:{_t['header_text']};text-decoration:none;padding:16px 48px;border-radius:6px;font-size:16px;font-weight:700;letter-spacing:0.5px;">
                    Consulter la soumission
                  </a>
                </td>
              </tr>
            </table>
            <p style="margin:0 0 8px;font-size:13px;color:#a0aec0;line-height:1.5;">
              Si le bouton ne fonctionne pas, copiez ce lien dans votre navigateur :
            </p>
            <p style="margin:0 0 24px;font-size:13px;color:{_t['primary']};word-break:break-all;">{safe_url}</p>
          </td>
        </tr>
        <tr>
          <td style="background:#f7fafc;padding:20px 40px;border-top:1px solid #e2e8f0;">
            <p style="margin:0;font-size:12px;color:#a0aec0;text-align:center;">
              {esc(enterprise_name)} &mdash; Envoye via Constructo AI
            </p>
          </td>
        </tr>
      </table>
    </td></tr>
  </table>
</body>
</html>"""

    # Apply tenant border color to the hardcoded gray borders in the template
    html_body = html_body.replace('#e2e8f0', _t['border'])

    text_body = (
        f"Bonjour,\n\n"
        f"Veuillez trouver ci-dessous votre soumission de la part de {enterprise_name}.\n\n"
        f"Soumission : {numero_devis}\n"
        f"Projet : {nom_projet}\n"
        f"Total TTC : {total_ttc}\n\n"
        f"Consultez votre soumission ici :\n{full_url}\n\n"
        f"---\n{enterprise_name} - Envoye via Constructo AI"
    )

    subject = f"Soumission {numero_devis} - {nom_projet}"

    try:
        msg = MIMEMultipart("alternative")
        msg["From"] = formataddr((SMTP_FROM_NAME, SMTP_USER))
        msg["To"] = to_email
        msg["Subject"] = subject
        msg["Date"] = formatdate(localtime=True)
        msg["Reply-To"] = SMTP_USER

        msg.attach(MIMEText(text_body, "plain", "utf-8"))
        msg.attach(MIMEText(html_body, "html", "utf-8"))

        if SMTP_USE_SSL:
            with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, timeout=30) as server:
                server.login(SMTP_USER, SMTP_PASSWORD)
                server.send_message(msg)
        else:
            with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=30) as server:
                server.starttls()
                server.login(SMTP_USER, SMTP_PASSWORD)
                server.send_message(msg)

        logger.info("Devis email sent to %s for %s", to_email, numero_devis)
        return True

    except Exception as exc:
        logger.error("Failed to send devis email to %s: %s", to_email, exc)
        return False


class DevisSendRequest(BaseModel):
    email: str

    @field_validator("email", mode="before")
    @classmethod
    def _validate_email(cls, v):
        s = (v or "").strip()
        if not _EMAIL_RE.match(s):
            raise ValueError("Adresse courriel invalide")
        return s


@router.post("/{devis_id}/send")
async def send_devis(devis_id: int, body: DevisSendRequest, user: ErpUser = Depends(get_current_user)):
    """Send a devis to a client via email and update status to ENVOYE."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")

    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        cursor = conn.cursor()

        # Fetch devis
        cursor.execute("SELECT id, validation_token, statut, nom_projet FROM devis WHERE id = %s", (devis_id,))
        devis = cursor.fetchone()
        if not devis:
            raise HTTPException(status_code=404, detail="Devis non trouvé")

        token = devis["validation_token"]
        if not token:
            nom_projet = devis.get("nom_projet") or "devis"
            token = _generate_readable_token(nom_projet)
            cursor.execute(
                "UPDATE devis SET validation_token = %s WHERE id = %s",
                (token, devis_id),
            )

        # Update status to Envoye
        cursor.execute(
            "UPDATE devis SET statut = 'Envoye' WHERE id = %s",
            (devis_id,),
        )

        # Store send date in metadonnees
        try:
            cursor.execute("SELECT metadonnees_json FROM devis WHERE id = %s", (devis_id,))
            row = cursor.fetchone()
            meta = row.get("metadonnees_json") or {} if row else {}
            if isinstance(meta, str):
                try:
                    meta = json.loads(meta)
                except (json.JSONDecodeError, TypeError):
                    meta = {}
            meta["sent_to"] = body.email
            meta["sent_at"] = datetime.now().isoformat()
            cursor.execute(
                "UPDATE devis SET metadonnees_json = %s WHERE id = %s",
                (json.dumps(meta), devis_id),
            )
        except Exception:
            pass


        # Register token in public lookup table
        _register_public_token(conn, token, user.schema, devis_id)

        # Build public URL
        public_url = f"/devis/public/{token}"

        # --- Send actual email ---
        email_sent = False
        try:
            db.set_tenant(conn, user.schema)
            cursor.execute(
                "SELECT numero_devis, nom_projet, investissement_total, "
                "total_avant_taxes, tps, tvq, prix_estime, total_travaux "
                "FROM devis WHERE id = %s",
                (devis_id,),
            )
            d = cursor.fetchone()

            from .html_utils import get_company_info, get_document_theme
            enterprise = get_company_info(cursor)
            ent_name = (enterprise or {}).get("nom", "") or "Votre entrepreneur"
            email_theme = get_document_theme(cursor)

            if d:
                ttc_val = float(d.get("investissement_total") or 0)
                if d.get("investissement_total") is None:
                    st = float(d.get("total_avant_taxes") or d.get("prix_estime") or d.get("total_travaux") or 0)
                    tps_v = float(d.get("tps")) if d.get("tps") is not None else round(st * 0.05, 2)
                    tvq_v = float(d.get("tvq")) if d.get("tvq") is not None else round(st * 0.09975, 2)
                    ttc_val = round(st + tps_v + tvq_v, 2)

                import asyncio
                email_sent = await asyncio.get_running_loop().run_in_executor(
                    None,
                    lambda: _send_devis_email(
                        to_email=body.email,
                        enterprise_name=ent_name,
                        numero_devis=d.get("numero_devis", ""),
                        nom_projet=d.get("nom_projet", ""),
                        total_ttc=_fmt_money(ttc_val),
                        public_url=public_url,
                        theme=email_theme,
                    ),
                )
        except Exception as mail_exc:
            logger.warning("Could not send devis email: %s", mail_exc)
            email_sent = False

        return {
            "sent": True,
            "email": body.email,
            "public_url": public_url,
            "token": token,
            "email_sent": email_sent,
            "message": (
                f"Soumission envoyee par courriel a {body.email}"
                if email_sent
                else "Soumission mise à jour (lien genere). Le courriel n'a pas pu etre envoye."
            ),
        }

    except HTTPException:
        raise
    except Exception as exc:
        logger.error("send_devis error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de l'envoi du devis")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


# ============================================
# PUBLIC DEVIS ENDPOINTS (NO AUTH)
# ============================================

@router.get("/public/{token}")
async def get_public_devis(token: str):
    """Public endpoint to view a devis by its validation token. No authentication required."""
    # Validate token format
    if not token or not re.match(r'^[a-zA-Z0-9\-_]{6,120}$', token):
        raise HTTPException(status_code=404, detail="Lien invalide")

    conn = db.get_conn()
    cursor = None
    try:
        # Fast path: look up from public tokens table
        lookup = _lookup_token(conn, token)
        if not lookup:
            # Fallback: scan all tenant schemas (backwards compatibility)
            lookup = _find_devis_by_token_fallback(conn, token)
        if not lookup:
            raise HTTPException(status_code=404, detail="Devis non trouvé")

        found_schema, devis_id = lookup

        db.set_tenant(conn, found_schema)
        cursor = conn.cursor()
        _ensure_visibility_columns(cursor)

        # Fetch devis
        cursor.execute("SELECT * FROM devis WHERE id = %s", (devis_id,))
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Devis non trouvé")
        devis = dict(row)

        # Only allow viewing if status permits
        if devis.get("statut") not in ("Envoye", "Accepte", "Refuse", "En attente"):
            raise HTTPException(status_code=404, detail="Devis non disponible")

        # Fetch lignes
        cursor.execute(
            "SELECT * FROM devis_lignes WHERE devis_id = %s ORDER BY sequence_ligne ASC",
            (devis_id,),
        )
        lignes = [dict(r) for r in cursor.fetchall()]

        # Fetch client company
        client_company = None
        if devis.get("client_company_id"):
            cursor.execute("SELECT * FROM companies WHERE id = %s", (devis["client_company_id"],))
            cli_row = cursor.fetchone()
            if cli_row:
                client_company = dict(cli_row)

        # Fetch client contact
        client_contact = None
        if devis.get("client_contact_id"):
            cursor.execute("SELECT * FROM contacts WHERE id = %s", (devis["client_contact_id"],))
            ct_row = cursor.fetchone()
            if ct_row:
                client_contact = dict(ct_row)

        # Fetch enterprise info
        enterprise = None
        try:
            from .html_utils import get_company_info
            enterprise = get_company_info(cursor)
        except Exception:
            pass
        entreprise_defaults = _get_entreprise_devis_defaults(cursor)
        from .html_utils import get_document_theme
        theme = get_document_theme(cursor)

        # Generate HTML
        html = _generate_devis_html(devis, lignes, client_company, enterprise, client_contact=client_contact, entreprise_defaults=entreprise_defaults, theme=theme)

        # Prepare response data
        for k in ("created_at", "date_prevu", "date_decision", "signature_date", "updated_at"):
            if devis.get(k) is not None:
                devis[k] = str(devis[k])
        for k in ("total_travaux", "tps", "tvq", "investissement_total",
                   "administration", "contingences", "profit", "total_avant_taxes"):
            if devis.get(k) is not None:
                devis[k] = float(devis[k])
        for l in lignes:
            for k in ("quantite", "prix_unitaire", "montant_ligne"):
                if l.get(k) is not None:
                    l[k] = float(l[k])
            # Strip internal pricing metadata before returning to the public
            # endpoint. Per-line markup % and MO/MAT split are business-internal
            # — already absent from the rendered HTML, but the JSON response
            # would otherwise leak them through DevTools. Defense-in-depth.
            for sensitive_key in (
                "admin_pct_ligne", "contingence_pct_ligne", "profit_pct_ligne",
                "mo_pct", "mat_pct",
            ):
                l.pop(sensitive_key, None)

        # Remove sensitive fields
        for key in ("validation_token", "notes", "metadonnees_json"):
            devis.pop(key, None)

        return {
            "devis": devis,
            "lignes": lignes,
            "html": html,
            "enterprise_name": enterprise.get("nom", "") if enterprise else "",
        }

    except HTTPException:
        raise
    except Exception as exc:
        logger.error("get_public_devis error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


class PublicDevisAcceptRequest(BaseModel):
    client_nom_signature: str
    signature_data: Optional[str] = None

    @field_validator("client_nom_signature")
    @classmethod
    def validate_nom(cls, v):
        if len(v.strip()) < 2:
            raise ValueError("Le nom doit contenir au moins 2 caracteres")
        if len(v) > 200:
            raise ValueError("Le nom ne doit pas depasser 200 caracteres")
        return v.strip()

    @field_validator("signature_data")
    @classmethod
    def validate_signature(cls, v):
        if v is None:
            return None
        if not isinstance(v, str):
            raise ValueError("signature_data doit etre une chaine")
        v = v.strip()
        if not v:
            return None  # whitespace-only treated as missing
        if len(v) > 500_000:
            raise ValueError("Les donnees de signature sont trop volumineuses")
        if not v.startswith("data:image/"):
            raise ValueError("signature_data doit etre une data URL image")
        return v


@router.post("/public/{token}/accept")
async def accept_public_devis(token: str, body: PublicDevisAcceptRequest):
    """Public endpoint to accept a devis. No authentication required.

    Flow (autocommit mode — each statement is its own transaction):
      1. Race-safe atomic UPDATE on devis.statut (only one concurrent request wins)
      2. Store signature metadata (best-effort)
      3. Create project from devis (best-effort — failure does NOT roll back acceptance)
      4. Link devis.project_id + update opportunity (best-effort)
      5. Copy attachments (best-effort)

    Once step 1 succeeds the endpoint always returns 200. Downstream failures
    are logged so an admin can reconcile — but the client gets confirmation
    that their legally-binding signature was captured. The previous flow had
    the opposite behaviour: if project creation crashed the whole acceptance
    rolled back and the user saw "Erreur" even though signing was successful.
    """
    if not token or not re.match(r'^[a-zA-Z0-9\-_]{6,120}$', token):
        raise HTTPException(status_code=404, detail="Lien invalide")

    conn = db.get_conn()
    cursor = None
    try:
        # Secure token lookup
        lookup = _lookup_token(conn, token)
        if not lookup:
            lookup = _find_devis_by_token_fallback(conn, token)
        if not lookup:
            raise HTTPException(status_code=404, detail="Devis non trouvé")

        found_schema, devis_id = lookup
        db.set_tenant(conn, found_schema)
        cursor = conn.cursor()

        # Step 1 — Race-safe atomic status change. The WHERE clause on statut
        # guarantees that even with concurrent requests (two tabs, double-click,
        # mobile retry), only the first request transitions Envoye→Accepte.
        # Subsequent requests get rowcount=0 and a 400 error.
        cursor.execute(
            "UPDATE devis SET statut = 'Accepte', signature_date = CURRENT_TIMESTAMP "
            "WHERE id = %s AND statut IN ('Envoye', 'En attente') "
            "RETURNING id",
            (devis_id,),
        )
        claimed = cursor.fetchone()
        if not claimed:
            # Either devis doesn't exist or is already in a final state.
            # Distinguish 404 vs 400 for a better error message.
            cursor.execute("SELECT statut FROM devis WHERE id = %s", (devis_id,))
            existing = cursor.fetchone()
            if not existing:
                raise HTTPException(status_code=404, detail="Devis non trouvé")
            raise HTTPException(
                status_code=400,
                detail="Ce devis ne peut plus etre accepte (statut actuel: {})".format(existing["statut"]),
            )

        # Step 2 — Signature metadata (best-effort; failure does not roll back acceptance)
        try:
            cursor.execute("SELECT metadonnees_json FROM devis WHERE id = %s", (devis_id,))
            meta_row = cursor.fetchone()
            meta = meta_row.get("metadonnees_json") or {} if meta_row else {}
            if isinstance(meta, str):
                try:
                    meta = json.loads(meta)
                except (json.JSONDecodeError, TypeError):
                    meta = {}
            meta["client_nom_signature"] = body.client_nom_signature
            meta["signature_date"] = datetime.now().isoformat()
            meta["accepted"] = True
            if body.signature_data:
                meta["signature_data"] = body.signature_data
            cursor.execute(
                "UPDATE devis SET metadonnees_json = %s WHERE id = %s",
                (json.dumps(meta), devis_id),
            )
        except Exception as meta_exc:
            logger.warning("Could not persist acceptance metadata for devis %s: %s", devis_id, meta_exc)

        # Step 3 — Create project (best-effort; admin can retry manually on failure)
        project_id = None
        try:
            cursor.execute("SELECT * FROM devis WHERE id = %s", (devis_id,))
            d = dict(cursor.fetchone())

            if not d.get("project_id"):
                # Compute budget (TTC)
                budget = float(d.get("investissement_total") or 0)
                if d.get("investissement_total") is None:
                    st = float(d.get("total_avant_taxes") or d.get("prix_estime") or d.get("total_travaux") or 0)
                    tps_v = float(d.get("tps") or 0)
                    tvq_v = float(d.get("tvq") or 0)
                    if d.get("tps") is None and st:
                        tps_v = round(st * 0.05, 2)
                    if d.get("tvq") is None and st:
                        tvq_v = round(st * 0.09975, 2)
                    budget = round(st + tps_v + tvq_v, 2)

                # Ensure all columns we're about to INSERT exist on this tenant (lazy migration)
                _ensure_projects_insert_columns(cursor, conn, found_schema)

                # Re-set tenant after ALTER TABLE commits/rollbacks
                db.set_tenant(conn, found_schema)

                # Fix sequence if out of sync (can happen after manual INSERTs).
                # Use GREATEST(max, 1) + 3-arg setval to avoid "value 0 is out of bounds"
                # when projects table is empty (PostgreSQL sequence range is 1..2^31-1).
                try:
                    cursor.execute(
                        "SELECT setval(pg_get_serial_sequence('projects', 'id'), "
                        "GREATEST(COALESCE((SELECT MAX(id) FROM projects), 0), 1), "
                        "(SELECT COUNT(*) > 0 FROM projects))"
                    )
                except Exception as seq_exc:
                    logger.warning("setval projects_id_seq failed: %s", seq_exc)
                    try:
                        db.set_tenant(conn, found_schema)
                    except Exception:
                        pass

                # Insert project — RETURNING id
                cursor.execute(
                    "INSERT INTO projects (nom_projet, client_company_id, client_contact_id, "
                    "client_nom_cache, po_client, statut, priorite, type_projet, "
                    "budget_total, date_prevu, description, "
                    "devis_id, devis_source_id, numero_devis, "
                    "created_at, updated_at) "
                    "VALUES (%s,%s,%s,%s,%s,'En cours',%s,'Construction',"
                    "%s,%s,%s,%s,%s,%s,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP) RETURNING id",
                    (
                        d.get("nom_projet") or f"Projet DEV-{devis_id}",
                        d.get("client_company_id"),
                        d.get("client_contact_id"),
                        d.get("client_nom_cache"),
                        d.get("po_client"),
                        d.get("priorite") or "Moyenne",
                        budget,
                        d.get("date_prevu"),
                        d.get("description") or d.get("notes") or f"Projet créé depuis devis #{devis_id}",
                        devis_id,
                        devis_id,
                        d.get("numero_devis"),
                    ),
                )
                project_id = cursor.fetchone()["id"]

                # Generate numero_projet (PROJ-YYYY-NNNNN) — try/except simple,
                # voir commentaire detaille au site #1 (~ligne 3710).
                try:
                    cursor.execute(
                        "UPDATE projects "
                        "SET numero_projet = 'PROJ-' || EXTRACT(YEAR FROM COALESCE(created_at, CURRENT_TIMESTAMP))::int "
                        "|| '-' || LPAD(id::text, 5, '0') "
                        "WHERE id = %s AND (numero_projet IS NULL OR numero_projet = '')",
                        (project_id,),
                    )
                except Exception as numproj_exc:
                    logger.error("Could not set numero_projet for project %s: %s — _backfill_numero_projet rattrapera au prochain list_projects", project_id, numproj_exc)

                # Link devis to project
                cursor.execute("UPDATE devis SET project_id = %s WHERE id = %s", (project_id, devis_id))

                # Update linked opportunity if exists. Defensive ALTER for old
                # tenants that predate the projet_id/converted_at columns.
                opportunity_id = d.get("opportunity_id")
                if opportunity_id:
                    for col, ctype in (
                        ("projet_id", "INTEGER"),
                        ("converted_at", "TIMESTAMP"),
                        ("updated_at", "TIMESTAMP"),
                    ):
                        try:
                            cursor.execute(
                                f"ALTER TABLE opportunities ADD COLUMN IF NOT EXISTS {col} {ctype}"
                            )
                        except Exception as opp_alter_exc:
                            logger.warning("ALTER opportunities ADD %s failed: %s", col, opp_alter_exc)
                    try:
                        cursor.execute(
                            "UPDATE opportunities SET projet_id = %s, statut = 'GAGNE', "
                            "converted_at = CURRENT_TIMESTAMP, updated_at = CURRENT_TIMESTAMP "
                            "WHERE id = %s",
                            (project_id, opportunity_id),
                        )
                    except Exception as opp_exc:
                        logger.warning("UPDATE opportunities failed for devis %s: %s", devis_id, opp_exc)
        except Exception as proj_exc:
            # Project creation failed AFTER devis was successfully accepted.
            # Log for admin reconciliation but return success to the client —
            # they signed, that part worked, project can be created manually.
            logger.error("Project creation failed for accepted devis %s: %s", devis_id, proj_exc)

        # Step 4 — Copy attachments (best-effort, non-critical)
        if project_id:
            att_cursor = None
            try:
                db.set_tenant(conn, found_schema)
                att_cursor = conn.cursor()
                att_cursor.execute(
                    "INSERT INTO project_attachments "
                    "(project_id, filename, original_filename, file_size, file_type, file_extension, "
                    "category, description, file_path, file_data, file_hash, upload_date, "
                    "uploaded_by, is_active) "
                    "SELECT %s, filename, original_filename, file_size, file_type, file_extension, "
                    "category, description, file_path, file_data, file_hash, upload_date, "
                    "uploaded_by, is_active "
                    "FROM devis_attachments WHERE devis_id = %s AND is_active = TRUE",
                    (project_id, devis_id),
                )
            except Exception as att_exc:
                logger.warning("Could not copy devis attachments on public accept: %s", att_exc)
            finally:
                if att_cursor:
                    att_cursor.close()

        logger.info("Public accept devis %s → project %s (schema: %s)", devis_id, project_id, found_schema)

        result = {
            "accepted": True,
            "project_created": project_id is not None,
            "message": "Devis accepte avec succes",
        }
        if project_id:
            result["project_id"] = project_id
            result["message"] = "Devis accepté — Projet créé automatiquement"
        return result

    except HTTPException:
        raise
    except Exception as exc:
        logger.error("accept_public_devis error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de l'acceptation du devis")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


class PublicDevisRefuseRequest(BaseModel):
    raison: Optional[str] = None

    @field_validator("raison")
    @classmethod
    def validate_raison(cls, v):
        if v and len(v) > 2000:
            raise ValueError("La raison ne doit pas depasser 2000 caracteres")
        return v


@router.post("/public/{token}/refuse")
async def refuse_public_devis(token: str, body: PublicDevisRefuseRequest):
    """Public endpoint to refuse a devis. No authentication required."""
    if not token or not re.match(r'^[a-zA-Z0-9\-_]{6,120}$', token):
        raise HTTPException(status_code=404, detail="Lien invalide")

    conn = db.get_conn()
    cursor = None
    try:
        # Secure token lookup
        lookup = _lookup_token(conn, token)
        if not lookup:
            lookup = _find_devis_by_token_fallback(conn, token)
        if not lookup:
            raise HTTPException(status_code=404, detail="Devis non trouvé")

        found_schema, devis_id = lookup
        db.set_tenant(conn, found_schema)
        cursor = conn.cursor()

        cursor.execute("SELECT id, statut, notes FROM devis WHERE id = %s", (devis_id,))
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Devis non trouvé")

        if row["statut"] not in ("Envoye", "En attente"):
            raise HTTPException(
                status_code=400,
                detail="Ce devis ne peut plus etre refuse (statut actuel: {})".format(row["statut"]),
            )

        # Add refusal reason to notes
        notes = row.get("notes") or ""
        if body.raison:
            notes += f"\n[REFUSE] Raison du client: {body.raison}"

        cursor.execute(
            "UPDATE devis SET statut = 'Refuse', notes = %s WHERE id = %s",
            (notes.strip(), devis_id),
        )

        # Store in metadonnees_json
        try:
            cursor.execute("SELECT metadonnees_json FROM devis WHERE id = %s", (devis_id,))
            meta_row = cursor.fetchone()
            meta = meta_row.get("metadonnees_json") or {} if meta_row else {}
            if isinstance(meta, str):
                try:
                    meta = json.loads(meta)
                except (json.JSONDecodeError, TypeError):
                    meta = {}
            meta["refused"] = True
            meta["refuse_raison"] = body.raison or ""
            meta["refuse_date"] = datetime.now().isoformat()
            cursor.execute(
                "UPDATE devis SET metadonnees_json = %s WHERE id = %s",
                (json.dumps(meta), devis_id),
            )
        except Exception:
            pass

        return {"refused": True, "message": "Devis refuse"}

    except HTTPException:
        raise
    except Exception as exc:
        logger.error("refuse_public_devis error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


# ============================================
# DEVIS ASSIGNMENTS
# ============================================

@router.get("/{devis_id}/assignments")
async def list_devis_assignments(devis_id: int, user: ErpUser = Depends(get_current_user)):
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        cursor = conn.cursor()
        # Ensure table exists
        cursor.execute(
            "CREATE TABLE IF NOT EXISTS devis_assignations ("
            "id SERIAL PRIMARY KEY, "
            "devis_id INT NOT NULL, "
            "employee_id INT NOT NULL, "
            "role VARCHAR(100), "
            "assigned_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"
        )
        cursor.execute(
            "SELECT da.id, da.devis_id, da.employee_id, da.role, da.assigned_at, "
            "e.prenom || ' ' || e.nom AS employe_nom "
            "FROM devis_assignations da "
            "LEFT JOIN employees e ON e.id = da.employee_id "
            "WHERE da.devis_id = %s ORDER BY da.assigned_at",
            (devis_id,),
        )
        items = []
        for row in cursor.fetchall():
            d = dict(row)
            if d.get("assigned_at"):
                d["assigned_at"] = str(d["assigned_at"])
            items.append(d)
        return {"items": items}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("list_devis_assignments error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.post("/{devis_id}/assignments")
async def add_devis_assignment(devis_id: int, body: DevisAssignmentCreate, user: ErpUser = Depends(get_current_user)):
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        cursor = conn.cursor()
        # Ensure table exists
        cursor.execute(
            "CREATE TABLE IF NOT EXISTS devis_assignations ("
            "id SERIAL PRIMARY KEY, "
            "devis_id INT NOT NULL, "
            "employee_id INT NOT NULL, "
            "role VARCHAR(100), "
            "assigned_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"
        )
        # Check duplicate
        cursor.execute(
            "SELECT id FROM devis_assignations WHERE devis_id = %s AND employee_id = %s",
            (devis_id, body.employee_id),
        )
        if cursor.fetchone():
            raise HTTPException(status_code=409, detail="Employé déjà assigné a ce devis")
        cursor.execute(
            "INSERT INTO devis_assignations (devis_id, employee_id, role, assigned_at) "
            "VALUES (%s, %s, %s, CURRENT_TIMESTAMP) RETURNING id",
            (devis_id, body.employee_id, body.role),
        )
        row = cursor.fetchone()
        return {"id": row["id"], "message": "Employe assigne au devis"}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("add_devis_assignment error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.delete("/{devis_id}/assignments/{assignment_id}")
async def remove_devis_assignment(devis_id: int, assignment_id: int, user: ErpUser = Depends(get_current_user)):
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user.schema)
        cursor = conn.cursor()
        cursor.execute(
            "DELETE FROM devis_assignations WHERE id = %s AND devis_id = %s",
            (assignment_id, devis_id),
        )
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="Assignation non trouvée")
        return {"message": "Assignation supprimée"}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("remove_devis_assignment error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()
