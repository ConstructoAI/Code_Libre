"""
ERP React - AI Router
Expert IA consultation, document analysis, usage tracking, credit management.
Based on expert_logic.py (858 lines) + assistant_ia_simple.py (10,184 lines)
+ ai_usage_tracker.py (1,569 lines) + ai_guard.py (189 lines) + document_analyzer.py (1,176 lines).
"""

import os
import sys
import io
import json
import base64
import logging
import time as time_module
from datetime import datetime, date, timedelta
from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from pydantic import BaseModel
from typing import Optional, List

from ..erp_auth import get_current_user, ErpUser
from .. import erp_database as db

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/ai", tags=["Intelligence Artificielle"])

# Import Anthropic client from project root
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
try:
    import anthropic
    _anthropic_client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY", ""))
except ImportError:
    _anthropic_client = None
    logger.warning("Anthropic SDK not installed - AI features unavailable")

try:
    import httpx
except ImportError:
    httpx = None

# AI Guard exempt company IDs
AI_GUARD_EXEMPT_IDS = {1, 105, 172}
AI_MODEL = "claude-sonnet-4-6"
AI_MAX_TOKENS = 32000


def _today_prompt_line() -> str:
    """Inject the current date into the system prompt so the model reasons
    correctly about overdue invoices, upcoming deadlines, etc. Without this,
    Claude falls back to its training-cutoff date and flags recent invoices
    as "overdue" when they are actually in the future."""
    today = datetime.now().strftime("%Y-%m-%d")
    return (
        f"DATE DU JOUR: {today} (format YYYY-MM-DD). "
        f"Utilise cette date comme reference absolue pour tout raisonnement "
        f"temporel (retards de paiement, echeances, dates limites, delais). "
        f"Ne te base JAMAIS sur ta date d'entrainement."
    )


# ============================================
# TOOL-USE: Outils BD pour l'Assistant IA
# ============================================

_AI_TOOLS = [
    {
        "name": "recherche_bd",
        "description": (
            "Execute une requete SQL SELECT en lecture seule sur la base de donnees du tenant. "
            "Utilise cet outil pour repondre a TOUTE question sur les donnees ERP. Retourne max 50 lignes.\n\n"
            "ASTUCE DECOUVERTE: Pour lister toutes les tables disponibles: "
            "SELECT table_name FROM information_schema.tables WHERE table_schema = current_schema() ORDER BY table_name\n"
            "Pour voir les colonnes d'une table: "
            "SELECT column_name, data_type FROM information_schema.columns WHERE table_schema = current_schema() AND table_name = 'NOM_TABLE'\n\n"
            "TABLES PAR MODULE:\n"
            "CORE: projects, employees, companies, contacts, users, work_centers, operations\n"
            "FORMULAIRES/BT: formulaires (type_formulaire='BON_TRAVAIL'|'ESTIMATION'), formulaire_lignes, "
            "bt_assignations (employee_id, role), bt_avancement, bt_comments\n"
            "COMPTABILITE: factures, facture_lignes, facture_paiements, depenses, paiements_recus, "
            "journal_entries, journal_lines, plan_comptable, grand_livre, periodes_comptables, "
            "annees_fiscales, cost_centers, budgets_projets\n"
            "DEVIS: devis, devis_lignes, devis_attachments, devis_envois\n"
            "ACHATS: bons_commande, bon_commande_lignes\n"
            "INVENTAIRE: produits, inventory_items, mouvements_stock, reservations_stock, produit_categories\n"
            "POINTAGE: time_entries (punch_in, punch_out, total_hours, employee_id, formulaire_bt_id, date_travail)\n"
            "CRM: opportunities, interactions, crm_activities\n"
            "CONFORMITE: licences_rbq, cartes_ccq, attestations_conformite, inspections_chantier\n"
            "LOGISTIQUE: logistics_deliveries, logistics_delivery_items, logistics_equipment, "
            "logistics_equipment_reservations, logistics_equipment_maintenance, logistics_vehicles, "
            "logistics_vehicle_trips, logistics_site_coordination, logistics_alerts\n"
            "IMMOBILIER: immo_terrains, immo_projets, immo_financement, immo_deblocages, "
            "immo_construction_phases, immo_unites, immo_commercialisation, immo_livraisons, "
            "immo_paiements, immo_documents, immo_inspections\n"
            "LOCATION: location_contrats, location_contrat_lignes, location_items, location_retours, location_contrats_employes, location_employes_heures, employee_location\n"
            "MAINTENANCE: maintenance_demandes, maintenance_interventions, maintenance_planification, "
            "maintenance_historique, maintenance_types, maintenance_pieces, maintenance_compteurs, maintenance_alertes\n"
            "SUBVENTIONS: subventions_categories, subventions_programmes, subventions_demandes, "
            "subventions_documents, subventions_eligibilite, subventions_alertes\n"
            "GPS: gps_locations, gps_vehicle_tracking, gps_routes, gps_geofences, gps_geofence_alerts\n"
            "PAIE: payroll_runs, payroll_entries, payroll_periods, payroll_config, tax_tables\n"
            "MESSAGERIE: conference_channels, conference_messages, conference_members, direct_messages\n"
            "NOTES: project_notes, project_note_files\n"
            "DOSSIERS: dossier_projets, dossier_documents, dossier_notes, dossier_etapes\n"
            "GANTT: gantt_dependencies (source_type/source_id->target_type/target_id, dependency_type IN ('finish_to_start','start_to_start','finish_to_finish','start_to_finish'), lag_days)\n"
            "METEO: alertes_meteo, previsions_meteo, historique_meteo_chantier\n"
            "CHARGE TRIBUTAIRE: charge_tributaire_calculs\n"
            "FONDS PREVOYANCE: fp_coproprietes, fp_composantes_batiment, fp_etudes, fp_projections, fp_carnet_entretien, fp_attestations_vente\n"
            "EMAIL: emails, email_threads, email_attachments\n"
            "CALENDRIER: calendar_events\n"
            "NOTIFICATIONS: notifications, broadcast_messages"
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "sql": {
                    "type": "string",
                    "description": "Requete SQL SELECT uniquement. Ex: SELECT e.prenom, e.nom, COUNT(te.id) FROM employees e JOIN time_entries te ON te.employee_id = e.id WHERE te.punch_in >= NOW() - INTERVAL '7 days' GROUP BY e.id",
                },
            },
            "required": ["sql"],
        },
    },
    {
        "name": "executer_action",
        "description": (
            "Execute une action de modification sur la base de donnees du tenant: INSERT, UPDATE ou DELETE. "
            "Utilise cet outil quand l'utilisateur demande de creer, modifier ou supprimer des donnees. "
            "IMPORTANT: Confirme toujours avec l'utilisateur avant de supprimer (DELETE). "
            "Applicable a TOUTES les tables du tenant: projects, employees, companies, contacts, "
            "formulaires (BT/BA), factures, devis, produits, bons_commande, opportunities, interactions, "
            "licences_rbq, cartes_ccq, logistics_*, immo_*, location_*, maintenance_*, subventions_*, "
            "project_notes, calendar_events, time_entries, journal_entries, depenses, etc. "
            "Utilise recherche_bd pour decouvrir les colonnes avant un INSERT. "
            "Toujours inclure les colonnes obligatoires et company_id quand applicable."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "sql": {
                    "type": "string",
                    "description": "Requete SQL INSERT, UPDATE ou DELETE.",
                },
                "description": {
                    "type": "string",
                    "description": "Description en francais de l'action effectuee pour le journal d'audit.",
                },
            },
            "required": ["sql", "description"],
        },
    },
]

# SQL interdit (securite)
_SQL_BLOCKED_KEYWORDS = {"DROP", "TRUNCATE", "ALTER", "CREATE", "GRANT", "REVOKE",
                          "SET ROLE", "SET SESSION", "COPY", "LOCK", "VACUUM"}


def _serialize_sql_value(v):
    """Convertit les types PostgreSQL non-serialisables en JSON."""
    import decimal
    if isinstance(v, decimal.Decimal):
        return float(v)
    if isinstance(v, datetime):
        return v.isoformat()
    if isinstance(v, date):
        return v.isoformat()
    if isinstance(v, timedelta):
        return str(v)
    if isinstance(v, bytes):
        return v.decode("utf-8", errors="replace")
    return v


def _strip_sql_comments(sql: str) -> str:
    """Strip SQL comments to prevent keyword bypass via DR/**/OP or DR--\\nOP."""
    import re
    sql = re.sub(r'/\*.*?\*/', ' ', sql, flags=re.DOTALL)  # block comments
    sql = re.sub(r'--[^\n]*', ' ', sql)  # line comments
    return sql


def _validate_sql_safe(sql_upper: str) -> str | None:
    """Verifie qu'une requete SQL ne contient pas de mots-cles dangereux. Retourne l'erreur ou None."""
    # Strip comments first to prevent bypass (DR/**/OP → DROP)
    sql_upper = _strip_sql_comments(sql_upper).upper()
    # Block semicolons to prevent multi-statement injection
    if ";" in sql_upper:
        return "Caractere interdit: point-virgule (;)"
    for kw in _SQL_BLOCKED_KEYWORDS:
        if kw in sql_upper:
            return f"Mot-cle interdit: {kw}"
    return None


def _execute_tenant_select(user_schema: str, sql: str) -> dict:
    """Execute un SELECT securise sur le schema du tenant. Retourne max 50 lignes."""
    sql_stripped = sql.strip().rstrip(";")  # strip trailing semicolons
    sql_upper = sql_stripped.upper()
    # Valider que c'est bien un SELECT ou WITH ... SELECT (CTE)
    if not (sql_upper.startswith("SELECT") or sql_upper.startswith("WITH")):
        return {"error": "Seules les requetes SELECT sont autorisees avec recherche_bd."}
    # Bloquer les mots-cles dangereux (strips comments first)
    err = _validate_sql_safe(sql_upper)
    if err:
        return {"error": err}
    # Bloquer INSERT/UPDATE/DELETE en standalone
    clean_upper = _strip_sql_comments(sql_upper).upper()
    for kw in ("INSERT INTO", "UPDATE ", "DELETE FROM"):
        if kw in clean_upper:
            return {"error": f"Mot-cle interdit dans un SELECT: {kw.strip()}"}

    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user_schema)
        # Le pool retourne autocommit=True (lecon #122) — sans bascule en
        # autocommit=False, SET TRANSACTION READ ONLY et SET LOCAL sont
        # silencieusement no-op (PG WARNING "can only be used in transaction
        # blocks"), ce qui DESACTIVE la protection read-only et le timeout 10s.
        conn.autocommit = False
        cursor = conn.cursor()
        # Read-only transaction + timeout for defense in depth
        cursor.execute("SET TRANSACTION READ ONLY")
        cursor.execute("SET LOCAL statement_timeout TO '10000'")
        # Ajouter LIMIT si absent
        if "LIMIT" not in sql_upper:
            sql_stripped = sql_stripped.rstrip(";") + " LIMIT 50"
        cursor.execute(sql_stripped)
        columns = [desc[0] for desc in cursor.description] if cursor.description else []
        rows = cursor.fetchall()
        # Rollback explicite — read-only, aucune ecriture ne devrait passer.
        # Termine proprement la transaction avant de restaurer autocommit.
        conn.rollback()
        results = [dict(r) for r in rows] if rows else []
        # Convertir les types non-serialisables
        for row in results:
            for k, v in row.items():
                row[k] = _serialize_sql_value(v)
        # Audit log SELECT
        logger.info("AI select user=%s schema=%s: %s | rows=%d",
                     getattr(cursor, '_user', '?'), user_schema, sql_stripped[:200], len(results))
        return {"columns": columns, "rows": results, "count": len(results)}
    except Exception as exc:
        try:
            conn.rollback()
        except Exception:
            pass
        return {"error": str(exc)}
    finally:
        if cursor:
            cursor.close()
        try:
            conn.autocommit = True  # Restore default before returning to pool
        except Exception:
            pass
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


def _execute_tenant_action(user_schema: str, sql: str, description: str, user: 'ErpUser') -> dict:
    """Execute un INSERT/UPDATE/DELETE securise sur le schema du tenant."""
    sql_stripped = sql.strip()
    sql_upper = sql_stripped.upper()
    # Valider le type de requete
    is_insert = sql_upper.startswith("INSERT")
    is_update = sql_upper.startswith("UPDATE")
    is_delete = sql_upper.startswith("DELETE")
    if not (is_insert or is_update or is_delete):
        return {"error": "Seules les requetes INSERT, UPDATE et DELETE sont autorisees."}
    # Bloquer les mots-cles dangereux
    err = _validate_sql_safe(sql_upper)
    if err:
        return {"error": err}

    conn = db.get_conn()
    cursor = None
    try:
        db.set_tenant(conn, user_schema)
        # Le pool retourne autocommit=True (lecon #122) — sans bascule en
        # autocommit=False, SET LOCAL statement_timeout est silencieusement
        # no-op (PG WARNING) et l'INSERT/UPDATE/DELETE auto-commit sans
        # protection contre les write-bombs.
        conn.autocommit = False
        cursor = conn.cursor()
        # Timeout requete: 10 secondes max
        cursor.execute("SET LOCAL statement_timeout TO '10000'")
        cursor.execute(sql_stripped)
        rowcount = cursor.rowcount
        # Recuperer l'id si INSERT RETURNING
        returning_row = None
        if is_insert and cursor.description:
            returning_row = cursor.fetchone()
            if returning_row:
                returning_row = dict(returning_row)
                for k, v in returning_row.items():
                    returning_row[k] = _serialize_sql_value(v)
        conn.commit()
        # Audit log
        logger.info("AI action [%s] user=%s schema=%s: %s | SQL: %s | rows=%d",
                     "INSERT" if is_insert else "UPDATE" if is_update else "DELETE",
                     getattr(user, 'username', '?'), user_schema, description, sql_stripped[:200], rowcount)
        result = {"success": True, "rows_affected": rowcount, "description": description}
        if returning_row:
            result["returning"] = returning_row
        return result
    except Exception as exc:
        try:
            conn.rollback()
        except Exception:
            pass
        return {"error": str(exc)}
    finally:
        if cursor:
            cursor.close()
        try:
            conn.autocommit = True  # Restore default before returning to pool
        except Exception:
            pass
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


def _call_claude(*, model=AI_MODEL, max_tokens=AI_MAX_TOKENS, system=None, messages, tools=None):
    """Call Claude via streaming to support long Opus requests (>10 min).

    Wrapped with retry on transient connection errors (Render edge proxy can
    close the chunked stream on long responses, surfacing as
    `httpx.RemoteProtocolError`). httpx errors are NOT subclasses of
    `anthropic.APIError`, so on final failure they are wrapped in
    `anthropic.APIConnectionError` so callers' typed handlers map them to a
    503 instead of a generic 500. Mirrors the helper in `devis.py`.
    """
    kwargs = {"model": model, "max_tokens": max_tokens, "messages": messages, "timeout": 600.0}
    if system:
        kwargs["system"] = system
    if tools:
        kwargs["tools"] = tools

    transient_excs = [anthropic.APIConnectionError]
    if httpx is not None:
        transient_excs.extend([httpx.RemoteProtocolError, httpx.ReadError, httpx.ReadTimeout])
    transient_excs = tuple(transient_excs)

    last_exc = None
    for attempt in range(2):
        try:
            with _anthropic_client.messages.stream(**kwargs) as stream:
                return stream.get_final_message()
        except transient_excs as exc:
            last_exc = exc
            logger.warning(
                "_call_claude transient error on attempt %d: %s — retrying",
                attempt + 1, exc,
            )
            time_module.sleep(0.5 * (attempt + 1))

    if isinstance(last_exc, anthropic.APIError):
        raise last_exc
    try:
        req = getattr(last_exc, "request", None)
    except Exception:
        req = None
    try:
        wrapped = anthropic.APIConnectionError(
            message=f"Anthropic connection error: {last_exc}",
            request=req,
        )
    except Exception:
        raise last_exc
    raise wrapped from last_exc


class ChatMessage(BaseModel):
    message: str
    profile: str = "expert_construction"
    context: Optional[str] = None
    conversation_id: Optional[int] = None


class DocumentAnalysisRequest(BaseModel):
    filename: str
    content_base64: str
    analysis_type: str = "general"


# ============================================
# AI PROFILES
# ============================================

AI_PROFILES = {
    "expert_construction": {
        "name": "Expert Construction",
        "system": "Tu es un expert en construction au Quebec. Tu reponds en francais avec des references aux normes RBQ, CCQ et CNB. Tu donnes des conseils pratiques et precis.",
    },
    "estimateur": {
        "name": "Estimateur",
        "system": "Tu es un estimateur professionnel en construction au Quebec. Tu aides a estimer les couts de projets avec precision, en tenant compte des materiaux, main-d'oeuvre et taxes (TPS 5%, TVQ 9.975%).",
    },
    "comptable": {
        "name": "Comptable Construction",
        "system": "Tu es un comptable specialise dans l'industrie de la construction au Quebec. Tu connais les DAS, TPS/TVQ, les retenues a la source et la paie CCQ.",
    },
    "juridique": {
        "name": "Conseiller Juridique",
        "system": "Tu es un conseiller juridique specialise en droit de la construction au Quebec. Tu references la Loi sur le batiment, le Code civil et les reglements RBQ.",
    },
    "securite": {
        "name": "Expert Securite",
        "system": "Tu es un expert en sante et securite sur les chantiers de construction au Quebec. Tu connais les normes CNESST et les reglements SST.",
    },
    "general": {
        "name": "Assistant Constructo AI",
        "system": (
            "Tu es l'assistant intelligent de Constructo AI, un ERP pour l'industrie de la construction au Quebec. "
            "Tu es polyvalent et tu adaptes automatiquement ton expertise selon le sujet de la conversation.\n\n"

            "EXPERTISES DISPONIBLES — adapte-toi selon les mots-cles de l'utilisateur:\n"
            "- ELECTRICITE: normes CSA C22.1, Code electrique Quebec, calibres de fils, charges, panneaux, circuits\n"
            "- PLOMBERIE: Code de plomberie Quebec, tuyauterie, drainage, ventilation, appareils sanitaires\n"
            "- ESTIMATION: couts materiaux, main-d'oeuvre, sous-traitants, prix du marche Quebec, TPS 5%, TVQ 9.975%\n"
            "- COMPTABILITE: DAS, TPS/TVQ, retenues a la source, paie CCQ, ecritures comptables, facturation\n"
            "- JURIDIQUE: Loi sur le batiment, Code civil, reglements RBQ, contrats, hypotheques legales\n"
            "- SECURITE/SST: normes CNESST, reglements SST, EPI, plans de securite chantier\n"
            "- STRUCTURE/BETON: calculs de charges, dalles, fondations, coffrage, normes CSA A23.3\n"
            "- TOITURE: bardeaux, membranes, pentes, solins, ventilation de toit\n"
            "- ISOLATION/ENVELOPPE: valeurs R, pare-vapeur, pare-air, Code energetique Quebec\n"
            "- SOUDURE: certifications CWB, normes CSA W47.1, procedes SMAW/GMAW/FCAW\n"
            "- GESTION ERP: projets, employes, factures, inventaire, devis, bons de travail, bons de commande\n\n"

            "DONNEES ERP: Tu as acces aux donnees temps reel de l'entreprise. "
            "Quand l'utilisateur pose une question sur ses projets, employes, factures, inventaire, etc., "
            "base-toi exclusivement sur les donnees fournies dans le contexte. Ne jamais inventer de donnees.\n\n"

            "REGLES:\n"
            "- Reponds en francais du Quebec\n"
            "- Sois concis et precis\n"
            "- Reference les normes applicables (RBQ, CCQ, CNB, CSA, CNESST) quand pertinent\n"
            "- Pour les questions ERP, utilise les donnees reelles du tenant\n"
            "- Pour les questions techniques, donne des conseils pratiques et actionnables\n"
            "- Si tu ne sais pas ou si l'info n'est pas dans le contexte, dis-le clairement"
        ),
    },
}


# ============================================
# AI GUARD
# ============================================

def check_ai_guard(user: ErpUser) -> tuple[bool, str]:
    """Check if the user is allowed to use AI features."""
    # Super-admin always allowed
    if user.user_type == "super_admin":
        return True, ""

    # Check exempt companies
    if user.schema:
        conn = db.get_conn()
        cursor = None
        try:
            cursor = conn.cursor()
            # entreprises.slug stores full prefix (e.g. "tenant_constructi_08e4ef")
            slug = user.schema if user.schema.startswith("tenant_") else f"tenant_{user.schema}"
            cursor.execute(
                "SELECT id FROM public.entreprises WHERE slug = %s", (slug,)
            )
            row = cursor.fetchone()
            if row and row["id"] in AI_GUARD_EXEMPT_IDS:
                return True, ""
        except Exception:
            pass
        finally:
            if cursor:
                cursor.close()
            conn.close()

    # For now, allow all authenticated users (credit check can be added later)
    return True, ""


def _ensure_ai_tables(cursor, conn):
    """Ensure all AI-related tables exist.
    If ai_prepaid_credits already exists (e.g., from erp_stripe.py with
    balance_cad/entreprise_id schema), we add our columns (tenant_slug,
    balance_usd etc.) without recreating the table.
    """
    # Usage tracking table
    cursor.execute(
        "SELECT EXISTS (SELECT FROM information_schema.tables "
        "WHERE table_schema = 'public' AND table_name = 'ai_usage_tracking')"
    )
    if not cursor.fetchone().get("exists", False):
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS public.ai_usage_tracking (
                id SERIAL PRIMARY KEY,
                user_id INTEGER,
                tenant_slug VARCHAR(100),
                feature VARCHAR(100),
                model VARCHAR(100),
                input_tokens INTEGER DEFAULT 0,
                output_tokens INTEGER DEFAULT 0,
                cost_usd NUMERIC(12,6) DEFAULT 0,
                duration_ms INTEGER DEFAULT 0,
                success BOOLEAN DEFAULT TRUE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.commit()
    else:
        # Add missing columns if table was created by Streamlit app with different schema
        tracking_cols_to_add = {
            "tenant_slug": "VARCHAR(100)",
            "input_tokens": "INTEGER DEFAULT 0",
            "output_tokens": "INTEGER DEFAULT 0",
            "cost_usd": "NUMERIC(12,6) DEFAULT 0",
            "duration_ms": "INTEGER DEFAULT 0",
            "success": "BOOLEAN DEFAULT TRUE",
            "model": "VARCHAR(100)",
        }
        for col, col_type in tracking_cols_to_add.items():
            cursor.execute(
                "SELECT column_name FROM information_schema.columns "
                "WHERE table_schema = 'public' AND table_name = 'ai_usage_tracking' "
                "AND column_name = %s",
                (col,),
            )
            if not cursor.fetchone():
                # Rename model_used → model if legacy column exists
                if col == "model":
                    cursor.execute(
                        "SELECT column_name FROM information_schema.columns "
                        "WHERE table_schema = 'public' AND table_name = 'ai_usage_tracking' "
                        "AND column_name = 'model_used'"
                    )
                    if cursor.fetchone():
                        try:
                            cursor.execute("ALTER TABLE public.ai_usage_tracking RENAME COLUMN model_used TO model")
                            conn.commit()
                            continue
                        except Exception:
                            conn.rollback()
                            # Rename failed — fall through to ADD COLUMN below
                try:
                    cursor.execute(f"ALTER TABLE public.ai_usage_tracking ADD COLUMN IF NOT EXISTS {col} {col_type}")
                    conn.commit()
                except Exception:
                    conn.rollback()

    # Prepaid credits table
    cursor.execute(
        "SELECT EXISTS (SELECT FROM information_schema.tables "
        "WHERE table_schema = 'public' AND table_name = 'ai_prepaid_credits')"
    )
    if not cursor.fetchone().get("exists", False):
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS public.ai_prepaid_credits (
                id SERIAL PRIMARY KEY,
                entreprise_id INTEGER NOT NULL,
                tenant_slug VARCHAR(100),
                balance_cad NUMERIC(12,4) DEFAULT 0.00,
                balance_usd NUMERIC(12,4) NOT NULL DEFAULT 0.00,
                monthly_limit_usd NUMERIC(12,4) DEFAULT 999999.99,
                auto_recharge BOOLEAN DEFAULT FALSE,
                recharge_amount_usd NUMERIC(12,4) DEFAULT 10.00,
                stripe_payment_method VARCHAR(255),
                total_consumed_usd NUMERIC(12,4) NOT NULL DEFAULT 0.00,
                total_charged_usd NUMERIC(12,4) NOT NULL DEFAULT 0.00,
                charges_count INTEGER NOT NULL DEFAULT 0,
                last_charge_stripe_id VARCHAR(255),
                last_charge_at TIMESTAMP,
                product_type VARCHAR(50) NOT NULL DEFAULT 'ERP',
                entreprise_nom VARCHAR(255),
                billing_year INTEGER NOT NULL,
                billing_month INTEGER NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                CONSTRAINT ai_prepaid_credits_entreprise_id_product_type_billing_year__key
                    UNIQUE (entreprise_id, product_type, billing_year, billing_month)
            )
        """)
        conn.commit()
    else:
        # Add missing columns to existing table (created by erp_stripe or earlier version)
        cols_to_add = {
            "tenant_slug": "VARCHAR(100)",
            "balance_usd": "NUMERIC(12,4) DEFAULT 0.00",
            "balance_cad": "NUMERIC(12,4) DEFAULT 0.00",
            "monthly_limit_usd": "NUMERIC(12,4) DEFAULT 999999.99",
            "recharge_amount_usd": "NUMERIC(12,4) DEFAULT 10.00",
            "auto_recharge": "BOOLEAN DEFAULT FALSE",
            "stripe_payment_method": "VARCHAR(255)",
            "total_consumed_usd": "NUMERIC(12,4) DEFAULT 0.00",
            "total_charged_usd": "NUMERIC(12,4) DEFAULT 0.00",
            "charges_count": "INTEGER DEFAULT 0",
            "last_charge_stripe_id": "VARCHAR(255)",
            "last_charge_at": "TIMESTAMP",
            "product_type": "VARCHAR(50) DEFAULT 'ERP'",
            "entreprise_nom": "VARCHAR(255)",
            "billing_year": "INTEGER",
            "billing_month": "INTEGER",
        }
        for col, col_type in cols_to_add.items():
            # ADD COLUMN IF NOT EXISTS avoids race between the SELECT above
            # and the ALTER below: two workers on a fresh tenant could each
            # observe the column as missing and then collide on ALTER.
            try:
                cursor.execute(
                    f"ALTER TABLE public.ai_prepaid_credits "
                    f"ADD COLUMN IF NOT EXISTS {col} {col_type}"
                )
                conn.commit()
            except Exception:
                conn.rollback()

        # Backfill NULL billing_year/billing_month from created_at
        try:
            cursor.execute(
                "UPDATE public.ai_prepaid_credits "
                "SET billing_year = EXTRACT(YEAR FROM COALESCE(updated_at, CURRENT_DATE))::int, "
                "    billing_month = EXTRACT(MONTH FROM COALESCE(updated_at, CURRENT_DATE))::int "
                "WHERE billing_year IS NULL OR billing_month IS NULL"
            )
        except Exception:
            pass

        # Ensure UNIQUE composite constraint exists for ON CONFLICT clauses
        try:
            cursor.execute(
                "SELECT 1 FROM pg_constraint WHERE conname = 'ai_prepaid_credits_entreprise_id_product_type_billing_year__key'"
            )
            if not cursor.fetchone():
                cursor.execute(
                    "ALTER TABLE public.ai_prepaid_credits "
                    "ADD CONSTRAINT ai_prepaid_credits_entreprise_id_product_type_billing_year__key "
                    "UNIQUE (entreprise_id, product_type, billing_year, billing_month)"
                )
                conn.commit()
        except Exception:
            conn.rollback()


def _get_tenant_slug(user: ErpUser) -> str:
    """Extract tenant slug from ErpUser schema."""
    if user.schema:
        return user.schema.replace("tenant_", "")
    return ""


PREPAID_RECHARGE_AMOUNT = float(os.getenv("AI_PREPAID_RECHARGE_AMOUNT", "10.00"))
# Minimum balance threshold: auto-recharge triggers when balance falls below this.
# Prevents the pattern where balance is $0.10, API call costs $2, and $1.90 is given free.
MIN_BALANCE_THRESHOLD = 0.10


def _get_entreprise_stripe_id(slug: str) -> tuple[int | None, str | None]:
    """Look up entreprise_id and stripe_customer_id from slug."""
    full_slug = slug if slug.startswith("tenant_") else f"tenant_{slug}"
    conn = db.get_conn()
    cursor = None
    try:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, stripe_customer_id FROM public.entreprises WHERE slug = %s",
            (full_slug,),
        )
        row = cursor.fetchone()
        if row:
            return row.get("id"), row.get("stripe_customer_id")
        return None, None
    except Exception:
        return None, None
    finally:
        if cursor:
            cursor.close()
        conn.close()


def _auto_recharge_credits(slug: str, cursor, conn) -> float:
    """Auto-recharge credits via Stripe when balance is 0.
    Returns new balance or 0.0 if recharge failed.
    Raises HTTPException(402) with a clear FR message if the card was
    declined (insufficient_funds, expired_card, etc.) so the user knows
    *why* their AI quota wasn't refilled instead of seeing a generic
    "Credits IA epuises" message.
    """
    from ..erp_stripe import (
        charge_ai_prepaid_credit,
        format_decline_message_fr,
    )

    entreprise_id, stripe_customer_id = _get_entreprise_stripe_id(slug)
    if not stripe_customer_id:
        logger.warning("Auto-recharge: no stripe_customer_id for slug=%s", slug)
        return 0.0

    # Charge $10 on client's Stripe card
    result = charge_ai_prepaid_credit(
        stripe_customer_id=stripe_customer_id,
        amount=PREPAID_RECHARGE_AMOUNT,
        currency="cad",
    )

    if not result or not result.get("paid"):
        logger.warning("Auto-recharge failed for slug=%s: %s", slug, result)
        # Si Stripe a rejete la carte, surfacer la raison exacte au lieu
        # d'un generique "Credits IA epuises". Le caller (_check_credits)
        # re-raise HTTPException intacte (lignes "except HTTPException: raise").
        if result and result.get("error_type") == "card_declined":
            decline_code = result.get("decline_code") or ""
            user_message = result.get("user_message") or ""
            detail = format_decline_message_fr(decline_code, user_message)
            raise HTTPException(
                status_code=402,
                detail=f"Recharge auto refusee — {detail}",
            )
        return 0.0

    # Add credits to balance + update tracking fields
    invoice_id = result.get("invoice_id", "")
    try:
        cursor.execute(
            "UPDATE public.ai_prepaid_credits "
            "SET balance_usd = COALESCE(balance_usd, 0) + %s, "
            "    balance_cad = COALESCE(balance_cad, 0) + %s, "
            "    total_charged_usd = COALESCE(total_charged_usd, 0) + %s, "
            "    charges_count = COALESCE(charges_count, 0) + 1, "
            "    last_charge_stripe_id = %s, "
            "    last_charge_at = CURRENT_TIMESTAMP, "
            "    updated_at = CURRENT_TIMESTAMP "
            "WHERE tenant_slug = %s AND product_type = 'ERP' "
            "AND billing_year = EXTRACT(YEAR FROM CURRENT_DATE)::int "
            "AND billing_month = EXTRACT(MONTH FROM CURRENT_DATE)::int",
            (PREPAID_RECHARGE_AMOUNT, PREPAID_RECHARGE_AMOUNT, PREPAID_RECHARGE_AMOUNT,
             invoice_id, slug),
        )
        if cursor.rowcount == 0 and entreprise_id:
            cursor.execute(
                "UPDATE public.ai_prepaid_credits "
                "SET balance_usd = COALESCE(balance_usd, 0) + %s, "
                "    balance_cad = COALESCE(balance_cad, 0) + %s, "
                "    total_charged_usd = COALESCE(total_charged_usd, 0) + %s, "
                "    charges_count = COALESCE(charges_count, 0) + 1, "
                "    last_charge_stripe_id = %s, "
                "    last_charge_at = CURRENT_TIMESTAMP, "
                "    updated_at = CURRENT_TIMESTAMP "
                "WHERE entreprise_id = %s AND product_type = 'ERP' "
                "AND billing_year = EXTRACT(YEAR FROM CURRENT_DATE)::int "
                "AND billing_month = EXTRACT(MONTH FROM CURRENT_DATE)::int",
                (PREPAID_RECHARGE_AMOUNT, PREPAID_RECHARGE_AMOUNT, PREPAID_RECHARGE_AMOUNT,
                 invoice_id, entreprise_id),
            )
            if cursor.rowcount == 0:
                # CRITICAL: Stripe charged but no DB row updated.
                # Attempt INSERT as last resort so credits are not lost.
                logger.critical(
                    "STRIPE CHARGED BUT NO DB ROW UPDATED: slug=%s entreprise_id=%s "
                    "invoice=%s amount=%.2f — attempting INSERT fallback",
                    slug, entreprise_id, invoice_id, PREPAID_RECHARGE_AMOUNT,
                )
                from datetime import date as _date
                _now = _date.today()
                try:
                    cursor.execute(
                        "INSERT INTO public.ai_prepaid_credits "
                        "(entreprise_id, tenant_slug, product_type, balance_usd, balance_cad, "
                        " total_charged_usd, charges_count, last_charge_stripe_id, "
                        " last_charge_at, billing_year, billing_month) "
                        "VALUES (%s, %s, 'ERP', %s, %s, %s, 1, %s, CURRENT_TIMESTAMP, %s, %s) "
                        "ON CONFLICT (entreprise_id, product_type, billing_year, billing_month) "
                        "DO UPDATE SET balance_usd = ai_prepaid_credits.balance_usd + %s, "
                        "             balance_cad = ai_prepaid_credits.balance_cad + %s, "
                        "             total_charged_usd = COALESCE(ai_prepaid_credits.total_charged_usd, 0) + %s, "
                        "             charges_count = COALESCE(ai_prepaid_credits.charges_count, 0) + 1, "
                        "             last_charge_stripe_id = %s, "
                        "             last_charge_at = CURRENT_TIMESTAMP",
                        (entreprise_id, slug, PREPAID_RECHARGE_AMOUNT, PREPAID_RECHARGE_AMOUNT,
                         PREPAID_RECHARGE_AMOUNT, invoice_id, _now.year, _now.month,
                         PREPAID_RECHARGE_AMOUNT, PREPAID_RECHARGE_AMOUNT, PREPAID_RECHARGE_AMOUNT,
                         invoice_id),
                    )
                except Exception as insert_exc:
                    logger.critical(
                        "INSERT FALLBACK ALSO FAILED: slug=%s invoice=%s error=%s",
                        slug, invoice_id, insert_exc,
                    )
        conn.commit()
        logger.info(
            "Auto-recharge OK: slug=%s amount=%.2f invoice=%s",
            slug, PREPAID_RECHARGE_AMOUNT, invoice_id,
        )
        return PREPAID_RECHARGE_AMOUNT
    except Exception as exc:
        logger.error("Auto-recharge DB update failed: %s", exc)
        conn.rollback()
        return 0.0


def _check_credits(user: ErpUser) -> tuple[bool, float]:
    """Check if the tenant has AI credits remaining.
    Returns (allowed, balance).
    Super-admins always allowed.
    If balance < MIN_BALANCE_THRESHOLD, attempts auto-recharge via Stripe ($10).
    Handles both tenant_slug-based and entreprise_id-based credit rows.
    """
    if user.user_type == "super_admin":
        return True, 999.99

    slug = _get_tenant_slug(user)
    if not slug:
        return True, 0.0

    conn = db.get_conn()
    cursor = None
    try:
        # Ensure autocommit is False so FOR UPDATE lock holds across
        # the Stripe charge call in _auto_recharge_credits().
        conn.autocommit = False
        cursor = conn.cursor()
        _ensure_ai_tables(cursor, conn)

        # Try tenant_slug first, then fall back to entreprise_id lookup
        # FOR UPDATE prevents race condition: two concurrent requests both seeing
        # balance=0 and triggering duplicate Stripe charges.
        cursor.execute(
            "SELECT balance_usd, balance_cad, monthly_limit_usd FROM public.ai_prepaid_credits "
            "WHERE tenant_slug = %s AND product_type = 'ERP' "
            "AND billing_year = EXTRACT(YEAR FROM CURRENT_DATE)::int "
            "AND billing_month = EXTRACT(MONTH FROM CURRENT_DATE)::int "
            "FOR UPDATE",
            (slug,),
        )
        row = cursor.fetchone()

        if not row:
            # Try matching via entreprise slug (entreprises.slug has "tenant_" prefix)
            full_slug = slug if slug.startswith("tenant_") else f"tenant_{slug}"
            cursor.execute(
                "SELECT apc.balance_usd, apc.balance_cad, apc.monthly_limit_usd "
                "FROM public.ai_prepaid_credits apc "
                "JOIN public.entreprises e ON apc.entreprise_id = e.id "
                "WHERE e.slug = %s AND apc.product_type = 'ERP' "
                "AND apc.billing_year = EXTRACT(YEAR FROM CURRENT_DATE)::int "
                "AND apc.billing_month = EXTRACT(MONTH FROM CURRENT_DATE)::int "
                "FOR UPDATE OF apc",
                (full_slug,),
            )
            row = cursor.fetchone()

        if not row:
            # Create credits entry with 0 balance — no free credits.
            # Client will be charged $10 via Stripe auto-recharge below.
            full_slug = slug if slug.startswith("tenant_") else f"tenant_{slug}"
            ent_id, _ = _get_entreprise_stripe_id(slug)
            if not ent_id:
                try:
                    cursor.execute("SELECT id FROM public.entreprises WHERE slug = %s", (full_slug,))
                    ent_row = cursor.fetchone()
                    ent_id = ent_row["id"] if ent_row else 0
                except Exception:
                    ent_id = 0
            from datetime import date
            now = date.today()
            try:
                cursor.execute(
                    "INSERT INTO public.ai_prepaid_credits "
                    "(entreprise_id, tenant_slug, product_type, balance_usd, balance_cad, "
                    " billing_year, billing_month) "
                    "VALUES (%s, %s, 'ERP', 0.00, 0.00, %s, %s) "
                    "ON CONFLICT (entreprise_id, product_type, billing_year, billing_month) DO NOTHING",
                    (ent_id, slug, now.year, now.month),
                )
                conn.commit()
            except Exception as e:
                logger.warning("Failed to create default credits: %s", e)
                conn.rollback()
            # Fall through to auto-recharge (balance = 0)
            balance = 0.0
        else:
            # Use balance_usd if available, fall back to balance_cad
            balance = float(row.get("balance_usd") or row.get("balance_cad") or 0)

        # Monthly spend cap removed by product decision: the tenant is never
        # hard-blocked on usage. Runaway-cost protection is delegated to the
        # Stripe auto-recharge loop (_auto_recharge_credits) and to per-tenant
        # Stripe subscription limits. The monthly_limit_usd column is kept on
        # ai_prepaid_credits for reporting/UI only.

        if balance < MIN_BALANCE_THRESHOLD:
            # Auto-recharge via Stripe before balance is fully drained.
            # This prevents the pattern: balance=0.05, call costs 2.00, 1.95 free.
            new_balance = _auto_recharge_credits(slug, cursor, conn)
            if new_balance > 0:
                return True, balance + new_balance
            if balance <= 0:
                return False, 0.0
            # Balance is positive but below threshold and recharge failed —
            # allow this last call but warn.
            logger.warning("Low balance %.2f below threshold for slug=%s, recharge failed", balance, slug)

        return True, balance
    except HTTPException:
        raise  # Monthly limit 402 must not be swallowed by the generic handler
    except Exception as exc:
        logger.error("_check_credits error: %s", exc)
        return False, 0.0  # Fail-closed: block AI on error to prevent free usage
    finally:
        try:
            conn.autocommit = True  # Restore default before returning to pool
        except Exception:
            pass
        if cursor:
            cursor.close()
        conn.close()


def _deduct_credits(user: ErpUser, cost_usd: float):
    """Deduct cost from tenant's prepaid credits.
    Updates both balance_usd and balance_cad for compatibility with erp_stripe.
    Allows negative balance so all usage is tracked accurately — auto-recharge
    will fire on the next _check_credits() call when balance < threshold.
    """
    if user.user_type == "super_admin":
        return

    slug = _get_tenant_slug(user)
    if not slug:
        return

    conn = db.get_conn()
    cursor = None
    try:
        cursor = conn.cursor()
        # Try updating by tenant_slug — NO GREATEST(0): allow negative balance
        # so the real consumed amount is always tracked. Auto-recharge triggers
        # on next _check_credits() when balance < MIN_BALANCE_THRESHOLD.
        cursor.execute(
            "UPDATE public.ai_prepaid_credits "
            "SET balance_usd = COALESCE(balance_usd, 0) - %s, "
            "    balance_cad = COALESCE(balance_cad, 0) - %s, "
            "    total_consumed_usd = COALESCE(total_consumed_usd, 0) + %s, "
            "    updated_at = CURRENT_TIMESTAMP "
            "WHERE tenant_slug = %s AND product_type = 'ERP' "
            "AND billing_year = EXTRACT(YEAR FROM CURRENT_DATE)::int "
            "AND billing_month = EXTRACT(MONTH FROM CURRENT_DATE)::int",
            (cost_usd, cost_usd, cost_usd, slug),
        )
        if cursor.rowcount == 0:
            # Fall back to entreprise lookup (entreprises.slug has "tenant_" prefix)
            full_slug = slug if slug.startswith("tenant_") else f"tenant_{slug}"
            cursor.execute(
                "UPDATE public.ai_prepaid_credits "
                "SET balance_usd = COALESCE(balance_usd, 0) - %s, "
                "    balance_cad = COALESCE(balance_cad, 0) - %s, "
                "    total_consumed_usd = COALESCE(total_consumed_usd, 0) + %s, "
                "    updated_at = CURRENT_TIMESTAMP "
                "WHERE entreprise_id = (SELECT id FROM public.entreprises WHERE slug = %s) "
                "AND product_type = 'ERP' "
                "AND billing_year = EXTRACT(YEAR FROM CURRENT_DATE)::int "
                "AND billing_month = EXTRACT(MONTH FROM CURRENT_DATE)::int",
                (cost_usd, cost_usd, cost_usd, full_slug),
            )
        conn.commit()
    except Exception as exc:
        logger.error("_deduct_credits error: %s", exc)
    finally:
        if cursor:
            cursor.close()
        conn.close()


def track_ai_usage(user: ErpUser, feature: str, input_tokens: int, output_tokens: int,
                    cost_usd: float, duration_ms: int, success: bool = True,
                    model: Optional[str] = None):
    """Track AI usage for billing. Pass model explicitly when using a different model than AI_MODEL."""
    conn = db.get_conn()
    cursor = None
    try:
        cursor = conn.cursor()
        _ensure_ai_tables(cursor, conn)

        slug = _get_tenant_slug(user)
        model_used = model or AI_MODEL
        cursor.execute(
            "INSERT INTO public.ai_usage_tracking (user_id, tenant_slug, feature, model, "
            "input_tokens, output_tokens, cost_usd, duration_ms, success, created_at) "
            "VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP)",
            (user.user_id, slug, feature, model_used, input_tokens, output_tokens,
             cost_usd, duration_ms, success),
        )
        conn.commit()
    except Exception as exc:
        logger.error("track_ai_usage error: %s", exc)
    finally:
        if cursor:
            cursor.close()
        conn.close()


# ============================================
# CHAT ENDPOINT
# ============================================

def _tables_in_schema(cursor) -> set:
    """Return the set of tables present in the current schema.

    Why: many optional/legacy tables (licences_rbq, cartes_ccq, logistics_*, immo_*,
    location_*, maintenance_*, subventions_*, etc.) do not exist on every tenant.
    Issuing a raw SELECT on a missing table or column floods PostgreSQL logs with
    ERROR entries even when the app catches the exception. Pre-checking existence
    once per request keeps the logs clean.
    """
    try:
        cursor.execute(
            "SELECT table_name FROM information_schema.tables "
            "WHERE table_schema = current_schema()"
        )
        rows = cursor.fetchall() or []
        out: set = set()
        for r in rows:
            if isinstance(r, dict):
                out.add(r.get("table_name"))
            else:
                try:
                    out.add(r[0])
                except Exception:
                    pass
        return {t for t in out if t}
    except Exception:
        return set()


def _build_tenant_context(user: ErpUser) -> str:
    """Build a context string with tenant ERP data for the AI assistant.
    Queries projects, employees, companies, inventory, invoices, devis, etc.
    Inspired by Streamlit's _fouiller_donnees_erp_completes().
    """
    if not user.schema:
        return ""
    conn = db.get_conn()
    cursor = None
    sections = []
    try:
        db.set_tenant(conn, user.schema)
        cursor = conn.cursor()

        # Snapshot of existing tables to skip optional ones cleanly and avoid
        # spamming PostgreSQL logs with "relation/column does not exist" errors.
        existing = _tables_in_schema(cursor)

        # 1. Stats globales projets
        if "projects" in existing:
            try:
                cursor.execute("""
                    SELECT COUNT(*) as total,
                        COUNT(CASE WHEN statut IN ('EN COURS','En cours') THEN 1 END) as en_cours,
                        COUNT(CASE WHEN statut IN ('TERMINE','Termine','Complété') THEN 1 END) as termines
                    FROM projects
                """)
                row = cursor.fetchone()
                if row:
                    d = dict(row)
                    sections.append(f"PROJETS: {d.get('total',0)} total, {d.get('en_cours',0)} en cours, {d.get('termines',0)} termines")
            except Exception:
                pass

        # 2. Projets actifs (max 20) — real columns: adresse_chantier, ville_chantier, date_prevu
        if "projects" in existing:
            try:
                cursor.execute("""
                    SELECT id, nom_projet, statut, adresse_chantier, ville_chantier, prix_estime, date_prevu
                    FROM projects
                    WHERE statut NOT IN ('ANNULE','Annule','ARCHIVE','Archive')
                    ORDER BY created_at DESC LIMIT 20
                """)
                rows = cursor.fetchall()
                if rows:
                    projets = []
                    for p in rows:
                        d = dict(p)
                        nom = d.get('nom_projet', 'Sans nom')
                        statut = d.get('statut', '?')
                        prix = d.get('prix_estime')
                        adresse = d.get('adresse_chantier') or ''
                        ville = d.get('ville_chantier') or ''
                        line = f"  - {nom} [{statut}]"
                        if prix:
                            line += f" budget: {prix}$"
                        lieu = " ".join(x for x in (adresse, ville) if x)
                        if lieu:
                            line += f" @ {lieu}"
                        projets.append(line)
                    sections.append("PROJETS ACTIFS:\n" + "\n".join(projets))
            except Exception:
                pass

        # 3. Employes actifs
        if "employees" in existing:
            try:
                cursor.execute("""
                    SELECT id, prenom, nom, poste, departement, taux_horaire, statut
                    FROM employees WHERE statut = 'ACTIF'
                    ORDER BY nom LIMIT 50
                """)
                rows = cursor.fetchall()
                if rows:
                    emps = []
                    for e in rows:
                        d = dict(e)
                        line = f"  - {d.get('prenom','')} {d.get('nom','')}"
                        if d.get('poste'):
                            line += f" ({d['poste']})"
                        if d.get('departement'):
                            line += f" dept: {d['departement']}"
                        if d.get('taux_horaire'):
                            line += f" {d['taux_horaire']}$/h"
                        emps.append(line)
                    sections.append(f"EMPLOYES ACTIFS ({len(rows)}):\n" + "\n".join(emps))
            except Exception:
                pass

        # 4. Entreprises/Clients
        if "companies" in existing:
            try:
                cursor.execute("""
                    SELECT id, nom, type_entreprise, ville, telephone, email
                    FROM companies ORDER BY created_at DESC LIMIT 30
                """)
                rows = cursor.fetchall()
                if rows:
                    clients = []
                    for c in rows:
                        d = dict(c)
                        line = f"  - {d.get('nom','?')}"
                        if d.get('type_entreprise'):
                            line += f" ({d['type_entreprise']})"
                        if d.get('ville'):
                            line += f" {d['ville']}"
                        clients.append(line)
                    sections.append(f"ENTREPRISES/CLIENTS ({len(rows)}):\n" + "\n".join(clients))
            except Exception:
                pass

        # 5. Inventaire produits
        if "produits" in existing:
            try:
                cursor.execute("""
                    SELECT id, nom, categorie, stock_disponible, stock_minimum, unite_vente, prix_unitaire
                    FROM produits WHERE active = TRUE
                    ORDER BY nom LIMIT 30
                """)
                rows = cursor.fetchall()
                if rows:
                    prods = []
                    for p in rows:
                        d = dict(p)
                        line = f"  - {d.get('nom','?')}"
                        if d.get('categorie'):
                            line += f" [{d['categorie']}]"
                        stock = d.get('stock_disponible', 0) or 0
                        seuil = d.get('stock_minimum', 0) or 0
                        unite = d.get('unite_vente', '')
                        line += f" stock: {stock} {unite}"
                        if stock <= seuil and seuil > 0:
                            line += " *** STOCK BAS ***"
                        prods.append(line)
                    sections.append(f"INVENTAIRE ({len(rows)} produits):\n" + "\n".join(prods))
            except Exception:
                pass

        # 6. Factures recentes
        if "factures" in existing:
            try:
                cursor.execute("""
                    SELECT id, numero_facture, client_nom, montant_total, statut, date_facture
                    FROM factures ORDER BY date_facture DESC LIMIT 15
                """)
                rows = cursor.fetchall()
                if rows:
                    facts = []
                    for f in rows:
                        d = dict(f)
                        line = f"  - {d.get('numero_facture','?')} {d.get('client_nom','')} {d.get('montant_total',0)}$ [{d.get('statut','?')}]"
                        facts.append(line)
                    sections.append(f"FACTURES RECENTES ({len(rows)}):\n" + "\n".join(facts))
            except Exception:
                pass

        # 7. Devis/Soumissions — real column: investissement_total (not montant_estime)
        if "devis" in existing:
            try:
                cursor.execute("""
                    SELECT id, nom_projet, investissement_total, total_avant_taxes, statut, created_at
                    FROM devis ORDER BY created_at DESC LIMIT 15
                """)
                rows = cursor.fetchall()
                if rows:
                    devis = []
                    for d_ in rows:
                        d = dict(d_)
                        montant = d.get('investissement_total') or d.get('total_avant_taxes') or 0
                        line = f"  - {d.get('nom_projet','?')} {montant}$ [{d.get('statut','?')}]"
                        devis.append(line)
                    sections.append(f"DEVIS/SOUMISSIONS ({len(rows)}):\n" + "\n".join(devis))
            except Exception:
                pass

        # 8. Bons de travail — real source: formulaires WHERE type_formulaire = 'BON_TRAVAIL'
        if "formulaires" in existing:
            try:
                cursor.execute("""
                    SELECT id, numero_document, nom, statut, priorite, date_creation, montant_total
                    FROM formulaires
                    WHERE type_formulaire = 'BON_TRAVAIL'
                    ORDER BY created_at DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    bts = []
                    for b in rows:
                        d = dict(b)
                        titre = d.get('nom') or d.get('numero_document') or '?'
                        line = f"  - {titre} [{d.get('statut','?')}] priorite: {d.get('priorite','?')}"
                        if d.get('montant_total'):
                            line += f" {d['montant_total']}$"
                        bts.append(line)
                    sections.append(f"BONS DE TRAVAIL ({len(rows)}):\n" + "\n".join(bts))
            except Exception:
                pass

        # 9. Mouvements stock recents
        if "mouvements_stock" in existing:
            try:
                cursor.execute("""
                    SELECT m.type_mouvement, m.quantite, m.reference_document, m.created_at, p.nom as produit_nom
                    FROM mouvements_stock m
                    LEFT JOIN produits p ON m.produit_id = p.id
                    ORDER BY m.created_at DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    mvts = []
                    for m in rows:
                        d = dict(m)
                        line = f"  - {d.get('type_mouvement','')} {d.get('produit_nom','?')} qte: {d.get('quantite',0)}"
                        if d.get('reference_document'):
                            line += f" ref: {d['reference_document']}"
                        mvts.append(line)
                    sections.append("MOUVEMENTS STOCK RECENTS:\n" + "\n".join(mvts))
            except Exception:
                pass

        # 10. Opportunites CRM
        if "opportunities" in existing:
            try:
                cursor.execute("""
                    SELECT id, nom, montant_estime, probabilite, statut, source, date_cloture_prevue
                    FROM opportunities ORDER BY created_at DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    opps = []
                    for o in rows:
                        d = dict(o)
                        line = f"  - {d.get('nom','?')} [{d.get('statut','?')}] {d.get('montant_estime',0)}$ prob: {d.get('probabilite',0)}%"
                        opps.append(line)
                    sections.append(f"OPPORTUNITES CRM ({len(rows)}):\n" + "\n".join(opps))
            except Exception:
                pass

        # 11. Licences RBQ (optionnel — legacy module, absent sur la plupart des tenants)
        if "licences_rbq" in existing:
            try:
                cursor.execute("""
                    SELECT id, numero_licence, titulaire, categorie, statut, date_expiration
                    FROM licences_rbq ORDER BY date_expiration DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    lics = []
                    for l in rows:
                        d = dict(l)
                        line = f"  - {d.get('numero_licence','?')} {d.get('titulaire','')} [{d.get('statut','?')}] exp: {d.get('date_expiration','?')}"
                        lics.append(line)
                    sections.append(f"LICENCES RBQ ({len(rows)}):\n" + "\n".join(lics))
            except Exception:
                pass

        # 12. Cartes CCQ (optionnel — legacy)
        if "cartes_ccq" in existing:
            try:
                cursor.execute("""
                    SELECT id, numero_carte, nom_complet, metier, statut, date_expiration
                    FROM cartes_ccq ORDER BY date_expiration DESC LIMIT 15
                """)
                rows = cursor.fetchall()
                if rows:
                    cartes = []
                    for c in rows:
                        d = dict(c)
                        line = f"  - {d.get('numero_carte','?')} {d.get('nom_complet','')} ({d.get('metier','?')}) [{d.get('statut','?')}]"
                        cartes.append(line)
                    sections.append(f"CARTES CCQ ({len(rows)}):\n" + "\n".join(cartes))
            except Exception:
                pass

        # 13. Inspections chantier (optionnel — colonne project_id pas projet_id)
        if "inspections_chantier" in existing:
            try:
                cursor.execute("""
                    SELECT id, type_inspection, project_id, statut, date_inspection, score_conformite
                    FROM inspections_chantier ORDER BY date_inspection DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    insps = []
                    for i in rows:
                        d = dict(i)
                        line = f"  - {d.get('type_inspection','?')} [{d.get('statut','?')}] score: {d.get('score_conformite','?')}"
                        insps.append(line)
                    sections.append(f"INSPECTIONS CHANTIER ({len(rows)}):\n" + "\n".join(insps))
            except Exception:
                pass

        # 14. Logistique - Livraisons (optionnel)
        if "logistics_deliveries" in existing:
            try:
                cursor.execute("""
                    SELECT id, reference, statut, date_prevue
                    FROM logistics_deliveries ORDER BY created_at DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    livs = []
                    for l in rows:
                        d = dict(l)
                        line = f"  - {d.get('reference','?')} [{d.get('statut','?')}] prevue: {d.get('date_prevue','?')}"
                        livs.append(line)
                    sections.append(f"LIVRAISONS ({len(rows)}):\n" + "\n".join(livs))
            except Exception:
                pass

        # 15. Logistique - Equipements (optionnel)
        if "logistics_equipment" in existing:
            try:
                cursor.execute("""
                    SELECT id, code, nom, categorie, statut, localisation_actuelle
                    FROM logistics_equipment ORDER BY nom LIMIT 15
                """)
                rows = cursor.fetchall()
                if rows:
                    equips = []
                    for e in rows:
                        d = dict(e)
                        line = f"  - {d.get('code','')} {d.get('nom','?')} [{d.get('statut','?')}] {d.get('localisation_actuelle','')}"
                        equips.append(line)
                    sections.append(f"EQUIPEMENTS ({len(rows)}):\n" + "\n".join(equips))
            except Exception:
                pass

        # 16. Logistique - Vehicules (optionnel)
        if "logistics_vehicles" in existing:
            try:
                cursor.execute("""
                    SELECT id, immatriculation, marque, modele, statut, type_vehicule
                    FROM logistics_vehicles ORDER BY immatriculation LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    vehs = []
                    for v in rows:
                        d = dict(v)
                        line = f"  - {d.get('immatriculation','?')} {d.get('marque','')} {d.get('modele','')} [{d.get('statut','?')}]"
                        vehs.append(line)
                    sections.append(f"VEHICULES ({len(rows)}):\n" + "\n".join(vehs))
            except Exception:
                pass

        # 17. Immobilier - Terrains (optionnel)
        if "immo_terrains" in existing:
            try:
                cursor.execute("""
                    SELECT id, numero_dossier, adresse, ville, statut, superficie_m2, prix_demande
                    FROM immo_terrains ORDER BY created_at DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    terrains = []
                    for t in rows:
                        d = dict(t)
                        line = f"  - {d.get('numero_dossier','?')} {d.get('adresse','')} {d.get('ville','')} [{d.get('statut','?')}]"
                        if d.get('prix_demande'):
                            line += f" {d['prix_demande']}$"
                        terrains.append(line)
                    sections.append(f"TERRAINS IMMOBILIER ({len(rows)}):\n" + "\n".join(terrains))
            except Exception:
                pass

        # 18. Immobilier - Projets (optionnel)
        if "immo_projets" in existing:
            try:
                cursor.execute("""
                    SELECT id, numero_projet, nom_projet, statut, budget_total
                    FROM immo_projets ORDER BY created_at DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    ipjs = []
                    for p in rows:
                        d = dict(p)
                        line = f"  - {d.get('numero_projet','?')} {d.get('nom_projet','')} [{d.get('statut','?')}] budget: {d.get('budget_total','')}$"
                        ipjs.append(line)
                    sections.append(f"PROJETS IMMOBILIER ({len(rows)}):\n" + "\n".join(ipjs))
            except Exception:
                pass

        # 19. Location - Contrats (optionnel)
        if "location_contrats" in existing:
            try:
                cursor.execute("""
                    SELECT id, numero_contrat, statut, date_debut, date_fin_prevue, date_fin_reelle, montant_total
                    FROM location_contrats ORDER BY created_at DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    locs = []
                    for l in rows:
                        d = dict(l)
                        date_fin = d.get('date_fin_reelle') or d.get('date_fin_prevue') or '?'
                        line = f"  - {d.get('numero_contrat','?')} [{d.get('statut','?')}] du {d.get('date_debut','?')} au {date_fin}"
                        if d.get('montant_total'):
                            line += f" {d['montant_total']}$"
                        locs.append(line)
                    sections.append(f"CONTRATS LOCATION ({len(rows)}):\n" + "\n".join(locs))
            except Exception:
                pass

        # 20. Maintenance (optionnel)
        if "maintenance_demandes" in existing:
            try:
                cursor.execute("""
                    SELECT id, titre, statut, priorite, type_maintenance, date_demande
                    FROM maintenance_demandes ORDER BY created_at DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    mains = []
                    for m in rows:
                        d = dict(m)
                        line = f"  - {d.get('titre','?')} [{d.get('statut','?')}] {d.get('type_maintenance','')} priorite: {d.get('priorite','?')}"
                        mains.append(line)
                    sections.append(f"MAINTENANCE ({len(rows)}):\n" + "\n".join(mains))
            except Exception:
                pass

        # 21. Subventions (optionnel — colonne programme_id pas programme)
        if "subventions_demandes" in existing:
            try:
                cursor.execute("""
                    SELECT id, programme_id, statut, montant, date_soumission
                    FROM subventions_demandes ORDER BY created_at DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    subs = []
                    for s in rows:
                        d = dict(s)
                        line = f"  - programme #{d.get('programme_id','?')} [{d.get('statut','?')}] {d.get('montant','')}$"
                        subs.append(line)
                    sections.append(f"SUBVENTIONS ({len(rows)}):\n" + "\n".join(subs))
            except Exception:
                pass

        # 22. Plan comptable
        if "plan_comptable" in existing:
            try:
                cursor.execute("""
                    SELECT id, code, nom, type, classe FROM plan_comptable
                    WHERE actif = TRUE ORDER BY code LIMIT 30
                """)
                rows = cursor.fetchall()
                if rows:
                    comptes = []
                    for c in rows:
                        d = dict(c)
                        comptes.append(f"  - {d.get('code','?')} {d.get('nom','')} ({d.get('type','')}/{d.get('classe','')})")
                    sections.append(f"PLAN COMPTABLE ({len(rows)} comptes):\n" + "\n".join(comptes))
            except Exception:
                pass

        # 23. Ecritures journal recentes
        if "journal_entries" in existing:
            try:
                cursor.execute("""
                    SELECT id, numero_ecriture, date_comptable, type_journal, description, statut
                    FROM journal_entries ORDER BY date_comptable DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    ecrs = []
                    for e in rows:
                        d = dict(e)
                        ecrs.append(f"  - {d.get('numero_ecriture','?')} {d.get('date_comptable','?')} [{d.get('type_journal','')}] {d.get('description','')}")
                    sections.append(f"ECRITURES JOURNAL ({len(rows)}):\n" + "\n".join(ecrs))
            except Exception:
                pass

        # 24. Bons de commande — real column: numero_bon (not numero_bc)
        if "bons_commande" in existing:
            try:
                cursor.execute("""
                    SELECT id, numero_bon, statut, montant_total, date_commande
                    FROM bons_commande ORDER BY created_at DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    bcs = []
                    for b in rows:
                        d = dict(b)
                        bcs.append(f"  - {d.get('numero_bon','?')} [{d.get('statut','?')}] {d.get('montant_total','')}$ date: {d.get('date_commande','?')}")
                    sections.append(f"BONS DE COMMANDE ({len(rows)}):\n" + "\n".join(bcs))
            except Exception:
                pass

        # 25. Notes chantier (optionnel — colonne project_id pas projet_id)
        if "project_notes" in existing:
            try:
                cursor.execute("""
                    SELECT id, titre, categorie, project_id, created_at
                    FROM project_notes ORDER BY created_at DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    notes = []
                    for n in rows:
                        d = dict(n)
                        notes.append(f"  - {d.get('titre','?')} [{d.get('categorie','')}] projet: {d.get('project_id','')}")
                    sections.append(f"NOTES CHANTIER ({len(rows)}):\n" + "\n".join(notes))
            except Exception:
                pass

        # 26. Dossiers — real table: dossiers (not dossier_projets), colonne: titre (not nom)
        if "dossiers" in existing:
            try:
                cursor.execute("""
                    SELECT id, numero_dossier, titre, statut
                    FROM dossiers ORDER BY created_at DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    doss = []
                    for d_ in rows:
                        d = dict(d_)
                        doss.append(f"  - {d.get('numero_dossier','?')} {d.get('titre','')} [{d.get('statut','?')}]")
                    sections.append(f"DOSSIERS ({len(rows)}):\n" + "\n".join(doss))
            except Exception:
                pass

        # 27. Calendrier — real columns: start_date, end_date (not start_time, end_time)
        if "calendar_events" in existing:
            try:
                cursor.execute("""
                    SELECT id, title, start_date, end_date
                    FROM calendar_events WHERE start_date >= CURRENT_DATE
                    ORDER BY start_date LIMIT 15
                """)
                rows = cursor.fetchall()
                if rows:
                    evts = []
                    for e in rows:
                        d = dict(e)
                        evts.append(f"  - {d.get('title','?')} du {d.get('start_date','?')} au {d.get('end_date','?')}")
                    sections.append(f"CALENDRIER ({len(rows)} evenements):\n" + "\n".join(evts))
            except Exception:
                pass

        # 28. Paiements recus (optionnel — table legacy)
        if "paiements_recus" in existing:
            try:
                cursor.execute("""
                    SELECT id, facture_id, date_paiement, montant, mode_paiement, reference
                    FROM paiements_recus ORDER BY date_paiement DESC LIMIT 10
                """)
                rows = cursor.fetchall()
                if rows:
                    paie = []
                    for p in rows:
                        d = dict(p)
                        paie.append(f"  - Facture #{d.get('facture_id','?')} {d.get('montant',0)}$ [{d.get('mode_paiement','')}] {d.get('date_paiement','?')}")
                    sections.append(f"PAIEMENTS RECUS ({len(rows)}):\n" + "\n".join(paie))
            except Exception:
                pass

    except Exception as exc:
        logger.warning("_build_tenant_context error: %s", exc)
    finally:
        if cursor:
            cursor.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()

    if not sections:
        return ""

    context = "DONNEES TEMPS REEL DE L'ERP DE L'ENTREPRISE:\n\n" + "\n\n".join(sections)
    # Truncate if too long (keep under ~40K chars to leave room for response)
    if len(context) > 40000:
        context = context[:40000] + "\n\n... [DONNEES TRONQUEES]"
    return context


def _detect_intention(message: str) -> str:
    """Detect user intention from keywords, inspired by Streamlit's _detecter_intention_conversation."""
    msg = message.lower()

    # Suppression (priority)
    if any(k in msg for k in ['supprime', 'supprimer', 'efface', 'effacer', 'delete', 'enleve', 'retirer']):
        return 'supprimer_element'

    # Opinion
    if any(k in msg for k in ['penses quoi', 'ton avis', 'que penses-tu', "qu'en penses-tu", 'opinion']):
        return 'demande_opinion'

    # Status
    if any(k in msg for k in ['comment va', 'etat', 'situation', 'avancement', 'ou en est']):
        return 'demande_status'

    # Problemes
    if any(k in msg for k in ['probleme', 'souci', 'risque', 'attention', 'inquiet', 'retard', 'urgent']):
        return 'detection_problemes'

    # Conseils
    if any(k in msg for k in ['conseil', 'recommande', 'suggere', 'que faire', 'devrais', 'comment faire']):
        return 'demande_conseil'

    # Finance
    if any(k in msg for k in ['budget', 'cout', 'prix', 'rentable', 'argent', 'marge', 'profit', 'facture', 'paiement']):
        return 'analyse_financiere'

    # Equipe
    if any(k in msg for k in ['equipe', 'employe', 'qui peut', 'competence', 'disponible', 'ressource']):
        return 'gestion_equipe'

    # Inventaire
    if any(k in msg for k in ['materiel', 'stock', 'inventaire', 'manque', 'commande', 'produit']):
        return 'gestion_materiel'

    # Planning
    if any(k in msg for k in ['planning', 'delai', 'echeance', 'calendrier', 'livraison', 'quand']):
        return 'gestion_temps'

    # Electricite
    if any(k in msg for k in ['electri', 'filage', 'panneau electri', 'circuit', 'ampere', 'volt', 'disjoncteur', 'prise', 'csa c22']):
        return 'expert_electricite'

    # Plomberie
    if any(k in msg for k in ['plomberi', 'tuyau', 'drain', 'robinet', 'sanitaire', 'egout', 'cuivre', 'pex', 'siphon']):
        return 'expert_plomberie'

    # Structure/Beton
    if any(k in msg for k in ['beton', 'fondation', 'coffrage', 'dalle', 'armature', 'mpa', 'structure', 'portee', 'poutre']):
        return 'expert_structure'

    # Toiture
    if any(k in msg for k in ['toiture', 'toit', 'bardeau', 'membrane', 'solin', 'pente', 'gouttiere']):
        return 'expert_toiture'

    # Isolation
    if any(k in msg for k in ['isolation', 'valeur r', 'pare-vapeur', 'pare-air', 'laine', 'styromousse', 'enveloppe']):
        return 'expert_isolation'

    # Securite SST
    if any(k in msg for k in ['securite', 'sst', 'cnesst', 'epi', 'casque', 'harnais', 'echafaud', 'accident']):
        return 'expert_securite'

    # Juridique/RBQ
    if any(k in msg for k in ['rbq', 'licence', 'permis', 'legal', 'juridique', 'code du batiment', 'hypotheque legale', 'contrat']):
        return 'expert_juridique'

    # Estimation
    if any(k in msg for k in ['estim', 'soumission', 'devis', 'evaluer', 'combien coute', 'prix de']):
        return 'expert_estimation'

    # Comptabilite
    if any(k in msg for k in ['comptab', 'tps', 'tvq', 'das', 'retenue', 'paie', 'ccq', 'ecriture']):
        return 'expert_comptabilite'

    # Creation
    if any(k in msg for k in ['creer', 'cree', 'nouveau', 'nouvelle', 'ajouter', 'ajoute']):
        return 'creer_element'

    # Modification
    if any(k in msg for k in ['modifier', 'modifie', 'changer', 'change', 'mettre a jour']):
        return 'modifier_element'

    return 'question_generale'


_INTENTION_INSTRUCTIONS = {
    'demande_opinion': "L'utilisateur te demande ton avis. Analyse les donnees reelles et donne ton opinion d'expert, sois franc mais constructif.",
    'demande_status': "Donne un point de situation naturel, comme avec un collegue. Mentionne les elements importants de facon conversationnelle.",
    'detection_problemes': "Analyse attentivement les donnees pour identifier les vrais problemes ou risques. Sois le collegue vigilant qui repere les soucis.",
    'demande_conseil': "Donne des conseils pratiques et actionnables bases sur les donnees reelles. Sois l'expert qui guide.",
    'analyse_financiere': "Concentre-toi sur l'analyse financiere: budgets, couts, rentabilite, marges. Utilise les donnees reelles.",
    'gestion_equipe': "Reponds en te concentrant sur l'equipe: disponibilite, competences, assignation. Utilise les donnees des employes.",
    'gestion_materiel': "Concentre-toi sur l'inventaire et les materiaux. Identifie les stocks bas ou les manques.",
    'gestion_temps': "Reponds sur les delais, le planning et les echeances. Identifie les retards potentiels.",
    'expert_electricite': "Reponds comme un maitre electricien du Quebec. Reference les normes CSA C22.1 et le Code electrique Quebec.",
    'expert_plomberie': "Reponds comme un maitre plombier du Quebec. Reference le Code de plomberie du Quebec.",
    'expert_structure': "Reponds comme un ingenieur en structure. Reference les normes CSA A23.3 et le CNB.",
    'expert_toiture': "Reponds comme un expert en toiture du Quebec. Couvre bardeaux, membranes, pentes, solins.",
    'expert_isolation': "Reponds comme un expert en enveloppe du batiment. Reference le Code energetique du Quebec et les valeurs R.",
    'expert_securite': "Reponds comme un expert SST. Reference les normes CNESST et les reglements de securite sur les chantiers.",
    'expert_juridique': "Reponds comme un conseiller juridique en construction. Reference la Loi sur le batiment, le Code civil et les reglements RBQ.",
    'expert_estimation': "Reponds comme un estimateur professionnel. Donne des couts realistes du marche Quebec avec TPS 5% et TVQ 9.975%.",
    'expert_comptabilite': "Reponds comme un comptable specialise construction Quebec. Couvre DAS, TPS/TVQ, retenues a la source, paie CCQ.",
}


@router.post("/chat")
async def ai_chat(body: ChatMessage, user: ErpUser = Depends(get_current_user)):
    """Send a message to the AI expert and get a response."""
    if not _anthropic_client:
        raise HTTPException(status_code=503, detail="Service IA non disponible")

    # Guard check
    allowed, error_msg = check_ai_guard(user)
    if not allowed:
        raise HTTPException(status_code=403, detail=error_msg or "Acces IA refuse")

    # Credit check
    credits_ok, balance = _check_credits(user)
    if not credits_ok:
        raise HTTPException(
            status_code=402,
            detail="Credits IA epuises. Veuillez recharger votre solde pour continuer.",
        )

    profile = AI_PROFILES.get(body.profile, AI_PROFILES["general"])

    # Detect intention and build adaptive prompt
    intention = _detect_intention(body.message)
    system_prompt = _today_prompt_line() + "\n\n" + profile["system"]

    # Add intention-specific instruction
    instruction = _INTENTION_INSTRUCTIONS.get(intention)
    if instruction:
        system_prompt += f"\n\nINTENTION DETECTEE: {intention}\nINSTRUCTION: {instruction}"

    # Add tenant database context
    tenant_context = _build_tenant_context(user)
    if tenant_context:
        system_prompt += f"\n\n{tenant_context}"

    system_prompt += "\n\nREGLE: Base-toi sur les donnees ERP reelles quand disponibles. Ne jamais inventer de donnees. Reponds en francais du Quebec, comme un collegue expert sympathique."
    system_prompt += "\n\nTu as acces a la base de donnees du tenant via les outils recherche_bd et executer_action. Utilise-les pour repondre aux questions avec des donnees reelles. Pour les questions sur les pointages, heures, employes, factures, projets, etc., fais TOUJOURS une requete BD au lieu de dire que tu n'as pas acces. Pour les actions de creation/modification/suppression, utilise executer_action."
    if body.context:
        system_prompt += f"\n\nContexte additionnel:\n{body.context}"

    # Load conversation history if conversation_id provided
    conversation_id = body.conversation_id
    existing_messages: list[dict] = []
    conversation_loaded = True
    if conversation_id and user.schema:
        existing_messages = _load_conversation_messages(user, conversation_id)
        if not existing_messages and conversation_id:
            # Could be empty conversation or load failure — flag to avoid overwriting
            conversation_loaded = False

    try:
        start = time_module.time()
        # Build messages: history + new user message
        messages = []
        for m in existing_messages:
            messages.append({"role": m["role"], "content": m["content"]})
        messages.append({"role": "user", "content": body.message})
        total_in = 0
        total_out = 0

        # Boucle tool-use: Claude peut appeler des outils, on execute et on renvoie les resultats
        max_iterations = 5  # Securite: max 5 appels d'outils par requete
        for _ in range(max_iterations):
            response = _call_claude(
                model=AI_MODEL,
                max_tokens=AI_MAX_TOKENS,
                system=system_prompt,
                messages=messages,
                tools=_AI_TOOLS if user.schema else None,
            )
            total_in += response.usage.input_tokens
            total_out += response.usage.output_tokens

            # Si Claude a termine (pas de tool_use), on sort de la boucle
            if response.stop_reason != "tool_use":
                break

            # Traiter les appels d'outils
            tool_results = []
            for block in response.content:
                if block.type == "tool_use":
                    tool_name = block.name
                    tool_input = block.input
                    if tool_name == "recherche_bd":
                        result = _execute_tenant_select(user.schema, tool_input.get("sql", ""))
                    elif tool_name == "executer_action":
                        result = _execute_tenant_action(
                            user.schema, tool_input.get("sql", ""),
                            tool_input.get("description", "Action IA"), user
                        )
                    else:
                        result = {"error": f"Outil inconnu: {tool_name}"}
                    tool_results.append({
                        "type": "tool_result",
                        "tool_use_id": block.id,
                        "content": str(result),
                    })

            # Ajouter la reponse de Claude (avec tool_use) et les resultats des outils
            messages.append({"role": "assistant", "content": response.content})
            messages.append({"role": "user", "content": tool_results})

        elapsed = time_module.time() - start

        # Extraire le texte final de la reponse
        content = ""
        for block in response.content:
            if hasattr(block, "text"):
                content += block.text

        # Si la boucle a epuise les iterations sans reponse texte
        if not content and response.stop_reason == "tool_use":
            content = "J'ai effectue plusieurs recherches mais la requete est trop complexe. Peux-tu reformuler ta question?"

        total_tokens = total_in + total_out
        duration_ms = int(elapsed * 1000)
        # Approximate cost (Claude Sonnet pricing: input $0.003/1K, output $0.015/1K)
        cost = (total_in * 0.003 + total_out * 0.015) / 1000 * 1.30  # 30% markup

        # Track usage and deduct credits
        track_ai_usage(user, f"chat_{body.profile}", total_in, total_out,
                        cost, duration_ms, success=True)
        _deduct_credits(user, cost)

        # Auto-save conversation (skip if history load failed to avoid overwriting)
        if user.schema and (conversation_loaded or not conversation_id):
            conversation_id = _save_conversation_after_chat(
                user, conversation_id, existing_messages,
                body.message, content, body.profile,
            )

        return {
            "response": content,
            "profile": profile["name"],
            "input_tokens": total_in,
            "output_tokens": total_out,
            "tokens_used": total_tokens,
            "cost_usd": round(cost, 6),
            "elapsed_seconds": round(elapsed, 2),
            "credit_balance": round(balance - cost, 4),
            "conversation_id": conversation_id,
        }
    except HTTPException:
        raise
    except anthropic.APIError as exc:
        logger.error("ai_chat API error: %s", exc)
        status = getattr(exc, "status_code", 0)
        exc_str = str(exc).lower()
        if status == 413 or (status == 400 and "too_large" in exc_str):
            raise HTTPException(
                status_code=413,
                detail="Conversation trop longue pour l'IA. Essayez de demarrer une nouvelle conversation.",
            )
        if status == 529 or "overloaded" in exc_str:
            raise HTTPException(status_code=503, detail="Le service IA est temporairement surcharge. Veuillez reessayer dans quelques secondes.")
        if status == 429 or "rate_limit" in exc_str:
            raise HTTPException(status_code=429, detail="Trop de demandes vers l'IA. Attendez quelques instants avant de reessayer.")
        raise HTTPException(status_code=503, detail="Service IA temporairement indisponible. Veuillez reessayer dans quelques instants.")
    except Exception as exc:
        logger.error("ai_chat error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur interne du service IA. Veuillez reessayer.")


# ============================================
# CONVERSATION PERSISTENCE (Assistant IA)
# ============================================


_conversations_last_updated_ensured: set = set()


def _ensure_conversations_last_updated_at(cur, schema: str = "") -> None:
    """Defensive migration: ensure `conversations.last_updated_at` is present
    and non-blocking (DEFAULT NOW(), nullable).

    Sur tenants legacy, la colonne existe parfois en NOT NULL sans default,
    ce qui fait planter les INSERT qui omettent la colonne avec
    "null value in column last_updated_at violates not-null constraint"
    (vu 9x/72h sur POST /devis/conversations).

    Memoized par schema pour eviter le lock ACCESS EXCLUSIVE repete sur
    chaque appel — pattern identique a `_email_tables_ensured_for`
    (emails.py:107).
    """
    global _conversations_last_updated_ensured
    if schema and schema in _conversations_last_updated_ensured:
        return
    try:
        cur.execute(
            "ALTER TABLE conversations ADD COLUMN IF NOT EXISTS last_updated_at TIMESTAMP DEFAULT NOW()"
        )
        cur.execute(
            "ALTER TABLE conversations ALTER COLUMN last_updated_at SET DEFAULT NOW()"
        )
        cur.execute(
            "ALTER TABLE conversations ALTER COLUMN last_updated_at DROP NOT NULL"
        )
    except Exception as alter_exc:
        logger.warning("conversations.last_updated_at defensive ALTER skipped: %s", alter_exc)
        return
    if schema:
        _conversations_last_updated_ensured.add(schema)


def _ensure_conversations_table(cur):
    """Create conversations table if it doesn't exist (defensive for new tenants).
    Also adds columns that may be missing if the table was created by mobile backend."""
    cur.execute("""
        CREATE TABLE IF NOT EXISTS conversations (
            id SERIAL PRIMARY KEY,
            name TEXT,
            devis_id INTEGER,
            project_id INTEGER,
            user_id INTEGER,
            subject TEXT,
            status TEXT DEFAULT 'active',
            messages TEXT,
            messages_json JSONB,
            metadata TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            last_updated_at TIMESTAMP DEFAULT NOW()
        )
    """)
    # Add columns that may be missing if table was created by mobile (6-column schema)
    for col, typ in [
        ("user_id", "INTEGER"),
        ("subject", "TEXT"),
        ("messages_json", "JSONB"),
        ("updated_at", "TIMESTAMP DEFAULT CURRENT_TIMESTAMP"),
        ("status", "TEXT DEFAULT 'active'"),
        ("devis_id", "INTEGER"),
        ("project_id", "INTEGER"),
    ]:
        try:
            cur.execute(f"ALTER TABLE conversations ADD COLUMN IF NOT EXISTS {col} {typ}")
        except Exception:
            pass


def _load_conversation_messages(user: ErpUser, conversation_id: int) -> list[dict]:
    """Load existing messages from a conversation."""
    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        cur = conn.cursor()
        _ensure_conversations_table(cur)
        conn.commit()
        cur.execute(
            "SELECT messages_json, messages FROM conversations "
            "WHERE id = %s AND user_id = %s AND subject = 'assistant_ia'",
            (conversation_id, user.user_id),
        )
        row = cur.fetchone()
        if not row:
            return []
        r = dict(row)
        msgs = []
        if r.get("messages_json"):
            msgs = r["messages_json"] if isinstance(r["messages_json"], list) else json.loads(r["messages_json"])
        elif r.get("messages"):
            try:
                msgs = json.loads(r["messages"])
            except (json.JSONDecodeError, TypeError):
                pass
        # Only return user/assistant messages (no system)
        return [m for m in msgs if m.get("role") in ("user", "assistant")]
    except Exception as exc:
        logger.warning("_load_conversation_messages error: %s", exc)
        return []
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


def _save_conversation_after_chat(
    user: ErpUser, conversation_id: Optional[int],
    existing_messages: list[dict], user_message: str,
    assistant_response: str, profile: str,
) -> Optional[int]:
    """Save or update conversation after a chat exchange. Returns conversation_id."""
    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        cur = conn.cursor()
        _ensure_conversations_table(cur)
        # Defensif: meme bug que devis.py:save_conversation. Sur tenants
        # legacy, conversations.last_updated_at est NOT NULL sans default.
        # Memoized par schema, no-op sur les calls subsequents.
        _ensure_conversations_last_updated_at(cur, user.schema or "")

        all_msgs = list(existing_messages)
        all_msgs.append({"role": "user", "content": user_message})
        all_msgs.append({"role": "assistant", "content": assistant_response})
        msgs_json = json.dumps(all_msgs, ensure_ascii=False)

        if conversation_id:
            cur.execute(
                "UPDATE conversations SET messages_json = %s::jsonb, messages = %s, "
                "updated_at = NOW(), last_updated_at = NOW() WHERE id = %s AND user_id = %s",
                (msgs_json, msgs_json, conversation_id, user.user_id),
            )
        else:
            # Generate name from first user message
            name_words = user_message.split()[:6]
            conv_name = " ".join(name_words)
            if len(user_message.split()) > 6:
                conv_name += "..."
            conv_name = conv_name[:80]

            # last_updated_at inclus explicitement avec NOW() pour les
            # tenants legacy ou la colonne est NOT NULL sans default.
            cur.execute(
                "INSERT INTO conversations "
                "(name, user_id, subject, status, messages_json, messages, created_at, updated_at, last_updated_at) "
                "VALUES (%s, %s, 'assistant_ia', 'active', %s::jsonb, %s, NOW(), NOW(), NOW()) RETURNING id",
                (conv_name, user.user_id, msgs_json, msgs_json),
            )
            row = cur.fetchone()
            if row:
                conversation_id = row["id"]

        conn.commit()
        return conversation_id
    except Exception as exc:
        logger.error("_save_conversation_after_chat error: %s", exc)
        try:
            conn.rollback()
        except Exception:
            pass
        return conversation_id
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.get("/conversations")
async def list_ai_conversations(user: ErpUser = Depends(get_current_user)):
    """List user's Assistant IA conversations."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        cur = conn.cursor()
        _ensure_conversations_table(cur)
        conn.commit()
        cur.execute(
            "SELECT id, name, created_at, updated_at, messages_json, messages "
            "FROM conversations WHERE user_id = %s AND subject = 'assistant_ia' "
            "ORDER BY updated_at DESC NULLS LAST LIMIT 30",
            (user.user_id,),
        )
        items = []
        for row in cur.fetchall():
            r = dict(row)
            # Count user/assistant messages
            msgs = []
            raw = r.get("messages_json") or r.get("messages")
            if raw:
                try:
                    msgs = raw if isinstance(raw, list) else json.loads(raw)
                except (json.JSONDecodeError, TypeError):
                    pass
            msg_count = len([m for m in msgs if m.get("role") in ("user", "assistant")])
            items.append({
                "id": r["id"],
                "name": r.get("name", ""),
                "messageCount": msg_count,
                "createdAt": str(r.get("created_at", "")),
                "updatedAt": str(r.get("updated_at", "")),
            })
        return {"items": items}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("list_ai_conversations error: %s", exc)
        return {"items": []}
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.get("/conversations/{conv_id}")
async def get_ai_conversation(conv_id: int, user: ErpUser = Depends(get_current_user)):
    """Get a single Assistant IA conversation with messages."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        cur = conn.cursor()
        _ensure_conversations_table(cur)
        conn.commit()
        cur.execute(
            "SELECT id, name, created_at, updated_at, messages_json, messages "
            "FROM conversations WHERE id = %s AND user_id = %s AND subject = 'assistant_ia'",
            (conv_id, user.user_id),
        )
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Conversation non trouvee")
        r = dict(row)
        msgs = []
        if r.get("messages_json"):
            msgs = r["messages_json"] if isinstance(r["messages_json"], list) else json.loads(r["messages_json"])
        elif r.get("messages"):
            try:
                msgs = json.loads(r["messages"])
            except (json.JSONDecodeError, TypeError):
                pass
        display_msgs = [m for m in msgs if m.get("role") in ("user", "assistant")]
        return {
            "id": r["id"],
            "name": r.get("name", ""),
            "messages": display_msgs,
            "createdAt": str(r.get("created_at", "")),
            "updatedAt": str(r.get("updated_at", "")),
        }
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("get_ai_conversation error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


@router.delete("/conversations/{conv_id}")
async def delete_ai_conversation(conv_id: int, user: ErpUser = Depends(get_current_user)):
    """Delete an Assistant IA conversation."""
    if not user.schema:
        raise HTTPException(status_code=400, detail="Contexte tenant manquant")
    conn = db.get_conn()
    cur = None
    try:
        db.set_tenant(conn, user.schema)
        cur = conn.cursor()
        cur.execute(
            "DELETE FROM conversations WHERE id = %s AND user_id = %s AND subject = 'assistant_ia'",
            (conv_id, user.user_id),
        )
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="Conversation non trouvee")
        try:
            conn.commit()
        except Exception:
            pass
        return {"message": "Conversation supprimee"}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("delete_ai_conversation error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur")
    finally:
        if cur:
            cur.close()
        try:
            db.reset_tenant(conn)
        except Exception:
            pass
        conn.close()


# ============================================
# PROFILES LIST
# ============================================

@router.get("/profiles")
async def list_ai_profiles(user: ErpUser = Depends(get_current_user)):
    """List available AI expert profiles."""
    return {
        "profiles": [
            {"id": k, "name": v["name"]}
            for k, v in AI_PROFILES.items()
        ]
    }


# ============================================
# USAGE STATS
# ============================================

@router.get("/usage")
async def get_ai_usage_stats(
    user: ErpUser = Depends(get_current_user),
    period_days: int = Query(30, ge=1, le=365),
):
    """Get AI usage statistics for the current tenant."""
    slug = _get_tenant_slug(user)
    conn = db.get_conn()
    cursor = None
    try:
        cursor = conn.cursor()

        cursor.execute(
            "SELECT EXISTS (SELECT FROM information_schema.tables "
            "WHERE table_schema = 'public' AND table_name = 'ai_usage_tracking')"
        )
        if not cursor.fetchone().get("exists", False):
            return {"total_requests": 0, "total_tokens": 0, "total_cost": 0, "by_feature": []}

        # Filter by tenant_slug — user_id is not unique cross-tenant.
        # Super-admins see the global totals (no filter).
        if user.user_type == "super_admin":
            where_clause = "created_at >= CURRENT_DATE - make_interval(days => %s)"
            params: tuple = (period_days,)
        else:
            where_clause = "tenant_slug = %s AND created_at >= CURRENT_DATE - make_interval(days => %s)"
            params = (slug, period_days)

        cursor.execute(
            "SELECT COUNT(*) as total_requests, "
            "COALESCE(SUM(COALESCE(input_tokens, tokens_input, 0) + COALESCE(output_tokens, tokens_output, 0)), 0) as total_tokens, "
            "COALESCE(SUM(COALESCE(cost_usd, estimated_cost_usd, 0)), 0) as total_cost "
            "FROM public.ai_usage_tracking "
            f"WHERE {where_clause}",
            params,
        )
        totals = cursor.fetchone()

        cursor.execute(
            "SELECT feature, COUNT(*) as requests, "
            "COALESCE(SUM(COALESCE(input_tokens, tokens_input, 0) + COALESCE(output_tokens, tokens_output, 0)), 0) as tokens, "
            "COALESCE(SUM(COALESCE(cost_usd, estimated_cost_usd, 0)), 0) as cost "
            "FROM public.ai_usage_tracking "
            f"WHERE {where_clause} "
            "GROUP BY feature ORDER BY cost DESC",
            params,
        )
        by_feature = []
        for row in cursor.fetchall():
            by_feature.append({
                "feature": row["feature"],
                "requests": row["requests"],
                "tokens": row["tokens"],
                "cost": round(float(row["cost"]), 6),
            })

        return {
            "total_requests": totals["total_requests"] if totals else 0,
            "total_tokens": totals["total_tokens"] if totals else 0,
            "total_cost": round(float(totals["total_cost"]), 4) if totals else 0,
            "by_feature": by_feature,
        }
    except Exception as exc:
        logger.error("get_ai_usage_stats error: %s", exc)
        return {"total_requests": 0, "total_tokens": 0, "total_cost": 0, "by_feature": []}
    finally:
        if cursor:
            cursor.close()
        conn.close()


# ============================================
# CREDITS
# ============================================

@router.get("/credits")
async def get_ai_credits(user: ErpUser = Depends(get_current_user)):
    """Get AI credit balance for the current tenant."""
    is_exempt = user.user_type == "super_admin"
    slug = _get_tenant_slug(user)

    if is_exempt or not slug:
        return {
            "balance_usd": 999.99,
            "monthly_limit_usd": 999.99,
            "auto_recharge": False,
            "recharge_amount_usd": 0,
            "is_exempt": True,
        }

    conn = db.get_conn()
    cursor = None
    try:
        cursor = conn.cursor()
        _ensure_ai_tables(cursor, conn)

        # Try tenant_slug first, then entreprise slug fallback
        cursor.execute(
            "SELECT balance_usd, balance_cad, monthly_limit_usd, auto_recharge, recharge_amount_usd "
            "FROM public.ai_prepaid_credits WHERE tenant_slug = %s AND product_type = 'ERP' "
            "AND billing_year = EXTRACT(YEAR FROM CURRENT_DATE)::int "
            "AND billing_month = EXTRACT(MONTH FROM CURRENT_DATE)::int",
            (slug,),
        )
        row = cursor.fetchone()

        if not row:
            # Try via entreprise (entreprises.slug has "tenant_" prefix)
            full_slug = slug if slug.startswith("tenant_") else f"tenant_{slug}"
            cursor.execute(
                "SELECT apc.balance_usd, apc.balance_cad, apc.monthly_limit_usd, "
                "apc.auto_recharge, apc.recharge_amount_usd "
                "FROM public.ai_prepaid_credits apc "
                "JOIN public.entreprises e ON apc.entreprise_id = e.id "
                "WHERE e.slug = %s AND apc.product_type = 'ERP' "
                "AND apc.billing_year = EXTRACT(YEAR FROM CURRENT_DATE)::int "
                "AND apc.billing_month = EXTRACT(MONTH FROM CURRENT_DATE)::int",
                (full_slug,),
            )
            row = cursor.fetchone()

        if not row:
            # Create entry with 0 balance — no free credits
            full_slug = slug if slug.startswith("tenant_") else f"tenant_{slug}"
            try:
                cursor.execute("SELECT id FROM public.entreprises WHERE slug = %s", (full_slug,))
                ent_row = cursor.fetchone()
                ent_id = ent_row["id"] if ent_row else 0
            except Exception:
                ent_id = 0
            from datetime import date
            now = date.today()
            try:
                cursor.execute(
                    "INSERT INTO public.ai_prepaid_credits "
                    "(entreprise_id, tenant_slug, product_type, balance_usd, balance_cad, "
                    " billing_year, billing_month) "
                    "VALUES (%s, %s, 'ERP', 0.00, 0.00, %s, %s) "
                    "ON CONFLICT (entreprise_id, product_type, billing_year, billing_month) DO NOTHING",
                    (ent_id, slug, now.year, now.month),
                )
                conn.commit()
            except Exception:
                conn.rollback()
            return {
                "balance_usd": 0.00,
                "monthly_limit_usd": 999999.99,
                "auto_recharge": False,
                "recharge_amount_usd": 10.00,
                "is_exempt": False,
            }

        balance = float(row.get("balance_usd") or row.get("balance_cad") or 0)
        limit_usd = float(row.get("monthly_limit_usd") or 999999.99)
        auto_rech = row.get("auto_recharge", False)
        rech_amt = float(row.get("recharge_amount_usd") or 10.00)

        # Get monthly usage (tenant_slug only — user_id is not cross-tenant unique)
        cursor.execute(
            "SELECT COALESCE(SUM(COALESCE(cost_usd, estimated_cost_usd, 0)), 0) as monthly_used "
            "FROM public.ai_usage_tracking "
            "WHERE tenant_slug = %s "
            "AND created_at >= date_trunc('month', CURRENT_DATE)",
            (slug,),
        )
        monthly_row = cursor.fetchone()
        monthly_used = float(monthly_row["monthly_used"]) if monthly_row else 0

        return {
            "balance_usd": balance,
            "monthly_limit_usd": limit_usd,
            "monthly_used_usd": round(monthly_used, 4),
            "auto_recharge": auto_rech,
            "recharge_amount_usd": rech_amt,
            "is_exempt": False,
        }
    except Exception as exc:
        logger.error("get_ai_credits error: %s", exc)
        return {
            "balance_usd": 0,
            "monthly_limit_usd": 999999.99,
            "auto_recharge": False,
            "recharge_amount_usd": 10.00,
            "is_exempt": is_exempt,
        }
    finally:
        if cursor:
            cursor.close()
        conn.close()


# ============================================
# QUOTA CHECK
# ============================================

@router.get("/quota")
async def get_ai_quota(user: ErpUser = Depends(get_current_user)):
    """Check if current user/tenant has AI credits.
    Returns allowed status, balance, and monthly usage.
    """
    is_exempt = user.user_type == "super_admin"
    slug = _get_tenant_slug(user)

    if is_exempt:
        return {
            "allowed": True,
            "balance": 999.99,
            "monthly_used": 0,
            "monthly_limit": 999.99,
            "is_exempt": True,
        }

    conn = db.get_conn()
    cursor = None
    try:
        cursor = conn.cursor()
        _ensure_ai_tables(cursor, conn)

        # Get balance
        cursor.execute(
            "SELECT balance_usd, monthly_limit_usd "
            "FROM public.ai_prepaid_credits WHERE tenant_slug = %s "
            "AND product_type = 'ERP' "
            "AND billing_year = EXTRACT(YEAR FROM CURRENT_DATE)::int "
            "AND billing_month = EXTRACT(MONTH FROM CURRENT_DATE)::int",
            (slug,),
        )
        credit_row = cursor.fetchone()
        balance = float(credit_row["balance_usd"]) if credit_row else 0.00
        monthly_limit = float(credit_row["monthly_limit_usd"]) if credit_row else 999999.99

        # Get monthly usage
        cursor.execute(
            "SELECT COALESCE(SUM(COALESCE(cost_usd, estimated_cost_usd, 0)), 0) as monthly_used "
            "FROM public.ai_usage_tracking "
            "WHERE tenant_slug = %s "
            "AND created_at >= date_trunc('month', CURRENT_DATE)",
            (slug,),
        )
        usage_row = cursor.fetchone()
        monthly_used = float(usage_row["monthly_used"]) if usage_row else 0

        # allowed=True always: the only gate at runtime is _check_credits,
        # which lets balance go briefly negative then auto-recharges via
        # Stripe. A restrictive `balance > 0` here would flash the UI into
        # a blocked state while the backend would still accept the call.
        return {
            "allowed": True,
            "balance": round(balance, 4),
            "monthly_used": round(monthly_used, 4),
            "monthly_limit": monthly_limit,
            "is_exempt": False,
        }
    except Exception as exc:
        logger.error("get_ai_quota error: %s", exc)
        return {"allowed": True, "balance": 0, "monthly_used": 0, "monthly_limit": 999999.99, "is_exempt": False}
    finally:
        if cursor:
            cursor.close()
        conn.close()


# ============================================
# DAILY USAGE BREAKDOWN
# ============================================

@router.get("/usage/daily")
async def get_ai_usage_daily(
    user: ErpUser = Depends(get_current_user),
    days: int = Query(30, ge=1, le=90),
):
    """Get daily AI usage breakdown for the last N days."""
    slug = _get_tenant_slug(user)
    conn = db.get_conn()
    cursor = None
    try:
        cursor = conn.cursor()
        _ensure_ai_tables(cursor, conn)

        # Build user filter: super-admin sees all, tenant users see their tenant
        if user.user_type == "super_admin":
            where_clause = "created_at >= CURRENT_DATE - make_interval(days => %s)"
            params = [days]
        else:
            where_clause = "tenant_slug = %s AND created_at >= CURRENT_DATE - make_interval(days => %s)"
            params = [slug, days]

        cursor.execute(
            f"SELECT DATE(created_at) as date, "
            f"COUNT(*) as total_requests, "
            f"COALESCE(SUM(COALESCE(input_tokens, tokens_input, 0)), 0) as total_input_tokens, "
            f"COALESCE(SUM(COALESCE(output_tokens, tokens_output, 0)), 0) as total_output_tokens, "
            f"COALESCE(SUM(COALESCE(cost_usd, estimated_cost_usd, 0)), 0) as total_cost_usd "
            f"FROM public.ai_usage_tracking "
            f"WHERE {where_clause} "
            f"GROUP BY DATE(created_at) "
            f"ORDER BY date DESC",
            params,
        )
        items = []
        for row in cursor.fetchall():
            items.append({
                "date": str(row["date"]),
                "total_requests": row["total_requests"],
                "total_input_tokens": int(row["total_input_tokens"]),
                "total_output_tokens": int(row["total_output_tokens"]),
                "total_cost_usd": round(float(row["total_cost_usd"]), 6),
            })
        return {"items": items, "days": days}
    except Exception as exc:
        logger.error("get_ai_usage_daily error: %s", exc)
        return {"items": [], "days": days}
    finally:
        if cursor:
            cursor.close()
        conn.close()


# ============================================
# MONTHLY USAGE BREAKDOWN
# ============================================

@router.get("/usage/monthly")
async def get_ai_usage_monthly(
    user: ErpUser = Depends(get_current_user),
    months: int = Query(6, ge=1, le=24),
):
    """Get monthly AI usage breakdown for the last N months."""
    slug = _get_tenant_slug(user)
    conn = db.get_conn()
    cursor = None
    try:
        cursor = conn.cursor()
        _ensure_ai_tables(cursor, conn)

        if user.user_type == "super_admin":
            where_clause = "created_at >= CURRENT_DATE - make_interval(months => %s)"
            params = [months]
        else:
            where_clause = "tenant_slug = %s AND created_at >= CURRENT_DATE - make_interval(months => %s)"
            params = [slug, months]

        cursor.execute(
            f"SELECT EXTRACT(YEAR FROM created_at)::int as annee, "
            f"EXTRACT(MONTH FROM created_at)::int as mois, "
            f"feature, "
            f"COUNT(*) as total_requests, "
            f"COALESCE(SUM(cost_usd), 0) as total_cost_usd "
            f"FROM public.ai_usage_tracking "
            f"WHERE {where_clause} "
            f"GROUP BY annee, mois, feature "
            f"ORDER BY annee DESC, mois DESC, total_cost_usd DESC",
            params,
        )
        items = []
        for row in cursor.fetchall():
            items.append({
                "annee": row["annee"],
                "mois": row["mois"],
                "feature": row["feature"],
                "total_requests": row["total_requests"],
                "total_cost_usd": round(float(row["total_cost_usd"]), 6),
            })
        return {"items": items, "months": months}
    except Exception as exc:
        logger.error("get_ai_usage_monthly error: %s", exc)
        return {"items": [], "months": months}
    finally:
        if cursor:
            cursor.close()
        conn.close()


# ============================================
# DOCUMENT ANALYSIS
# ============================================

ALLOWED_DOC_EXTENSIONS = {
    ".pdf", ".docx", ".xlsx", ".csv", ".txt", ".md", ".json", ".html",
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff",
}

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff"}


def _extract_extension(filename: str) -> str:
    """Extract lowercase file extension."""
    return os.path.splitext(filename.lower())[1]


# Protection ZIP bomb sur fichiers Office (xlsx, docx — conteneurs ZIP).
_ZIP_MAX_TOTAL_UNCOMPRESSED = 200 * 1024 * 1024  # 200 MB
_ZIP_MAX_COMPRESSION_RATIO = 100
_ZIP_MAX_ENTRIES = 1000


def _validate_office_zip_safe(content: bytes) -> None:
    """Vérifie qu'un fichier Office (xlsx/docx) n'est pas une ZIP bomb.

    Lève ValueError si suspect.
    """
    import zipfile

    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            entries = zf.infolist()
            if len(entries) > _ZIP_MAX_ENTRIES:
                raise ValueError(f"Trop d'entrées ZIP ({len(entries)})")
            total_uncompressed = sum(e.file_size for e in entries)
            if total_uncompressed > _ZIP_MAX_TOTAL_UNCOMPRESSED:
                raise ValueError("Décompression refusée (taille)")
            total_compressed = sum(e.compress_size for e in entries) or 1
            if total_uncompressed / total_compressed > _ZIP_MAX_COMPRESSION_RATIO:
                raise ValueError("Ratio de compression suspect")
    except zipfile.BadZipFile:
        raise ValueError("ZIP invalide")


def _extract_text_from_pdf(content: bytes) -> tuple[str, int]:
    """Extract text from PDF bytes. Returns (text, page_count)."""
    try:
        # pypdf (anciennement PyPDF2) — versions récentes corrigent de nombreuses CVE.
        try:
            import pypdf as _pypdf
        except ImportError:
            import PyPDF2 as _pypdf  # type: ignore[no-redef]
        reader = _pypdf.PdfReader(io.BytesIO(content))
        pages = len(reader.pages)
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
        return "\n\n".join(text_parts), pages
    except ImportError:
        logger.warning("pypdf not installed - PDF text extraction unavailable")
        return "[Extraction PDF non disponible - pypdf requis]", 0
    except Exception as exc:
        logger.error("PDF extraction error: %s", exc)
        return f"[Erreur extraction PDF: {exc}]", 0


def _extract_text_from_docx(content: bytes) -> tuple[str, int]:
    """Extract text from DOCX bytes. Returns (text, page_count_estimate)."""
    try:
        _validate_office_zip_safe(content)
        import docx
        doc = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        # Rough page estimate: ~3000 chars per page
        pages = max(1, len(text) // 3000)
        return text, pages
    except ValueError as exc:
        logger.warning("DOCX rejected (zip-bomb guard): %s", exc)
        return f"[DOCX refusé: {exc}]", 0
    except ImportError:
        logger.warning("python-docx not installed - DOCX extraction unavailable")
        return "[Extraction DOCX non disponible - python-docx requis]", 0
    except Exception as exc:
        logger.error("DOCX extraction error: %s", exc)
        return f"[Erreur extraction DOCX: {exc}]", 0


def _extract_text_from_xlsx(content: bytes) -> tuple[str, int]:
    """Extract text from XLSX/CSV bytes. Returns (text, sheet_count)."""
    try:
        _validate_office_zip_safe(content)
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        sheets = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(c) if c is not None else "" for c in row)
                if row_text.strip(" |"):
                    rows.append(row_text)
            if rows:
                sheets.append(f"=== Feuille: {sheet_name} ===\n" + "\n".join(rows))
        wb.close()
        return "\n\n".join(sheets), len(wb.sheetnames)
    except ValueError as exc:
        logger.warning("XLSX rejected (zip-bomb guard): %s", exc)
        return f"[XLSX refusé: {exc}]", 0
    except ImportError:
        # Fallback: try reading as CSV text
        try:
            text = content.decode("utf-8", errors="replace")
            return text, 1
        except Exception:
            return "[Extraction XLSX non disponible - openpyxl requis]", 0
    except Exception as exc:
        logger.error("XLSX extraction error: %s", exc)
        return f"[Erreur extraction XLSX: {exc}]", 0


def _detect_media_type_from_bytes(content: bytes) -> Optional[str]:
    """Detect MIME type from magic bytes. Returns None if unknown."""
    if not content:
        return None
    if content[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if content[:3] == b"\xff\xd8\xff":
        return "image/jpeg"
    if content[:6] in (b"GIF87a", b"GIF89a"):
        return "image/gif"
    if len(content) >= 12 and content[:4] == b"RIFF" and content[8:12] == b"WEBP":
        return "image/webp"
    if content[:2] == b"BM":
        return "image/bmp"
    return None


def _resize_image_base64(content: bytes, max_dim: int = 1568) -> tuple[str, str]:
    """Resize image and return (base64, media_type) tuple.

    Le chemin nominal (PIL dispo) reencode TOUJOURS en JPEG — le media_type
    retourne doit alors etre "image/jpeg" pour matcher le payload, sinon
    l'API Anthropic rejette avec 400 "image was specified using the image/png
    media type, but the image appears to be a image/jpeg image" (cas vu
    quand un upload .png etait en realite un JPEG renomme, ou quand PIL
    reencodait un PNG en JPEG).

    En fallback (ImportError ou erreur PIL), on retourne les bytes bruts et
    on detecte le media_type via les magic bytes. Si les bytes ne matchent
    AUCUN type image connu, on raise une HTTPException 400 plutot que
    d'envoyer un media_type fictif a l'API Anthropic (qui rejetterait avec
    un 400 generique difficile a debugger).
    """
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(content))
        if img.width > max_dim or img.height > max_dim:
            ratio = min(max_dim / img.width, max_dim / img.height)
            new_size = (int(img.width * ratio), int(img.height * ratio))
            img = img.resize(new_size, Image.LANCZOS)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        return base64.b64encode(buf.getvalue()).decode("utf-8"), "image/jpeg"
    except ImportError:
        detected = _detect_media_type_from_bytes(content)
        if not detected:
            raise HTTPException(
                status_code=400,
                detail="Format image non reconnu. Formats supportes: PNG, JPEG, GIF, WebP, BMP.",
            )
        return base64.b64encode(content).decode("utf-8"), detected
    except Exception as exc:
        logger.error("Image resize error: %s", exc)
        detected = _detect_media_type_from_bytes(content)
        if not detected:
            raise HTTPException(
                status_code=400,
                detail="Format image invalide ou corrompu. Formats supportes: PNG, JPEG, GIF, WebP, BMP.",
            )
        return base64.b64encode(content).decode("utf-8"), detected


def _detect_media_type(ext: str) -> str:
    """Map file extension to MIME media type (fallback lorsque les bytes ne sont pas disponibles)."""
    mapping = {
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".png": "image/png", ".gif": "image/gif",
        ".bmp": "image/bmp", ".webp": "image/webp",
        ".tiff": "image/tiff",
    }
    return mapping.get(ext, "image/jpeg")


@router.post("/analyze-document")
async def analyze_document(
    file: UploadFile = File(...),
    prompt: Optional[str] = Form(None),
    user: ErpUser = Depends(get_current_user),
):
    """Analyze an uploaded document using Claude AI."""
    if not _anthropic_client:
        raise HTTPException(status_code=503, detail="Service IA non disponible")

    allowed, error_msg = check_ai_guard(user)
    if not allowed:
        raise HTTPException(status_code=403, detail=error_msg or "Acces IA refuse")

    credits_ok, balance = _check_credits(user)
    if not credits_ok:
        raise HTTPException(status_code=402, detail="Credits IA epuises. Veuillez recharger votre solde.")

    filename = file.filename or "document"
    ext = _extract_extension(filename)
    if ext not in ALLOWED_DOC_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Type de fichier non supporte: {ext}. Formats acceptes: {', '.join(sorted(ALLOWED_DOC_EXTENSIONS))}",
        )

    content = await file.read()
    if len(content) > 50 * 1024 * 1024:  # 50 MB limit
        raise HTTPException(status_code=413, detail="Fichier trop volumineux (max 50 Mo)")

    pages = 1
    document_type = "text"
    messages_content = []

    # Determine extraction strategy based on file type
    if ext in IMAGE_EXTENSIONS:
        document_type = "image"
        img_b64, media_type = _resize_image_base64(content)
        messages_content.append({
            "type": "image",
            "source": {"type": "base64", "media_type": media_type, "data": img_b64},
        })
        user_text = prompt or "Analyse ce document image en detail. Identifie le type de document, son contenu et les informations cles."
        messages_content.append({"type": "text", "text": user_text})
    else:
        # Text-based extraction
        if ext == ".pdf":
            document_type = "pdf"
            extracted, pages = _extract_text_from_pdf(content)
        elif ext == ".docx":
            document_type = "docx"
            extracted, pages = _extract_text_from_docx(content)
        elif ext in (".xlsx",):
            document_type = "xlsx"
            extracted, pages = _extract_text_from_xlsx(content)
        elif ext == ".csv":
            document_type = "csv"
            extracted = content.decode("utf-8", errors="replace")
            pages = max(1, extracted.count("\n") // 50)
        else:
            # TXT, MD, JSON, HTML — read directly
            document_type = ext.lstrip(".")
            extracted = content.decode("utf-8", errors="replace")
            pages = max(1, len(extracted) // 3000)

        # Truncate if too long (Claude context window consideration)
        if len(extracted) > 100000:
            extracted = extracted[:100000] + "\n\n[... document tronque a 100 000 caracteres ...]"

        user_text = prompt or "Analyse ce document en detail. Identifie le type de document, son contenu et les informations cles pour la construction."
        messages_content.append({
            "type": "text",
            "text": f"Document: {filename}\nType: {document_type}\n\n--- CONTENU DU DOCUMENT ---\n{extracted}\n--- FIN DU DOCUMENT ---\n\n{user_text}",
        })

    system_prompt = (
        f"{_today_prompt_line()}\n\n"
        "Tu es un expert en construction au Quebec avec 40 ans d'experience. "
        "Tu analyses des documents lies a la construction: devis, plans, contrats, "
        "specifications techniques, bons de commande, factures, rapports d'inspection, etc. "
        "Tu reponds en francais avec precision et references aux normes RBQ, CCQ et CNB quand applicable. "
        "Structure ta reponse avec des sections claires."
    )

    try:
        start = time_module.time()
        response = _call_claude(
            model=AI_MODEL,
            max_tokens=AI_MAX_TOKENS,
            system=system_prompt,
            messages=[{"role": "user", "content": messages_content}],
        )
        elapsed = time_module.time() - start

        analysis = response.content[0].text if response.content else ""
        tokens_in = response.usage.input_tokens
        tokens_out = response.usage.output_tokens
        total_tokens = tokens_in + tokens_out
        duration_ms = int(elapsed * 1000)
        cost = (tokens_in * 0.003 + tokens_out * 0.015) / 1000 * 1.30  # Sonnet 4.6 pricing $3/$15 per MTok + 30% markup

        track_ai_usage(user, "analyze_document", tokens_in, tokens_out, cost, duration_ms, success=True)
        _deduct_credits(user, cost)

        return {
            "analysis": analysis,
            "document_type": document_type,
            "pages": pages,
            "tokens_used": total_tokens,
            "cost_usd": round(cost, 6),
            "elapsed_seconds": round(elapsed, 2),
        }
    except HTTPException:
        raise
    except anthropic.APIError as exc:
        logger.error("analyze_document API error: %s", exc)
        status = getattr(exc, "status_code", 0)
        exc_str = str(exc).lower()
        if status == 413 or (status == 400 and "too_large" in exc_str):
            raise HTTPException(
                status_code=413,
                detail="Document trop volumineux pour l'analyse IA. Essayez un fichier plus petit ou un PDF avec moins de pages.",
            )
        if status == 529 or "overloaded" in exc_str:
            raise HTTPException(status_code=503, detail="Le service IA est temporairement surcharge. Veuillez reessayer dans quelques secondes.")
        if status == 429 or "rate_limit" in exc_str:
            raise HTTPException(status_code=429, detail="Trop de demandes vers l'IA. Attendez quelques instants avant de reessayer.")
        raise HTTPException(status_code=503, detail="Service IA temporairement indisponible. Veuillez reessayer dans quelques instants.")
    except Exception as exc:
        logger.error("analyze_document error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de l'analyse du document")


# ============================================
# PLAN ANALYSIS (Vision)
# ============================================

@router.post("/analyze-plan")
async def analyze_plan(
    files: List[UploadFile] = File(...),
    user: ErpUser = Depends(get_current_user),
):
    """Analyze construction plan images using Claude Vision."""
    if not _anthropic_client:
        raise HTTPException(status_code=503, detail="Service IA non disponible")

    allowed, error_msg = check_ai_guard(user)
    if not allowed:
        raise HTTPException(status_code=403, detail=error_msg or "Acces IA refuse")

    credits_ok, balance = _check_credits(user)
    if not credits_ok:
        raise HTTPException(status_code=402, detail="Credits IA epuises. Veuillez recharger votre solde.")

    if not files or len(files) == 0:
        raise HTTPException(status_code=400, detail="Aucun fichier fourni")

    if len(files) > 10:
        raise HTTPException(status_code=400, detail="Maximum 10 images par analyse")

    messages_content = []

    for f in files:
        file_content = await f.read()
        if len(file_content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail=f"Fichier {f.filename} trop volumineux (max 50 Mo)")

        ext = _extract_extension(f.filename or "image.jpg")
        if ext not in IMAGE_EXTENSIONS and ext != ".pdf":
            raise HTTPException(
                status_code=400,
                detail=f"Fichier {f.filename}: seuls les formats image (JPG, PNG) et PDF sont acceptes pour l'analyse de plans.",
            )

        if ext == ".pdf":
            # For PDF plans: extract first page as image if possible, otherwise use text
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(stream=file_content, filetype="pdf")
                for page_num in range(min(doc.page_count, 5)):  # Max 5 pages
                    page = doc[page_num]
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_bytes = pix.tobytes("jpeg")
                    img_b64, media_type = _resize_image_base64(img_bytes, max_dim=1568)
                    messages_content.append({
                        "type": "image",
                        "source": {"type": "base64", "media_type": media_type, "data": img_b64},
                    })
                doc.close()
            except ImportError:
                # Fallback: extract text from PDF
                text, _ = _extract_text_from_pdf(file_content)
                messages_content.append({
                    "type": "text",
                    "text": f"[Plan PDF - texte extrait de {f.filename}]\n{text}",
                })
            except HTTPException:
                # _resize_image_base64 peut raise 400 explicit pour bytes invalides.
                # Ne PAS catcher ici — propager au caller pour preserver le code 400.
                raise
            except Exception as exc:
                logger.warning("PDF plan image extraction failed: %s", exc)
                text, _ = _extract_text_from_pdf(file_content)
                messages_content.append({
                    "type": "text",
                    "text": f"[Plan PDF - texte extrait de {f.filename}]\n{text}",
                })
        else:
            # Image file
            img_b64, media_type = _resize_image_base64(file_content, max_dim=1568)
            messages_content.append({
                "type": "image",
                "source": {"type": "base64", "media_type": media_type, "data": img_b64},
            })

    # Add the analysis instruction
    messages_content.append({
        "type": "text",
        "text": (
            "Analyse ces plans de construction en detail. Pour chaque plan, identifie:\n"
            "1. Type de plan (architectural, structural, mecanique, electrique, plomberie, etc.)\n"
            "2. Dimensions principales (longueur, largeur, hauteur, superficies)\n"
            "3. Materiaux visibles ou specifies\n"
            "4. Quantites estimees pour chaque materiau\n"
            "5. Notes importantes, specifications particulieres\n"
            "6. Conformite potentielle aux normes RBQ/CNB\n\n"
            "Structure ta reponse en sections claires avec des estimations chiffrees."
        ),
    })

    system_prompt = (
        f"{_today_prompt_line()}\n\n"
        "Tu es un estimateur de construction avec 40 ans d'experience au Quebec. "
        "Tu analyses des plans de construction et tu identifies avec precision: "
        "le type de plan, les dimensions, les materiaux, les quantites estimees, "
        "et les specifications techniques. Tu references les normes RBQ, CCQ et CNB. "
        "Tu donnes des estimations realistes basees sur les prix du marche au Quebec."
    )

    try:
        start = time_module.time()
        response = _call_claude(
            model=AI_MODEL,
            max_tokens=AI_MAX_TOKENS,
            system=system_prompt,
            messages=[{"role": "user", "content": messages_content}],
        )
        elapsed = time_module.time() - start

        analysis_text = response.content[0].text if response.content else ""
        tokens_in = response.usage.input_tokens
        tokens_out = response.usage.output_tokens
        total_tokens = tokens_in + tokens_out
        duration_ms = int(elapsed * 1000)
        cost = (tokens_in * 0.003 + tokens_out * 0.015) / 1000 * 1.30  # Sonnet 4.6 pricing $3/$15 per MTok + 30% markup

        track_ai_usage(user, "analyze_plan", tokens_in, tokens_out, cost, duration_ms, success=True)
        _deduct_credits(user, cost)

        # Parse structured data from analysis if possible
        plan_type = "Plan de construction"
        if "architectural" in analysis_text.lower():
            plan_type = "Plan architectural"
        elif "structural" in analysis_text.lower() or "structure" in analysis_text.lower():
            plan_type = "Plan structural"
        elif "electrique" in analysis_text.lower():
            plan_type = "Plan electrique"
        elif "mecanique" in analysis_text.lower() or "cvac" in analysis_text.lower():
            plan_type = "Plan mecanique/CVAC"
        elif "plomberie" in analysis_text.lower():
            plan_type = "Plan plomberie"

        return {
            "plan_type": plan_type,
            "analysis": analysis_text,
            "dimensions": {},
            "materials": [],
            "files_analyzed": len(files),
            "tokens_used": total_tokens,
            "cost_usd": round(cost, 6),
            "elapsed_seconds": round(elapsed, 2),
        }
    except HTTPException:
        raise
    except anthropic.APIError as exc:
        logger.error("analyze_plan API error: %s", exc)
        status = getattr(exc, "status_code", 0)
        exc_str = str(exc).lower()
        if status == 413 or (status == 400 and "too_large" in exc_str):
            raise HTTPException(
                status_code=413,
                detail="Plan trop volumineux pour l'analyse IA. Essayez avec moins de fichiers ou des fichiers plus petits.",
            )
        if status == 529 or "overloaded" in exc_str:
            raise HTTPException(status_code=503, detail="Le service IA est temporairement surcharge. Veuillez reessayer dans quelques secondes.")
        if status == 429 or "rate_limit" in exc_str:
            raise HTTPException(status_code=429, detail="Trop de demandes vers l'IA. Attendez quelques instants avant de reessayer.")
        raise HTTPException(status_code=503, detail="Service IA temporairement indisponible. Veuillez reessayer dans quelques instants.")
    except Exception as exc:
        logger.error("analyze_plan error: %s", exc)
        raise HTTPException(status_code=500, detail="Erreur lors de l'analyse du plan")
